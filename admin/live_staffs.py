from flask import render_template, request, jsonify, Response
from bson import ObjectId
from datetime import datetime
import json
import json as _cjson
import json as _cjson2
import base64
import csv
import io
import re
import re as _re
import os
import threading

from database import db
from . import admin_bp
from admin.views import admin_required


# ── Helpers ──────────────────────────────────────────────────────────

# ── Google Cloud Storage helpers ──────────────────────────────────────

def _gcs_client():
    """Return an authenticated GCS client.
    Priority:
      1. GCS_CREDENTIALS_JSON env var (full JSON as string) — recommended
      2. GCS_KEY_FILE env var (path to JSON key file)
      3. Application Default Credentials (ADC) — if on Google Cloud VM
    """
    from google.cloud import storage as _gcs
    import json as _json

    creds_json = os.environ.get('GCS_CREDENTIALS_JSON', '').strip()
    if creds_json:
        from google.oauth2 import service_account
        info  = _json.loads(creds_json)
        creds = service_account.Credentials.from_service_account_info(
            info,
            scopes=["https://www.googleapis.com/auth/cloud-platform"]
        )
        return _gcs.Client(credentials=creds, project=info.get('project_id'))

    key_path = os.environ.get('GCS_KEY_FILE', '').strip()
    if key_path and os.path.exists(key_path):
        return _gcs.Client.from_service_account_json(key_path)

    # Fallback — Application Default Credentials (works on GCE/Cloud Run)
    return _gcs.Client()


def _gcs_bucket():
    return _gcs_client().bucket(os.environ.get('GCS_BUCKET_NAME', ''))


def _gcs_upload(blob_name, data_bytes, content_type='application/octet-stream'):
    """Upload bytes to GCS and return the blob name."""
    bucket = _gcs_bucket()
    blob   = bucket.blob(blob_name)
    blob.upload_from_string(data_bytes, content_type=content_type)
    return blob_name


def _gcs_download(blob_name):
    """Download bytes from GCS blob."""
    bucket = _gcs_bucket()
    blob   = bucket.blob(blob_name)
    return blob.download_as_bytes()


def _gcs_signed_url(blob_name, expiry_minutes=60):
    """
    Generate a signed URL for a GCS blob (time-limited, no public access needed).
    Requires the service account to have roles/iam.serviceAccountTokenCreator.
    Falls back to a direct Flask download URL if signing fails.
    """
    import datetime as _dt
    try:
        bucket = _gcs_bucket()
        blob   = bucket.blob(blob_name)
        url    = blob.generate_signed_url(
            expiration=_dt.timedelta(minutes=expiry_minutes),
            method='GET',
            version='v4',
        )
        return url
    except Exception:
        # Return None — caller will use internal download route instead
        return None



# ── Doc webhook collection ────────────────────────────────────────────

def _doc_webhook_col():
    return db.doc_webhook


# ── DOCX → PDF converter ──────────────────────────────────────────────

def _docx_to_pdf_bytes(docx_bytes):
    """
    Convert DOCX bytes to PDF bytes.
    Uses LibreOffice (soffice) via subprocess — available on the server.
    Falls back to docx2pdf if soffice not found.
    """
    import subprocess, tempfile, pathlib

    with tempfile.TemporaryDirectory() as tmp:
        docx_path = pathlib.Path(tmp) / 'input.docx'
        pdf_path  = pathlib.Path(tmp) / 'input.pdf'
        docx_path.write_bytes(docx_bytes)

        # Try LibreOffice first (preferred)
        for soffice in ['soffice', 'libreoffice',
                        '/usr/bin/soffice', '/usr/bin/libreoffice']:
            try:
                result = subprocess.run(
                    [soffice, '--headless', '--convert-to', 'pdf',
                     '--outdir', tmp, str(docx_path)],
                    capture_output=True, timeout=60
                )
                if result.returncode == 0 and pdf_path.exists():
                    return pdf_path.read_bytes()
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue

        raise RuntimeError(
            'LibreOffice not found. Install with: sudo apt install libreoffice'
        )


# ── Background HSE document upload ───────────────────────────────────

HSE_DOC_TYPES = {
    'cv':          'hse_cv',
    'interview':   'interview_notes',
    'appform':     'application_form',
}


def _resolve_xn_staff_id(staff_mongo_id, email):
    """
    Get the XN Portal staff_id to use in the HSE document upload API.

    Priority:
      1. live_staffs.staff_id field (already stored) — use directly
      2. live_staffs.xn_staff_id field (previously fetched) — use directly
      3. Call XN Portal user-document-list API with email → get data.id
         → store in live_staffs.xn_staff_id for future use

    Returns the resolved staff_id string, or None if not found.
    """
    import requests as _req

    col = _staffs_col()

    # 1. Check live_staffs.staff_id field first (primary source)
    doc = col.find_one(
        {"_id": ObjectId(staff_mongo_id)},
        {"staff_id": 1, "xn_staff_id": 1, "email": 1}
    )
    if doc and doc.get('staff_id'):
        return str(doc['staff_id'])

    # 2. Check previously resolved xn_staff_id
    if doc and doc.get('xn_staff_id'):
        return str(doc['xn_staff_id'])

    # 3. Fetch from XN Portal API
    if not email:
        return None

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')

    if not base_url:
        return None

    try:
        resp = _req.post(
            f"{base_url}/ai/recruitments/user-document-list",
            json={"email": email},
            headers={
                "Api-Key":       api_key,
                "X-App-Country": app_country,
                "Content-Type":  "application/json",
                "Accept":        "application/json",
            },
            timeout=30,
        )
        # Handle 405 — retry with GET
        if resp.status_code == 405:
            resp = _req.get(
                f"{base_url}/ai/recruitments/user-document-list",
                params={"email": email},
                headers={
                    "Api-Key":       api_key,
                    "X-App-Country": app_country,
                    "Accept":        "application/json",
                },
                timeout=30,
            )

        data     = resp.json()
        api_data = data.get('data') or {}

        # Handle both list and dict response
        if isinstance(api_data, list):
            api_data = api_data[0] if api_data else {}

        xn_id = str(api_data.get('id', '')).strip()
        if xn_id:
            # Store as xn_staff_id and also as staff_id for future use
            col.update_one(
                {"_id": ObjectId(staff_mongo_id)},
                {"$set": {
                    "xn_staff_id": xn_id,
                    "staff_id":    xn_id,
                }}
            )
            return xn_id

    except Exception:
        pass

    return None


def _push_hse_document_background(staff_id_str, doc_type_key,
                                   docx_bytes, staff_name='',
                                   mongo_id=None, email=''):
    """
    Fire-and-forget background task.
    Converts DOCX → PDF, POSTs to HSE document upload API,
    saves full response to doc_webhook collection.
    """
    def _run():
        import requests as _req

        base_url    = os.environ.get('DOC_BASE_URL', '').rstrip('/')
        api_key     = os.environ.get('DOC_API_KEY', '')
        app_country = os.environ.get('XN_APP_COUNTRY', '')

        if not base_url:
            _doc_webhook_col().insert_one({
                "staff_id":     staff_id_str,
                "staff_name":   staff_name,
                "doc_type":     doc_type_key,
                "status":       "error",
                "error":        "DOC_BASE_URL not set in environment",
                "triggered_at": datetime.utcnow(),
            })
            return

        # Resolve xn_staff_id — fetch from XN Portal if missing
        resolved_staff_id = staff_id_str
        if mongo_id and email:
            xn_id = _resolve_xn_staff_id(mongo_id, email)
            if xn_id:
                resolved_staff_id = xn_id

        endpoint      = f"{base_url}/api/admin/staff/hse-document-upload"
        hse_type      = HSE_DOC_TYPES.get(doc_type_key, doc_type_key)
        pdf_bytes     = None
        convert_error = None

        # Convert DOCX → PDF
        try:
            pdf_bytes = _docx_to_pdf_bytes(docx_bytes)
        except Exception as e:
            convert_error = str(e)

        if not pdf_bytes:
            _doc_webhook_col().insert_one({
                "staff_id":     staff_id_str,
                "staff_name":   staff_name,
                "doc_type":     doc_type_key,
                "hse_type":     hse_type,
                "status":       "error",
                "error":        f"PDF conversion failed: {convert_error}",
                "triggered_at": datetime.utcnow(),
            })
            return

        # POST to HSE API
        try:
            resp = _req.post(
                endpoint,
                data={
                    "staff_id":          resolved_staff_id,
                    "hse_document_type": hse_type,
                },
                files={
                    "file": (f"{hse_type}.pdf", pdf_bytes, "application/pdf"),
                },
                headers={
                    "Api-Key":       api_key,
                    "X-App-Country": app_country,
                    "Accept":        "application/json",
                },
                timeout=60,
            )

            # Try to parse JSON response
            try:
                resp_json = resp.json()
            except Exception:
                resp_json = {"raw": resp.text[:500]}

            _doc_webhook_col().insert_one({
                "staff_id":           staff_id_str,
                "xn_staff_id":        resolved_staff_id,
                "staff_name":         staff_name,
                "doc_type":           doc_type_key,
                "hse_type":           hse_type,
                "endpoint":           endpoint,
                "status":             "success" if resp.status_code < 300 else "api_error",
                "http_status":        resp.status_code,
                "response":           resp_json,
                "triggered_at":       datetime.utcnow(),
            })

        except Exception as e:
            _doc_webhook_col().insert_one({
                "staff_id":     staff_id_str,
                "staff_name":   staff_name,
                "doc_type":     doc_type_key,
                "hse_type":     hse_type,
                "endpoint":     endpoint,
                "status":       "error",
                "error":        str(e),
                "triggered_at": datetime.utcnow(),
            })

    # Launch in background thread — does not block the HTTP response
    t = threading.Thread(target=_run, daemon=True)
    t.start()


def _staffs_col():
    return db.live_staffs


def _serialize(doc):
    """Recursively convert ObjectId / datetime to JSON-safe types."""
    if isinstance(doc, list):
        return [_serialize(i) for i in doc]
    if isinstance(doc, dict):
        return {k: _serialize(v) for k, v in doc.items()}
    if isinstance(doc, ObjectId):
        return str(doc)
    if isinstance(doc, datetime):
        return doc.isoformat()
    return doc


def _get_all(search, page, per_page):
    match = {}
    if search:
        pattern = re.compile(re.escape(search), re.IGNORECASE)
        match = {"$or": [
            {"section_1_personal_details.full_name": pattern},
            {"email": pattern},
            {"employee_code": pattern},
            {"section_1_personal_details.nationality": pattern},
            {"user_type": pattern},
        ]}
    col   = _staffs_col()
    total = col.count_documents(match)

    # Aggregation: add a sort key so records with extracted_cv come first
    pipeline = []
    if match:
        pipeline.append({"$match": match})
    pipeline += [
        {"$addFields": {
            "_cv_filled": {
                "$cond": {
                    "if": {
                        "$and": [
                            {"$ifNull": ["$extracted_cv", False]},
                            {"$ne": ["$extracted_cv", ""]},
                            {"$ne": ["$extracted_cv", None]},
                        ]
                    },
                    "then": 0,   # has CV — sort first
                    "else": 1    # no CV — sort after
                }
            }
        }},
        {"$sort": {
            "_cv_filled": 1,
            "section_1_personal_details.full_name": 1
        }},
        {"$skip":  (page - 1) * per_page},
        {"$limit": per_page},
    ]

    items = list(col.aggregate(pipeline))
    # Serialize BEFORE passing to template so tojson never sees ObjectId
    return [_serialize(doc) for doc in items], total


def _parse_json_content(content):
    """
    Handle all JSON variants that can come from the export pipeline:
      1. Standard JSON array  [ {...}, ... ]
      2. Standard JSON object { "records": [ ... ] }
      3. Bare fragment        "records": [ ... ]   ← missing outer braces
      4. JSONL                {...}\n{...}\n
      5. Concatenated objects {...}{...}
    """
    content = content.strip()

    # 1 & 2 — standard JSON
    try:
        raw = json.loads(content)
        return raw if isinstance(raw, list) else raw.get('records', [raw])
    except json.JSONDecodeError:
        pass

    # 3 — bare fragment (missing outer braces)
    try:
        raw = json.loads('{' + content + '}')
        if 'records' in raw:
            return raw['records']
    except json.JSONDecodeError:
        pass

    # 4 — JSONL
    try:
        lines = [l for l in content.splitlines() if l.strip()]
        records = [json.loads(l) for l in lines]
        if records:
            return records
    except json.JSONDecodeError:
        pass

    # 5 — concatenated objects
    try:
        records = []
        decoder = json.JSONDecoder()
        idx = 0
        while idx < len(content):
            while idx < len(content) and content[idx] in ' \t\r\n,':
                idx += 1
            if idx >= len(content):
                break
            obj, end = decoder.raw_decode(content, idx)
            records.append(obj)
            idx = end
        if records:
            return records
    except json.JSONDecodeError:
        pass

    raise ValueError("Could not parse JSON — unrecognised format.")


# ── Routes ───────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs')
@admin_required
def live_staffs():
    page     = int(request.args.get('page', 1))
    search   = request.args.get('search', '').strip()
    per_page = 20

    items, total = _get_all(search, page, per_page)

    return render_template(
        'admin/live_staffs.html',
        staffs=items,
        page=page,
        total=total,
        per_page=per_page,
        search=search,
    )


@admin_bp.route('/live-staffs/get')
@admin_required
def live_staff_get():
    """Return a single staff record as JSON — used by view/edit modals."""
    staff_id = (request.args.get('id') or '').strip()
    if not staff_id:
        return jsonify({"success": False, "error": "Missing id"}), 400
    try:
        doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not doc:
            return jsonify({"success": False, "error": "Record not found"}), 404
        return jsonify({"success": True, "record": _serialize(doc)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/edit-extracted-cv', methods=['POST'])
@admin_required
def live_staff_edit_extracted_cv():
    """
    Save edited extracted_cv text for a staff member.
    Also resets experience_analysed_at so AI analysis re-runs on next cron call.

    POST /admin/live-staffs/edit-extracted-cv
    Body: {"staff_id": "...", "extracted_cv": "..."}
    """
    data      = request.get_json() or {}
    staff_id  = (data.get('staff_id') or '').strip()
    new_text  = data.get('extracted_cv', '')
    user_type = (data.get('user_type') or '').strip()

    if not staff_id:
        return jsonify({"success": False, "error": "Missing staff_id"}), 400
    try:
        update_fields = {
            "extracted_cv":               new_text,
            "extracted_cv_at":            datetime.utcnow(),
            "extracted_cv_edited":        True,
            # Reset AI analysis so it re-runs with updated text/type
            "experience_analysed_at":     None,
            "experience_list_at":         None,
            "experience_list_processing": False,
        }
        if user_type:
            update_fields["user_type"] = user_type

        result = _staffs_col().update_one(
            {"_id": ObjectId(staff_id)},
            {"$set": update_fields}
        )
        if result.matched_count == 0:
            return jsonify({"success": False, "error": "Staff not found"}), 404
        return jsonify({"success": True, "message": "Changes saved successfully"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@admin_bp.route('/live-staffs/points')
@admin_required
def live_staff_points():
    """
    Get points for a staff member by email address.

    GET /admin/live-staffs/points?email=someone@example.com

    Returns:
      {
        "success": true,
        "email": "someone@example.com",
        "name": "Sherin Augustine",
        "points": 8
      }
    """
    email = (request.args.get('email') or '').strip().lower()
    if not email:
        return jsonify({"success": False, "error": "Missing email parameter"}), 400
    try:
        doc = _staffs_col().find_one(
            {"email": email},
            {"points": 1, "section_1_personal_details": 1, "email": 1}
        )
        if not doc:
            return jsonify({"success": False, "error": f"No staff found with email: {email}"}), 404

        s1   = doc.get('section_1_personal_details') or {}
        name = _v(s1.get('full_name') or '')

        return jsonify({
            "success": True,
            "email":   email,
            "name":    name,
            "points":  doc.get('points'),
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@admin_bp.route('/live-staffs/experience', methods=['GET', 'POST'])
def live_staff_experience():
    """
    Analyse a staff member's extracted_cv using Gemini AI and return
    total years and months of experience as digits.

    Accepts X-API-Key header (no session required) OR admin session cookie.

    GET  /admin/live-staffs/experience?email=someone@example.com
    POST /admin/live-staffs/experience  body: {"staff_id": "68abc..."} or {"email": "..."}

    Response:
      {
        "success": true,
        "staff_name": "Jane Smith",
        "email": "jane@example.com",
        "years": 4,
        "months": 6,
        "total_months": 54,
        "summary": "4 years and 6 months",
        "source": "extracted_cv"   // or "section_5" if CV text unavailable
      }
    """
    # Accept either API key or admin session
    api_key_provided = request.headers.get('X-API-Key', '').strip()
    if api_key_provided:
        ok, err = _validate_api_key()
        if not ok:
            return jsonify({"success": False, "error": err}), 401
    else:
        from admin.views import admin_required as _admin_required
        from flask import session, redirect, url_for
        if not session.get('admin_logged_in'):
            return jsonify({"success": False,
                            "error": "Unauthorised — provide X-API-Key header or login"}), 401

    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set on server"}), 500

    # ── Resolve staff record ──────────────────────────────────────────
    if request.method == 'POST':
        data     = request.get_json(silent=True) or {}
        staff_id = (data.get('staff_id') or '').strip()
        email    = (data.get('email') or '').strip().lower()
    else:
        staff_id = (request.args.get('id') or '').strip()
        email    = (request.args.get('email') or '').strip().lower()

    if staff_id:
        # 1. Try live_staffs.staff_id field (XN Portal ID)
        doc = _staffs_col().find_one({"staff_id": staff_id})
        # 2. Try live_staffs.xn_staff_id field
        if not doc:
            doc = _staffs_col().find_one({"xn_staff_id": staff_id})
        # 3. Try MongoDB _id
        if not doc:
            try:
                doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
            except Exception:
                pass
    elif email:
        # Search email in top-level field and section_1 sub-field
        doc = _staffs_col().find_one({"$or": [
            {"email": email},
            {"section_1_personal_details.email_address": email},
        ]})
    else:
        return jsonify({"success": False,
                        "error": "Provide staff_id or email (query param or JSON body)"}), 400

    if not doc:
        return jsonify({"success": False, "error": "Staff record not found"}), 404

    s1        = doc.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email_out = _v(doc.get('email') or s1.get('email_address') or email)
    user_type = _v(doc.get('user_type') or '')

    # ── Return cached result if already analysed ─────────────────────
    if doc.get('experience_analysed_at') and doc.get('experience_years') is not None:
        years        = int(doc.get('experience_years') or 0)
        months       = int(doc.get('experience_months') or 0)
        total_months = int(doc.get('experience_total_months') or 0)
        if total_months == 0:
            total_months = years * 12 + months
        parts   = []
        if years:  parts.append(f"{years} year{'s' if years != 1 else ''}")
        if months: parts.append(f"{months} month{'s' if months != 1 else ''}")
        summary = ' and '.join(parts) if parts else '0 months'
        return jsonify({
            "success":      True,
            "staff_name":   full_name,
            "email":        email_out,
            "years":        years,
            "months":       months,
            "total_months": total_months,
            "summary":      summary,
            "note":         _v(doc.get('experience_note') or ''),
            "source":       _v(doc.get('experience_source') or 'cached'),
            "cached":       True,
        })

    # ── Determine text to analyse ─────────────────────────────────────
    extracted_cv = _v(doc.get('extracted_cv') or '')
    has_cv_text  = (
        extracted_cv and
        not extracted_cv.startswith('[') and
        extracted_cv != 'No doc found'
    )

    # Fallback: use section_5 employment history total_experience string
    s5           = doc.get('section_5_employment_history') or {}
    total_exp_db = _v(s5.get('total_experience') or '')

    if not has_cv_text and not total_exp_db:
        return jsonify({
            "success":      True,
            "staff_name":   full_name,
            "email":        email_out,
            "years":        None,
            "months":       None,
            "total_months": None,
            "summary":      "No experience data available",
            "source":       "none",
        })

    # ── Call Gemini ───────────────────────────────────────────────────
    try:
        from google import genai as google_genai

        # Determine role filter based on user_type
        ut_lower = (user_type or '').lower()
        if 'nurse' in ut_lower or 'nursing' in ut_lower:
            role_rule = (
                "IMPORTANT — Count ONLY nursing-related work experience, regardless of the country where it was gained. "
                "This includes: Registered Nurse, Staff Nurse, Clinical Nurse, ICU Nurse, Theatre Nurse, "
                "Community Nurse, Mental Health Nurse, Nursing Home Nurse, or any role with Nurse or Nursing in the title. "
                "DO NOT count non-nursing roles such as healthcare assistant, carer, support worker, admin, or any other role."
            )
        elif 'hca' in ut_lower or 'healthcare assistant' in ut_lower or 'health care assistant' in ut_lower:
            role_rule = (
                "IMPORTANT — Count ONLY Healthcare Assistant (HCA) work experience, regardless of the country where it was gained. "
                "This includes: Healthcare Assistant, HCA, Care Assistant, Care Worker, Support Worker in a clinical/care setting, "
                "or any role with Healthcare Assistant or HCA in the title. "
                "DO NOT count nursing roles (Registered Nurse, Staff Nurse, etc.) or non-care roles such as admin, retail, or hospitality."
            )
        else:
            role_rule = (
                "Count only direct healthcare or care-related work experience, regardless of country. "
                "Exclude non-healthcare roles such as admin, retail, hospitality, or general support roles "
                "unless they are clearly in a clinical or care setting."
            )

        if has_cv_text:
            source = 'extracted_cv'
            prompt = f"""You are a professional CV analyser specialising in Irish healthcare staffing.

Candidate role: {user_type}

Read the CV text below and calculate the candidate's TOTAL relevant work experience.

{role_rule}

Calculation Rules:
- Only count roles that match the role filter above.
- Experience gained in ANY country counts — not just Ireland.
- If a matching role has no end date, assume it is still ongoing (use today's date to calculate).
- If two matching roles overlap in time, count the overlapping period only once.
- Ignore all non-matching roles entirely — do not add them.
- Return ONLY a JSON object with these exact keys — nothing else, no markdown, no explanation:
  {{"years": <integer>, "months": <integer 0-11>, "total_months": <integer>, "note": "<one sentence summary of which roles were counted and why>"}}

CV TEXT:
{extracted_cv[:10000]}
"""
        else:
            source = 'section_5'
            prompt = f"""You are a professional CV analyser specialising in Irish healthcare staffing.

Candidate role: {user_type}

The candidate's total experience is described as: "{total_exp_db}"

{role_rule}

Extract the relevant years and months from this description, applying the role filter above.
Return ONLY a JSON object — nothing else, no markdown:
{{"years": <integer>, "months": <integer 0-11>, "total_months": <integer>, "note": "<one sentence summary>"}}
"""

        client   = google_genai.Client(api_key=gemini_key)
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        raw = (response.text or '').strip()

        # Strip markdown code fences if Gemini wraps it
        raw = _re.sub(r'^```(?:json)?\s*', '', raw, flags=_re.MULTILINE)
        raw = _re.sub(r'```\s*$', '', raw, flags=_re.MULTILINE).strip()

        result = _json.loads(raw)

        years        = int(result.get('years', 0) or 0)
        months       = int(result.get('months', 0) or 0)
        total_months = int(result.get('total_months', 0) or 0)
        note         = _v(result.get('note', ''))

        # Recalculate total_months as sanity check
        if total_months == 0:
            total_months = years * 12 + months

        # Build human-readable summary
        parts = []
        if years:  parts.append(f"{years} year{'s' if years != 1 else ''}")
        if months: parts.append(f"{months} month{'s' if months != 1 else ''}")
        summary = ' and '.join(parts) if parts else '0 months'

        # Save to DB
        _staffs_col().update_one(
            {"_id": doc['_id']},
            {"$set": {
                "experience_years":        years,
                "experience_months":       months,
                "experience_total_months": total_months,
                "experience_note":         note,
                "experience_source":       source,
                "experience_analysed_at":  datetime.utcnow(),
            }}
        )

        return jsonify({
            "success":      True,
            "staff_name":   full_name,
            "email":        email_out,
            "years":        years,
            "months":       months,
            "total_months": total_months,
            "summary":      summary,
            "note":         note,
            "source":       source,
            "cached":       False,
        })

    except _cjson.JSONDecodeError:
        return jsonify({
            "success": False,
            "error":   "Gemini returned non-JSON response",
            "raw":     raw[:300],
        }), 500
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/api/experience', methods=['POST'])
def api_experience():
    """
    External API — analyse experience from extracted_cv via Gemini.

    Headers:
      X-API-Key: <LIVE_STAFF_API_KEY>
      Content-Type: application/json

    Body:
      {"staff_id": "68abc123..."} or {"email": "jane@example.com"}
    """
    ok, err = _validate_api_key()
    if not ok:
        return jsonify({"success": False, "error": err}), 401

    data     = request.get_json(silent=True) or {}
    staff_id = (data.get('staff_id') or '').strip()
    email    = (data.get('email') or '').strip().lower()

    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set on server"}), 500

    if staff_id:
        # 1. Try live_staffs.staff_id field (XN Portal ID)
        doc = _staffs_col().find_one({"staff_id": staff_id})
        # 2. Try live_staffs.xn_staff_id field
        if not doc:
            doc = _staffs_col().find_one({"xn_staff_id": staff_id})
        # 3. Try MongoDB _id
        if not doc:
            try:
                doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
            except Exception:
                pass
    elif email:
        # Search email in top-level field and section_1 sub-field
        doc = _staffs_col().find_one({"$or": [
            {"email": email},
            {"section_1_personal_details.email_address": email},
        ]})
    else:
        return jsonify({"success": False,
                        "error": "Provide staff_id or email in JSON body"}), 400

    if not doc:
        return jsonify({"success": False, "error": "Staff record not found"}), 404

    s1        = doc.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email_out = _v(doc.get('email') or s1.get('email_address') or email)
    user_type = _v(doc.get('user_type') or '')

    extracted_cv = _v(doc.get('extracted_cv') or '')
    has_cv_text  = (
        extracted_cv and
        not extracted_cv.startswith('[') and
        extracted_cv != 'No doc found'
    )

    s5           = doc.get('section_5_employment_history') or {}
    total_exp_db = _v(s5.get('total_experience') or '')

    if not has_cv_text and not total_exp_db:
        return jsonify({
            "success":      True,
            "staff_name":   full_name,
            "email":        email_out,
            "years":        None,
            "months":       None,
            "total_months": None,
            "summary":      "No experience data available",
            "source":       "none",
        })

    # ── Return cached result if already analysed ──────────────────────
    if doc.get('experience_analysed_at') and doc.get('experience_years') is not None:
        years        = int(doc.get('experience_years') or 0)
        months       = int(doc.get('experience_months') or 0)
        total_months = int(doc.get('experience_total_months') or 0)
        if total_months == 0:
            total_months = years * 12 + months
        parts   = []
        if years:  parts.append(f"{years} year{'s' if years != 1 else ''}")
        if months: parts.append(f"{months} month{'s' if months != 1 else ''}")
        summary = ' and '.join(parts) if parts else '0 months'
        return jsonify({
            "success":      True,
            "staff_name":   full_name,
            "email":        email_out,
            "years":        years,
            "months":       months,
            "total_months": total_months,
            "summary":      summary,
            "note":         _v(doc.get('experience_note') or ''),
            "source":       _v(doc.get('experience_source') or 'cached'),
            "cached":       True,
        })

    try:
        from google import genai as google_genai

        # Determine role filter based on user_type
        ut_lower = (user_type or '').lower()
        if 'nurse' in ut_lower or 'nursing' in ut_lower:
            role_rule = (
                "IMPORTANT — Count ONLY nursing-related work experience, regardless of the country where it was gained. "
                "This includes: Registered Nurse, Staff Nurse, Clinical Nurse, ICU Nurse, Theatre Nurse, "
                "Community Nurse, Mental Health Nurse, Nursing Home Nurse, or any role with Nurse or Nursing in the title. "
                "DO NOT count non-nursing roles such as healthcare assistant, carer, support worker, admin, or any other role."
            )
        elif 'hca' in ut_lower or 'healthcare assistant' in ut_lower or 'health care assistant' in ut_lower:
            role_rule = (
                "IMPORTANT — Count ONLY Healthcare Assistant (HCA) work experience, regardless of the country where it was gained. "
                "This includes: Healthcare Assistant, HCA, Care Assistant, Care Worker, Support Worker in a clinical/care setting, "
                "or any role with Healthcare Assistant or HCA in the title. "
                "DO NOT count nursing roles (Registered Nurse, Staff Nurse, etc.) or non-care roles such as admin, retail, or hospitality."
            )
        else:
            role_rule = (
                "Count only direct healthcare or care-related work experience, regardless of country. "
                "Exclude non-healthcare roles such as admin, retail, hospitality, or general support roles "
                "unless they are clearly in a clinical or care setting."
            )

        if has_cv_text:
            source = 'extracted_cv'
            prompt = f"""You are a professional CV analyser specialising in Irish healthcare staffing.

Candidate role: {user_type}

Read the CV text below and calculate the candidate's TOTAL relevant work experience.

{role_rule}

Calculation Rules:
- Only count roles that match the role filter above.
- Experience gained in ANY country counts — not just Ireland.
- If a matching role has no end date, assume it is still ongoing (use today's date to calculate).
- If two matching roles overlap in time, count the overlapping period only once.
- Ignore all non-matching roles entirely — do not add them.
- Return ONLY a JSON object with these exact keys — nothing else, no markdown, no explanation:
  {{"years": <integer>, "months": <integer 0-11>, "total_months": <integer>, "note": "<one sentence summary of which roles were counted>"}}

CV TEXT:
{extracted_cv[:10000]}
"""
        else:
            source = 'section_5'
            prompt = f"""You are a professional CV analyser specialising in Irish healthcare staffing.

Candidate role: {user_type}

The candidate's total experience is described as: "{total_exp_db}"

{role_rule}

Extract the relevant years and months from this description, applying the role filter above.
Return ONLY a JSON object — nothing else, no markdown:
{{"years": <integer>, "months": <integer 0-11>, "total_months": <integer>, "note": "<one sentence summary>"}}
"""


        client   = google_genai.Client(api_key=gemini_key)
        response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
        raw      = (response.text or '').strip()
        raw      = _re.sub(r'^```(?:json)?\s*', '', raw, flags=_re.MULTILINE)
        raw      = _re.sub(r'```\s*$', '', raw, flags=_re.MULTILINE).strip()

        result       = _json.loads(raw)
        years        = int(result.get('years', 0) or 0)
        months       = int(result.get('months', 0) or 0)
        total_months = int(result.get('total_months', 0) or 0)
        note         = _v(result.get('note', ''))
        if total_months == 0:
            total_months = years * 12 + months

        parts   = []
        if years:  parts.append(f"{years} year{'s' if years != 1 else ''}")
        if months: parts.append(f"{months} month{'s' if months != 1 else ''}")
        summary = ' and '.join(parts) if parts else '0 months'

        _staffs_col().update_one(
            {"_id": doc['_id']},
            {"$set": {
                "experience_years":        years,
                "experience_months":       months,
                "experience_total_months": total_months,
                "experience_analysed_at":  datetime.utcnow(),
            }}
        )

        return jsonify({
            "success":      True,
            "staff_name":   full_name,
            "email":        email_out,
            "years":        years,
            "months":       months,
            "total_months": total_months,
            "summary":      summary,
            "note":         note,
            "source":       source,
        })

    except _cjson.JSONDecodeError:
        return jsonify({"success": False, "error": "Gemini returned non-JSON response", "raw": raw[:300]}), 500
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/import-last-point-scale', methods=['POST'])
@admin_required
def live_staff_import_last_point_scale():
    """
    Seed last_point_scale for a staff member by email.
    Also resets points_checked so the cron will re-check them.

    POST /admin/live-staffs/import-last-point-scale
    Body: {"email": "...", "last_point_scale": 9}
    """
    data     = request.get_json() or {}
    email    = (data.get('email') or '').strip().lower()
    last_ps  = data.get('last_point_scale')

    if not email:
        return jsonify({"success": False, "error": "Missing email"}), 400
    if last_ps is None:
        return jsonify({"success": False, "error": "Missing last_point_scale"}), 400

    try:
        result = _staffs_col().update_one(
            {"email": email},
            {"$set": {
                "last_point_scale": last_ps,
                "points_checked":   False,   # reset so cron re-checks
            }}
        )
        if result.matched_count == 0:
            return jsonify({"success": False,
                            "error": f"No staff found with email: {email}"}), 404
        return jsonify({
            "success":  True,
            "email":    email,
            "last_point_scale": last_ps,
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── AI CV Collection helper ───────────────────────────────────────────

def _ai_cvs_col():
    return db.live_staff_ai_cvs

def _ai_interviews_col():
    return db.live_staff_ai_interviews


def _ai_appforms_col():
    return db.live_staff_ai_appforms


@admin_bp.route('/live-staffs/ai-appform/reset-all', methods=['POST'])
@admin_required
def live_staff_ai_appform_reset_all():
    """
    Delete all saved application forms from MongoDB so the cron
    will regenerate them with the latest template.
    POST /admin/live-staffs/ai-appform/reset-all
    """
    try:
        result = _ai_appforms_col().delete_many({})
        return jsonify({
            "success": True,
            "deleted": result.deleted_count,
            "message": f"Cleared {result.deleted_count} application forms — cron will regenerate them.",
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Generate AI Application Form ──────────────────────────────────────

@admin_bp.route('/live-staffs/ai-appform/generate', methods=['POST'])
@admin_required
def live_staff_ai_appform_generate():
    """
    Build a filled Xpress Health Application Form .docx for a staff member.
    Uses actual DB data — no AI hallucination, no Gemini needed.
    Uploads to GCS and saves metadata to MongoDB.
    """
    data     = request.get_json()
    staff_id = (data.get('staff_id') or '').strip()
    if not staff_id:
        return jsonify({"success": False, "error": "Missing staff_id"}), 400

    try:
        doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not doc:
            return jsonify({"success": False, "error": "Staff record not found"}), 404

        s1        = doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff')
        email     = _v(doc.get('email'))
        emp_code  = _v(doc.get('employee_code') or '')

        # Download signature from GCS if available
        signature_bytes = None
        sig_blob = _v(doc.get('signature_gcs_blob') or '')
        if sig_blob:
            try:
                signature_bytes = _gcs_download(sig_blob)
            except Exception:
                signature_bytes = None

        docx_bytes = _build_appform_docx(doc, signature_bytes=signature_bytes)
        safe_name = full_name.replace(' ', '_').replace('/', '_')
        filename  = f"AppForm_{safe_name}.docx"
        gcs_blob  = f"appforms/{filename}"

        _gcs_upload(
            gcs_blob, docx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        col      = _ai_appforms_col()
        existing = col.find_one({"staff_id": str(doc['_id'])})
        rec = {
            "staff_id":      str(doc['_id']),
            "staff_name":    full_name,
            "employee_code": emp_code,
            "filename":      filename,
            "gcs_blob":      gcs_blob,
            "generated_at":  datetime.utcnow(),
        }
        if existing:
            col.update_one({"_id": existing["_id"]}, {"$set": rec})
            rec_id = str(existing["_id"])
        else:
            result = col.insert_one(rec)
            rec_id = str(result.inserted_id)

        # Background: push to HSE document API
        _push_hse_document_background(
            staff_id_str=staff_id,
            doc_type_key='appform',
            docx_bytes=docx_bytes,
            staff_name=full_name,
            mongo_id=staff_id,
            email=email,
        )

        return jsonify({
            "success":      True,
            "appform_id":   rec_id,
            "staff_name":   full_name,
            "message":      f"Application form generated for {full_name}",
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/ai-appform/download/<appform_id>')
@admin_required
def live_staff_ai_appform_download(appform_id):
    """Serve saved application form DOCX from Google Cloud Storage."""
    try:
        rec = _ai_appforms_col().find_one({"_id": ObjectId(appform_id)})
        if not rec:
            return "Application form not found", 404
        gcs_blob = rec.get('gcs_blob', '')
        if not gcs_blob:
            return "File not found in storage — please regenerate", 404
        name       = (rec.get('staff_name') or 'staff').replace(' ', '_')
        docx_bytes = _gcs_download(gcs_blob)
        return Response(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f'attachment; filename="AppForm_{name}.docx"'}
        )
    except Exception as e:
        return str(e), 500


@admin_bp.route('/live-staffs/ai-appform/saved/<staff_id>')
@admin_required
def live_staff_ai_appform_saved(staff_id):
    """Check if a saved application form exists for this staff member."""
    try:
        rec = _ai_appforms_col().find_one({"staff_id": staff_id})
        if not rec:
            return jsonify({"success": True, "found": False})
        return jsonify({
            "success":      True,
            "found":        True,
            "appform_id":   str(rec["_id"]),
            "generated_at": rec["generated_at"].strftime("%d %b %Y %H:%M"),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/ai-appform/upload/<staff_id>', methods=['POST'])
@admin_required
def live_staff_ai_appform_upload(staff_id):
    """Replace the saved application form with an edited .docx upload."""
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file provided"}), 400
    file = request.files['file']
    if not file.filename.lower().endswith('.docx'):
        return jsonify({"success": False, "error": "Only .docx files are accepted"}), 400
    try:
        col = _ai_appforms_col()
        rec = col.find_one({"staff_id": staff_id})
        if not rec:
            return jsonify({"success": False,
                            "error": "No saved application form found for this staff member"}), 404
        gcs_blob = rec.get('gcs_blob', '')
        if not gcs_blob:
            doc2 = _staffs_col().find_one({"_id": ObjectId(staff_id)})
            s1   = (doc2.get('section_1_personal_details') or {}) if doc2 else {}
            name = _v(s1.get('full_name') or 'staff').replace(' ', '_').replace('/', '_')
            gcs_blob = f"appforms/AppForm_{name}.docx"
        data_bytes = file.read()
        _gcs_upload(
            gcs_blob, data_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        col.update_one(
            {"_id": rec["_id"]},
            {"$set": {
                "gcs_blob":      gcs_blob,
                "filename":      os.path.basename(gcs_blob),
                "last_uploaded": datetime.utcnow(),
                "uploaded_by":   "admin",
            }}
        )
        # Background: push updated form to HSE document API
        try:
            _push_hse_document_background(
                staff_id_str=staff_id,
                doc_type_key='appform',
                docx_bytes=data_bytes,
                staff_name=(rec.get('staff_name') or ''),
                mongo_id=staff_id,
                email=_v((_staffs_col().find_one({"_id": ObjectId(staff_id)}) or {}).get('email') or ''),
            )
        except Exception:
            pass

        return jsonify({
            "success":  True,
            "message":  "Application form replaced successfully",
            "filename": os.path.basename(gcs_blob),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Build Application Form DOCX ───────────────────────────────────────

def _build_appform_docx(doc, signature_bytes=None):
    """
    Build a filled Xpress Health Application Form matching the uploaded template exactly.
    Sections: Personal Details, Identity Verification, Qualification and Experience,
              Declaration, Signature.
    All data pulled from live_staffs MongoDB document — no hallucination.
    If signature_bytes is provided, embeds the actual signature image.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io as _io

    BLACK  = RGBColor(0x00, 0x00, 0x00)
    NAVY   = RGBColor(0x1B, 0x3A, 0x6B)
    GREEN  = RGBColor(0x2E, 0x9E, 0x44)
    GRAY   = RGBColor(0x55, 0x55, 0x55)
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

    # ── Extract data from doc ──────────────────────────────────────────
    s1   = doc.get('section_1_personal_details') or {}
    s2   = doc.get('section_2_identity_verification') or {}
    s3   = doc.get('section_3_professional_registration') or {}
    s4   = doc.get('section_4_qualifications') or {}
    s5   = doc.get('section_5_employment_history') or {}
    visa = s1.get('work_permit_visa_status') or {}
    docs = s2.get('documents_submitted') or {}

    full_name   = _v(s1.get('full_name'))
    email       = _v(doc.get('email'))
    user_type   = _v(doc.get('user_type'))
    address     = _v(s1.get('address'))
    eircode     = _v(s1.get('eircode_postcode'))
    mobile      = _v(s1.get('mobile_number'))
    pps         = _v(s1.get('pps_number'))
    perm_work   = _v(visa.get('permission_to_work'))
    total_exp   = _v(s5.get('total_experience'))
    nmbi_pin    = _v(s3.get('registration_number_pin'))
    divisions   = s3.get('divisions_registered_in') or []

    is_nurse  = 'nurse' in user_type.lower() if user_type else bool(divisions or nmbi_pin)
    is_hca    = not is_nurse

    d = DocxDocument()
    for sec in d.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    normal = d.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    def add_border_bottom(para, color='1B3A6B', size=12):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    str(size))
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color)
        pBdr.append(bot)
        pPr.append(pBdr)

    def add_doc_title():
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(14)
        r = p.add_run('Xpress Health Application Form')
        r.bold = True
        r.font.size = Pt(18)
        r.font.name = 'Calibri'
        r.font.color.rgb = NAVY

    def add_section_heading(title):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = 'Calibri'
        r.font.color.rgb = NAVY
        add_border_bottom(p, color='2E9E44', size=8)

    def add_field(label, value):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(label + '  ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.color.rgb = NAVY
        r2 = p.add_run(value or '')
        r2.font.name = 'Calibri'
        r2.font.color.rgb = BLACK

    def _add_tick_run(para, checked):
        """
        Add a tick/box using Unicode ballot box characters with DejaVu Sans font.
        U+2611 ☑ = ballot box with check (ticked)
        U+2610 ☐ = ballot box (empty)
        DejaVu Sans is bundled with LibreOffice and renders these correctly in PDF.
        """
        r = para.add_run('☑' if checked else '☐')
        r.font.name = 'DejaVu Sans'
        r.font.size = Pt(12)
        r.font.color.rgb = BLACK
        return r

    def add_checkbox_line(label, checked):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _add_tick_run(p, checked)
        r2 = p.add_run(f'  {label}')
        r2.font.name = 'Calibri'
        r2.font.size = Pt(11)
        r2.font.color.rgb = BLACK

    def add_checkbox_row(items):
        """Multiple checkboxes on one line: [(label, checked), ...]"""
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        for i, (label, checked) in enumerate(items):
            _add_tick_run(p, checked)
            run = p.add_run(f'  {label}')
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK
            if i < len(items) - 1:
                spacer = p.add_run('       ')
                spacer.font.name = 'Calibri'

    def add_spacer(pts=6):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = Pt(pts)

    # ── Build document ─────────────────────────────────────────────────

    add_doc_title()

    # ── Section 1: Personal Details ────────────────────────────────────
    add_section_heading('Personal Details')
    add_field('Full Name:', full_name)
    add_field('Email:', email)
    add_field('Role:', user_type)
    add_field('Address:', address)
    add_field('Eircode/Postcode:', eircode)
    add_field('Mobile Number:', mobile)
    add_field('Work Permit / Visa Status:', perm_work or ('Yes' if visa.get('visa_type') else ''))

    # ── Section 2: Identity Verification ─────────────────────────────
    add_section_heading('Identity Verification')
    p_id = d.add_paragraph()
    p_id.paragraph_format.space_before = Pt(4)
    p_id.paragraph_format.space_after  = Pt(4)
    r_id = p_id.add_run('ID Proof:')
    r_id.bold = True
    r_id.font.name = 'Calibri'
    r_id.font.color.rgb = NAVY

    add_checkbox_row([
        ('Passport',            bool(docs.get('passport'))),
        ('Birth Certificate',   bool(docs.get('birth_certificate'))),
        ('Driving Licence',     bool(docs.get('driving_licence'))),
        ('Proof of Address',    bool(docs.get('proof_of_address'))),
    ])

    # ── Section 3: Qualification and Experience ───────────────────────
    add_section_heading('Qualification and Experience')

    add_checkbox_row([
        ('Nurse (NMBI):', is_nurse),
        ('HCA (QQI L5):', is_hca),
    ])

    add_spacer(4)
    add_field('Total years of experience:', total_exp)

    # NMBI PIN if nurse
    if is_nurse and nmbi_pin:
        add_field('NMBI PIN:', nmbi_pin)
    if divisions:
        add_field('Divisions:', ', '.join(divisions))

    # ── Declaration ────────────────────────────────────────────────────
    add_section_heading('Declaration')
    p_decl = d.add_paragraph()
    p_decl.paragraph_format.space_before = Pt(4)
    p_decl.paragraph_format.space_after  = Pt(12)
    r_decl = p_decl.add_run(
        'I declare that the information provided in this application form is true and accurate '
        'to the best of my knowledge. I understand that any false or misleading information may '
        'result in the withdrawal of an offer of employment or termination of employment.'
    )
    r_decl.font.name = 'Calibri'
    r_decl.font.size = Pt(10)
    r_decl.font.color.rgb = GRAY
    r_decl.italic = True

    # ── Signature ──────────────────────────────────────────────────────
    p_sig_lbl = d.add_paragraph()
    p_sig_lbl.paragraph_format.space_before = Pt(6)
    p_sig_lbl.paragraph_format.space_after  = Pt(2)
    r_sig = p_sig_lbl.add_run('Applicant Signature:')
    r_sig.bold = True
    r_sig.font.name = 'Calibri'
    r_sig.font.color.rgb = NAVY

    if signature_bytes:
        # Embed the actual signature image
        try:
            import tempfile as _tmp, pathlib as _pl
            with _tmp.NamedTemporaryFile(suffix='.png', delete=False) as tf:
                tf.write(signature_bytes)
                tf_path = tf.name
            from docx.shared import Inches as _Inches
            p_img = d.add_paragraph()
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after  = Pt(4)
            run_img = p_img.add_run()
            run_img.add_picture(tf_path, width=_Inches(2.0))
            import os as _os
            _os.unlink(tf_path)
        except Exception:
            # Fall back to blank line if image embedding fails
            p_blank = d.add_paragraph()
            p_blank.add_run('   _______________________________')
    else:
        p_blank = d.add_paragraph()
        p_blank.paragraph_format.space_before = Pt(0)
        p_blank.add_run('   _______________________________')

    buf = _io.BytesIO()
    d.save(buf)
    return buf.getvalue()




# ── Generate AI CV ────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/ai-cv/generate', methods=['POST'])
@admin_required
def live_staff_ai_cv_generate():
    """Call Gemini to write a personalised CV, render to DOCX, upload to Google Cloud Storage."""
    data     = request.get_json()
    staff_id = (data.get('staff_id') or '').strip()
    if not staff_id:
        return jsonify({"success": False, "error": "Missing staff_id"}), 400

    try:
        doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not doc:
            return jsonify({"success": False, "error": "Staff record not found"}), 404

        s1   = doc.get('section_1_personal_details') or {}
        s3   = doc.get('section_3_professional_registration') or {}
        s4   = doc.get('section_4_qualifications') or {}
        s5   = doc.get('section_5_employment_history') or {}
        s8   = doc.get('section_8_garda_vetting_police_clearance') or {}
        s9   = doc.get('section_9_occupational_health') or {}
        s10  = doc.get('section_10_mandatory_training') or {}
        visa = s1.get('work_permit_visa_status') or {}

        def _vv(val):
            if val is None: return ''
            return str(val).strip()

        full_name   = _vv(s1.get('full_name'))
        user_type   = _vv(doc.get('user_type'))
        address     = _vv(s1.get('address'))
        mobile      = _vv(s1.get('mobile_number'))
        email       = _vv(doc.get('email'))
        dob         = _vv(s1.get('date_of_birth'))
        nationality = _vv(s1.get('nationality'))
        emp_code    = _vv(doc.get('employee_code'))
        total_exp   = _vv(s5.get('total_experience'))
        divisions   = ', '.join(s3.get('divisions_registered_in') or [])
        reg_pin     = _vv(s3.get('registration_number_pin'))
        reg_exp     = _vv(s3.get('registration_expiry_date'))
        nmbi        = 'Yes' if s3.get('nmbi_active_declaration') else 'No'
        visa_type   = _vv(visa.get('visa_type'))
        perm_work   = _vv(visa.get('permission_to_work'))
        garda       = 'Yes' if s8.get('garda_vetting_submitted') else 'No'
        fit         = 'Yes' if s9.get('fit_for_nursing_duties') else 'No'

        qual_lines = []
        for qk in ['nursing_degree', 'postgraduate_qualification', 'other_qualification']:
            q = s4.get(qk) or {}
            if q.get('qualification') or q.get('institution'):
                qual_lines.append(
                    f"  - {_vv(q.get('qualification'))} | "
                    f"{_vv(q.get('institution'))} | "
                    f"{_vv(q.get('year_completed'))}"
                )
        # Also include NMBI/QQI numbers as qualification context
        nmbi_num = _vv(doc.get('nmbi_number') or s3.get('registration_number_pin') or '')
        qqi_num  = _vv(doc.get('qqi_number') or '')
        if nmbi_num and not any('nmbi' in l.lower() or 'registration' in l.lower() for l in qual_lines):
            qual_lines.append(f"  - NMBI Registration PIN: {nmbi_num}")
        if qqi_num and not any('qqi' in l.lower() for l in qual_lines):
            qual_lines.append(f"  - QQI Level 5 Certificate No: {qqi_num}")

        entries = [e for e in (s5.get('entries') or [])
                   if e.get('employer') or e.get('position')]
        exp_lines = []
        for e in entries:
            exp_lines.append(
                f"  - {_vv(e.get('position'))} at {_vv(e.get('employer'))} "
                f"({_vv(e.get('from'))} - {_vv(e.get('to') or 'Present')})"
            )

        TLABELS = {
            'manual_handling': 'Manual Handling',
            'cpr_bls': 'CPR / BLS',
            'fire_safety': 'Fire Safety',
            'infection_prevention_control': 'Infection Prevention & Control',
            'hand_hygiene': 'Hand Hygiene',
            'safeguarding': 'Safeguarding',
            'children_first': 'Children First',
            'cyber_security': 'Cyber Security',
            'dignity_at_work': 'Dignity at Work',
            'open_disclosure': 'Open Disclosure',
            'mapa_pmav': 'MAPA / PMAV',
        }
        certs = [label for k, label in TLABELS.items() if s10.get(k)][:6]

        # Pull extracted_cv text from DB if available
        extracted_cv = _v(doc.get('extracted_cv') or '')
        has_extracted_cv = (
            extracted_cv and
            not extracted_cv.startswith('[') and
            extracted_cv != '[no CV document found]'
        )

        data_summary = f"""
Candidate: {full_name}
Role / User Type: {user_type}
Employee Code: {emp_code}
Address: {address}
Mobile: {mobile}
Email: {email}
Nationality: {nationality}
Total Experience: {total_exp}
Divisions / Speciality: {divisions}
Registration PIN: {reg_pin}
Registration Expiry: {reg_exp}
NMBI Active Declaration: {nmbi}
Permission to Work: {perm_work}
Visa / Stamp Type: {visa_type}
Garda Vetted: {garda}
Fit for Nursing Duties: {fit}

Qualifications:
{chr(10).join(qual_lines) if qual_lines else '  None recorded'}

Employment History (from profile):
{chr(10).join(exp_lines) if exp_lines else '  None recorded'}

Training & Certifications (on file):
{chr(10).join('  - ' + c for c in certs) if certs else '  None recorded'}
""".strip()

        # Build extracted CV section for prompt
        extracted_cv_section = f"""

EXTRACTED CV TEXT (use this as the PRIMARY source for PROFESSIONAL EXPERIENCE, TRAINING & CERTIFICATIONS and KEY SKILLS — copy the actual duties, skills and certifications directly from this text, preserving the candidate's own words):
{extracted_cv[:8000]}
""" if has_extracted_cv else ""

        prompt = f"""You are an expert professional CV writer specialising in Irish healthcare staffing.

STRICT RULE — NO HALLUCINATION:
You MUST use ONLY the exact facts provided in the CANDIDATE DATA and EXTRACTED CV TEXT below.
Do NOT invent, assume, or add any information that is not explicitly stated.
If a field is empty or says "None recorded", skip it.

SECTION SOURCE RULES:
- PERSONAL DETAILS, PROFESSIONAL PROFILE: use CANDIDATE DATA.
- EDUCATION & QUALIFICATIONS: MANDATORY — ALWAYS include this section. Use qualifications from CANDIDATE DATA.
  If Qualifications section says "None recorded", infer from their role:
    * Nurse → "Bachelor of Nursing Science (or equivalent) | [University based on nationality/location] | [estimated year based on experience]"
    * Healthcare Assistant → "QQI Level 5 in Healthcare Support | [College of Further Education] | [estimated year]"
  Include NMBI PIN or QQI certificate number if provided in CANDIDATE DATA.
- PROFESSIONAL EXPERIENCE: {"Extract directly from EXTRACTED CV TEXT — copy the actual job titles, employers, dates, and duties word-for-word as written by the candidate. Do not rewrite or invent." if has_extracted_cv else "Use Employment History from CANDIDATE DATA. Write 5-6 appropriate duties per role."}
- TRAINING & CERTIFICATIONS: {"Extract directly from EXTRACTED CV TEXT — list only the certifications the candidate actually listed." if has_extracted_cv else "Use Training & Certifications from CANDIDATE DATA only."}
- KEY SKILLS: {"Extract directly from EXTRACTED CV TEXT — use the candidate's own skills list exactly as written." if has_extracted_cv else "Write 8-10 bullet points from their role and certifications in the data."}

Structure the CV exactly as follows (EXACT section headings in UPPERCASE on their own line):

PERSONAL DETAILS
PROFESSIONAL PROFILE
EDUCATION & QUALIFICATIONS
PROFESSIONAL EXPERIENCE
TRAINING & CERTIFICATIONS
KEY SKILLS
ADDITIONAL INFORMATION

Section format rules:
- PERSONAL DETAILS: "Label: Value" per line. Skip blank fields.
- PROFESSIONAL PROFILE: 2 paragraphs, FIRST PERSON ("I am", "I have", "I bring"). Genuine personal statement.
- EDUCATION & QUALIFICATIONS: ALWAYS write at least one entry. Format: Qualification Name | Institution | Year
  Include registration numbers (NMBI PIN, QQI Certificate No) as a separate line if available.
- PROFESSIONAL EXPERIENCE: One block per role:
    Job Title: [title]
    Employer: [employer]
    Dates: [from] - [to]
    Duties:
    - [duty]
- TRAINING & CERTIFICATIONS: Bullet list of certifications only.
- KEY SKILLS: 8-10 bullet points.
- ADDITIONAL INFORMATION: Include ONLY these two lines, nothing else:
Driving Licence: No
Own Transport: No

---
CANDIDATE DATA:
{data_summary}
{extracted_cv_section}
---

Output CV text only. No preamble, no explanation, no markdown symbols like ** or ##. Plain text with section headings and dash bullet points.
"""

        gemini_key = os.environ.get('GEMINI_API_KEY', '')
        if not gemini_key:
            return jsonify({"success": False, "error": "GEMINI_API_KEY not set in environment"}), 500

        from google import genai as google_genai
        client   = google_genai.Client(api_key=gemini_key)
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        cv_text = response.text.strip()

        docx_bytes = _build_ai_cv_docx(doc, cv_text)

        safe_name   = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        cv_filename = f"{safe_name}.docx"
        gcs_blob    = f"cv/{cv_filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        col      = _ai_cvs_col()
        existing = col.find_one({"staff_id": str(doc['_id'])})
        ai_doc = {
            "staff_id":      str(doc['_id']),
            "staff_name":    full_name,
            "employee_code": emp_code,
            "cv_text":       cv_text,
            "cv_filename":   cv_filename,
            "gcs_blob":      gcs_blob,
            "generated_at":  datetime.utcnow(),
        }
        if existing:
            col.update_one({"_id": existing["_id"]}, {"$set": ai_doc})
            ai_id = str(existing["_id"])
        else:
            result = col.insert_one(ai_doc)
            ai_id  = str(result.inserted_id)

        # Background: push to HSE document API
        _push_hse_document_background(
            staff_id_str=staff_id,
            doc_type_key='cv',
            docx_bytes=docx_bytes,
            staff_name=full_name,
            mongo_id=staff_id,
            email=email,
        )

        return jsonify({
            "success":     True,
            "ai_cv_id":    ai_id,
            "cv_filename": cv_filename,
            "staff_name":  full_name,
            "message":     f"AI CV generated for {full_name}"
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/ai-cv/download/<ai_cv_id>')
@admin_required
def live_staff_ai_cv_download(ai_cv_id):
    """Serve the saved AI CV DOCX from Google Cloud Storage."""
    try:
        rec = _ai_cvs_col().find_one({"_id": ObjectId(ai_cv_id)})
        if not rec:
            return "AI CV not found", 404
        gcs_blob = rec.get('gcs_blob', '')
        if not gcs_blob:
            return "CV not found in storage — please regenerate", 404
        name       = (rec.get('staff_name') or 'staff').replace(' ', '_')
        filename   = f"{name}.docx"
        docx_bytes = _gcs_download(gcs_blob)
        return Response(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return str(e), 500


@admin_bp.route('/live-staffs/ai-cv/saved/<staff_id>')
@admin_required
def live_staff_ai_cv_saved(staff_id):
    """Check if a saved AI CV exists for this staff member."""
    try:
        rec = _ai_cvs_col().find_one(
            {"staff_id": staff_id},
            {"cv_text": 0}
        )
        if not rec:
            return jsonify({"success": True, "found": False})
        return jsonify({
            "success":      True,
            "found":        True,
            "ai_cv_id":     str(rec["_id"]),
            "cv_filename":  rec.get("cv_filename", ""),
            "generated_at": rec["generated_at"].strftime("%d %b %Y %H:%M"),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/ai-cv/upload/<staff_id>', methods=['POST'])
@admin_required
def live_staff_ai_cv_upload(staff_id):
    """
    Replace the saved AI CV file with an edited version uploaded by the user.
    Accepts a .docx file upload and overwrites the existing file on disk.
    """
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file provided"}), 400

    file = request.files['file']
    if not file.filename.lower().endswith('.docx'):
        return jsonify({"success": False, "error": "Only .docx files are accepted"}), 400

    try:
        col = _ai_cvs_col()
        rec = col.find_one({"staff_id": staff_id})

        if not rec:
            return jsonify({"success": False, "error": "No saved CV found for this staff member"}), 404

        # Determine GCS blob name
        gcs_blob = rec.get('gcs_blob', '')
        if not gcs_blob:
            doc2 = _staffs_col().find_one({"_id": ObjectId(staff_id)})
            s1   = (doc2.get('section_1_personal_details') or {}) if doc2 else {}
            name = _v(s1.get('full_name') or 'staff').replace(' ', '_').replace('/', '_')
            gcs_blob = f"cv/{name}.docx"

        # Upload to GCS — overwrites existing blob
        data_bytes = file.read()
        _gcs_upload(gcs_blob, data_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        col.update_one(
            {"_id": rec["_id"]},
            {"$set": {
                "gcs_blob":      gcs_blob,
                "cv_filename":   os.path.basename(gcs_blob),
                "last_uploaded": datetime.utcnow(),
                "uploaded_by":   "admin",
            }}
        )

        # Background: push updated CV to HSE document API
        try:
            _push_hse_document_background(
                staff_id_str=staff_id,
                doc_type_key='cv',
                docx_bytes=data_bytes,
                staff_name=(rec.get('staff_name') or ''),
                mongo_id=staff_id,
                email=_v((_staffs_col().find_one({"_id": ObjectId(staff_id)}) or {}).get('email') or ''),
            )
        except Exception:
            pass

        return jsonify({
            "success":  True,
            "message":  "CV replaced successfully with uploaded version",
            "filename": os.path.basename(gcs_blob),
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Generate AI Interview Notes ───────────────────────────────────────

@admin_bp.route('/live-staffs/ai-interview/generate', methods=['POST'])
@admin_required
def live_staff_ai_interview_generate():
    """
    Call Gemini to write realistic interview notes for a staff member
    using the exact structure of the Nurse Interview Template.
    Saves DOCX to Google Cloud Storage and metadata to MongoDB.
    """
    data     = request.get_json()
    staff_id = (data.get('staff_id') or '').strip()
    if not staff_id:
        return jsonify({"success": False, "error": "Missing staff_id"}), 400

    try:
        doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not doc:
            return jsonify({"success": False, "error": "Staff record not found"}), 404

        s1   = doc.get('section_1_personal_details') or {}
        s3   = doc.get('section_3_professional_registration') or {}
        s5   = doc.get('section_5_employment_history') or {}
        s8   = doc.get('section_8_garda_vetting_police_clearance') or {}
        s9   = doc.get('section_9_occupational_health') or {}
        s10  = doc.get('section_10_mandatory_training') or {}
        visa = s1.get('work_permit_visa_status') or {}

        full_name   = _v(s1.get('full_name'))
        email       = _v(doc.get('email'))
        user_type   = _v(doc.get('user_type'))
        address     = _v(s1.get('address'))
        nationality = _v(s1.get('nationality'))
        reg_pin     = _v(s3.get('registration_number_pin'))
        visa_type   = _v(visa.get('visa_type'))
        divisions   = ', '.join(s3.get('divisions_registered_in') or [])
        total_exp   = _v(s5.get('total_experience'))
        entries     = [e for e in (s5.get('entries') or [])
                       if e.get('employer') or e.get('position')]
        nmbi        = 'Yes' if s3.get('nmbi_active_declaration') else 'No'
        garda       = 'Yes' if s8.get('garda_vetting_submitted') else 'No'
        bls         = 'Yes' if s10.get('cpr_bls') else 'No'
        manual      = 'Yes' if s10.get('manual_handling') else 'No'
        fit         = 'Yes' if s9.get('fit_for_nursing_duties') else 'No'

        # Preferred county from address
        county = ''
        if address:
            parts = [p.strip() for p in address.replace(',', ' ').split()]
            for p in parts:
                if p.lower().startswith('co.') or p.lower() == 'county':
                    idx = parts.index(p)
                    if idx + 1 < len(parts):
                        county = parts[idx + 1]
                    break
            if not county:
                county = parts[-1] if parts else ''

        # Build experience summary for prompt
        exp_lines = []
        for e in entries[:5]:
            pos = _v(e.get('position')); emp = _v(e.get('employer'))
            d_from = _v(e.get('from')); d_to = _v(e.get('to'))
            if pos or emp:
                exp_lines.append(
                    f"  - {pos} at {emp} ({d_from} – {d_to or 'Present'})"
                )

        TLABELS = {
            'manual_handling': 'Manual Handling',
            'cpr_bls': 'BLS/CPR',
            'safeguarding': 'Safeguarding',
            'fire_safety': 'Fire Safety',
            'infection_prevention_control': 'Infection Prevention & Control',
        }
        certs = [label for k, label in TLABELS.items() if s10.get(k)]

        data_summary = f"""
Name: {full_name}
Role / User Type: {user_type}
Address / Location: {address}
Nationality: {nationality}
Visa / Stamp Type: {visa_type}
NMBI Registration PIN: {reg_pin}
NMBI Registration Active: {nmbi}
Divisions / Speciality: {divisions}
Total Experience: {total_exp}
Garda Vetted: {garda}
BLS/CPR on file: {bls}
Manual Handling on file: {manual}
Fit for Duties: {fit}

Employment History:
{chr(10).join(exp_lines) if exp_lines else '  None recorded'}

Certifications on file: {', '.join(certs) if certs else 'None recorded'}
""".strip()

        prompt = f"""You are an experienced nursing recruitment consultant at Xpress Health, Ireland.

Using ONLY the verified candidate data below, complete a realistic, professional nurse interview notes template.
Answers must be written as if the candidate themselves just answered each question in a live phone/video interview.
Write naturally — conversational but professional. First person where appropriate ("I have", "I work", "I currently").

STRICT RULES — NO HALLUCINATION:
- Use ONLY the facts provided in CANDIDATE DATA. Do not invent employers, dates, locations, or qualifications.
- If data is missing for a field, write a realistic professional answer appropriate to their role and experience level without inventing specific names.
- Clinical question answers must be clinically appropriate for a {user_type}.
- Assessment scores: pick a random realistic score for each between 3.5 and 5.0 in 0.5 increments (e.g. 3.5/5, 4/5, 4.5/5, 5/5). Vary the three scores — do not give the same score to all three.
- Do NOT add any text outside the template structure below.

Output ONLY the completed template below — no preamble, no explanations, no markdown symbols:

---
Completed {user_type} Interview

Name: [full name]
Location: [county/city from address]
NMBI PIN: [registration pin or N/A]
Visa Status: [visa type]

Experience

1. Tell me about your nursing experience.
[Write a 4–6 sentence answer in first person describing their experience, speciality, and current/most recent role. Use only the data provided.]

2. How many years in Ireland?
[Write a realistic answer based on employment history dates. If Ireland-based work is evident, state it clearly.]

3. Acute, Nursing Home, Community, or Mental Health?
[Based on employment history, state the most relevant care setting.]

Clinical Questions

1. How would you manage a deteriorating patient?
[Write a clinically accurate 4–5 sentence answer appropriate for a {user_type}. Use recognised frameworks (ABCDE, NEWS2, ISBAR) where appropriate.]

2. What would you do if you witnessed a medication error?
[Write a clinically accurate 4–5 sentence answer covering patient safety, reporting, documentation, and prevention.]

Compliance
NMBI Registration: [Yes/No based on data]
BLS/CPR: [Yes/No based on data]
Manual Handling: [Yes/No based on data]
Garda Vetting: [Yes/No based on data]
References: Yes

Availability
Preferred counties: [county from address, or nearest city]
Day/Night/Both: Both
Earliest start date: Immediate

Assessment
Communication: [3.5/5 or 4/5 or 4.5/5 or 5/5 — vary randomly]
Clinical Knowledge: [3.5/5 or 4/5 or 4.5/5 or 5/5 — vary randomly]
Experience: [3.5/5 or 4/5 or 4.5/5 or 5/5 — vary randomly]
Suitable: Yes
---

CANDIDATE DATA (use ONLY this):
{data_summary}
"""

        gemini_key = os.environ.get('GEMINI_API_KEY', '')
        if not gemini_key:
            return jsonify({"success": False,
                            "error": "GEMINI_API_KEY not set"}), 500

        from google import genai as google_genai
        client   = google_genai.Client(api_key=gemini_key)
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        interview_text = response.text.strip()
        # Strip leading/trailing --- if Gemini included them
        interview_text = interview_text.strip('-').strip()

        # Build DOCX
        docx_bytes = _build_interview_docx(doc, interview_text)

        # Upload to Google Cloud Storage
        safe_name    = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        filename     = f"Interview_{safe_name}_{staff_id}.docx"
        gcs_blob  = f"interviews/{filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        # Save metadata to MongoDB
        col      = _ai_interviews_col()
        existing = col.find_one({"staff_id": str(doc['_id'])})
        rec = {
            "staff_id":       str(doc['_id']),
            "staff_name":     full_name,
            "employee_code":  _v(doc.get('employee_code')),
            "interview_text": interview_text,
            "filename":       filename,
            "gcs_blob":       gcs_blob,
            "generated_at":   datetime.utcnow(),
        }
        if existing:
            col.update_one({"_id": existing["_id"]}, {"$set": rec})
            rec_id = str(existing["_id"])
        else:
            result = col.insert_one(rec)
            rec_id = str(result.inserted_id)

        # Background: push to HSE document API
        _push_hse_document_background(
            staff_id_str=staff_id,
            doc_type_key='interview',
            docx_bytes=docx_bytes,
            staff_name=full_name,
            mongo_id=staff_id,
            email=email,
        )

        return jsonify({
            "success":      True,
            "interview_id": rec_id,
            "staff_name":   full_name,
            "message":      f"Interview notes generated for {full_name}"
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/ai-interview/download/<interview_id>')
@admin_required
def live_staff_ai_interview_download(interview_id):
    """Serve saved interview DOCX from Google Cloud Storage."""
    try:
        rec = _ai_interviews_col().find_one({"_id": ObjectId(interview_id)})
        if not rec:
            return "Interview notes not found", 404
        gcs_blob = rec.get('gcs_blob', '')
        if not gcs_blob:
            return "File not found in storage — please regenerate", 404
        name       = (rec.get('staff_name') or 'staff').replace(' ', '_')
        docx_bytes = _gcs_download(gcs_blob)
        return Response(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition":
                     f'attachment; filename="Interview_{name}.docx"'}
        )
    except Exception as e:
        return str(e), 500


@admin_bp.route('/live-staffs/ai-interview/saved/<staff_id>')
@admin_required
def live_staff_ai_interview_saved(staff_id):
    """Check if saved interview notes exist for this staff member."""
    try:
        rec = _ai_interviews_col().find_one(
            {"staff_id": staff_id},
            {"interview_text": 0}
        )
        if not rec:
            return jsonify({"success": True, "found": False})
        return jsonify({
            "success":      True,
            "found":        True,
            "interview_id": str(rec["_id"]),
            "generated_at": rec["generated_at"].strftime("%d %b %Y %H:%M"),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/ai-interview/upload/<staff_id>', methods=['POST'])
@admin_required
def live_staff_ai_interview_upload(staff_id):
    """
    Replace the saved interview notes file with an edited version uploaded by the user.
    Accepts a .docx file and overwrites the existing file on disk.
    """
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file provided"}), 400

    file = request.files['file']
    if not file.filename.lower().endswith('.docx'):
        return jsonify({"success": False, "error": "Only .docx files are accepted"}), 400

    try:
        col = _ai_interviews_col()
        rec = col.find_one({"staff_id": staff_id})

        if not rec:
            return jsonify({"success": False,
                            "error": "No saved interview notes found for this staff member"}), 404

        gcs_blob = rec.get('gcs_blob', '')
        if not gcs_blob:
            doc2 = _staffs_col().find_one({"_id": ObjectId(staff_id)})
            s1   = (doc2.get('section_1_personal_details') or {}) if doc2 else {}
            name = _v(s1.get('full_name') or 'staff').replace(' ', '_').replace('/', '_')
            gcs_blob = f"interviews/Interview_{name}_{staff_id}.docx"

        data_bytes = file.read()
        _gcs_upload(gcs_blob, data_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        col.update_one(
            {"_id": rec["_id"]},
            {"$set": {
                "gcs_blob":      gcs_blob,
                "filename":      os.path.basename(gcs_blob),
                "last_uploaded": datetime.utcnow(),
                "uploaded_by":   "admin",
            }}
        )

        # Background: push updated interview to HSE document API
        try:
            _push_hse_document_background(
                staff_id_str=staff_id,
                doc_type_key='interview',
                docx_bytes=data_bytes,
                staff_name=(rec.get('staff_name') or ''),
                mongo_id=staff_id,
                email=_v((_staffs_col().find_one({"_id": ObjectId(staff_id)}) or {}).get('email') or ''),
            )
        except Exception:
            pass

        return jsonify({
            "success":  True,
            "message":  "Interview notes replaced successfully with uploaded version",
            "filename": os.path.basename(gcs_blob),
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Build Interview Notes PDF ─────────────────────────────────────────



def _build_interview_docx(doc, interview_text):
    """
    Render AI interview notes as a Word doc matching the original
    Completed Nurse Interview PDF design:
    - Plain white background throughout
    - Bold black section headings with a simple bottom border line
    - Bold label + plain value for fields
    - Numbered questions in bold, answers as plain indented paragraphs
    - No coloured boxes, no navy/green fills
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io as _io, re as _re

    BLACK  = RGBColor(0x00, 0x00, 0x00)
    DKGRAY = RGBColor(0x22, 0x22, 0x22)

    s1_d      = doc.get('section_1_personal_details') or {}
    full_name = _v(s1_d.get('full_name')) or 'Candidate'
    user_type = _v(doc.get('user_type')) or 'Nurse'

    d = DocxDocument()

    # ── Margins — match original template ────────────────────────────
    for sec in d.sections:
        sec.top_margin    = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin   = Cm(2.54)
        sec.right_margin  = Cm(2.54)

    # ── Default Normal style ──────────────────────────────────────────
    normal = d.styles['Normal']
    normal.font.name  = 'Calibri'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = BLACK

    # ── Helpers ───────────────────────────────────────────────────────

    def add_para_border_bottom(para):
        """Add a thin black bottom border to a paragraph (section divider)."""
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')       # 0.75pt
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '000000')
        pBdr.append(bot)
        pPr.append(pBdr)

    def add_section_heading(title):
        """Bold heading + bottom border line — plain black on white."""
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = BLACK
        run.font.name = 'Calibri'
        add_para_border_bottom(p)

    def add_field(label, value):
        """Bold label followed by plain value on the same line."""
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(label + ' ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.color.rgb = BLACK
        r2 = p.add_run(value or '')
        r2.font.name = 'Calibri'
        r2.font.color.rgb = BLACK

    def add_numbered_question(text):
        """Bold numbered question."""
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = BLACK

    def add_answer_text(text):
        """Plain answer paragraph — indented slightly."""
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Inches(0.2)
        run = p.add_run(text or '')
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK

    def add_compliance_field(label, value):
        """Bold label + bold value (no colour change — plain black)."""
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(label + ' ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.color.rgb = BLACK
        r2 = p.add_run(value or '')
        r2.bold = True
        r2.font.name = 'Calibri'
        r2.font.color.rgb = BLACK

    def add_score_field(label, value):
        """Bold label + bold score value."""
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(label + ' ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.color.rgb = BLACK
        r2 = p.add_run(value or '')
        r2.bold = True
        r2.font.name = 'Calibri'
        r2.font.color.rgb = BLACK

    # ── Parse interview text ──────────────────────────────────────────
    parsed = {'header': {}, 'experience': {}, 'clinical': {},
              'compliance': {}, 'availability': {}, 'assessment': {}}
    current = 'header'
    cur_q   = None
    cur_ans = []

    def flush():
        nonlocal cur_q, cur_ans
        if cur_q and cur_ans:
            parsed[current][cur_q] = ' '.join(cur_ans).strip()
            cur_q = None
            cur_ans = []

    for line in interview_text.splitlines():
        sl = line.strip()
        if sl == 'Experience':
            flush(); current = 'experience'; continue
        elif sl == 'Clinical Questions':
            flush(); current = 'clinical'; continue
        elif sl == 'Compliance':
            flush(); current = 'compliance'; continue
        elif sl == 'Availability':
            flush(); current = 'availability'; continue
        elif sl == 'Assessment':
            flush(); current = 'assessment'; continue

        if current == 'header':
            if ':' in sl:
                k, v = sl.split(':', 1)
                parsed['header'][k.strip()] = v.strip()
        elif current in ('experience', 'clinical'):
            if _re.match(r'^[0-9]+\.\s+.+$', sl):
                flush(); cur_q = sl; cur_ans = []
            elif cur_q and sl:
                cur_ans.append(sl)
        elif current in ('compliance', 'availability', 'assessment'):
            if ':' in sl:
                k, v = sl.split(':', 1)
                parsed[current][k.strip()] = v.strip()
    flush()

    # ── Build document ────────────────────────────────────────────────

    # Document title — centred bold
    title_p = d.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(16)
    t_run = title_p.add_run(f'Completed {user_type} Interview')
    t_run.bold = True
    t_run.font.size = Pt(16)
    t_run.font.name = 'Calibri'
    t_run.font.color.rgb = BLACK

    # Header fields (Name / Location / NMBI PIN / Visa Status)
    hdr = parsed['header']
    for key in ['Name', 'Location', 'NMBI PIN', 'Visa Status']:
        add_field(f'{key}:', hdr.get(key, ''))

    # ── Experience ────────────────────────────────────────────────────
    add_section_heading('Experience')
    for q_text, answer in parsed['experience'].items():
        add_numbered_question(q_text)
        add_answer_text(answer)

    # ── Clinical Questions ────────────────────────────────────────────
    add_section_heading('Clinical Questions')
    for q_text, answer in parsed['clinical'].items():
        add_numbered_question(q_text)
        add_answer_text(answer)

    # ── Compliance ────────────────────────────────────────────────────
    add_section_heading('Compliance')
    for field in ['NMBI Registration', 'BLS/CPR', 'Manual Handling',
                  'Garda Vetting', 'References']:
        add_compliance_field(
            f'{field}:', parsed['compliance'].get(field, '')
        )

    # ── Availability ──────────────────────────────────────────────────
    add_section_heading('Availability')
    for field in ['Preferred counties', 'Day/Night/Both', 'Earliest start date']:
        add_field(f'{field}:', parsed['availability'].get(field, ''))

    # ── Assessment ────────────────────────────────────────────────────
    add_section_heading('Assessment')
    for field in ['Communication', 'Clinical Knowledge', 'Experience']:
        add_score_field(f'{field}:', parsed['assessment'].get(field, ''))
    add_compliance_field('Suitable:', parsed['assessment'].get('Suitable', 'Yes'))

    buf = _io.BytesIO()
    d.save(buf)
    return buf.getvalue()



@admin_bp.route('/live-staffs/cv/<staff_id>')
@admin_required
def live_staff_cv(staff_id):
    """Generate and download a filled HSE CV PDF for a staff member."""
    try:
        doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not doc:
            return "Staff record not found", 404
        pdf_bytes = _build_cv_pdf(doc)
        s1   = (doc.get('section_1_personal_details') or {})
        name = (s1.get('full_name') or 'staff').replace(' ', '_')
        filename = f"CV_{name}.pdf"
        return Response(
            pdf_bytes,
            mimetype='application/pdf',
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return str(e), 500


def _v(val):
    """Return value as string, or empty string if None/empty."""
    if val is None:
        return ''
    return str(val).strip()






def _build_ai_cv_docx(doc, cv_text):
    """
    Render AI-generated CV text as a clean Word document (.docx).
    Plain black text, Calibri font, section headings with bottom border —
    matches the clean professional style of the interview docx.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io as _io, re as _re

    BLACK = RGBColor(0x00, 0x00, 0x00)
    GRAY  = RGBColor(0x44, 0x44, 0x44)

    s1_d      = doc.get('section_1_personal_details') or {}
    full_name = _v(s1_d.get('full_name')) or 'Candidate'
    mobile    = _v(s1_d.get('mobile_number'))
    email     = _v(doc.get('email'))
    address   = _v(s1_d.get('address'))

    d = DocxDocument()
    for sec in d.sections:
        sec.top_margin    = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin   = Cm(2.54)
        sec.right_margin  = Cm(2.54)

    normal = d.styles['Normal']
    normal.font.name  = 'Calibri'
    normal.font.size  = Pt(11)

    def add_border_bottom(para):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '000000')
        pBdr.append(bot)
        pPr.append(pBdr)

    def add_section_heading(title):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(13)
        r.font.name = 'Calibri'
        r.font.color.rgb = BLACK
        add_border_bottom(p)

    def add_name_header():
        # Name — large centred
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(full_name)
        r.bold = True
        r.font.size = Pt(18)
        r.font.name = 'Calibri'
        r.font.color.rgb = BLACK
        # Contact line
        contact = '   |   '.join(x for x in [mobile, email, address] if x)
        if contact:
            cp = d.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(12)
            cr = cp.add_run(contact)
            cr.font.size = Pt(9)
            cr.font.name = 'Calibri'
            cr.font.color.rgb = GRAY

    def add_body(text):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = BLACK

    def add_field(label, value):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(label + '  ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.color.rgb = BLACK
        r2 = p.add_run(value or '')
        r2.font.name = 'Calibri'
        r2.font.color.rgb = BLACK

    def add_role_heading(title, dates=''):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(title)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r1.font.color.rgb = BLACK
        if dates:
            r2 = p.add_run(f'   {dates}')
            r2.font.size = Pt(9)
            r2.font.name = 'Calibri'
            r2.font.color.rgb = GRAY

    def add_sub(text):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        r.font.color.rgb = GRAY

    def add_bullet(text):
        clean = text.lstrip('- •	').strip()
        if not clean:
            return
        p = d.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.25)
        r = p.add_run(clean)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = BLACK

    # ── Parse sections ────────────────────────────────────────────────
    HEADINGS = [
        'PERSONAL DETAILS', 'PROFESSIONAL PROFILE',
        'EDUCATION & QUALIFICATIONS', 'PROFESSIONAL EXPERIENCE',
        'TRAINING & CERTIFICATIONS', 'KEY SKILLS', 'ADDITIONAL INFORMATION',
    ]
    sections = {}
    current  = '__pre__'
    sections[current] = []
    for line in cv_text.splitlines():
        matched = next((h for h in HEADINGS if line.strip().upper() == h), None)
        if matched:
            current = matched
            sections[current] = []
        else:
            sections.setdefault(current, []).append(line)

    # ── Build document ────────────────────────────────────────────────
    add_name_header()

    for heading in HEADINGS:
        lines = [l for l in sections.get(heading, []) if l.strip()]
        if not lines:
            continue

        add_section_heading(heading)

        if heading == 'PERSONAL DETAILS':
            for line in lines:
                if ':' in line:
                    parts = line.split(':', 1)
                    lbl = parts[0].strip() + ':'
                    val = parts[1].strip()
                    if val:
                        add_field(lbl, val)

        elif heading == 'PROFESSIONAL PROFILE':
            para_buf = []
            for line in lines:
                if line.strip() == '':
                    if para_buf:
                        add_body(' '.join(para_buf))
                        para_buf = []
                else:
                    para_buf.append(line.strip())
            if para_buf:
                add_body(' '.join(para_buf))

        elif heading == 'EDUCATION & QUALIFICATIONS':
            for line in lines:
                s = line.strip().lstrip('- ').strip()
                if not s:
                    continue
                parts = [p.strip() for p in s.split('|')]
                qual  = parts[0] if parts else s
                inst  = parts[1] if len(parts) > 1 else ''
                year  = parts[2] if len(parts) > 2 else ''
                p = d.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(1)
                r = p.add_run(qual + (f'  ({year})' if year else ''))
                r.bold = True
                r.font.name = 'Calibri'
                r.font.color.rgb = BLACK
                if inst:
                    add_sub(inst)

        elif heading == 'PROFESSIONAL EXPERIENCE':
            roles = []
            cur   = []
            for line in lines:
                if line.strip().lower().startswith('job title:') and cur:
                    roles.append(cur); cur = [line]
                else:
                    cur.append(line)
            if cur:
                roles.append(cur)

            for role_lines in roles:
                jt = en = ds = ''
                duties = []
                for rl in role_lines:
                    sl = rl.strip(); sll = sl.lower()
                    if not sl: continue
                    if sll.startswith('job title:'):  jt = sl.split(':',1)[1].strip()
                    elif sll.startswith('employer:'): en = sl.split(':',1)[1].strip()
                    elif sll.startswith('dates:'):    ds = sl.split(':',1)[1].strip()
                    elif sll.startswith('duties'):    pass
                    elif sl.startswith('-') or sl.startswith('•'):
                        duties.append(sl.lstrip('- •').strip())
                if not jt and not en:
                    continue
                add_role_heading(jt or 'Role', ds)
                if en:
                    add_sub(en)
                for duty in duties:
                    if duty:
                        add_bullet(duty)
                d.add_paragraph()

        elif heading in ('TRAINING & CERTIFICATIONS', 'KEY SKILLS'):
            for line in lines:
                s = line.strip()
                if s:
                    add_bullet(s)

        elif heading == 'ADDITIONAL INFORMATION':
            # Only render Driving Licence and Own Transport lines
            for line in lines:
                s = line.strip()
                if not s:
                    continue
                sl = s.lower()
                if not (sl.startswith('driving licence') or sl.startswith('own transport')):
                    continue
                if ':' in s:
                    parts = s.split(':', 1)
                    lbl   = parts[0].strip() + ':'
                    val   = parts[1].strip()
                    p = d.add_paragraph()
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after  = Pt(1)
                    r = p.add_run(f'{lbl}  {val}')
                    r.font.name  = 'Calibri'
                    r.font.size  = Pt(11)
                    r.font.color.rgb = BLACK

    buf = _io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _build_ai_cv_pdf(doc, cv_text):
    """
    4 visually distinct ATS-friendly CV designs, all black text.
    Theme chosen by md5(staff_id) % 4 — same staff always gets same design.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
        Image as RLImage, Table, TableStyle
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
    import io as _io, hashlib

    BLACK     = colors.HexColor('#000000')
    NEAR_BLK  = colors.HexColor('#111111')
    DARK_GRAY = colors.HexColor('#333333')
    LT_GRAY   = colors.HexColor('#CCCCCC')
    WHITE     = colors.white

    W, H = A4
    id_str  = str(doc.get('_id', ''))
    theme_n = int(hashlib.md5(id_str.encode()).hexdigest(), 16) % 4

    LAYOUTS = [
        {'lm':22*mm,'rm':22*mm,'tm':18*mm,'bm':15*mm,'bf':'Helvetica','bfb':'Helvetica-Bold','bfi':'Helvetica-Oblique','name_size':22,'name_align':TA_LEFT,'sec_size':10,'body_size':10,'contact_align':TA_LEFT},
        {'lm':25*mm,'rm':25*mm,'tm':20*mm,'bm':15*mm,'bf':'Times-Roman','bfb':'Times-Bold','bfi':'Times-Italic','name_size':24,'name_align':TA_CENTER,'sec_size':11,'body_size':10,'contact_align':TA_CENTER},
        {'lm':18*mm,'rm':18*mm,'tm':15*mm,'bm':12*mm,'bf':'Helvetica','bfb':'Helvetica-Bold','bfi':'Helvetica-Oblique','name_size':20,'name_align':TA_LEFT,'sec_size':9,'body_size':9.5,'contact_align':TA_LEFT},
        {'lm':28*mm,'rm':28*mm,'tm':22*mm,'bm':18*mm,'bf':'Times-Roman','bfb':'Times-Bold','bfi':'Times-Italic','name_size':26,'name_align':TA_CENTER,'sec_size':11,'body_size':10.5,'contact_align':TA_CENTER},
    ]
    L      = LAYOUTS[theme_n]
    PAGE_W = W - L['lm'] - L['rm']
    BF     = L['bf']
    BFB    = L['bfb']
    BFI    = L['bfi']

    def ps(name, **kw):
        d = dict(fontName=BF, fontSize=L['body_size'], textColor=NEAR_BLK,
                 spaceAfter=2, leading=L['body_size'] * 1.5)
        d.update(kw)
        return ParagraphStyle(name, **d)

    S = {
        'name'    : ps('name',    fontName=BFB, fontSize=L['name_size'],
                       textColor=BLACK, alignment=L['name_align'],
                       spaceAfter=3, leading=L['name_size'] * 1.3),
        'contact' : ps('contact', fontSize=L['body_size'] - 1,
                       textColor=DARK_GRAY, alignment=L['contact_align'],
                       spaceAfter=0, leading=14),
        'sec'     : ps('sec',     fontName=BFB, fontSize=L['sec_size'],
                       textColor=BLACK, spaceAfter=0,
                       leading=L['sec_size'] * 1.4, tracking=30),
        'body'    : ps('body',    alignment=TA_JUSTIFY, spaceAfter=4,
                       leading=L['body_size'] * 1.55),
        'val'     : ps('val',     spaceAfter=1),
        'role'    : ps('role',    fontName=BFB, fontSize=L['body_size'] + 1,
                       textColor=BLACK, spaceAfter=1,
                       leading=(L['body_size'] + 1) * 1.4),
        'employer': ps('employer',fontName=BFI, fontSize=L['body_size'],
                       textColor=DARK_GRAY, spaceAfter=1, leading=14),
        'dates'   : ps('dates',   fontSize=L['body_size'] - 1,
                       textColor=DARK_GRAY, spaceAfter=2, leading=13),
        'bullet'  : ps('bullet',  leftIndent=10, spaceAfter=3,
                       leading=L['body_size'] * 1.5),
        'qual_q'  : ps('qual_q',  fontName=BFB, fontSize=L['body_size'],
                       textColor=BLACK, spaceAfter=1),
        'qual_i'  : ps('qual_i',  fontName=BFI, fontSize=L['body_size'] - 1,
                       textColor=DARK_GRAY, spaceAfter=0, leading=13),
    }

    sp   = lambda n=3: Spacer(1, n * mm)
    thin = lambda: HRFlowable(width=PAGE_W, color=LT_GRAY, thickness=0.5, spaceAfter=2)

    def sec_heading(title):
        if theme_n == 0:
            return [Paragraph(title.upper(), S['sec']),
                    HRFlowable(width=PAGE_W, color=BLACK, thickness=0.8, spaceAfter=3)]
        elif theme_n == 1:
            return [HRFlowable(width=PAGE_W, color=BLACK, thickness=0.4, spaceAfter=2),
                    Paragraph(title.upper(), S['sec']),
                    HRFlowable(width=PAGE_W, color=BLACK, thickness=1.2, spaceAfter=4)]
        elif theme_n == 2:
            t = Table([[Paragraph(f'  {title.upper()}', S['sec'])]], colWidths=[PAGE_W])
            t.setStyle(TableStyle([
                ('LINEBEFORE',(0,0),(0,-1),3.5,BLACK),('LINEBELOW',(0,0),(-1,-1),0.4,LT_GRAY),
                ('TOPPADDING',(0,0),(-1,-1),3),('BOTTOMPADDING',(0,0),(-1,-1),4),
                ('LEFTPADDING',(0,0),(-1,-1),6)]))
            return [t]
        else:
            t = Table([[Paragraph(f'  {title.upper()}  ', S['sec'])]], colWidths=[PAGE_W])
            t.setStyle(TableStyle([
                ('BOX',(0,0),(-1,-1),0.8,BLACK),('TOPPADDING',(0,0),(-1,-1),4),
                ('BOTTOMPADDING',(0,0),(-1,-1),4),('LEFTPADDING',(0,0),(-1,-1),8)]))
            return [t]

    def bullet_p(text):
        clean = text.lstrip('- \u2022\t').strip()
        return Paragraph(f'\u2022\u2003{clean}', S['bullet']) if clean else None

    HEADINGS = [
        'PERSONAL DETAILS','PROFESSIONAL PROFILE','EDUCATION & QUALIFICATIONS',
        'PROFESSIONAL EXPERIENCE','TRAINING & CERTIFICATIONS','KEY SKILLS',
        'ADDITIONAL INFORMATION',
    ]
    sections = {}
    current  = '__pre__'
    sections[current] = []
    for line in cv_text.splitlines():
        matched = next((h for h in HEADINGS if line.strip().upper() == h), None)
        if matched:
            current = matched
            sections[current] = []
        else:
            sections.setdefault(current, []).append(line)

    s1_d      = doc.get('section_1_personal_details') or {}
    full_name = _v(s1_d.get('full_name')) or 'Candidate'
    mobile    = _v(s1_d.get('mobile_number'))
    email     = _v(doc.get('email'))
    address   = _v(s1_d.get('address'))

    logo_path = None
    for c in [
        os.path.join(os.path.dirname(__file__), '..', 'static', 'images', 'logo.png'),
        os.path.join(os.path.dirname(__file__), '..', 'static', 'img', 'logo.png'),
        os.path.join(os.path.dirname(__file__), '..', 'static', 'logo.png'),
        'static/images/logo.png', 'static/img/logo.png', 'static/logo.png',
    ]:
        if os.path.exists(c):
            logo_path = c
            break

    buf   = _io.BytesIO()
    story = []

    # Header
    if logo_path:
        logo_w   = 30*mm
        logo_img = RLImage(logo_path, width=logo_w, height=logo_w*94/316)
        if L['name_align'] == TA_CENTER:
            story.append(logo_img); story.append(sp(2))
        else:
            hdr_t = Table([[Paragraph('', S['body']), logo_img]],
                          colWidths=[PAGE_W - logo_w - 2*mm, logo_w + 2*mm])
            hdr_t.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'TOP'),
                                       ('TOPPADDING',(0,0),(-1,-1),0),
                                       ('BOTTOMPADDING',(0,0),(-1,-1),0),
                                       ('LEFTPADDING',(0,0),(-1,-1),0),
                                       ('RIGHTPADDING',(0,0),(-1,-1),0)]))
            story.append(hdr_t)

    story.append(Paragraph(full_name, S['name']))
    contact_parts = [p for p in [mobile, email, address] if p]
    if contact_parts:
        sep = '   |   ' if L['contact_align'] == TA_CENTER else '  •  '
        story.append(Paragraph(sep.join(contact_parts), S['contact']))

    story.append(sp(3))
    if theme_n == 0:
        story.append(HRFlowable(width=PAGE_W,color=BLACK,thickness=1.5,spaceAfter=0))
        story.append(HRFlowable(width=PAGE_W,color=BLACK,thickness=0.4,spaceAfter=4))
    elif theme_n == 1:
        story.append(HRFlowable(width=PAGE_W,color=BLACK,thickness=1.2,spaceAfter=4))
    elif theme_n == 2:
        story.append(HRFlowable(width=PAGE_W,color=LT_GRAY,thickness=0.5,spaceAfter=4))
    else:
        story.append(HRFlowable(width=PAGE_W,color=BLACK,thickness=0.6,spaceAfter=2))
        story.append(HRFlowable(width=PAGE_W,color=BLACK,thickness=0.6,spaceAfter=4))
    story.append(sp(2))

    for heading in HEADINGS:
        lines = [l for l in sections.get(heading, []) if l.strip()]
        if not lines:
            continue
        story += sec_heading(heading)
        story.append(sp(2))

        if heading == 'PERSONAL DETAILS':
            for line in lines:
                if ':' in line:
                    parts = line.split(':', 1)
                    lbl_t = parts[0].strip() + ':'
                    val_t = parts[1].strip()
                    if val_t:
                        story.append(Paragraph(f'<b>{lbl_t}</b> {val_t}', S['val']))
                        story.append(sp(1))
            story.append(sp(3))

        elif heading == 'PROFESSIONAL PROFILE':
            pb = []
            for line in lines:
                if line.strip() == '':
                    if pb:
                        story.append(Paragraph(' '.join(pb), S['body']))
                        story.append(sp(2))
                        pb = []
                else:
                    pb.append(line.strip())
            if pb:
                story.append(Paragraph(' '.join(pb), S['body']))
            story.append(sp(4))

        elif heading == 'EDUCATION & QUALIFICATIONS':
            for line in lines:
                s = line.strip().lstrip('- ').strip()
                if not s:
                    continue
                parts = [p.strip() for p in s.split('|')]
                qual  = parts[0] if parts else s
                inst  = parts[1] if len(parts) > 1 else ''
                year  = parts[2] if len(parts) > 2 else ''
                yr_txt = f' ({year})' if year else ''
                story.append(Paragraph(f'<b>{qual}</b>{yr_txt}', S['qual_q']))
                if inst:
                    story.append(Paragraph(inst, S['qual_i']))
                story += [sp(2), thin(), sp(2)]
            story.append(sp(3))

        elif heading == 'PROFESSIONAL EXPERIENCE':
            roles = []
            cur   = []
            for line in lines:
                if line.strip().lower().startswith('job title:') and cur:
                    roles.append(cur); cur = [line]
                else:
                    cur.append(line)
            if cur:
                roles.append(cur)

            for ri, role_lines in enumerate(roles):
                job_title = emp_name = dates_str = ''
                duties    = []
                for rl in role_lines:
                    sl    = rl.strip()
                    sl_lo = sl.lower()
                    if not sl: continue
                    if sl_lo.startswith('job title:'):   job_title = sl.split(':',1)[1].strip()
                    elif sl_lo.startswith('employer:'):  emp_name  = sl.split(':',1)[1].strip()
                    elif sl_lo.startswith('dates:'):     dates_str = sl.split(':',1)[1].strip()
                    elif sl_lo.startswith('duties'):     pass
                    elif sl.startswith('-') or sl.startswith('\u2022'):
                        duties.append(sl.lstrip('- \u2022').strip())

                if not job_title and not emp_name:
                    continue

                if theme_n in (1, 3) and dates_str:
                    rt = Table([[Paragraph(f'<b>{job_title}</b>', S['role']),
                                 Paragraph(dates_str, S['dates'])]],
                               colWidths=[PAGE_W*0.65, PAGE_W*0.35])
                    rt.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'BOTTOM'),
                                            ('TOPPADDING',(0,0),(-1,-1),0),
                                            ('BOTTOMPADDING',(0,0),(-1,-1),2),
                                            ('LEFTPADDING',(0,0),(-1,-1),0),
                                            ('RIGHTPADDING',(0,0),(-1,-1),0),
                                            ('ALIGN',(1,0),(1,-1),'RIGHT')]))
                    story.append(rt)
                else:
                    story.append(Paragraph(job_title or 'Role', S['role']))
                    if dates_str:
                        story.append(Paragraph(dates_str, S['dates']))

                if emp_name:
                    story.append(Paragraph(emp_name, S['employer']))
                story.append(sp(2))
                for d in duties:
                    if d:
                        bp = bullet_p(d)
                        if bp: story.append(bp)
                story.append(sp(3))
                if ri < len(roles) - 1:
                    story.append(thin()); story.append(sp(2))
            story.append(sp(2))

        elif heading in ('TRAINING & CERTIFICATIONS', 'KEY SKILLS'):
            if theme_n in (0, 2):
                bitems = [bullet_p(l.strip()) for l in lines if bullet_p(l.strip())]
                pairs  = []
                for i in range(0, len(bitems), 2):
                    left  = bitems[i]
                    right = bitems[i+1] if i+1 < len(bitems) else Paragraph('', S['body'])
                    pairs.append([left, right])
                if pairs:
                    col_w = PAGE_W / 2 - 3*mm
                    bt = Table(pairs, colWidths=[col_w, col_w])
                    bt.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'TOP'),
                                            ('TOPPADDING',(0,0),(-1,-1),1),
                                            ('BOTTOMPADDING',(0,0),(-1,-1),1),
                                            ('LEFTPADDING',(0,0),(-1,-1),0),
                                            ('RIGHTPADDING',(0,0),(-1,-1),4)]))
                    story.append(bt)
            else:
                for line in lines:
                    bp = bullet_p(line.strip())
                    if bp: story.append(bp)
            story.append(sp(4))

        elif heading == 'ADDITIONAL INFORMATION':
            for line in lines:
                s = line.strip()
                if not s or s.lower().startswith('reference'): continue
                if ':' in s:
                    parts = s.split(':', 1)
                    story.append(Paragraph(
                        f'<b>{parts[0].strip()}:</b> {parts[1].strip()}', S['val']))
                    story.append(sp(1))
                else:
                    story.append(Paragraph(s, S['body']))
            story.append(sp(4))

    pdf_doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=L['lm'], rightMargin=L['rm'],
        topMargin=L['tm'],  bottomMargin=L['bm'],
    )
    pdf_doc.build(story)
    return buf.getvalue()



def _build_cv_pdf(doc):
    """
    Build a rich individual Xpress Health CV PDF.
    Mirrors the Abidemi Aluko CV structure:
      1. Personal Details
      2. Professional Profile  (auto-generated flowing paragraph)
      3. Education & Qualifications  (entry per qual)
      4. Professional Experience  (one card per role with full duties)
      5. Training & Certifications  (bullet list, max 6)
      6. Key Skills  (bullet list)
      7. Additional Information  (Driving / Transport / References / Date)
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table,
        TableStyle, HRFlowable, Image as RLImage, ListFlowable, ListItem
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    import io as _io, os

    # ── Brand palette ─────────────────────────────────────────────────
    NAVY      = colors.HexColor('#1B3A6B')
    XH_GREEN  = colors.HexColor('#2E9E44')
    LIGHT_BG  = colors.HexColor('#EFF6FF')
    STRIPE    = colors.HexColor('#F0FDF4')
    MID_GRAY  = colors.HexColor('#CBD5E1')
    TEXT_DARK = colors.HexColor('#1E293B')
    TEXT_GRAY = colors.HexColor('#475569')
    WHITE     = colors.white

    W, H   = A4
    PAGE_W = W - 30 * mm   # 15 mm margins each side

    # ── Styles ────────────────────────────────────────────────────────
    def ps(name, **kw):
        d = dict(fontName='Helvetica', fontSize=10, textColor=TEXT_GRAY,
                 spaceAfter=2, leading=15)
        d.update(kw)
        return ParagraphStyle(name, **d)

    S = {
        'cv_title'  : ps('cv_title',   fontName='Helvetica-Bold', fontSize=20,
                         textColor=WHITE, alignment=TA_CENTER, spaceAfter=0, leading=24),
        'cv_name'   : ps('cv_name',    fontName='Helvetica-Bold', fontSize=13,
                         textColor=NAVY, alignment=TA_CENTER, spaceAfter=0, leading=18),
        'sec_head'  : ps('sec_head',   fontName='Helvetica-Bold', fontSize=10,
                         textColor=WHITE, spaceAfter=0, leading=14),
        'lbl'       : ps('lbl',        fontName='Helvetica-Bold', fontSize=10,
                         textColor=NAVY, spaceAfter=0, leading=14),
        'val'       : ps('val',        fontSize=10, textColor=TEXT_GRAY,
                         spaceAfter=0, leading=14),
        'body'      : ps('body',       fontSize=10, textColor=TEXT_GRAY,
                         alignment=TA_JUSTIFY, spaceAfter=4, leading=15),
        'exp_title' : ps('exp_title',  fontName='Helvetica-Bold', fontSize=11,
                         textColor=NAVY, spaceAfter=0, leading=15),
        'exp_sub'   : ps('exp_sub',    fontName='Helvetica-Oblique', fontSize=10,
                         textColor=XH_GREEN, spaceAfter=0, leading=14),
        'exp_date'  : ps('exp_date',   fontName='Helvetica-Bold', fontSize=9,
                         textColor=WHITE, alignment=TA_CENTER, spaceAfter=0, leading=12),
        'duty'      : ps('duty',       fontSize=10, textColor=TEXT_GRAY,
                         leftIndent=8, spaceAfter=3, leading=15),
        'bullet'    : ps('bullet',     fontSize=10, textColor=TEXT_GRAY,
                         leftIndent=8, spaceAfter=3, leading=15),
        'footer'    : ps('footer',     fontSize=7,  textColor=MID_GRAY,
                         alignment=TA_CENTER, spaceAfter=0),
        'qual_title': ps('qual_title', fontName='Helvetica-Bold', fontSize=10,
                         textColor=NAVY, spaceAfter=1, leading=14),
        'qual_sub'  : ps('qual_sub',   fontName='Helvetica-Oblique', fontSize=9,
                         textColor=TEXT_GRAY, spaceAfter=0, leading=13),
    }

    sp = lambda n=3: Spacer(1, n * mm)

    # ── Helpers ───────────────────────────────────────────────────────
    def sec(title):
        t = Table([[Paragraph(title, S['sec_head'])]], colWidths=[PAGE_W])
        t.setStyle(TableStyle([
            ('BACKGROUND',    (0,0), (-1,-1), NAVY),
            ('TOPPADDING',    (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('LEFTPADDING',   (0,0), (-1,-1), 10),
            ('LINEBELOW',     (0,0), (-1,-1), 2, XH_GREEN),
        ]))
        return t

    def lv(label, value, lw=55*mm):
        val_text = value if value else '—'
        t = Table(
            [[Paragraph(label, S['lbl']), Paragraph(val_text, S['val'])]],
            colWidths=[lw, PAGE_W - lw]
        )
        t.setStyle(TableStyle([
            ('TOPPADDING',    (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ('LEFTPADDING',   (0,0), (0,0),   6),
            ('LEFTPADDING',   (1,0), (1,0),   4),
            ('LINEBELOW',     (0,0), (-1,-1), 0.3, MID_GRAY),
        ]))
        return t

    def date_badge(text):
        """Green pill badge for date range."""
        t = Table([[Paragraph(text, S['exp_date'])]], colWidths=[None])
        t.setStyle(TableStyle([
            ('BACKGROUND',    (0,0), (-1,-1), XH_GREEN),
            ('TOPPADDING',    (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ('LEFTPADDING',   (0,0), (-1,-1), 8),
            ('RIGHTPADDING',  (0,0), (-1,-1), 8),
            ('ROUNDEDCORNERS',(0,0), (-1,-1), 4),
        ]))
        return t

    def bullet_item(text):
        return Paragraph(f'\u2022\u2003{text}', S['bullet'])

    def duty_item(text):
        return Paragraph(f'\u2022\u2003{text}', S['duty'])

    def profile_box(text):
        t = Table([[Paragraph(text, S['body'])]], colWidths=[PAGE_W])
        t.setStyle(TableStyle([
            ('BACKGROUND',    (0,0), (-1,-1), LIGHT_BG),
            ('BOX',           (0,0), (-1,-1), 0.5, MID_GRAY),
            ('LINEBEFORE',    (0,0), (0,-1),  4,   XH_GREEN),
            ('TOPPADDING',    (0,0), (-1,-1), 10),
            ('BOTTOMPADDING', (0,0), (-1,-1), 10),
            ('LEFTPADDING',   (0,0), (-1,-1), 12),
            ('RIGHTPADDING',  (0,0), (-1,-1), 12),
        ]))
        return t

    # ── Data ─────────────────────────────────────────────────────────
    s1   = doc.get('section_1_personal_details') or {}
    s2   = doc.get('section_2_identity_verification') or {}
    s3   = doc.get('section_3_professional_registration') or {}
    s4   = doc.get('section_4_qualifications') or {}
    s5   = doc.get('section_5_employment_history') or {}
    s7   = doc.get('section_7_references') or {}
    s8   = doc.get('section_8_garda_vetting_police_clearance') or {}
    s9   = doc.get('section_9_occupational_health') or {}
    s10  = doc.get('section_10_mandatory_training') or {}
    s12  = doc.get('section_12_declaration') or {}
    visa     = s1.get('work_permit_visa_status') or {}
    docs_sub = s2.get('documents_submitted') or {}

    full_name  = _v(s1.get('full_name'))
    emp_code   = _v(doc.get('employee_code'))
    user_type  = _v(doc.get('user_type'))
    address    = _v(s1.get('address'))
    mobile     = _v(s1.get('mobile_number'))
    email      = _v(doc.get('email'))
    dob        = _v(s1.get('date_of_birth'))
    nationality= _v(s1.get('nationality'))
    reg_pin    = _v(s3.get('registration_number_pin'))
    reg_exp    = _v(s3.get('registration_expiry_date'))
    divisions  = ', '.join(s3.get('divisions_registered_in') or [])
    nmbi       = s3.get('nmbi_active_declaration')
    perm_work  = _v(visa.get('permission_to_work'))
    visa_type  = _v(visa.get('visa_type'))
    total_exp  = _v(s5.get('total_experience'))
    entries    = [e for e in (s5.get('entries') or [])
                  if e.get('employer') or e.get('position')]

    # ── Logo ─────────────────────────────────────────────────────────
    logo_path = None
    for c in [
        os.path.join(os.path.dirname(__file__), '..', 'static', 'images', 'logo.png'),
        os.path.join(os.path.dirname(__file__), '..', 'static', 'img', 'logo.png'),
        os.path.join(os.path.dirname(__file__), '..', 'static', 'logo.png'),
        'static/images/logo.png', 'static/img/logo.png', 'static/logo.png',
    ]:
        if os.path.exists(c):
            logo_path = c
            break

    # ── Build story ───────────────────────────────────────────────────
    buf   = _io.BytesIO()
    story = []

    # ════════════════════════════════════════════════════════════════
    # HEADER — Logo + "CURRICULUM VITAE" + candidate name
    # ════════════════════════════════════════════════════════════════
    title_rows = [
        [Paragraph('CURRICULUM VITAE', S['cv_title'])],
        [Paragraph(full_name or 'Candidate', S['cv_name'])],   # name under title
    ]
    title_w   = PAGE_W - (55*mm if logo_path else 0)
    title_tbl = Table(title_rows, colWidths=[title_w])
    title_tbl.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), NAVY),
        ('TOPPADDING',    (0,0), (-1,-1), 12),
        ('BOTTOMPADDING', (0,0), (-1,-1), 12),
        ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
        # name row slightly lighter bg
        ('BACKGROUND',    (0,1), (-1,1),  colors.HexColor('#162F58')),
        ('TOPPADDING',    (0,1), (-1,1),  6),
        ('BOTTOMPADDING', (0,1), (-1,1),  8),
    ]))

    if logo_path:
        logo_img  = RLImage(logo_path, width=45*mm, height=45*mm*94/316)
        logo_cell = Table([[logo_img]], colWidths=[55*mm])
        logo_cell.setStyle(TableStyle([
            ('BACKGROUND',    (0,0), (-1,-1), WHITE),
            ('TOPPADDING',    (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('LEFTPADDING',   (0,0), (-1,-1), 6),
            ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
        ]))
        banner = Table([[logo_cell, title_tbl]],
                       colWidths=[55*mm, PAGE_W - 55*mm])
    else:
        banner = Table([[title_tbl]], colWidths=[PAGE_W])

    banner.setStyle(TableStyle([
        ('TOPPADDING',    (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ('LEFTPADDING',   (0,0), (-1,-1), 0),
        ('RIGHTPADDING',  (0,0), (-1,-1), 0),
        ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
        ('LINEBELOW',     (0,0), (-1,-1), 3, XH_GREEN),
    ]))
    story += [banner, sp(5)]

    # ════════════════════════════════════════════════════════════════
    # 1. PERSONAL DETAILS
    # ════════════════════════════════════════════════════════════════
    story += [sec('PERSONAL DETAILS'), sp(3)]
    for label, value in [
        ('Full Name:',     full_name),
        ('Address:',       address),
        ('Mobile Number:', mobile),
        ('Email Address:', email),
        ('Nationality:',   nationality),
    ]:
        if value:
            story += [lv(label, value), sp(1)]
    story.append(sp(4))

    # ════════════════════════════════════════════════════════════════
    # 2. PROFESSIONAL PROFILE — rich individual paragraph
    # ════════════════════════════════════════════════════════════════
    story += [sec('PROFESSIONAL PROFILE'), sp(3)]

    # Build a natural multi-sentence profile from the data
    para_sentences = []

    # Opening — who they are + experience + speciality
    if full_name and user_type:
        opener = f"{full_name} is a compassionate and dedicated {user_type}"
        if divisions:
            opener += f" specialising in {divisions}"
        if total_exp:
            opener += f", with {total_exp} of professional healthcare experience"
        para_sentences.append(opener)

    # Most recent role
    if entries:
        latest = entries[0]
        pos_l  = _v(latest.get('position'))
        emp_l  = _v(latest.get('employer'))
        d_from = _v(latest.get('from'))
        d_to   = _v(latest.get('to'))
        if pos_l and emp_l:
            role_s = f"Most recently working as {pos_l} at {emp_l}"
            if d_from:
                role_s += f" from {d_from}"
                role_s += f" to {d_to}" if d_to else " to present"
            para_sentences.append(role_s)

    # Registration / professional status
    if reg_pin or reg_exp or nmbi:
        reg_s = "Professionally registered"
        if reg_pin:
            reg_s += f" (PIN: {reg_pin})"
        if reg_exp:
            reg_s += f" with registration valid until {reg_exp}"
        if nmbi:
            reg_s += ", holding an active NMBI declaration"
        para_sentences.append(reg_s)

    # Work authorisation
    if perm_work == 'Yes' and visa_type:
        para_sentences.append(
            f"Fully authorised to work in Ireland ({visa_type})"
        )
    elif perm_work == 'Yes':
        para_sentences.append("Fully authorised to work in Ireland")

    # Occupational health
    if s9.get('fit_for_nursing_duties'):
        para_sentences.append(
            "Confirmed fit for nursing duties with up-to-date occupational health clearance"
        )

    # Garda vetting
    if s8.get('garda_vetting_submitted'):
        para_sentences.append("Garda vetted and cleared to work with vulnerable adults")

    # Qualities closing line
    qualities = (
        "Known for excellent communication, a caring and professional manner, "
        "and a genuine commitment to promoting client dignity, independence, and wellbeing"
    )
    para_sentences.append(qualities)

    profile_text = '. '.join(para_sentences) + '.'

    story.append(profile_box(profile_text))
    story.append(sp(4))

    # ════════════════════════════════════════════════════════════════
    # 3. EDUCATION & QUALIFICATIONS
    # ════════════════════════════════════════════════════════════════
    story += [sec('EDUCATION & QUALIFICATIONS'), sp(3)]

    qual_keys = ['nursing_degree', 'postgraduate_qualification', 'other_qualification']
    qual_found = False
    for qk in qual_keys:
        q     = s4.get(qk) or {}
        qname = _v(q.get('qualification'))
        qinst = _v(q.get('institution'))
        qyear = _v(q.get('year_completed'))
        if not (qname or qinst):
            continue
        qual_found = True

        # Heading row: qual name left, year right
        h_left  = Paragraph(f'<b>{qname}</b>' if qname else '<b>Qualification</b>',
                             S['qual_title'])
        h_right = Paragraph(qyear, S['exp_date'])
        yr_cell = Table([[h_right]], colWidths=[30*mm])
        yr_cell.setStyle(TableStyle([
            ('BACKGROUND',    (0,0), (-1,-1), XH_GREEN),
            ('TOPPADDING',    (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ('LEFTPADDING',   (0,0), (-1,-1), 6),
            ('RIGHTPADDING',  (0,0), (-1,-1), 6),
        ]))
        head_row = Table([[h_left, yr_cell]],
                         colWidths=[PAGE_W - 34*mm, 34*mm])
        head_row.setStyle(TableStyle([
            ('TOPPADDING',    (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ('LEFTPADDING',   (0,0), (-1,-1), 0),
            ('RIGHTPADDING',  (0,0), (-1,-1), 0),
            ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(head_row)
        if qinst:
            story.append(Paragraph(qinst, S['qual_sub']))
        story += [sp(2), HRFlowable(width=PAGE_W, color=MID_GRAY, thickness=0.4), sp(3)]

    if not qual_found:
        story.append(Paragraph('No qualifications recorded.', S['body']))
        story.append(sp(3))

    story.append(sp(1))

    # ════════════════════════════════════════════════════════════════
    # 4. PROFESSIONAL EXPERIENCE — one card per role with full duties
    # ════════════════════════════════════════════════════════════════
    story += [sec('PROFESSIONAL EXPERIENCE'), sp(3)]

    if entries:
        for i, e in enumerate(entries):
            pos     = _v(e.get('position'))
            emp     = _v(e.get('employer'))
            loc     = _v(e.get('location', ''))   # location field if present
            d_from  = _v(e.get('from'))
            d_to    = _v(e.get('to'))
            leaving = _v(e.get('reason_for_leaving'))

            # Date range string
            if d_from and d_to:
                date_str = f"{d_from} \u2013 {d_to}"
            elif d_from:
                date_str = f"{d_from} \u2013 Present"
            elif d_to:
                date_str = f"Until {d_to}"
            else:
                date_str = ''

            # ── Role heading: title left, date badge right ────────
            t_para  = Paragraph(f'<b>{pos}</b>' if pos else '<b>Role</b>', S['exp_title'])
            if date_str:
                d_badge = Table([[Paragraph(date_str, S['exp_date'])]],
                                colWidths=[None])
                d_badge.setStyle(TableStyle([
                    ('BACKGROUND',    (0,0), (-1,-1), XH_GREEN),
                    ('TOPPADDING',    (0,0), (-1,-1), 4),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                    ('LEFTPADDING',   (0,0), (-1,-1), 10),
                    ('RIGHTPADDING',  (0,0), (-1,-1), 10),
                ]))
                head_t = Table([[t_para, d_badge]],
                               colWidths=[PAGE_W * 0.60, PAGE_W * 0.40])
            else:
                head_t = Table([[t_para]], colWidths=[PAGE_W])

            head_t.setStyle(TableStyle([
                ('TOPPADDING',    (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
                ('LEFTPADDING',   (0,0), (-1,-1), 0),
                ('RIGHTPADDING',  (0,0), (-1,-1), 0),
                ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
            ]))
            story.append(head_t)

            # Employer + location sub-line
            sub_parts = []
            if emp: sub_parts.append(emp)
            if loc: sub_parts.append(loc)
            if sub_parts:
                story.append(Paragraph(' \u2022 '.join(sub_parts), S['exp_sub']))

            story.append(sp(2))

            # ── Description paragraph ─────────────────────────────
            # Build a rich descriptive paragraph for this role
            desc_parts = []
            if pos and emp:
                desc = f"Worked as <b>{pos}</b> at {emp}"
                if loc:
                    desc += f", based in {loc}"
                if d_from and d_to:
                    desc += f", from {d_from} to {d_to}"
                elif d_from:
                    desc += f" from {d_from} to present"
                desc_parts.append(desc)

            # Add responsibilities based on user_type keywords
            ut_lower = user_type.lower() if user_type else ''
            if 'nurse' in ut_lower or 'nursing' in ut_lower:
                role_duties = [
                    "Assisted residents and clients with all aspects of personal care including personal hygiene, dressing, and grooming",
                    "Supported safe mobility and transfers, assisting with walking, wheelchair use, and repositioning",
                    "Observed and reported changes in residents' condition — including skin integrity, pain, and behaviour — to the nursing team",
                    "Assisted with medication administration under the direct supervision of qualified nursing staff",
                    "Maintained accurate records and contributed to care planning in line with individual care plans",
                    "Built positive and respectful therapeutic relationships with residents and their families",
                    "Worked effectively within multidisciplinary teams, supporting a safe and caring environment",
                ]
            elif 'healthcare' in ut_lower or 'hca' in ut_lower or 'assistant' in ut_lower:
                role_duties = [
                    "Provided high-quality, person-centred care and support tailored to each individual client's needs",
                    "Assisted clients with all activities of daily living including personal care, meal preparation, and mobility support",
                    "Observed and reported changes in clients' physical or emotional wellbeing to the supervising care team",
                    "Promoted client independence, dignity, and choice throughout all aspects of care delivery",
                    "Maintained comprehensive and accurate care records in line with organisational policies",
                    "Collaborated effectively with colleagues, families, and multidisciplinary teams to ensure continuity of care",
                    "Followed safe working practices, infection control procedures, and moving and handling guidelines at all times",
                ]
            else:
                role_duties = [
                    "Delivered high standards of professional care and support in line with organisational policies and procedures",
                    "Maintained clear and accurate records and communicated effectively with the wider team",
                    "Promoted the dignity, independence, and wellbeing of all clients and residents at all times",
                ]

            if desc_parts:
                story.append(Paragraph('. '.join(desc_parts) + '.', S['body']))
                story.append(sp(2))

            # Duties heading
            story.append(Paragraph('<b>Duties &amp; Responsibilities</b>', S['lbl']))
            story.append(sp(1))
            for duty in role_duties:
                story.append(duty_item(duty))
            story.append(sp(2))

            if leaving:
                story.append(Paragraph(
                    f'<i>Reason for leaving: {leaving}</i>', S['qual_sub']
                ))
                story.append(sp(2))

            if i < len(entries) - 1:
                story.append(HRFlowable(width=PAGE_W, color=MID_GRAY, thickness=0.5))
                story.append(sp(3))

    else:
        story.append(Paragraph('No employment history recorded.', S['body']))
        story.append(sp(3))

    if total_exp:
        story += [sp(2), lv('Total Experience:', total_exp, lw=55*mm), sp(1)]
    story.append(sp(4))

    # ════════════════════════════════════════════════════════════════
    # 5. TRAINING & CERTIFICATIONS — bullet list, max 6
    # ════════════════════════════════════════════════════════════════
    story += [sec('TRAINING & CERTIFICATIONS'), sp(3)]
    TLABELS = {
        'manual_handling':              'Manual Handling',
        'cpr_bls':                      'CPR / Basic Life Support',
        'fire_safety':                  'Fire Safety',
        'infection_prevention_control': 'Infection Prevention & Control',
        'hand_hygiene':                 'Hand Hygiene',
        'safeguarding':                 'Safeguarding Vulnerable Adults',
        'children_first':               'Children First',
        'cyber_security':               'Cyber Security Awareness',
        'dignity_at_work':              'Dignity at Work',
        'open_disclosure':              'Open Disclosure',
        'mapa_pmav':                    'MAPA / PMAV (De-escalation)',
    }
    cert_labels = []
    for key, label in TLABELS.items():
        if s10.get(key):
            cert_labels.append(label)
        if len(cert_labels) == 6:
            break

    if cert_labels:
        for label in cert_labels:
            story.append(bullet_item(label))
        story.append(sp(4))
    else:
        story.append(Paragraph('No training certifications recorded.', S['body']))
        story.append(sp(4))

    # ════════════════════════════════════════════════════════════════
    # 6. KEY SKILLS — bullets from health/registration data
    # ════════════════════════════════════════════════════════════════
    story += [sec('KEY SKILLS'), sp(3)]
    ut_lower = user_type.lower() if user_type else ''
    if 'nurse' in ut_lower or 'nursing' in ut_lower:
        skills = [
            'Medication administration (under nursing supervision)',
            'Patient assessment and observation',
            'Personal and person-centred care',
            'Patient moving and handling / safe mobility support',
            'Communication and interpersonal skills',
            'Observation, monitoring, and reporting of patient condition',
            'Record keeping and report writing',
            'Teamwork and collaboration with multidisciplinary teams',
            'Compassion, empathy, and patience',
            'Promoting patient dignity and independence',
        ]
    else:
        skills = [
            'Person-centred care and support',
            'Assistance with all activities of daily living',
            'Patient moving and handling / safe mobility support',
            'Communication and interpersonal skills',
            'Observation, monitoring, and reporting of client condition',
            'Record keeping and report writing',
            'Teamwork and collaboration',
            'Compassion, empathy, and patience',
            'Promoting client dignity and independence',
        ]

    for skill in skills:
        story.append(bullet_item(skill))
    story.append(sp(4))

    # ════════════════════════════════════════════════════════════════
    # 7. ADDITIONAL INFORMATION
    # ════════════════════════════════════════════════════════════════
    story += [sec('ADDITIONAL INFORMATION'), sp(3)]
    # Generate a random date between 01 Jan 2024 and 31 Dec 2026
    import random as _rand
    from datetime import date as _date, timedelta as _td
    _d_start = _date(2024, 1, 1); _d_end = _date(2026, 12, 31)
    _rand_date = _d_start + _td(days=_rand.randint(0, (_d_end - _d_start).days))
    _cv_date   = _rand_date.strftime('%d %B %Y')
    for label, value in [
        ('Driving Licence:', 'No'),
        ('Own Transport:',   'No'),
    ]:
        story += [lv(label, value), sp(1)]
    story.append(sp(4))

    # ── Render ────────────────────────────────────────────────────────
    pdf_doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=15*mm, rightMargin=15*mm,
        topMargin=10*mm,  bottomMargin=15*mm,
    )
    pdf_doc.build(story)
    return buf.getvalue()




@admin_bp.route('/live-staffs/add', methods=['POST'])
@admin_required
def live_staff_add():
    data = request.get_json()

    email = (data.get('email') or '').strip().lower()
    if not email:
        return jsonify({"success": False, "error": "Email is required"}), 400

    if _staffs_col().count_documents({"email": email}) > 0:
        return jsonify({"success": False, "error": f'Email "{email}" already exists'}), 400

    doc = _build_doc(data)
    doc["created_at"] = datetime.utcnow()

    try:
        _staffs_col().insert_one(doc)
        return jsonify({"success": True, "message": "Staff record created"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/edit', methods=['POST'])
@admin_required
def live_staff_edit():
    data     = request.get_json()
    staff_id = (data.get('staff_id') or '').strip()

    if not staff_id:
        return jsonify({"success": False, "error": "Missing staff_id"}), 400

    email = (data.get('email') or '').strip().lower()
    if not email:
        return jsonify({"success": False, "error": "Email is required"}), 400

    if _staffs_col().count_documents({"email": email, "_id": {"$ne": ObjectId(staff_id)}}) > 0:
        return jsonify({"success": False, "error": f'Email "{email}" already exists'}), 400

    try:
        col     = _staffs_col()
        current = col.find_one({"_id": ObjectId(staff_id)})
        if not current:
            return jsonify({"success": False, "error": "Staff record not found"}), 404

        doc = _build_doc(data)
        col.update_one({"_id": ObjectId(staff_id)}, {"$set": doc})
        return jsonify({"success": True, "message": "Staff record updated"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/delete', methods=['POST'])
@admin_required
def live_staff_delete():
    data     = request.get_json()
    staff_id = (data.get('staff_id') or '').strip()

    if not staff_id:
        return jsonify({"success": False, "error": "Missing staff_id"}), 400
    try:
        _staffs_col().delete_one({"_id": ObjectId(staff_id)})
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Import ────────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/import', methods=['POST'])
@admin_required
def live_staff_import():
    """Accept a JSON file upload; upsert on email. Handles all JSON variants."""
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file provided"}), 400

    file = request.files['file']
    if not file.filename.endswith('.json'):
        return jsonify({"success": False, "error": "Only .json files are accepted"}), 400

    try:
        content = file.read().decode('utf-8')
        records = _parse_json_content(content)
    except Exception as e:
        return jsonify({"success": False, "error": f"Could not parse file: {e}"}), 400

    inserted = updated = skipped = 0
    errors   = []

    for idx, rec in enumerate(records):
        email = (rec.get('email') or '').strip().lower()
        if not email:
            skipped += 1
            continue
        try:
            doc = _map_import_record(rec)
            result = _staffs_col().update_one(
                {"email": email},
                {"$set": doc, "$setOnInsert": {"created_at": datetime.utcnow()}},
                upsert=True
            )
            if result.upserted_id:
                inserted += 1
            else:
                updated += 1
        except Exception as e:
            errors.append(f"Row {idx + 1} ({email}): {e}")

    return jsonify({
        "success": True,
        "inserted": inserted,
        "updated":  updated,
        "skipped":  skipped,
        "errors":   errors,
        "message":  f"Import complete — {inserted} added, {updated} updated, {skipped} skipped"
    })


# ── Export ────────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/export')
@admin_required
def live_staff_export():
    fmt   = request.args.get('format', 'json').lower()
    items = list(_staffs_col().find({}))

    if fmt == 'csv':
        return _export_csv(items)
    return _export_json(items)


@admin_bp.route('/live-staffs/export/xlsx')
@admin_required
def live_staff_export_xlsx():
    """
    Export all live_staffs records to an Excel (.xlsx) file.
    Columns: Sno | Name | Email | Points
    Sorted alphabetically by name.

    GET /admin/live-staffs/export/xlsx
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        # ── Fetch all staff ───────────────────────────────────────────
        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1, "points": 1}
        ))

        # Sort by name
        def _get_name(d):
            s1 = d.get('section_1_personal_details') or {}
            return _v(s1.get('full_name') or '').lower()

        docs.sort(key=_get_name)

        # ── Build workbook ────────────────────────────────────────────
        wb = Workbook()
        ws = wb.active
        ws.title = 'Staff Points'

        # Styles
        NAVY_HEX  = '1B3A6B'
        GREEN_HEX = '2E9E44'
        WHITE_HEX = 'FFFFFF'

        header_font  = Font(name='Arial', bold=True, color=WHITE_HEX, size=11)
        header_fill  = PatternFill('solid', start_color=NAVY_HEX, end_color=NAVY_HEX)
        header_align = Alignment(horizontal='center', vertical='center')

        body_font    = Font(name='Arial', size=10)
        alt_fill     = PatternFill('solid', start_color='EFF6FF', end_color='EFF6FF')
        center_align = Alignment(horizontal='center', vertical='center')
        left_align   = Alignment(horizontal='left', vertical='center')

        thin_side = Side(style='thin', color='CCCCCC')
        thin_border = Border(
            left=thin_side, right=thin_side,
            top=thin_side,  bottom=thin_side
        )

        # ── Headers ───────────────────────────────────────────────────
        headers = ['Sno', 'Name', 'Email', 'Points']
        col_widths = [6, 35, 40, 10]

        for col_idx, (header, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font      = header_font
            cell.fill      = header_fill
            cell.alignment = header_align
            cell.border    = thin_border
            ws.column_dimensions[cell.column_letter].width = width

        ws.row_dimensions[1].height = 22

        # ── Bottom border on header in green ─────────────────────────
        for col_idx in range(1, 5):
            cell = ws.cell(row=1, column=col_idx)
            cell.border = Border(
                left=thin_side, right=thin_side,
                top=thin_side,
                bottom=Side(style='medium', color=GREEN_HEX)
            )

        # ── Data rows ─────────────────────────────────────────────────
        for row_idx, doc in enumerate(docs, start=2):
            s1     = doc.get('section_1_personal_details') or {}
            name   = _v(s1.get('full_name') or '')
            email  = _v(doc.get('email') or '')
            points = doc.get('points')

            row_fill = alt_fill if row_idx % 2 == 0 else None

            values = [row_idx - 1, name, email, points]
            aligns = [center_align, left_align, left_align, center_align]

            for col_idx, (val, align) in enumerate(zip(values, aligns), start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.font      = body_font
                cell.alignment = align
                cell.border    = thin_border
                if row_fill:
                    cell.fill = row_fill

            ws.row_dimensions[row_idx].height = 18

        # Freeze header row
        ws.freeze_panes = 'A2'

        # ── Auto-filter ───────────────────────────────────────────────
        ws.auto_filter.ref = f'A1:D{len(docs) + 1}'

        # ── Render to bytes ───────────────────────────────────────────
        buf = _io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        xlsx_bytes = buf.getvalue()

        return Response(
            xlsx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={
                'Content-Disposition':
                    f'attachment; filename="staff_points_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'
            }
        )

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



@admin_bp.route('/live-staffs/export/vetting-xlsx')
@admin_required
def live_staff_export_vetting_xlsx():
    """
    Export 4 Excel files in one ZIP, each filtered by vetting status:

    ?filter=no_garda         - garda_vetting == 0
    ?filter=garda_expired    - garda_vetting_expired == 1
    ?filter=no_police        - police_clearance == 0
    ?filter=police_expired   - police_clearance_expired == 1
    ?filter=all              - all staff with all 4 status columns (default)

    Columns: Sno | Name | Email | Phone | Garda Vetting | Garda Expired | Police Clearance | Police Expired
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        filter_type = request.args.get('filter', 'all').lower()

        FILTERS = {
            'no_garda':       {"garda_vetting":          0},
            'garda_expired':  {"garda_vetting_expired":  1},
            'no_police':      {"police_clearance":       0},
            'police_expired': {"police_clearance_expired": 1},
            'all':            {},
        }

        TITLES = {
            'no_garda':       'No Garda Vetting',
            'garda_expired':  'Garda Vetting Expired',
            'no_police':      'No Police Clearance',
            'police_expired': 'Police Clearance Expired',
            'all':            'All Staff — Vetting Status',
        }

        if filter_type not in FILTERS:
            return jsonify({"success": False,
                            "error": f"Invalid filter. Use: {', '.join(FILTERS.keys())}"}), 400

        query = FILTERS[filter_type]
        docs  = list(_staffs_col().find(
            query,
            {"section_1_personal_details": 1, "email": 1,
             "garda_vetting": 1, "garda_vetting_expired": 1,
             "police_clearance": 1, "police_clearance_expired": 1}
        ))

        # Sort by name
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        # ── Styles ────────────────────────────────────────────────────
        NAVY   = '1B3A6B'; GREEN = '2E9E44'; RED = 'CC0000'
        WHITE  = 'FFFFFF'; ALT   = 'EFF6FF'; WARN = 'FFF3CD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=11)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        c_align = Alignment(horizontal='center', vertical='center')
        l_align = Alignment(horizontal='left', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)

        def status_label(val, field_type):
            if val is None:
                return 'Not Checked'
            if field_type == 'vetting':
                return 'Approved' if val == 1 else 'No Garda Vetting'
            if field_type == 'vetting_exp':
                return 'Expired' if val == 1 else 'Valid'
            if field_type == 'police':
                return 'Approved' if val == 1 else 'No Police Clearance'
            if field_type == 'police_exp':
                return 'Expired' if val == 1 else 'Valid'
            return str(val)

        def cell_color(val, is_expired=False):
            if val is None:
                return PatternFill('solid', start_color='EEEEEE', end_color='EEEEEE')
            if is_expired:
                return PatternFill('solid', start_color=WARN, end_color=WARN) if val == 1 else None
            return None if val == 1 else PatternFill('solid', start_color='FFDDDD', end_color='FFDDDD')

        # ── Build workbook ────────────────────────────────────────────
        wb    = Workbook()
        ws    = wb.active
        ws.title = TITLES[filter_type][:31]

        # Title row
        title_cell = ws.cell(row=1, column=1, value=TITLES[filter_type])
        title_cell.font = Font(name='Arial', bold=True, size=13, color=NAVY)
        ws.merge_cells('A1:H1')
        ws.row_dimensions[1].height = 24

        # Header row
        headers    = ['Sno', 'Name', 'Email', 'Phone',
                      'Garda Vetting', 'Garda Expired',
                      'Police Clearance', 'Police Expired']
        col_widths = [6, 32, 38, 18, 20, 18, 22, 18]

        for col_idx, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=2, column=col_idx, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = border
            ws.column_dimensions[cell.column_letter].width = width
            # Green bottom border
            cell.border = Border(
                left=Side(style='thin', color='CCCCCC'),
                right=Side(style='thin', color='CCCCCC'),
                top=Side(style='thin', color='CCCCCC'),
                bottom=Side(style='medium', color=GREEN)
            )
        ws.row_dimensions[2].height = 22
        ws.freeze_panes = 'A3'
        ws.auto_filter.ref = f'A2:H{len(docs) + 2}'

        # Data rows
        for row_idx, doc in enumerate(docs, start=3):
            s1     = doc.get('section_1_personal_details') or {}
            name   = _v(s1.get('full_name') or '')
            email  = _v(doc.get('email') or '')
            phone  = _v(s1.get('mobile_number') or '')
            gv     = doc.get('garda_vetting')
            ge     = doc.get('garda_vetting_expired')
            pc     = doc.get('police_clearance')
            pe     = doc.get('police_clearance_expired')

            alt = PatternFill('solid', start_color=ALT, end_color=ALT) if row_idx % 2 == 0 else None

            row_data = [
                (row_idx - 2,              c_align, None),
                (name,                     l_align, None),
                (email,                    l_align, None),
                (phone,                    l_align, None),
                (status_label(gv,  'vetting'),     c_align, cell_color(gv)),
                (status_label(ge,  'vetting_exp'), c_align, cell_color(ge, True)),
                (status_label(pc,  'police'),      c_align, cell_color(pc)),
                (status_label(pe,  'police_exp'),  c_align, cell_color(pe, True)),
            ]

            for col_idx, (val, align, fill) in enumerate(row_data, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.font      = b_font
                cell.alignment = align
                cell.border    = border
                if fill:
                    cell.fill = fill
                elif alt:
                    cell.fill = alt

            ws.row_dimensions[row_idx].height = 17

        # Summary row
        summary_row = len(docs) + 3
        ws.cell(row=summary_row, column=1,
                value=f'Total: {len(docs)} staff').font = Font(name='Arial', bold=True, size=10)

        buf = _io.BytesIO()
        wb.save(buf)
        xlsx_bytes = buf.getvalue()

        date_str  = datetime.utcnow().strftime('%Y%m%d')
        file_name = f"vetting_{filter_type}_{date_str}.xlsx"

        return Response(
            xlsx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition": f'attachment; filename="{file_name}"'}
        )

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500





# ── External API: Generate documents via static API key ──────────────
# Header: X-API-Key: <LIVE_STAFF_API_KEY env var>
# Body:   {"staff_id": "<mongo _id>"}
# ─────────────────────────────────────────────────────────────────────

def _validate_api_key():
    """Validate X-API-Key header against LIVE_STAFF_API_KEY env var."""
    expected = os.environ.get('LIVE_STAFF_API_KEY', '').strip()
    if not expected:
        return False, "LIVE_STAFF_API_KEY not configured on server"
    provided = (request.headers.get('X-API-Key') or '').strip()
    if not provided:
        return False, "Missing X-API-Key header"
    if provided != expected:
        return False, "Invalid API key"
    return True, None


@admin_bp.route('/live-staffs/api/generate-cv', methods=['POST'])
def api_generate_cv():
    """
    External API — generate AI CV for a staff member.

    Headers:
      X-API-Key: <LIVE_STAFF_API_KEY>
      Content-Type: application/json

    Body:
      {"staff_id": "68abc123..."}

    Response:
      {"success": true, "staff_name": "...", "cv_filename": "...", "gcs_blob": "..."}
    """
    ok, err = _validate_api_key()
    if not ok:
        return jsonify({"success": False, "error": err}), 401

    data     = request.get_json(silent=True) or {}
    staff_id = (data.get('staff_id') or '').strip()
    email    = (data.get('email') or '').strip().lower()

    if not staff_id and not email:
        return jsonify({"success": False,
                        "error": "Provide staff_id or email in request body"}), 400

    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set on server"}), 500

    try:
        doc = None
        if staff_id:
            # Search by staff_id (XN Portal ID), xn_staff_id, then MongoDB _id
            doc = _staffs_col().find_one({"staff_id": staff_id})
            if not doc:
                doc = _staffs_col().find_one({"xn_staff_id": staff_id})
            if not doc:
                try:
                    doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
                except Exception:
                    pass
        if not doc and email:
            doc = _staffs_col().find_one({"$or": [
                {"email": email},
                {"section_1_personal_details.email_address": email},
            ]})
        if not doc:
            identifier = staff_id or email
            return jsonify({"success": False,
                            "error": f"Staff not found: {identifier}"}), 404

        s1        = doc.get('section_1_personal_details') or {}
        s3        = doc.get('section_3_professional_registration') or {}
        s4        = doc.get('section_4_qualifications') or {}
        s5        = doc.get('section_5_employment_history') or {}
        s8        = doc.get('section_8_garda_vetting_police_clearance') or {}
        s9        = doc.get('section_9_occupational_health') or {}
        s10       = doc.get('section_10_mandatory_training') or {}
        visa      = s1.get('work_permit_visa_status') or {}

        def _vv(v): return '' if v is None else str(v).strip()

        full_name   = _vv(s1.get('full_name'))
        email       = _vv(doc.get('email'))
        user_type   = _vv(doc.get('user_type'))
        emp_code    = _vv(doc.get('employee_code'))
        address     = _vv(s1.get('address'))
        mobile      = _vv(s1.get('mobile_number'))
        nationality = _vv(s1.get('nationality'))
        total_exp   = _vv(s5.get('total_experience'))
        divisions   = ', '.join(s3.get('divisions_registered_in') or [])
        reg_pin     = _vv(s3.get('registration_number_pin'))
        reg_exp     = _vv(s3.get('registration_expiry_date'))
        nmbi        = 'Yes' if s3.get('nmbi_active_declaration') else 'No'
        visa_type   = _vv(visa.get('visa_type'))
        perm_work   = _vv(visa.get('permission_to_work'))
        garda       = 'Yes' if s8.get('garda_vetting_submitted') else 'No'
        fit         = 'Yes' if s9.get('fit_for_nursing_duties') else 'No'

        qual_lines = []
        for qk in ['nursing_degree', 'postgraduate_qualification', 'other_qualification']:
            q = s4.get(qk) or {}
            if q.get('qualification') or q.get('institution'):
                qual_lines.append(
                    f"  - {_vv(q.get('qualification'))} | "
                    f"{_vv(q.get('institution'))} | "
                    f"{_vv(q.get('year_completed'))}"
                )
        # Also include NMBI/QQI numbers as qualification context
        nmbi_num = _vv(doc.get('nmbi_number') or s3.get('registration_number_pin') or '')
        qqi_num  = _vv(doc.get('qqi_number') or '')
        if nmbi_num and not any('nmbi' in l.lower() or 'registration' in l.lower() for l in qual_lines):
            qual_lines.append(f"  - NMBI Registration PIN: {nmbi_num}")
        if qqi_num and not any('qqi' in l.lower() for l in qual_lines):
            qual_lines.append(f"  - QQI Level 5 Certificate No: {qqi_num}")

        entries   = [e for e in (s5.get('entries') or []) if e.get('employer') or e.get('position')]
        exp_lines = [
            f"  - {_vv(e.get('position'))} at {_vv(e.get('employer'))} "
            f"({_vv(e.get('from'))} - {_vv(e.get('to') or 'Present')})"
            for e in entries
        ]

        TLABELS = {
            'manual_handling': 'Manual Handling', 'cpr_bls': 'CPR / BLS',
            'fire_safety': 'Fire Safety',
            'infection_prevention_control': 'Infection Prevention & Control',
            'hand_hygiene': 'Hand Hygiene', 'safeguarding': 'Safeguarding',
        }
        certs = [label for k, label in TLABELS.items() if s10.get(k)][:6]

        extracted_cv = _v(doc.get('extracted_cv') or '')
        has_extracted = (
            extracted_cv and
            not extracted_cv.startswith('[') and
            extracted_cv != 'No doc found'
        )

        data_summary = f"""
Candidate: {full_name}
Role / User Type: {user_type}
Employee Code: {emp_code}
Address: {address}
Mobile: {mobile}
Email: {email}
Nationality: {nationality}
Total Experience: {total_exp}
Divisions / Speciality: {divisions}
Registration PIN: {reg_pin}
Registration Expiry: {reg_exp}
NMBI Active Declaration: {nmbi}
Permission to Work: {perm_work}
Visa / Stamp Type: {visa_type}
Garda Vetted: {garda}
Fit for Nursing Duties: {fit}

Qualifications:
{chr(10).join(qual_lines) if qual_lines else "  None recorded"}

Employment History:
{chr(10).join(exp_lines) if exp_lines else "  None recorded"}

Training & Certifications:
{chr(10).join("  - " + c for c in certs) if certs else "  None recorded"}
""".strip()

        extracted_cv_section = f"""
EXTRACTED CV TEXT (use as PRIMARY source for PROFESSIONAL EXPERIENCE, TRAINING & CERTIFICATIONS and KEY SKILLS):
{extracted_cv[:8000]}
""" if has_extracted else ""

        prompt = f"""You are an expert professional CV writer specialising in Irish healthcare staffing.

STRICT RULE — NO HALLUCINATION: Use ONLY the exact facts in CANDIDATE DATA.

Structure (EXACT UPPERCASE headings on their own line):
PERSONAL DETAILS
PROFESSIONAL PROFILE
EDUCATION & QUALIFICATIONS
PROFESSIONAL EXPERIENCE
TRAINING & CERTIFICATIONS
KEY SKILLS
ADDITIONAL INFORMATION

Rules:
- PERSONAL DETAILS: "Label: Value" per line. Skip blank fields.
- PROFESSIONAL PROFILE: 2 paragraphs, FIRST PERSON.
- EDUCATION & QUALIFICATIONS: Qualification Name | Institution | Year
- PROFESSIONAL EXPERIENCE: Job Title: / Employer: / Dates: / Duties: / - duty
- TRAINING & CERTIFICATIONS: Bullet list only.
- KEY SKILLS: 8-10 bullets.
- ADDITIONAL INFORMATION: Include ONLY these two lines:
Driving Licence: No
Own Transport: No

---
CANDIDATE DATA:
{data_summary}
{extracted_cv_section}
---

Output CV text only. No preamble, no markdown symbols.
"""

        from google import genai as google_genai
        client   = google_genai.Client(api_key=gemini_key)
        response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
        cv_text  = response.text.strip()

        docx_bytes  = _build_ai_cv_docx(doc, cv_text)
        safe_name   = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        cv_filename = f"{safe_name}.docx"
        gcs_blob    = f"cv/{cv_filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        col      = _ai_cvs_col()
        existing = col.find_one({"staff_id": str(doc['_id'])})
        ai_doc   = {
            "staff_id":      str(doc['_id']),
            "staff_name":    full_name,
            "employee_code": emp_code,
            "cv_text":       cv_text,
            "cv_filename":   cv_filename,
            "gcs_blob":      gcs_blob,
            "generated_at":  datetime.utcnow(),
        }
        if existing:
            col.update_one({"_id": existing["_id"]}, {"$set": ai_doc})
            ai_id = str(existing["_id"])
        else:
            ai_id = str(col.insert_one(ai_doc).inserted_id)

        _push_hse_document_background(
            staff_id_str=staff_id, doc_type_key='cv',
            docx_bytes=docx_bytes, staff_name=full_name,
            mongo_id=staff_id, email=email,
        )

        download_url = _gcs_signed_url(gcs_blob) or ''
        return jsonify({
            "success":      True,
            "staff_id":     staff_id,
            "staff_name":   full_name,
            "cv_id":        ai_id,
            "cv_filename":  cv_filename,
            "gcs_blob":     gcs_blob,
            "download_url": download_url,
            "generated_at": datetime.utcnow().isoformat(),
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/api/generate-interview', methods=['POST'])
def api_generate_interview():
    """
    External API — generate AI Interview Notes for a staff member.

    Headers:
      X-API-Key: <LIVE_STAFF_API_KEY>
      Content-Type: application/json

    Body:
      {"staff_id": "68abc123..."}
    """
    ok, err = _validate_api_key()
    if not ok:
        return jsonify({"success": False, "error": err}), 401

    data     = request.get_json(silent=True) or {}
    staff_id = (data.get('staff_id') or '').strip()
    email    = (data.get('email') or '').strip().lower()

    if not staff_id and not email:
        return jsonify({"success": False,
                        "error": "Provide staff_id or email in request body"}), 400

    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set on server"}), 500

    try:
        doc = None
        if staff_id:
            # Search by staff_id (XN Portal ID), xn_staff_id, then MongoDB _id
            doc = _staffs_col().find_one({"staff_id": staff_id})
            if not doc:
                doc = _staffs_col().find_one({"xn_staff_id": staff_id})
            if not doc:
                try:
                    doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
                except Exception:
                    pass
        if not doc and email:
            doc = _staffs_col().find_one({"$or": [
                {"email": email},
                {"section_1_personal_details.email_address": email},
            ]})
        if not doc:
            identifier = staff_id or email
            return jsonify({"success": False,
                            "error": f"Staff not found: {identifier}"}), 404

        s1        = doc.get('section_1_personal_details') or {}
        s3        = doc.get('section_3_professional_registration') or {}
        s5        = doc.get('section_5_employment_history') or {}
        s8        = doc.get('section_8_garda_vetting_police_clearance') or {}
        s9        = doc.get('section_9_occupational_health') or {}
        s10       = doc.get('section_10_mandatory_training') or {}
        visa      = s1.get('work_permit_visa_status') or {}

        def _vv(v): return '' if v is None else str(v).strip()

        full_name   = _vv(s1.get('full_name'))
        email       = _vv(doc.get('email'))
        user_type   = _vv(doc.get('user_type') or 'Nurse')
        emp_code    = _vv(doc.get('employee_code'))
        address     = _vv(s1.get('address'))
        nationality = _vv(s1.get('nationality'))
        reg_pin     = _vv(s3.get('registration_number_pin'))
        visa_type   = _vv(visa.get('visa_type'))
        divisions   = ', '.join(s3.get('divisions_registered_in') or [])
        total_exp   = _vv(s5.get('total_experience'))
        nmbi        = 'Yes' if s3.get('nmbi_active_declaration') else 'No'
        garda       = 'Yes' if s8.get('garda_vetting_submitted') else 'No'
        bls         = 'Yes' if s10.get('cpr_bls') else 'No'
        manual      = 'Yes' if s10.get('manual_handling') else 'No'
        fit         = 'Yes' if s9.get('fit_for_nursing_duties') else 'No'

        entries   = [e for e in (s5.get('entries') or []) if e.get('employer') or e.get('position')]
        exp_lines = [
            f"  - {_vv(e.get('position'))} at {_vv(e.get('employer'))} "
            f"({_vv(e.get('from'))} – {_vv(e.get('to') or 'Present')})"
            for e in entries[:5]
        ]

        TLABELS = {
            'manual_handling': 'Manual Handling', 'cpr_bls': 'BLS/CPR',
            'safeguarding': 'Safeguarding', 'fire_safety': 'Fire Safety',
            'infection_prevention_control': 'Infection Prevention & Control',
        }
        certs = [label for k, label in TLABELS.items() if s10.get(k)]

        county = ''
        if address:
            parts = [p.strip() for p in address.replace(',', ' ').split()]
            for p in parts:
                if p.lower().startswith('co.') or p.lower() == 'county':
                    idx = parts.index(p)
                    county = parts[idx + 1] if idx + 1 < len(parts) else ''
                    break
            if not county and parts:
                county = parts[-1]

        data_summary = f"""
Name: {full_name}
Role / User Type: {user_type}
Address / Location: {address}
Nationality: {nationality}
Visa / Stamp Type: {visa_type}
NMBI Registration PIN: {reg_pin}
NMBI Registration Active: {nmbi}
Divisions / Speciality: {divisions}
Total Experience: {total_exp}
Garda Vetted: {garda}
BLS/CPR on file: {bls}
Manual Handling on file: {manual}
Fit for Duties: {fit}

Employment History:
{chr(10).join(exp_lines) if exp_lines else "  None recorded"}

Certifications on file: {", ".join(certs) if certs else "None recorded"}
""".strip()

        prompt = f"""You are an experienced nursing recruitment consultant at Xpress Health, Ireland.

Complete a realistic professional interview notes template using ONLY the candidate data below.
Write answers in first person, conversational but professional. NO HALLUCINATION.
Assessment scores: pick varied random scores 3.5–5/5 in 0.5 increments.

Output ONLY the completed template — no preamble, no markdown:

---
Completed {user_type} Interview

Name: [full name]
Location: [county/city]
NMBI PIN: [pin or N/A]
Visa Status: [visa type]

Experience

1. Tell me about your nursing experience.
[4–6 sentence first-person answer using only data provided]

2. How many years in Ireland?
[realistic answer from employment history]

3. Acute, Nursing Home, Community, or Mental Health?
[most relevant care setting from employment history]

Clinical Questions

1. How would you manage a deteriorating patient?
[4–5 sentence clinically accurate answer for a {user_type}, using ABCDE/NEWS2/ISBAR]

2. What would you do if you witnessed a medication error?
[4–5 sentence answer covering patient safety, reporting, documentation, prevention]

Compliance
NMBI Registration: [Yes/No]
BLS/CPR: [Yes/No]
Manual Handling: [Yes/No]
Garda Vetting: [Yes/No]
References: Yes

Availability
Preferred counties: [county from address]
Day/Night/Both: Both
Earliest start date: Immediate

Assessment
Communication: [score/5]
Clinical Knowledge: [score/5]
Experience: [score/5]
Suitable: Yes
---

CANDIDATE DATA:
{data_summary}
"""

        from google import genai as google_genai
        client         = google_genai.Client(api_key=gemini_key)
        response       = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
        interview_text = response.text.strip().strip('-').strip()

        docx_bytes = _build_interview_docx(doc, interview_text)
        safe_name  = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        filename   = f"Interview_{safe_name}_{staff_id}.docx"
        gcs_blob   = f"interviews/{filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        col      = _ai_interviews_col()
        existing = col.find_one({"staff_id": str(doc['_id'])})
        rec = {
            "staff_id":       str(doc['_id']),
            "staff_name":     full_name,
            "employee_code":  emp_code,
            "interview_text": interview_text,
            "filename":       filename,
            "gcs_blob":       gcs_blob,
            "generated_at":   datetime.utcnow(),
        }
        if existing:
            col.update_one({"_id": existing["_id"]}, {"$set": rec})
            rec_id = str(existing["_id"])
        else:
            rec_id = str(col.insert_one(rec).inserted_id)

        _push_hse_document_background(
            staff_id_str=staff_id, doc_type_key='interview',
            docx_bytes=docx_bytes, staff_name=full_name,
            mongo_id=staff_id, email=email,
        )

        download_url = _gcs_signed_url(gcs_blob) or ''
        return jsonify({
            "success":       True,
            "staff_id":      staff_id,
            "staff_name":    full_name,
            "interview_id":  rec_id,
            "filename":      filename,
            "gcs_blob":      gcs_blob,
            "download_url":  download_url,
            "generated_at":  datetime.utcnow().isoformat(),
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/api/generate-appform', methods=['POST'])
def api_generate_appform():
    """
    External API — generate Application Form for a staff member.

    Headers:
      X-API-Key: <LIVE_STAFF_API_KEY>
      Content-Type: application/json

    Body:
      {"staff_id": "68abc123..."}
    """
    ok, err = _validate_api_key()
    if not ok:
        return jsonify({"success": False, "error": err}), 401

    data     = request.get_json(silent=True) or {}
    staff_id = (data.get('staff_id') or '').strip()
    email    = (data.get('email') or '').strip().lower()

    if not staff_id and not email:
        return jsonify({"success": False,
                        "error": "Provide staff_id or email in request body"}), 400

    try:
        doc = None
        if staff_id:
            doc = _staffs_col().find_one({"staff_id": staff_id})
            if not doc:
                doc = _staffs_col().find_one({"xn_staff_id": staff_id})
            if not doc:
                try:
                    doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
                except Exception:
                    pass
        if not doc and email:
            doc = _staffs_col().find_one({"$or": [
                {"email": email},
                {"section_1_personal_details.email_address": email},
            ]})
        if not doc:
            identifier = staff_id or email
            return jsonify({"success": False,
                            "error": f"Staff not found: {identifier}"}), 404

        s1        = doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff')
        email     = _v(doc.get('email'))
        emp_code  = _v(doc.get('employee_code') or '')

        # Download signature from GCS if available
        signature_bytes = None
        sig_blob = _v(doc.get('signature_gcs_blob') or '')
        if sig_blob:
            try:
                signature_bytes = _gcs_download(sig_blob)
            except Exception:
                signature_bytes = None

        docx_bytes = _build_appform_docx(doc, signature_bytes=signature_bytes)
        safe_name  = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        filename   = f"AppForm_{safe_name}.docx"
        gcs_blob   = f"appforms/{filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        col      = _ai_appforms_col()
        existing = col.find_one({"staff_id": str(doc['_id'])})
        rec = {
            "staff_id":      str(doc['_id']),
            "staff_name":    full_name,
            "employee_code": emp_code,
            "filename":      filename,
            "gcs_blob":      gcs_blob,
            "has_signature": bool(signature_bytes),
            "generated_at":  datetime.utcnow(),
        }
        if existing:
            col.update_one({"_id": existing["_id"]}, {"$set": rec})
            rec_id = str(existing["_id"])
        else:
            rec_id = str(col.insert_one(rec).inserted_id)

        _push_hse_document_background(
            staff_id_str=staff_id, doc_type_key='appform',
            docx_bytes=docx_bytes, staff_name=full_name,
            mongo_id=staff_id, email=email,
        )

        # Convert DOCX → PDF and upload to GCS for direct download link
        pdf_blob     = gcs_blob.replace('.docx', '.pdf')
        download_url = ''
        try:
            pdf_bytes_dl = _docx_to_pdf_bytes(docx_bytes)
            _gcs_upload(pdf_blob, pdf_bytes_dl, content_type='application/pdf')
            download_url = _gcs_signed_url(pdf_blob) or ''
        except Exception:
            # Fall back to DOCX signed URL if PDF conversion fails
            download_url = _gcs_signed_url(gcs_blob) or ''
            pdf_blob     = ''

        return jsonify({
            "success":       True,
            "staff_id":      staff_id,
            "staff_name":    full_name,
            "appform_id":    rec_id,
            "filename":      filename,
            "gcs_blob":      gcs_blob,
            "pdf_blob":      pdf_blob,
            "has_signature": bool(signature_bytes),
            "download_url":  download_url,
            "generated_at":  datetime.utcnow().isoformat(),
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Cron: Sync document list from XN Portal ───────────────────────────

@admin_bp.route('/live-staffs/cron/sync-documents', methods=['GET', 'POST'])
def live_staff_cron_sync_documents():
    """
    Cron job — processes ONE staff member per call.

    Logic:
      1. Find the first live_staffs record where extracted_cv is missing/empty.
      2. Call the XN Portal API to get their document list.
      3. If a document_type_name == "Cv" has a URL, download + extract text via Gemini.
      4. Save documents[] and extracted_cv back to MongoDB.
      5. Return result with how many staff still need processing (remaining_count).

    Call every N minutes via cron — it will work through all staff automatically,
    one at a time, until everyone has an extracted_cv.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req

    # ── Auth ──────────────────────────────────────────────────────────
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    # ── Env ───────────────────────────────────────────────────────────
    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')

    if not base_url:
        return jsonify({"success": False,
                        "error": "LIVE_STAFF_URL not set in environment"}), 500

    endpoint = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    col = _staffs_col()

    # ── Find next staff without extracted_cv ──────────────────────────
    staff = col.find_one(
        {"$or": [
            {"extracted_cv": {"$exists": False}},
            {"extracted_cv": None},
            {"extracted_cv": ""},
        ]},
        {"email": 1, "section_1_personal_details": 1}
    )

    # Count how many still need processing (including this one)
    remaining_total = col.count_documents(
        {"$or": [
            {"extracted_cv": {"$exists": False}},
            {"extracted_cv": None},
            {"extracted_cv": ""},
        ]}
    )

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff already have extracted_cv — nothing to do.",
            "remaining_count": 0,
            "processed":       None,
        })

    email = _v(
        (staff.get('section_1_personal_details') or {}).get('email_address') or
        staff.get('email') or ''
    )

    if not email:
        # Mark as attempted so it doesn't block the queue forever
        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {"extracted_cv": "[skipped — no email]",
                      "extracted_cv_at": datetime.utcnow()}}
        )
        return jsonify({
            "success":         True,
            "message":         "Skipped — staff record has no email address.",
            "remaining_count": remaining_total - 1,
            "processed":       str(staff.get("_id")),
        })

    # ── Call XN Portal API ────────────────────────────────────────────
    status_code   = None
    response_text = ''
    try:
        resp = _req.post(
            endpoint,
            json={"email": email},
            headers=api_headers,
            timeout=30
        )
        # If POST not allowed, retry with GET + query param
        if resp.status_code == 405:
            resp = _req.get(
                endpoint,
                params={"email": email},
                headers=api_headers,
                timeout=30
            )
        status_code   = resp.status_code
        response_text = resp.text[:500] if resp.text else ''
        resp.raise_for_status()
        data = resp.json()
    except _req.exceptions.ConnectionError as api_err:
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"Connection error — cannot reach XN Portal",
            "detail":          str(api_err)[:300],
            "endpoint":        endpoint,
            "check":           "Verify LIVE_STAFF_URL env var is correct and server is reachable",
            "remaining_count": remaining_total,
        })
    except _req.exceptions.Timeout:
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           "Timeout — XN Portal did not respond within 30s",
            "endpoint":        endpoint,
            "remaining_count": remaining_total,
        })
    except _req.exceptions.HTTPError as api_err:
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"HTTP {status_code} error from XN Portal",
            "response_body":   response_text,
            "endpoint":        endpoint,
            "remaining_count": remaining_total,
        })
    except Exception as api_err:
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"{type(api_err).__name__}: {api_err}",
            "endpoint":        endpoint,
            "remaining_count": remaining_total,
        })

    if not data.get('success'):
        # Mark attempted so cron moves on next time
        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {"extracted_cv": f"[API error: {data.get('message', 'unknown')}]",
                      "extracted_cv_at": datetime.utcnow()}}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           data.get('message', 'API returned success=false'),
            "api_raw":         data,
            "remaining_count": remaining_total - 1,
        })

    api_data = data.get('data')

    # API may return data as a list (empty []) or a dict with documents key
    if isinstance(api_data, list):
        documents = api_data          # data itself is the document list
    elif isinstance(api_data, dict):
        documents = api_data.get('documents') or []
    else:
        documents = []

    # No documents found — save "No doc found" and move on
    if not documents:
        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {
                "extracted_cv":    "No doc found",
                "extracted_cv_at": datetime.utcnow(),
            }}
        )
        return jsonify({
            "success":         True,
            "email":           email,
            "documents_found": 0,
            "cv_extracted":    False,
            "extracted_cv":    "No doc found",
            "remaining_count": max(0, remaining_total - 1),
            "message":         f"No documents returned by API for {email} — marked as 'No doc found'.",
        })

    # ── Process documents ─────────────────────────────────────────────
    extracted_cv_text = None
    cv_error          = None
    cv_url_found      = None

    for doc_item in documents:
        # Extract only when document_type_name is exactly "Cv"
        doc_name = (doc_item.get('document_type_name') or '').strip()
        doc_url  = doc_item.get('url') or ''

        if doc_name == 'Cv' and doc_url and not extracted_cv_text:
            cv_url_found = doc_url
            try:
                extracted_cv_text = _extract_text_from_url(doc_url, api_headers)
            except Exception as cv_err:
                cv_error = str(cv_err)
                extracted_cv_text = f"[CV extraction failed: {cv_err}]"

    # ── Save only extracted_cv to MongoDB ─────────────────────────────
    col.update_one(
        {"_id": staff["_id"]},
        {"$set": {
            "extracted_cv":      extracted_cv_text or "[no CV document found]",
            "extracted_cv_at":   datetime.utcnow(),
        }}
    )

    return jsonify({
        "success":         True,
        "email":           email,
        "cv_extracted":    bool(extracted_cv_text and not cv_error),
        "cv_url_found":    cv_url_found,
        "cv_error":        cv_error,
        "remaining_count": max(0, remaining_total - 1),
        "message": (
            f"Processed {email} — "
            f"{max(0, remaining_total - 1)} staff still need extraction."
        ),
        "synced_at": datetime.utcnow().isoformat(),
    })


def _extract_text_from_url(url, headers=None):
    """
    Download a CV document from URL, then use Gemini AI to extract
    and structure the full text content.

    Strategy:
      1. Download the file (PDF or DOCX).
      2. Get raw text via pdfplumber / python-docx as a pre-extraction step.
      3. Send that raw text to Gemini 2.5 Flash to clean, structure,
         and return a well-formatted plain-text CV extraction.
      4. If Gemini is unavailable, fall back to raw text only.
    """
    import requests as _req
    import io as _io

    # ── Download file ─────────────────────────────────────────────────
    dl_headers = dict(headers or {})
    dl_headers.pop('Content-Type', None)

    resp = _req.get(url, headers=dl_headers, timeout=60)
    resp.raise_for_status()

    content_type = resp.headers.get('Content-Type', '').lower()
    raw          = resp.content
    url_lower    = url.lower().split('?')[0]

    # ── Step 1: raw text extraction ───────────────────────────────────
    raw_text = ''

    # PDF
    if 'pdf' in content_type or url_lower.endswith('.pdf'):
        try:
            import pdfplumber
            with pdfplumber.open(_io.BytesIO(raw)) as pdf:
                raw_text = chr(10).join( page.extract_text() or '' for page in pdf.pages ).strip()
        except Exception:
            try:
                import PyPDF2
                reader   = PyPDF2.PdfReader(_io.BytesIO(raw))
                raw_text = chr(10).join( page.extract_text() or '' for page in reader.pages ).strip()
            except Exception:
                pass

    # DOCX
    elif ('wordprocessingml' in content_type or
          url_lower.endswith('.docx') or url_lower.endswith('.doc')):
        try:
            from docx import Document as _DocxDoc
            d        = _DocxDoc(_io.BytesIO(raw))
            raw_text = chr(10).join(p.text for p in d.paragraphs).strip()
        except Exception:
            pass

    # Plain text
    elif 'text' in content_type:
        raw_text = raw.decode('utf-8', errors='replace').strip()

    # Last resort — try PDF
    if not raw_text:
        try:
            import pdfplumber
            with pdfplumber.open(_io.BytesIO(raw)) as pdf:
                raw_text = chr(10).join( page.extract_text() or '' for page in pdf.pages ).strip()
        except Exception:
            pass

    if not raw_text:
        raise RuntimeError(
            f"Could not extract any text from document (content-type: {content_type})"
        )

    # ── Step 2: Gemini AI extraction & structuring ────────────────────
    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        # No Gemini key — return raw text as-is
        return raw_text

    try:
        from google import genai as _genai
        client = _genai.Client(api_key=gemini_key)

        prompt = f"""You are a professional CV parser.

The text below was extracted from a candidate's CV document (PDF or DOCX).
The text may be messy, have formatting issues, or be partially garbled from extraction.

Your task:
1. Read the raw extracted text carefully.
2. Identify and structure all CV content into clean, readable plain text.
3. Preserve ALL factual information exactly as stated — do NOT add, invent, or change any facts.
4. Format it with clear section headings (PERSONAL DETAILS, PROFESSIONAL PROFILE, EDUCATION & QUALIFICATIONS, PROFESSIONAL EXPERIENCE, TRAINING & CERTIFICATIONS, KEY SKILLS, ADDITIONAL INFORMATION) where the content exists.
5. If a section's content is not present in the raw text, omit that section entirely.
6. Return ONLY the clean structured CV text — no preamble, no commentary.

RAW EXTRACTED TEXT:
{raw_text[:12000]}
"""

        response   = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        gemini_out = (response.text or '').strip()

        # Return Gemini-structured text; fall back to raw if empty
        return gemini_out if gemini_out else raw_text

    except Exception as gemini_err:
        # Gemini failed — return raw text with a note
        return f"[Gemini extraction failed: {gemini_err}]\n\n{raw_text}"



# ── Cron: Generate AI CV one staff at a time ──────────────────────────

@admin_bp.route('/live-staffs/cron/generate-cv', methods=['GET', 'POST'])
def live_staff_cron_generate_cv():
    """
    Cron job — generates an AI CV for ONE staff member per call.

    Logic:
      1. Find the first live_staffs record where gcs_blob is missing
         in live_staff_ai_cvs (i.e. no CV generated yet).
      2. Call Gemini to write the CV.
      3. Build DOCX, upload to GCS, save metadata.
      4. Trigger HSE document upload in background.
      5. Return remaining_count so you know how many are left.

    Call every N seconds via cron until remaining_count == 0.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req

    # ── Auth ──────────────────────────────────────────────────────────
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        return jsonify({"success": False,
                        "error": "GEMINI_API_KEY not set"}), 500

    staffs_col  = _staffs_col()
    ai_cvs_col  = _ai_cvs_col()

    # ── Find next staff without a generated CV ────────────────────────
    # Get set of staff_ids that already have a CV
    existing_ids = set(
        str(r['staff_id'])
        for r in ai_cvs_col.find({}, {"staff_id": 1})
        if r.get('staff_id')
    )

    staff = staffs_col.find_one(
        {"_id": {"$nin": [ObjectId(i) for i in existing_ids if len(i) == 24]}},
        {"email": 1, "section_1_personal_details": 1, "user_type": 1,
         "employee_code": 1, "extracted_cv": 1}
    )

    remaining_total = staffs_col.count_documents({}) - len(existing_ids)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff already have AI CVs — nothing to do.",
            "remaining_count": 0,
            "processed":       None,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or '')

    if not full_name and not email:
        # Mark as skipped so it won't block the queue
        ai_cvs_col.insert_one({
            "staff_id":     staff_id,
            "staff_name":   '',
            "cv_text":      '[skipped — no name or email]',
            "gcs_blob":     '',
            "generated_at": datetime.utcnow(),
        })
        return jsonify({
            "success":         True,
            "message":         f"Skipped {staff_id} — no name or email",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Build full doc object for the CV builder ───────────────────────
    full_doc = staffs_col.find_one({"_id": staff['_id']})
    if not full_doc:
        return jsonify({"success": False, "error": "Staff record disappeared"}), 500

    # ── Call Gemini ───────────────────────────────────────────────────
    try:
        s1_f = full_doc.get('section_1_personal_details') or {}
        s3   = full_doc.get('section_3_professional_registration') or {}
        s4   = full_doc.get('section_4_qualifications') or {}
        s5   = full_doc.get('section_5_employment_history') or {}
        s8   = full_doc.get('section_8_garda_vetting_police_clearance') or {}
        s9   = full_doc.get('section_9_occupational_health') or {}
        s10  = full_doc.get('section_10_mandatory_training') or {}
        visa = s1_f.get('work_permit_visa_status') or {}

        def _vv(val):
            if val is None: return ''
            return str(val).strip()

        user_type   = _vv(full_doc.get('user_type'))
        emp_code    = _vv(full_doc.get('employee_code'))
        address     = _vv(s1_f.get('address'))
        mobile      = _vv(s1_f.get('mobile_number'))
        nationality = _vv(s1_f.get('nationality'))
        total_exp   = _vv(s5.get('total_experience'))
        divisions   = ', '.join(s3.get('divisions_registered_in') or [])
        reg_pin     = _vv(s3.get('registration_number_pin'))
        reg_exp     = _vv(s3.get('registration_expiry_date'))
        nmbi        = 'Yes' if s3.get('nmbi_active_declaration') else 'No'
        visa_type   = _vv(visa.get('visa_type'))
        perm_work   = _vv(visa.get('permission_to_work'))
        garda       = 'Yes' if s8.get('garda_vetting_submitted') else 'No'
        fit         = 'Yes' if s9.get('fit_for_nursing_duties') else 'No'
        dob         = _vv(s1_f.get('date_of_birth'))

        qual_lines = []
        for qk in ['nursing_degree', 'postgraduate_qualification', 'other_qualification']:
            q = s4.get(qk) or {}
            if q.get('qualification') or q.get('institution'):
                qual_lines.append(
                    f"  - {_vv(q.get('qualification'))} | "
                    f"{_vv(q.get('institution'))} | "
                    f"{_vv(q.get('year_completed'))}"
                )
        # Also include NMBI/QQI numbers as qualification context
        nmbi_num = _vv(doc.get('nmbi_number') or s3.get('registration_number_pin') or '')
        qqi_num  = _vv(doc.get('qqi_number') or '')
        if nmbi_num and not any('nmbi' in l.lower() or 'registration' in l.lower() for l in qual_lines):
            qual_lines.append(f"  - NMBI Registration PIN: {nmbi_num}")
        if qqi_num and not any('qqi' in l.lower() for l in qual_lines):
            qual_lines.append(f"  - QQI Level 5 Certificate No: {qqi_num}")

        entries = [e for e in (s5.get('entries') or [])
                   if e.get('employer') or e.get('position')]
        exp_lines = []
        for e in entries:
            exp_lines.append(
                f"  - {_vv(e.get('position'))} at {_vv(e.get('employer'))} "
                f"({_vv(e.get('from'))} - {_vv(e.get('to') or 'Present')})"
            )

        TLABELS = {
            'manual_handling': 'Manual Handling', 'cpr_bls': 'CPR / BLS',
            'fire_safety': 'Fire Safety',
            'infection_prevention_control': 'Infection Prevention & Control',
            'hand_hygiene': 'Hand Hygiene', 'safeguarding': 'Safeguarding',
            'children_first': 'Children First', 'cyber_security': 'Cyber Security',
            'dignity_at_work': 'Dignity at Work', 'open_disclosure': 'Open Disclosure',
            'mapa_pmav': 'MAPA / PMAV',
        }
        certs = [label for k, label in TLABELS.items() if s10.get(k)][:6]

        extracted_cv = _v(full_doc.get('extracted_cv') or '')
        has_extracted = (
            extracted_cv and
            not extracted_cv.startswith('[') and
            extracted_cv != 'No doc found'
        )

        data_summary = f"""
Candidate: {full_name}
Role / User Type: {user_type}
Employee Code: {emp_code}
Address: {address}
Mobile: {mobile}
Email: {email}
Nationality: {nationality}
Total Experience: {total_exp}
Divisions / Speciality: {divisions}
Registration PIN: {reg_pin}
Registration Expiry: {reg_exp}
NMBI Active Declaration: {nmbi}
Permission to Work: {perm_work}
Visa / Stamp Type: {visa_type}
Garda Vetted: {garda}
Fit for Nursing Duties: {fit}

Qualifications:
{chr(10).join(qual_lines) if qual_lines else '  None recorded'}

Employment History:
{chr(10).join(exp_lines) if exp_lines else '  None recorded'}

Training & Certifications:
{chr(10).join('  - ' + c for c in certs) if certs else '  None recorded'}
""".strip()

        extracted_cv_section = f"""

EXTRACTED CV TEXT (use as PRIMARY source for PROFESSIONAL EXPERIENCE, TRAINING & CERTIFICATIONS and KEY SKILLS):
{extracted_cv[:8000]}
""" if has_extracted else ""

        prompt = f"""You are an expert professional CV writer specialising in Irish healthcare staffing.

STRICT RULE — NO HALLUCINATION:
Use ONLY the exact facts in CANDIDATE DATA. Do not invent anything.

SECTION SOURCE RULES:
- PERSONAL DETAILS, PROFESSIONAL PROFILE: use CANDIDATE DATA.
- EDUCATION & QUALIFICATIONS: MANDATORY — ALWAYS include this section. Use qualifications from CANDIDATE DATA.
  If Qualifications says "None recorded", infer from their role:
    * Nurse → "Bachelor of Nursing Science | [University based on nationality] | [estimated year]"
    * Healthcare Assistant → "QQI Level 5 in Healthcare Support | [College of Further Education] | [estimated year]"
  Include NMBI PIN or QQI certificate number as a separate line if provided.
- PROFESSIONAL EXPERIENCE: {"Extract directly from EXTRACTED CV TEXT — copy job titles, employers, dates, and duties word-for-word." if has_extracted else "Use Employment History from CANDIDATE DATA. Write 5-6 appropriate duties per role."}
- TRAINING & CERTIFICATIONS: {"Extract directly from EXTRACTED CV TEXT." if has_extracted else "Use Training & Certifications from CANDIDATE DATA only."}
- KEY SKILLS: {"Extract directly from EXTRACTED CV TEXT." if has_extracted else "Write 8-10 bullet points from their role and certifications."}

Structure (EXACT UPPERCASE headings on their own line):
PERSONAL DETAILS
PROFESSIONAL PROFILE
EDUCATION & QUALIFICATIONS
PROFESSIONAL EXPERIENCE
TRAINING & CERTIFICATIONS
KEY SKILLS
ADDITIONAL INFORMATION

Rules:
- PERSONAL DETAILS: "Label: Value" per line. Skip blank fields.
- PROFESSIONAL PROFILE: 2 paragraphs, FIRST PERSON. Genuine personal statement.
- EDUCATION & QUALIFICATIONS: ALWAYS write at least one entry. Format: Qualification Name | Institution | Year
  Include registration numbers (NMBI PIN, QQI Certificate No) on a separate line if available.
- PROFESSIONAL EXPERIENCE: Job Title: / Employer: / Dates: / Duties: / - duty
- TRAINING & CERTIFICATIONS: Bullet list only.
- KEY SKILLS: 8-10 bullets.
- ADDITIONAL INFORMATION: Include ONLY these two lines, nothing else:
Driving Licence: No
Own Transport: No

---
CANDIDATE DATA:
{data_summary}
{extracted_cv_section}
---

Output CV text only. No preamble, no markdown symbols.
"""

        from google import genai as google_genai
        client   = google_genai.Client(api_key=gemini_key)
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        cv_text = response.text.strip()

    except Exception as e:
        # Mark as attempted so cron moves on
        ai_cvs_col.insert_one({
            "staff_id":     staff_id,
            "staff_name":   full_name,
            "cv_text":      f"[Gemini error: {e}]",
            "gcs_blob":     '',
            "generated_at": datetime.utcnow(),
        })
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"Gemini error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Build DOCX and upload to GCS ─────────────────────────────────
    try:
        docx_bytes  = _build_ai_cv_docx(full_doc, cv_text)
        safe_name   = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        cv_filename = f"{safe_name}.docx"
        gcs_blob    = f"cv/{cv_filename}"
        _gcs_upload(
            gcs_blob, docx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        ai_cvs_col.insert_one({
            "staff_id":     staff_id,
            "staff_name":   full_name,
            "cv_text":      cv_text,
            "gcs_blob":     '',
            "generated_at": datetime.utcnow(),
        })
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"DOCX/GCS error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Save to MongoDB ───────────────────────────────────────────────
    ai_cvs_col.insert_one({
        "staff_id":      staff_id,
        "staff_name":    full_name,
        "employee_code": emp_code,
        "cv_text":       cv_text,
        "cv_filename":   cv_filename,
        "gcs_blob":      gcs_blob,
        "generated_at":  datetime.utcnow(),
    })

    # ── Background HSE document push ─────────────────────────────────
    _push_hse_document_background(
        staff_id_str=staff_id,
        doc_type_key='cv',
        docx_bytes=docx_bytes,
        staff_name=full_name,
        mongo_id=staff_id,
        email=email,
    )

    return jsonify({
        "success":         True,
        "email":           email,
        "staff_name":      full_name,
        "cv_filename":     cv_filename,
        "remaining_count": max(0, remaining_total - 1),
        "message": (
            f"CV generated for {full_name} — "
            f"{max(0, remaining_total - 1)} staff still need a CV."
        ),
        "generated_at": datetime.utcnow().isoformat(),
    })



# ── Cron: Generate AI Interview Notes one staff at a time ─────────────

@admin_bp.route('/live-staffs/cron/generate-interview', methods=['GET', 'POST'])
def live_staff_cron_generate_interview():
    """
    Cron job — generates AI Interview Notes for ONE staff member per call.

    Logic:
      1. Find the first live_staffs record with no entry in live_staff_ai_interviews.
      2. Call Gemini 2.5 Flash to write the interview notes.
      3. Build DOCX, upload to GCS interviews/ folder, save metadata.
      4. Trigger HSE background upload.
      5. Return remaining_count.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """

    # ── Auth ──────────────────────────────────────────────────────────
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    staffs_col     = _staffs_col()
    interviews_col = _ai_interviews_col()

    # ── Find next staff without interview notes ───────────────────────
    existing_ids = set(
        str(r['staff_id'])
        for r in interviews_col.find({}, {"staff_id": 1})
        if r.get('staff_id')
    )

    staff = staffs_col.find_one(
        {"_id": {"$nin": [ObjectId(i) for i in existing_ids if len(i) == 24]}},
    )

    remaining_total = staffs_col.count_documents({}) - len(existing_ids)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff already have interview notes — nothing to do.",
            "remaining_count": 0,
            "processed":       None,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    s3        = staff.get('section_3_professional_registration') or {}
    s5        = staff.get('section_5_employment_history') or {}
    s8        = staff.get('section_8_garda_vetting_police_clearance') or {}
    s9        = staff.get('section_9_occupational_health') or {}
    s10       = staff.get('section_10_mandatory_training') or {}
    visa      = s1.get('work_permit_visa_status') or {}

    full_name   = _v(s1.get('full_name') or '')
    email       = _v(staff.get('email') or '')
    user_type   = _v(staff.get('user_type') or 'Nurse')
    emp_code    = _v(staff.get('employee_code') or '')
    address     = _v(s1.get('address') or '')
    nationality = _v(s1.get('nationality') or '')
    reg_pin     = _v(s3.get('registration_number_pin') or '')
    visa_type   = _v(visa.get('visa_type') or '')
    divisions   = ', '.join(s3.get('divisions_registered_in') or [])
    total_exp   = _v(s5.get('total_experience') or '')
    nmbi        = 'Yes' if s3.get('nmbi_active_declaration') else 'No'
    garda       = 'Yes' if s8.get('garda_vetting_submitted') else 'No'
    bls         = 'Yes' if s10.get('cpr_bls') else 'No'
    manual      = 'Yes' if s10.get('manual_handling') else 'No'
    fit         = 'Yes' if s9.get('fit_for_nursing_duties') else 'No'

    entries = [e for e in (s5.get('entries') or [])
               if e.get('employer') or e.get('position')]
    exp_lines = []
    for e in entries[:5]:
        pos    = _v(e.get('position'))
        emp    = _v(e.get('employer'))
        d_from = _v(e.get('from'))
        d_to   = _v(e.get('to'))
        if pos or emp:
            exp_lines.append(f"  - {pos} at {emp} ({d_from} – {d_to or 'Present'})")

    TLABELS = {
        'manual_handling': 'Manual Handling', 'cpr_bls': 'BLS/CPR',
        'safeguarding': 'Safeguarding', 'fire_safety': 'Fire Safety',
        'infection_prevention_control': 'Infection Prevention & Control',
    }
    certs = [label for k, label in TLABELS.items() if s10.get(k)]

    # County from address
    county = ''
    if address:
        parts = [p.strip() for p in address.replace(',', ' ').split()]
        for p in parts:
            if p.lower().startswith('co.') or p.lower() == 'county':
                idx = parts.index(p)
                if idx + 1 < len(parts):
                    county = parts[idx + 1]
                break
        if not county and parts:
            county = parts[-1]

    data_summary = f"""
Name: {full_name}
Role / User Type: {user_type}
Address / Location: {address}
Nationality: {nationality}
Visa / Stamp Type: {visa_type}
NMBI Registration PIN: {reg_pin}
NMBI Registration Active: {nmbi}
Divisions / Speciality: {divisions}
Total Experience: {total_exp}
Garda Vetted: {garda}
BLS/CPR on file: {bls}
Manual Handling on file: {manual}
Fit for Duties: {fit}

Employment History:
{chr(10).join(exp_lines) if exp_lines else '  None recorded'}

Certifications on file: {', '.join(certs) if certs else 'None recorded'}
""".strip()

    prompt = f"""You are an experienced nursing recruitment consultant at Xpress Health, Ireland.

Using ONLY the verified candidate data below, complete a realistic, professional nurse interview notes template.
Answers must be written as if the candidate themselves just answered each question in a live phone/video interview.
Write naturally — conversational but professional. First person where appropriate ("I have", "I work", "I currently").

STRICT RULES — NO HALLUCINATION:
- Use ONLY the facts provided in CANDIDATE DATA. Do not invent employers, dates, locations, or qualifications.
- If data is missing for a field, write a realistic professional answer appropriate to their role and experience level without inventing specific names.
- Clinical question answers must be clinically appropriate for a {user_type}.
- Assessment scores: pick a random realistic score for each between 3.5 and 5.0 in 0.5 increments (e.g. 3.5/5, 4/5, 4.5/5, 5/5). Vary the three scores — do not give the same score to all three.
- Do NOT add any text outside the template structure below.

Output ONLY the completed template below — no preamble, no explanations, no markdown symbols:

---
Completed {user_type} Interview

Name: [full name]
Location: [county/city from address]
NMBI PIN: [registration pin or N/A]
Visa Status: [visa type]

Experience

1. Tell me about your nursing experience.
[Write a 4–6 sentence answer in first person describing their experience, speciality, and current/most recent role. Use only the data provided.]

2. How many years in Ireland?
[Write a realistic answer based on employment history dates. If Ireland-based work is evident, state it clearly.]

3. Acute, Nursing Home, Community, or Mental Health?
[Based on employment history, state the most relevant care setting.]

Clinical Questions

1. How would you manage a deteriorating patient?
[Write a clinically accurate 4–5 sentence answer appropriate for a {user_type}. Use recognised frameworks (ABCDE, NEWS2, ISBAR) where appropriate.]

2. What would you do if you witnessed a medication error?
[Write a clinically accurate 4–5 sentence answer covering patient safety, reporting, documentation, and prevention.]

Compliance
NMBI Registration: [Yes/No based on data]
BLS/CPR: [Yes/No based on data]
Manual Handling: [Yes/No based on data]
Garda Vetting: [Yes/No based on data]
References: Yes

Availability
Preferred counties: [county from address, or nearest city]
Day/Night/Both: Both
Earliest start date: Immediate

Assessment
Communication: [3.5/5 or 4/5 or 4.5/5 or 5/5 — vary randomly]
Clinical Knowledge: [3.5/5 or 4/5 or 4.5/5 or 5/5 — vary randomly]
Experience: [3.5/5 or 4/5 or 4.5/5 or 5/5 — vary randomly]
Suitable: Yes
---

CANDIDATE DATA (use ONLY this):
{data_summary}
"""

    if not full_name and not email:
        interviews_col.insert_one({
            "staff_id":       staff_id,
            "staff_name":     '',
            "interview_text": '[skipped — no name or email]',
            "gcs_blob":       '',
            "generated_at":   datetime.utcnow(),
        })
        return jsonify({
            "success":         True,
            "message":         f"Skipped {staff_id} — no name or email",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Call Gemini ───────────────────────────────────────────────────
    try:
        from google import genai as google_genai
        client         = google_genai.Client(api_key=gemini_key)
        response       = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        interview_text = response.text.strip().strip('-').strip()

    except Exception as e:
        interviews_col.insert_one({
            "staff_id":       staff_id,
            "staff_name":     full_name,
            "interview_text": f"[Gemini error: {e}]",
            "gcs_blob":       '',
            "generated_at":   datetime.utcnow(),
        })
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"Gemini error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Build DOCX and upload to GCS ─────────────────────────────────
    try:
        docx_bytes = _build_interview_docx(staff, interview_text)
        safe_name  = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        filename   = f"Interview_{safe_name}_{staff_id}.docx"
        gcs_blob   = f"interviews/{filename}"
        _gcs_upload(
            gcs_blob, docx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        interviews_col.insert_one({
            "staff_id":       staff_id,
            "staff_name":     full_name,
            "interview_text": interview_text,
            "gcs_blob":       '',
            "generated_at":   datetime.utcnow(),
        })
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"DOCX/GCS error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Save to MongoDB ───────────────────────────────────────────────
    interviews_col.insert_one({
        "staff_id":       staff_id,
        "staff_name":     full_name,
        "employee_code":  emp_code,
        "interview_text": interview_text,
        "filename":       filename,
        "gcs_blob":       gcs_blob,
        "generated_at":   datetime.utcnow(),
    })

    # ── Background HSE document push ──────────────────────────────────
    _push_hse_document_background(
        staff_id_str=staff_id,
        doc_type_key='interview',
        docx_bytes=docx_bytes,
        staff_name=full_name,
        mongo_id=staff_id,
        email=email,
    )

    return jsonify({
        "success":         True,
        "email":           email,
        "staff_name":      full_name,
        "filename":        filename,
        "remaining_count": max(0, remaining_total - 1),
        "message": (
            f"Interview notes generated for {full_name} — "
            f"{max(0, remaining_total - 1)} staff still need interview notes."
        ),
        "generated_at": datetime.utcnow().isoformat(),
    })



# ── Cron: Fetch signature + points one staff at a time ───────────────

@admin_bp.route('/live-staffs/cron/fetch-signature', methods=['GET', 'POST'])
def live_staff_cron_fetch_signature():
    """
    Cron job — fetches signature URL and points for ONE staff member per call.

    Logic:
      1. Find first live_staffs record where:
         - signature_url is missing/empty AND
         - staff_id (xn_staff_id) is set (needed for the API call)
      2. POST to {LIVE_STAFF_URL}/ai/recruitments/detail with {"_id": staff_id}
      3. Download signature image from signature_url
      4. Upload image to GCS under signatures/ folder
      5. Save GCS URL + points to live_staffs collection
      6. Return remaining_count

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req

    # ── Auth ──────────────────────────────────────────────────────────
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')

    if not base_url:
        return jsonify({"success": False,
                        "error": "LIVE_STAFF_URL not set in environment"}), 500

    col = _staffs_col()

    # ── Count total needing signature ─────────────────────────────────
    needs_sig_query = {
        "$and": [
            # Missing signature_gcs_blob (not yet saved to GCS)
            {"$or": [
                {"signature_gcs_blob": {"$exists": False}},
                {"signature_gcs_blob": None},
                {"signature_gcs_blob": ""},
            ]},
            # Skip records already marked as no signature or errors
            {"$or": [
                {"signature_url": {"$exists": False}},
                {"signature_url": None},
                {"signature_url": ""},
                # Has a real URL (not a placeholder error string)
                {"signature_url": {"$not": {"$regex": "^\\["}}},
            ]},
            # Must have a staff_id to call the API
            {"$or": [
                {"staff_id":    {"$exists": True, "$ne": None, "$ne": ""}},
                {"xn_staff_id": {"$exists": True, "$ne": None, "$ne": ""}},
            ]},
        ]
    }
    remaining_total = col.count_documents(needs_sig_query)

    # ── Find next staff ───────────────────────────────────────────────
    staff = col.find_one(needs_sig_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff signatures already fetched — nothing to do.",
            "remaining_count": 0,
        })

    mongo_id  = str(staff['_id'])
    xn_id     = _v(staff.get('staff_id') or staff.get('xn_staff_id') or '')
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or staff.get('email') or mongo_id)
    email     = _v(staff.get('email') or '')

    if not xn_id:
        # No staff_id — skip to avoid blocking queue
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {"signature_url": "[skipped — no staff_id]"}}
        )
        return jsonify({
            "success":         True,
            "message":         f"Skipped {full_name} — no staff_id available",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Call XN Portal detail API ─────────────────────────────────────
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(
            f"{base_url}/ai/recruitments/detail",
            json={"_id": xn_id},
            headers=api_headers,
            timeout=30,
        )
        # Retry with GET if POST not allowed
        if resp.status_code == 405:
            resp = _req.get(
                f"{base_url}/ai/recruitments/detail",
                params={"_id": xn_id},
                headers=api_headers,
                timeout=30,
            )
        resp.raise_for_status()
        data = resp.json()

    except Exception as api_err:
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {"signature_url": f"[API error: {api_err}]"}}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"API call failed: {api_err}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {"signature_url": f"[API error: {data.get('message', 'unknown')}]"}}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           data.get('message', 'API returned success=false'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data      = data.get('data') or {}
    signature_url = (api_data.get('signature_url') or '').strip()
    points        = api_data.get('points')

    if not signature_url:
        # No signature — mark and move on
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {
                "signature_url": "[no signature found]",
                "points":        points,
            }}
        )
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "signature_saved": False,
            "points":          points,
            "message":         f"No signature URL returned for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Download signature image ───────────────────────────────────────
    try:
        img_resp = _req.get(signature_url, timeout=30)
        img_resp.raise_for_status()
        img_bytes = img_resp.content

        # Detect extension from URL or content-type
        content_type = img_resp.headers.get('Content-Type', 'image/png').lower()
        if 'jpeg' in content_type or 'jpg' in content_type:
            ext = 'jpg'
        elif 'webp' in content_type:
            ext = 'webp'
        else:
            ext = 'png'

        # Also try to get extension from URL
        url_path = signature_url.split('?')[0].lower()
        if url_path.endswith('.jpg') or url_path.endswith('.jpeg'):
            ext = 'jpg'
        elif url_path.endswith('.webp'):
            ext = 'webp'
        elif url_path.endswith('.png'):
            ext = 'png'

    except Exception as dl_err:
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {
                "signature_url": f"[download error: {dl_err}]",
                "points":        points,
            }}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"Signature download failed: {dl_err}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Upload to GCS ─────────────────────────────────────────────────
    try:
        safe_name  = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        gcs_blob   = f"signatures/{safe_name}_{mongo_id}.{ext}"
        gcs_ct     = f"image/{ext}"
        _gcs_upload(gcs_blob, img_bytes, content_type=gcs_ct)

        # Build a GCS public-style path to store (signed URL generated on demand)
        gcs_signature_path = gcs_blob

    except Exception as gcs_err:
        # GCS failed — still save original URL as fallback
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {
                "signature_url":      signature_url,
                "signature_gcs_blob": "",
                "points":             points,
                "signature_synced_at": datetime.utcnow(),
            }}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"GCS upload failed: {gcs_err} — original URL saved",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Save to live_staffs ───────────────────────────────────────────
    col.update_one(
        {"_id": staff['_id']},
        {"$set": {
            "signature_url":       signature_url,       # original CDN URL
            "signature_gcs_blob":  gcs_signature_path,  # our GCS copy
            "points":              points,
            "signature_synced_at": datetime.utcnow(),
        }}
    )

    return jsonify({
        "success":         True,
        "email":           email,
        "staff_name":      full_name,
        "signature_saved": True,
        "gcs_blob":        gcs_signature_path,
        "points":          points,
        "remaining_count": max(0, remaining_total - 1),
        "message": (
            f"Signature saved for {full_name} — "
            f"{max(0, remaining_total - 1)} staff still need signatures."
        ),
        "synced_at": datetime.utcnow().isoformat(),
    })



# ── Cron: Check points vs Last PointScale and flag mismatches ─────────

@admin_bp.route('/live-staffs/cron/check-points', methods=['GET', 'POST'])
def live_staff_cron_check_points():
    """
    Cron job — checks one staff member's current points against their
    Last PointScale from the spreadsheet and flags mismatches.

    Logic:
      1. Find first live_staffs record where points_checked is not True.
      2. Call XN Portal detail API to get current points.
      3. Compare with last_point_scale stored in live_staffs.
      4. Save: current_points, points_mismatch (bool), points_checked (True).

    Import last_point_scale into live_staffs first using the import route
    or via the spreadsheet (email + last_point_scale fields).

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req

    # ── Auth ──────────────────────────────────────────────────────────
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')

    if not base_url:
        return jsonify({"success": False,
                        "error": "LIVE_STAFF_URL not set"}), 500

    col = _staffs_col()

    # Count remaining
    remaining_total = col.count_documents(
        {"$or": [
            {"points_checked": {"$exists": False}},
            {"points_checked": False},
            {"points_checked": None},
        ]}
    )

    # Find next unchecked
    staff = col.find_one(
        {"$or": [
            {"points_checked": {"$exists": False}},
            {"points_checked": False},
            {"points_checked": None},
        ]}
    )

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff points checked — nothing to do.",
            "remaining_count": 0,
        })

    mongo_id  = str(staff['_id'])
    xn_id     = _v(staff.get('staff_id') or staff.get('xn_staff_id') or '')
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or staff.get('email') or mongo_id)
    email     = _v(staff.get('email') or '')
    last_ps   = staff.get('last_point_scale')   # stored from spreadsheet import

    if not xn_id:
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {
                "points_checked": True,
                "points_mismatch": None,
                "points_check_note": "skipped — no staff_id",
            }}
        )
        return jsonify({
            "success":         True,
            "message":         f"Skipped {full_name} — no staff_id",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Call XN Portal detail API ──────────────────────────────────────
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(
            f"{base_url}/ai/recruitments/detail",
            json={"_id": xn_id},
            headers=api_headers,
            timeout=30,
        )
        if resp.status_code == 405:
            resp = _req.get(
                f"{base_url}/ai/recruitments/detail",
                params={"_id": xn_id},
                headers=api_headers,
                timeout=30,
            )
        resp.raise_for_status()
        data = resp.json()

    except Exception as api_err:
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {
                "points_checked":    True,
                "points_mismatch":   None,
                "points_check_note": f"API error: {api_err}",
            }}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"API error: {api_err}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {
                "points_checked":    True,
                "points_mismatch":   None,
                "points_check_note": f"API error: {data.get('message')}",
            }}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           data.get('message', 'API returned success=false'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data       = data.get('data') or {}
    current_points = api_data.get('points')

    # ── Compare ────────────────────────────────────────────────────────
    mismatch = None
    if last_ps is not None and current_points is not None:
        try:
            mismatch = float(current_points) != float(last_ps)
        except (TypeError, ValueError):
            mismatch = None

    col.update_one(
        {"_id": staff['_id']},
        {"$set": {
            "points":             current_points,
            "last_point_scale":   last_ps,
            "points_mismatch":    mismatch,
            "points_checked":     True,
            "points_checked_at":  datetime.utcnow(),
        }}
    )

    return jsonify({
        "success":          True,
        "email":            email,
        "staff_name":       full_name,
        "last_point_scale": last_ps,
        "current_points":   current_points,
        "mismatch":         mismatch,
        "remaining_count":  max(0, remaining_total - 1),
        "message": (
            f"{'⚠ MISMATCH' if mismatch else '✓ Match'} — "
            f"{full_name}: last={last_ps}, current={current_points} — "
            f"{max(0, remaining_total - 1)} remaining."
        ),
        "checked_at": datetime.utcnow().isoformat(),
    })



# ── Cron: Generate Application Form one staff at a time ───────────────

@admin_bp.route('/live-staffs/cron/generate-appform', methods=['GET', 'POST'])
def live_staff_cron_generate_appform():
    """
    Cron job — generates an Application Form for ONE staff member per call.

    Logic:
      1. Find first live_staffs record with no entry in live_staff_ai_appforms.
      2. If staff has a signature_gcs_blob, download it from GCS.
      3. Build the Application Form DOCX (with real signature image if available).
      4. Upload to GCS appforms/ folder, save metadata to MongoDB.
      5. Fire HSE document background upload.
      6. Return remaining_count.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """

    # ── Auth ──────────────────────────────────────────────────────────
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    staffs_col   = _staffs_col()
    appforms_col = _ai_appforms_col()

    # ── Find next staff without an application form ───────────────────
    existing_ids = set(
        str(r['staff_id'])
        for r in appforms_col.find({}, {"staff_id": 1})
        if r.get('staff_id')
    )

    staff = staffs_col.find_one(
        {"_id": {"$nin": [ObjectId(i) for i in existing_ids if len(i) == 24]}}
    )

    remaining_total = staffs_col.count_documents({}) - len(existing_ids)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff already have application forms — nothing to do.",
            "remaining_count": 0,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or '')
    emp_code  = _v(staff.get('employee_code') or '')

    if not full_name and not email:
        appforms_col.insert_one({
            "staff_id":     staff_id,
            "staff_name":   '',
            "gcs_blob":     '',
            "generated_at": datetime.utcnow(),
        })
        return jsonify({
            "success":         True,
            "message":         f"Skipped {staff_id} — no name or email",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Download signature from GCS if available ──────────────────────
    signature_bytes = None
    sig_blob        = _v(staff.get('signature_gcs_blob') or '')
    if sig_blob:
        try:
            signature_bytes = _gcs_download(sig_blob)
        except Exception as sig_err:
            # Non-fatal — form will have blank signature line
            signature_bytes = None

    # ── Build Application Form DOCX ───────────────────────────────────
    try:
        docx_bytes = _build_appform_docx(staff, signature_bytes=signature_bytes)
        safe_name  = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        filename   = f"AppForm_{safe_name}.docx"
        gcs_blob   = f"appforms/{filename}"
        _gcs_upload(
            gcs_blob, docx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        appforms_col.insert_one({
            "staff_id":     staff_id,
            "staff_name":   full_name,
            "gcs_blob":     '',
            "generated_at": datetime.utcnow(),
        })
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"DOCX/GCS error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Save to MongoDB ───────────────────────────────────────────────
    appforms_col.insert_one({
        "staff_id":      staff_id,
        "staff_name":    full_name,
        "employee_code": emp_code,
        "filename":      filename,
        "gcs_blob":      gcs_blob,
        "has_signature": bool(signature_bytes),
        "generated_at":  datetime.utcnow(),
    })

    # ── Background HSE document push ──────────────────────────────────
    _push_hse_document_background(
        staff_id_str=staff_id,
        doc_type_key='appform',
        docx_bytes=docx_bytes,
        staff_name=full_name,
        mongo_id=staff_id,
        email=email,
    )

    return jsonify({
        "success":         True,
        "email":           email,
        "staff_name":      full_name,
        "filename":        filename,
        "has_signature":   bool(signature_bytes),
        "remaining_count": max(0, remaining_total - 1),
        "message": (
            f"Application form generated for {full_name} "
            f"({'with' if signature_bytes else 'without'} signature) — "
            f"{max(0, remaining_total - 1)} staff still need forms."
        ),
        "generated_at": datetime.utcnow().isoformat(),
    })



# ── Cron: Sync Garda Vetting & Police Clearance one staff at a time ───

@admin_bp.route('/live-staffs/cron/sync-vetting', methods=['GET', 'POST'])
def live_staff_cron_sync_vetting():
    """
    Cron job — processes ONE staff member per call.

    For each staff member calls the XN Portal document-list API and checks:
      - document_type_name == "Garda Vetting Document"
        * status approved  → garda_vetting = 1
        * status rejected/pending → garda_vetting = 0
        * expiry_date passed today → garda_vetting_expired = 1 else 0
      - document_type_name == "Police Clearance Certificate ( From Country Of Birth )"
        * same logic → police_clearance and police_clearance_expired

    Both checked in a single API call per staff member.
    Run one staff at a time until all processed (remaining_count == 0).

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req
    from datetime import date as _date

    # ── Auth ──────────────────────────────────────────────────────────
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')

    if not base_url:
        return jsonify({"success": False,
                        "error": "LIVE_STAFF_URL not set"}), 500

    col = _staffs_col()

    # ── Find next staff without vetting synced ────────────────────────
    pending_query = {
        "$or": [
            {"garda_vetting_synced": {"$exists": False}},
            {"garda_vetting_synced": False},
            {"garda_vetting_synced": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)

    staff = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff vetting documents synced — nothing to do.",
            "remaining_count": 0,
        })

    mongo_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or '')

    if not email:
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {"garda_vetting_synced": True,
                      "garda_vetting_sync_note": "skipped — no email"}}
        )
        return jsonify({
            "success":         True,
            "message":         f"Skipped {mongo_id} — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Call XN Portal API ────────────────────────────────────────────
    endpoint   = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(
            endpoint,
            json={"email": email},
            headers=api_headers,
            timeout=30,
        )
        if resp.status_code == 405:
            resp = _req.get(
                endpoint,
                params={"email": email},
                headers=api_headers,
                timeout=30,
            )
        resp.raise_for_status()
        data = resp.json()
    except Exception as api_err:
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {"garda_vetting_synced": True,
                      "garda_vetting_sync_note": f"API error: {api_err}"}}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"API error: {api_err}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {"garda_vetting_synced": True,
                      "garda_vetting_sync_note": f"API error: {data.get('message')}"}}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           data.get('message', 'API returned success=false'),
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Parse documents ────────────────────────────────────────────────
    api_data  = data.get('data') or {}
    if isinstance(api_data, list):
        documents = api_data
    elif isinstance(api_data, dict):
        documents = api_data.get('documents') or []
    else:
        documents = []

    today = _date.today()

    def _parse_expiry(expiry_str):
        """Parse expiry date string, return date object or None."""
        if not expiry_str:
            return None
        for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y', '%Y/%m/%d'):
            try:
                from datetime import datetime as _dt
                return _dt.strptime(str(expiry_str)[:10], fmt).date()
            except ValueError:
                continue
        return None

    def _check_doc(doc_type_name):
        """
        Find a document by type name and return (status_int, expired_int).
        status: 1=approved, 0=rejected/pending/missing
        expired: 1=expired, 0=valid or no expiry
        """
        found = [
            d for d in documents
            if (d.get('document_type_name') or '').strip().lower()
            == doc_type_name.strip().lower()
        ]
        if not found:
            return None, None   # document not present at all

        # Use the most recent / approved one if multiple
        doc = found[0]
        for d in found:
            if (d.get('status') or '').lower() == 'approved':
                doc = d
                break

        status_str = (doc.get('status') or '').strip().lower()
        approved   = 1 if status_str == 'approved' else 0

        expiry_date = _parse_expiry(doc.get('expiry_date'))
        if expiry_date:
            expired = 1 if expiry_date < today else 0
        else:
            expired = 0

        return approved, expired

    garda_status,  garda_expired  = _check_doc('Garda Vetting Document')
    police_status, police_expired = _check_doc('Police Clearance Certificate ( From Country Of Birth )')

    # ── Save to live_staffs ───────────────────────────────────────────
    update_fields = {
        "garda_vetting_synced":  True,
        "garda_vetting_synced_at": datetime.utcnow(),
    }

    if garda_status is not None:
        update_fields["garda_vetting"]         = garda_status
        update_fields["garda_vetting_expired"] = garda_expired
    else:
        update_fields["garda_vetting"]         = 0
        update_fields["garda_vetting_expired"] = 0

    if police_status is not None:
        update_fields["police_clearance"]         = police_status
        update_fields["police_clearance_expired"] = police_expired
    else:
        update_fields["police_clearance"]         = 0
        update_fields["police_clearance_expired"] = 0

    col.update_one({"_id": staff['_id']}, {"$set": update_fields})

    return jsonify({
        "success":               True,
        "email":                 email,
        "staff_name":            full_name,
        "garda_vetting":         update_fields.get("garda_vetting"),
        "garda_vetting_expired": update_fields.get("garda_vetting_expired"),
        "police_clearance":         update_fields.get("police_clearance"),
        "police_clearance_expired": update_fields.get("police_clearance_expired"),
        "remaining_count":       max(0, remaining_total - 1),
        "message": (
            f"Vetting synced for {full_name} — "
            f"{max(0, remaining_total - 1)} staff remaining."
        ),
        "synced_at": datetime.utcnow().isoformat(),
    })



# ── Cron: Analyse experience one staff at a time ──────────────────────

@admin_bp.route('/live-staffs/cron/analyse-experience', methods=['GET', 'POST'])
def live_staff_cron_analyse_experience():
    """
    Cron job — analyses experience for ONE staff member per call.

    Gemini runs in a background thread so the HTTP response returns
    immediately — prevents 504 gateway timeouts on slow Gemini calls.

    Logic:
      1. Find first record where experience_analysed_at is missing AND
         experience_processing is not True (not already in-flight).
      2. Mark experience_processing = True immediately.
      3. Fire Gemini analysis in background thread.
      4. Return 200 right away with remaining_count.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$and": [
            {"$or": [
                {"experience_analysed_at": {"$exists": False}},
                {"experience_analysed_at": None},
            ]},
            # Skip records currently being processed in background
            {"$or": [
                {"experience_processing": {"$exists": False}},
                {"experience_processing": False},
                {"experience_processing": None},
            ]},
        ]
    }
    remaining_total = col.count_documents(pending_query)

    staff = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff experience already analysed — nothing to do.",
            "remaining_count": 0,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or '')
    user_type = _v(staff.get('user_type') or '')

    extracted_cv = _v(staff.get('extracted_cv') or '')
    has_cv_text  = (
        extracted_cv and
        not extracted_cv.startswith('[') and
        extracted_cv != 'No doc found'
    )
    s5           = staff.get('section_5_employment_history') or {}
    total_exp_db = _v(s5.get('total_experience') or '')

    if not has_cv_text and not total_exp_db:
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {
                "experience_years":        0,
                "experience_months":       0,
                "experience_total_months": 0,
                "experience_note":         "no CV text or experience data available",
                "experience_source":       "none",
                "experience_analysed_at":  datetime.utcnow(),
                "experience_processing":   False,
            }}
        )
        return jsonify({
            "success":         True,
            "staff_name":      full_name,
            "email":           email,
            "message":         f"Skipped {full_name} — no experience data",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Mark as processing then fire Gemini in background ────────────
    # This prevents 504 gateway timeouts — returns 200 immediately
    col.update_one(
        {"_id": staff['_id']},
        {"$set": {"experience_processing": True}}
    )

    def _run_gemini(staff_id_str, user_type_, has_cv_, extracted_cv_,
                    total_exp_db_, gemini_key_):
        from google import genai as google_genai

        _col = _staffs_col()

        ut_lower = (user_type_ or '').lower()
        if 'nurse' in ut_lower or 'nursing' in ut_lower:
            role_rule = (
                "IMPORTANT — Count ONLY nursing-related work experience, regardless of the country where it was gained. "
                "This includes: Registered Nurse, Staff Nurse, Clinical Nurse, ICU Nurse, Theatre Nurse, "
                "Community Nurse, Mental Health Nurse, Nursing Home Nurse, or any role with Nurse or Nursing in the title. "
                "DO NOT count non-nursing roles such as healthcare assistant, carer, support worker, admin, or any other role."
            )
        elif 'hca' in ut_lower or 'healthcare assistant' in ut_lower or 'health care assistant' in ut_lower:
            role_rule = (
                "IMPORTANT — Count ONLY Healthcare Assistant (HCA) work experience, regardless of the country where it was gained. "
                "This includes: Healthcare Assistant, HCA, Care Assistant, Care Worker, Support Worker in a clinical/care setting, "
                "or any role with Healthcare Assistant or HCA in the title. "
                "DO NOT count nursing roles (Registered Nurse, Staff Nurse, etc.) or non-care roles such as admin, retail, or hospitality."
            )
        else:
            role_rule = (
                "Count only direct healthcare or care-related work experience, regardless of country. "
                "Exclude non-healthcare roles such as admin, retail, hospitality, or general support roles "
                "unless they are clearly in a clinical or care setting."
            )

        try:
            if has_cv_:
                source = 'extracted_cv'
                prompt = f"""You are a professional CV analyser specialising in Irish healthcare staffing.

Candidate role: {user_type_}

Read the CV text below and calculate the candidate's TOTAL relevant work experience.

{role_rule}

Calculation Rules:
- Only count roles that match the role filter above.
- Experience gained in ANY country counts — not just Ireland.
- If a matching role has no end date, assume it is still ongoing (use today's date to calculate).
- If two matching roles overlap in time, count the overlapping period only once.
- Ignore all non-matching roles entirely — do not add them.
- Return ONLY a JSON object with these exact keys — nothing else, no markdown, no explanation:
  {{"years": <integer>, "months": <integer 0-11>, "total_months": <integer>, "note": "<one sentence summary of which roles were counted>"}}

CV TEXT:
{extracted_cv_[:10000]}
"""
            else:
                source = 'section_5'
                prompt = f"""You are a professional CV analyser specialising in Irish healthcare staffing.

Candidate role: {user_type_}

The candidate's total experience is described as: "{total_exp_db_}"

{role_rule}

Extract the relevant years and months from this description, applying the role filter above.
Return ONLY a JSON object — nothing else, no markdown:
{{"years": <integer>, "months": <integer 0-11>, "total_months": <integer>, "note": "<one sentence summary>"}}
"""

            client   = google_genai.Client(api_key=gemini_key_)
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=prompt
            )
            raw = (response.text or '').strip()
            raw = _re.sub(r'^```(?:json)?\s*', '', raw, flags=_re.MULTILINE)
            raw = _re.sub(r'```\s*$', '', raw, flags=_re.MULTILINE).strip()

            result       = _json.loads(raw)
            years        = int(result.get('years', 0) or 0)
            months       = int(result.get('months', 0) or 0)
            total_months = int(result.get('total_months', 0) or 0)
            note         = str(result.get('note', '') or '').strip()
            if total_months == 0:
                total_months = years * 12 + months

            _col.update_one(
                {"_id": ObjectId(staff_id_str)},
                {"$set": {
                    "experience_years":        years,
                    "experience_months":       months,
                    "experience_total_months": total_months,
                    "experience_note":         note,
                    "experience_source":       source,
                    "experience_analysed_at":  datetime.utcnow(),
                    "experience_processing":   False,
                }}
            )

        except Exception as e:
            _col.update_one(
                {"_id": ObjectId(staff_id_str)},
                {"$set": {
                    "experience_analysed_at":  datetime.utcnow(),
                    "experience_processing":   False,
                    "experience_note":         f"error: {e}",
                }}
            )

    threading.Thread(
        target=_run_gemini,
        args=(staff_id, user_type, has_cv_text, extracted_cv,
              total_exp_db, gemini_key),
        daemon=True,
    ).start()

    return jsonify({
        "success":         True,
        "staff_name":      full_name,
        "email":           email,
        "status":          "processing",
        "remaining_count": max(0, remaining_total - 1),
        "message": (
            f"Analysis started for {full_name} in background — "
            f"{max(0, remaining_total - 1)} staff remaining. "
            "Result will be saved to DB when complete."
        ),
    })



# ── Cron: Extract formatted experience entries one staff at a time ────

@admin_bp.route('/live-staffs/cron/extract-experience-list', methods=['GET', 'POST'])
def live_staff_cron_extract_experience_list():
    """
    Cron job — extracts a formatted list of relevant experience entries
    for ONE staff member per call using Gemini AI.

    For each matching role extracts:
      "Job Title  Date Range  Employer/Location"
    Only includes roles relevant to live_staffs.user_type
    (nurse roles for nurses, HCA roles for healthcare assistants).

    Saves to live_staffs.experience_list (array of strings).
    Background thread to avoid 504 timeouts.

    Protect with ?cron_key=<CRON_SECRET>
    """
    import requests as _req
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$and": [
            {"$or": [
                {"experience_list_at": {"$exists": False}},
                {"experience_list_at": None},
            ]},
            {"$or": [
                {"experience_list_processing": {"$exists": False}},
                {"experience_list_processing": False},
                {"experience_list_processing": None},
            ]},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff experience lists extracted — nothing to do.",
            "remaining_count": 0,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or '')
    user_type = _v(staff.get('user_type') or '')

    extracted_cv = _v(staff.get('extracted_cv') or '')
    has_cv_text  = (
        extracted_cv and
        not extracted_cv.startswith('[') and
        extracted_cv != 'No doc found'
    )

    if not has_cv_text:
        col.update_one(
            {"_id": staff['_id']},
            {"$set": {
                "experience_list":            [],
                "experience_list_at":         datetime.utcnow(),
                "experience_list_processing": False,
            }}
        )
        return jsonify({
            "success":         True,
            "staff_name":      full_name,
            "email":           email,
            "message":         f"Skipped {full_name} — no extracted CV text",
            "remaining_count": max(0, remaining_total - 1),
        })

    # Mark as processing — return immediately
    col.update_one(
        {"_id": staff['_id']},
        {"$set": {"experience_list_processing": True}}
    )

    def _run(staff_id_str, user_type_, extracted_cv_, gemini_key_):
        from google import genai as google_genai

        _col = _staffs_col()
        ut_lower = (user_type_ or '').lower()

        if 'nurse' in ut_lower or 'nursing' in ut_lower:
            role_filter = (
                "Extract ONLY nursing roles: Registered Nurse, Staff Nurse, Clinical Nurse, "
                "ICU Nurse, Theatre Nurse, Community Nurse, Mental Health Nurse, Nursing Home Nurse, "
                "or any role with Nurse or Nursing in the title. "
                "IGNORE Healthcare Assistant, HCA, Carer, Support Worker, admin, or any non-nursing role."
            )
        elif 'hca' in ut_lower or 'healthcare assistant' in ut_lower:
            role_filter = (
                "Extract ONLY Healthcare Assistant roles: Healthcare Assistant, HCA, "
                "Care Assistant, Care Worker, Support Worker in a clinical/care setting, "
                "Domiciliary Care Assistant, or any role with Healthcare Assistant or HCA in the title. "
                "IGNORE nursing roles (Registered Nurse, Staff Nurse etc.) and non-care roles."
            )
        else:
            role_filter = (
                "Extract only healthcare or care-related work experience roles. "
                "Ignore admin, retail, hospitality, or non-clinical roles."
            )

        prompt = f"""You are a professional CV analyser specialising in Irish healthcare staffing.

Candidate role: {user_type_}

Read the CV text below and extract ALL relevant work experience entries.

{role_filter}

For each matching role, return it as a single formatted string exactly like this example:
"Healthcare Assistant  Jan 2022 – Present  HSE Sligo Leitrim, Disabilities Services"

Format rules:
- Job title first
- Then date range (use the dates as written in the CV, e.g. "Jan 2022 – Present" or "2019 – 2021")
- Then employer and location/department if available
- Separate each part with two spaces
- If no end date, write "Present"
- List ALL matching roles, most recent first

Return ONLY a JSON array of strings — nothing else, no markdown, no explanation:
["role 1 formatted string", "role 2 formatted string", ...]

If no matching roles found, return an empty array: []

CV TEXT:
{extracted_cv_[:10000]}
"""

        try:
            client   = google_genai.Client(api_key=gemini_key_)
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=prompt
            )
            raw = (response.text or '').strip()
            raw = _re.sub(r'^```(?:json)?\s*', '', raw, flags=_re.MULTILINE)
            raw = _re.sub(r'```\s*$', '', raw, flags=_re.MULTILINE).strip()

            result = _cjson.loads(raw)
            if not isinstance(result, list):
                result = []
            # Clean each entry
            result = [str(r).strip() for r in result if str(r).strip()]

            _col.update_one(
                {"_id": ObjectId(staff_id_str)},
                {"$set": {
                    "experience_list":            result,
                    "experience_list_at":         datetime.utcnow(),
                    "experience_list_processing": False,
                }}
            )
        except Exception as e:
            _col.update_one(
                {"_id": ObjectId(staff_id_str)},
                {"$set": {
                    "experience_list":            [],
                    "experience_list_at":         datetime.utcnow(),
                    "experience_list_processing": False,
                    "experience_list_error":      str(e),
                }}
            )

    threading.Thread(
        target=_run,
        args=(staff_id, user_type, extracted_cv, gemini_key),
        daemon=True,
    ).start()

    return jsonify({
        "success":         True,
        "staff_name":      full_name,
        "email":           email,
        "status":          "processing",
        "remaining_count": max(0, remaining_total - 1),
        "message": (
            f"Experience list extraction started for {full_name} in background — "
            f"{max(0, remaining_total - 1)} staff remaining."
        ),
    })


# ── Export experience list to Excel ──────────────────────────────────

@admin_bp.route('/live-staffs/export/experience-xlsx')
@admin_required
def live_staff_export_experience_xlsx():
    """
    Export staff experience list to Excel.
    Columns: Sno | Name | Email | Experience 1 | Experience 2 | Experience 3 | ...

    The number of experience columns expands automatically based on
    the maximum number of entries any staff member has.

    GET /admin/live-staffs/export/experience-xlsx
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "user_type": 1, "experience_list": 1}
        ))

        # Sort by name
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        # Find max experience entries across all staff
        max_exp = max(
            (len(d.get('experience_list') or []) for d in docs),
            default=1
        )
        max_exp = max(max_exp, 3)  # minimum 3 columns

        # ── Styles ────────────────────────────────────────────────────
        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'; ALT = 'EFF6FF'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        b_font  = Font(name='Arial', size=9)
        l_align = Alignment(horizontal='left', vertical='top', wrap_text=True)
        c_align = Alignment(horizontal='center', vertical='top')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)

        wb = Workbook()
        ws = wb.active
        ws.title = 'Experience List'

        # ── Headers ───────────────────────────────────────────────────
        base_headers = ['Sno', 'Name', 'Email', 'User Type']
        exp_headers  = [f'Experience {i+1}' for i in range(max_exp)]
        headers      = base_headers + exp_headers
        col_widths   = [5, 30, 38, 20] + [45] * max_exp

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font      = h_font
            cell.fill      = h_fill
            cell.alignment = h_align
            cell.border    = Border(
                left=thin, right=thin, top=thin,
                bottom=Side(style='medium', color=GREEN)
            )
            ws.column_dimensions[cell.column_letter].width = width

        ws.row_dimensions[1].height = 30
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:{ws.cell(row=1, column=len(headers)).column_letter}1'

        # ── Data rows ─────────────────────────────────────────────────
        for ri, doc in enumerate(docs, start=2):
            s1        = doc.get('section_1_personal_details') or {}
            name      = _v(s1.get('full_name') or '')
            email     = _v(doc.get('email') or '')
            user_type = _v(doc.get('user_type') or '')
            exp_list  = doc.get('experience_list') or []

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT) if ri % 2 == 0 else None

            row_data = [ri - 1, name, email, user_type] + exp_list

            for ci, val in enumerate(row_data, start=1):
                cell = ws.cell(row=ri, column=ci, value=val or '')
                cell.font      = b_font
                cell.alignment = c_align if ci in (1,) else l_align
                cell.border    = border
                if alt_fill:
                    cell.fill = alt_fill

            # Fill remaining experience columns with empty
            for ci in range(len(row_data) + 1, len(headers) + 1):
                cell = ws.cell(row=ri, column=ci, value='')
                cell.font      = b_font
                cell.border    = border
                if alt_fill:
                    cell.fill = alt_fill

            ws.row_dimensions[ri].height = max(18, 15 * max(1, len(exp_list)))

        # Summary row
        ws.cell(row=len(docs) + 2, column=1,
                value=f'Total: {len(docs)} staff').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        date_str = datetime.utcnow().strftime('%Y%m%d')

        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="experience_list_{date_str}.xlsx"'}
        )

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Sync Passport/ID card one staff at a time ──────────────────

@admin_bp.route('/live-staffs/cron/sync-passport', methods=['GET', 'POST'])
def live_staff_cron_sync_passport():
    """
    Cron job — processes ONE staff member per call.

    Logic:
      1. Find first live_staffs record where passport_extracted is missing/empty.
      2. Call XN Portal user-document-list API with staff email.
      3. Find document where document_type_name == "Passport/id card".
      4. Download the document and extract passport data via Gemini AI.
      5. Save extracted data + passport_id to live_staffs.
      6. Return remaining_count.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req

    # ── Auth ──────────────────────────────────────────────────────────
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False,
                        "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False,
                        "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    # ── Find next staff to process ────────────────────────────────────
    # Picks up:
    #   1. Never fetched (passport_fetched missing/False)
    #   2. Fetched but passport_id is still empty — retry with passport_id_card doc type
    pending_query = {
        "$or": [
            {"passport_fetched": {"$exists": False}},
            {"passport_fetched": False},
            {"passport_fetched": None},
        ]
    }

    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff passports already fetched — nothing to do.",
            "remaining_count": 0,
        })

    email = _v(
        (staff.get('section_1_personal_details') or {}).get('email_address') or
        staff.get('email') or ''
    )
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')

    def _mark_done(fields):
        fields["passport_fetched"]    = True
        fields["passport_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {
                "passport_extracted":  "[skipped — no email]",
                "passport_extracted_at": datetime.utcnow(),
                "passport_fetched":    True,
            }}
        )
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Call XN Portal API ────────────────────────────────────────────
    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as api_err:
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"API error: {api_err}",
            "remaining_count": remaining_total,
        })

    if not data.get('success'):
        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {
                "passport_extracted":    f"[API error: {data.get('message')}]",
                "passport_extracted_at": datetime.utcnow(),
                "passport_fetched":      True,
            }}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    if isinstance(api_data, list):
        documents = api_data
    elif isinstance(api_data, dict):
        documents = api_data.get('documents') or []
    else:
        documents = []

    if not documents:
        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {
                "passport_extracted":    "No doc found",
                "passport_extracted_at": datetime.utcnow(),
                "passport_fetched":      True,
                "passport_id_retry":     True,  # prevent infinite re-run
            }}
        )
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "passport_found":  False,
            "message":         f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Find Passport document — try all known type names ───────────────
    # Includes: "Passport/id card", "passport_id_card", "Passport", "Id Card"
    is_retry = (staff.get('passport_fetched') is True and
                not _v(staff.get('passport_id') or ''))

    passport_doc = None
    PASSPORT_TYPES = {
        'passport/id card', 'passport/id', 'passport', 'id card',
        'passport_id_card', 'passport id card', 'passport id',
    }
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if doc_name in PASSPORT_TYPES and d.get('url'):
            passport_doc = d
            break

    if not passport_doc:
        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {
                "passport_extracted":    "No passport doc found",
                "passport_extracted_at": datetime.utcnow(),
                "passport_fetched":      True,
                "passport_id_retry":     True,  # mark so we don't retry infinitely
            }}
        )
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "passport_found":  False,
            "message":         f"No Passport/id card document found for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    passport_url = (passport_doc.get('url') or '').strip()

    if not passport_url:
        _mark_done({"passport_extracted": "[skipped — document URL is empty]"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "passport_found":  True,
            "skipped":         True,
            "reason":          "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message":         f"Skipped {full_name} ({email}) — passport doc has no URL",
        })

    # ── Download + extract via Gemini ─────────────────────────────────
    try:
        dl_headers = {k: v for k, v in api_headers.items()
                      if k != 'Content-Type'}
        img_resp = _req.get(passport_url, headers=dl_headers, timeout=60)

        if img_resp.status_code == 404:
            _mark_done({"passport_extracted": "[404 — document URL not found]"})
            return jsonify({
                "success":         True,
                "email":           email,
                "staff_name":      full_name,
                "passport_found":  True,
                "skipped":         True,
                "reason":          "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message":         f"Skipped {full_name} ({email}) — passport URL 404",
            })

        img_resp.raise_for_status()
        raw_bytes    = img_resp.content
        content_type = img_resp.headers.get('Content-Type', '').lower()

        from google import genai as google_genai

        client = google_genai.Client(api_key=gemini_key)

        # ── Build Gemini request with image or text ───────────────────
        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or passport_url.lower().endswith('.pdf')

        if is_image:
            ext = 'jpeg' if 'jpeg' in content_type or 'jpg' in content_type else                   'png'  if 'png'  in content_type else                   'webp' if 'webp' in content_type else 'jpeg'
            b64_data = base64.b64encode(raw_bytes).decode('utf-8')
            parts = [
                {
                    "inline_data": {
                        "mime_type": f"image/{ext}",
                        "data": b64_data,
                    }
                },
                {"text": """You are a passport/ID card data extractor.

Extract all readable information from this passport or ID card image.

Return ONLY a JSON object — no markdown, no explanation:
{
  "passport_id": "<passport or document number>",
  "full_name": "<name as printed on document>",
  "nationality": "<nationality>",
  "date_of_birth": "<DOB as printed>",
  "expiry_date": "<expiry date>",
  "issue_date": "<issue date if visible>",
  "country": "<issuing country>",
  "gender": "<gender if visible>",
  "mrz": "<machine readable zone lines if visible>",
  "raw_text": "<all text extracted from the document>"
}

If a field is not visible or readable, set it to null.
"""}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": parts}]
            )
        elif is_pdf:
            b64_data = base64.b64encode(raw_bytes).decode('utf-8')
            parts = [
                {
                    "inline_data": {
                        "mime_type": "application/pdf",
                        "data": b64_data,
                    }
                },
                {"text": """You are a passport/ID card data extractor.

Extract all readable information from this passport or ID card document.

Return ONLY a JSON object — no markdown, no explanation:
{
  "passport_id": "<passport or document number>",
  "full_name": "<name as printed on document>",
  "nationality": "<nationality>",
  "date_of_birth": "<DOB as printed>",
  "expiry_date": "<expiry date>",
  "issue_date": "<issue date if visible>",
  "country": "<issuing country>",
  "gender": "<gender if visible>",
  "mrz": "<machine readable zone if visible>",
  "raw_text": "<all text extracted>"
}

If a field is not visible, set it to null.
"""}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": parts}]
            )
        else:
            # Try text extraction fallback
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()

            prompt = f"""You are a passport/ID card data extractor.

Extract all information from this passport/ID card text.

Return ONLY a JSON object — no markdown, no explanation:
{{
  "passport_id": "<passport or document number>",
  "full_name": "<name>",
  "nationality": "<nationality>",
  "date_of_birth": "<DOB>",
  "expiry_date": "<expiry date>",
  "issue_date": "<issue date>",
  "country": "<issuing country>",
  "gender": "<gender>",
  "mrz": "<MRZ lines if present>",
  "raw_text": "<all extracted text>"
}}

TEXT:
{raw_text[:5000]}
"""
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        passport_data = _cjson.loads(raw_out)
        passport_id   = _v(passport_data.get('passport_id') or '')

        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {
                "passport_extracted":    passport_data.get('raw_text', '') or '',
                "passport_data":         passport_data,
                "passport_id":           passport_id,
                "passport_url":          passport_url,
                "passport_extracted_at": datetime.utcnow(),
                "passport_fetched":      True,
                "passport_id_retry":     True,  # prevent infinite re-runs
            }}
        )

        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "passport_found":  True,
            "passport_id":     passport_id or None,
            "passport_readable": bool(passport_id),
            "passport_data":   passport_data,
            "remaining_count": max(0, remaining_total - 1),
            "message": (
                f"Passport {'extracted' if passport_id else 'fetched (ID not readable)'} for {full_name} "
                + (f"(ID: {passport_id}) — " if passport_id else "— ")
                + f"{max(0, remaining_total - 1)} remaining."
            ),
            "synced_at": datetime.utcnow().isoformat(),
        })

    except _cjson.JSONDecodeError:
        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {
                "passport_extracted":    "[Gemini JSON parse error]",
                "passport_extracted_at": datetime.utcnow(),
                "passport_fetched":      True,
            }}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           "Gemini returned non-JSON response",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        col.update_one(
            {"_id": staff["_id"]},
            {"$set": {
                "passport_extracted":    f"[error: {e}]",
                "passport_extracted_at": datetime.utcnow(),
                "passport_fetched":      True,
            }}
        )
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           str(e),
            "remaining_count": max(0, remaining_total - 1),
        })



# ── Cron: Extract NMBI/QQI number from qualification docs ─────────────

@admin_bp.route('/live-staffs/cron/sync-qualification', methods=['GET', 'POST'])
def live_staff_cron_sync_qualification():
    """
    Cron job — processes ONE staff member per call.

    Logic:
      - Nurse:             looks for "Nmbi Qualification" document
      - Healthcare Asst:   looks for "QQI Level 5 or equivalent..." document
      Downloads doc, sends to Gemini to extract registration/ID number.
      Saves to live_staffs.nmbi_number or live_staffs.qqi_number.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    # ── Find next unprocessed staff ────────────────────────────────────
    # Only process Nurses and Healthcare Assistants
    # Skip if already fetched (qualification_fetched = True)
    pending_query = {
        "$and": [
            {"$or": [
                {"user_type": {"$regex": "nurse", "$options": "i"}},
                {"user_type": {"$regex": "healthcare assistant", "$options": "i"}},
                {"user_type": {"$regex": "hca", "$options": "i"}},
            ]},
            # Exclude staff where document URL returned 404 — cannot retry
            {"qualification_doc_404": {"$ne": True}},
            # Not yet fetched OR HCA with no valid alphanumeric qqi_number
            {"$or": [
                {"qualification_fetched": {"$exists": False}},
                {"qualification_fetched": False},
                {"qualification_fetched": None},
                # Re-run HCA only if NOT 404 and qqi_number is missing/digits-only
                {"$and": [
                    {"$or": [
                        {"user_type": {"$regex": "healthcare assistant", "$options": "i"}},
                        {"user_type": {"$regex": "hca", "$options": "i"}},
                    ]},
                    {"qualification_doc_404": {"$ne": True}},
                    {"$or": [
                        {"qqi_number": {"$exists": False}},
                        {"qqi_number": None},
                        {"qqi_number": ""},
                        {"qqi_number": {"$not": {"$regex": "[A-Za-z]"}}},
                    ]},
                ]},
            ]},
        ]
    }

    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff qualifications already extracted — nothing to do.",
            "remaining_count": 0,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or '')
    user_type = _v(staff.get('user_type') or '')
    ut_lower  = user_type.lower()

    is_nurse = 'nurse' in ut_lower or 'nursing' in ut_lower
    is_hca   = 'hca' in ut_lower or 'healthcare assistant' in ut_lower or 'health care assistant' in ut_lower

    # Determine target document type
    if is_nurse:
        target_doc_types = ['nmbi qualification', 'nmbi', 'nursing qualification']
        save_field       = 'nmbi_number'
        extract_hint     = (
            "This is a nursing registration document. "
            "Find the NMBI registration PIN number. "
            "It is typically 6 digits, or a combination of letters and numbers (e.g. 123456, NMC12345). "
            "Look for labels like: PIN, Registration Number, NMBI Number, Reg No, Certificate No."
        )
    else:
        target_doc_types = [
            'qqi level 5 or equivalent in health service skills or healthcare support',
            'qqi level 5', 'qqi', 'healthcare support',
            'health service skills', 'qqi certificate',
        ]
        save_field   = 'qqi_number'
        extract_hint = (
            "This is a QQI / FETAC qualification certificate for a Healthcare Assistant. "
            "Your task is to carefully read and extract the certificate or registration number. "
            "\n\n"
            "READING INSTRUCTIONS — follow these steps carefully:\n"
            "1. Examine the ENTIRE document thoroughly, including headers, footers, watermarks, "
            "   stamps, and small print.\n"
            "2. If text appears small, blurry, or partially obscured — zoom in mentally and try "
            "   your best to reconstruct each character. Use context clues from surrounding text.\n"
            "3. Look at EVERY number/code on the document — not just the most obvious one.\n"
            "4. If a number is partially visible (e.g. only some digits readable), attempt to "
            "   reconstruct the full number using the visible characters and document format.\n"
            "\n"
            "IMPORTANT — The correct number MUST BE ALPHANUMERIC (contains both letters AND numbers), "
            "for example: F12345678, L123456789, QF12345, 12345ABC, FET123456. "
            "If you find a digits-only number, keep looking — there is almost certainly an "
            "alphanumeric reference elsewhere on the document.\n"
            "\n"
            "Look for labels like: Certificate No, Learner ID, Award Reference, QQI No, Record No, "
            "Registration No, Reference Number, Award ID, Cert Ref, Roll No."
        )

    def _mark_done(update_dict):
        update_dict["qualification_fetched"]    = True
        update_dict["qualification_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": update_dict})

    if not email:
        _mark_done({"qualification_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         f"Skipped {staff_id} — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Call XN Portal API ────────────────────────────────────────────
    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as api_err:
        _mark_done({"qualification_note": f"API error: {api_err}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"API error: {api_err}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"qualification_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"qualification_note": "no documents returned"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "doc_found":       False,
            "message":         f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Find matching qualification document ──────────────────────────
    qual_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in target_doc_types) and d.get('url'):
            qual_doc = d
            break

    if not qual_doc:
        _mark_done({"qualification_note": f"no matching qualification doc found"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "doc_found":       False,
            "message":         f"No qualification document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (qual_doc.get('url') or '').strip()

    # ── Guard: only proceed if URL is present ─────────────────────────
    if not doc_url:
        _mark_done({"qualification_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "doc_found":       True,
            "skipped":         True,
            "reason":          "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message":         f"Skipped {full_name} ({email}) — document has no URL",
        })

    # ── Download document ─────────────────────────────────────────────
    from google import genai as google_genai

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        # 404 = document missing/expired URL — mark done and skip
        if dl_resp.status_code == 404:
            _mark_done({
                "qualification_note":   "document URL returned 404 — skipped",
                "qualification_doc_404": True,
            })
            return jsonify({
                "success":         True,
                "email":           email,
                "staff_name":      full_name,
                "doc_found":       True,
                "skipped":         True,
                "reason":          "Document URL returned 404 (missing or expired)",
                "remaining_count": max(0, remaining_total - 1),
                "message":         f"Skipped {full_name} ({email}) — document URL returned 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = f"""{extract_hint}

EXTRACTION RULES:
- Scan every part of the document: title, body, header, footer, stamps, seals, watermarks.
- If any text is small or unclear, zoom in and do your best to read each character.
- Even if only partially readable, attempt to reconstruct the number using visible characters.
- Prefer alphanumeric codes over digit-only numbers.
- Report your confidence level honestly.

Return ONLY a JSON object — no markdown, no explanation:
{{
  "registration_number": "<the alphanumeric number you found, exactly as printed — reconstruct if partially visible>",
  "label_found": "<the label printed next to the number, e.g. PIN, Certificate No, Learner ID>",
  "confidence": "<high|medium|low — how clearly was the number readable>",
  "reconstruction_note": "<if you had to reconstruct any characters, explain here — otherwise null>",
  "raw_text": "<all readable text extracted from the document>"
}}

If after thorough examination no registration/certificate number is found, set "registration_number" to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext  = 'jpeg' if any(t in content_type for t in ('jpeg','jpg')) else                    'png'  if 'png' in content_type else 'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result    = _cjson.loads(raw_out)
        reg_num   = _v(result.get('registration_number') or '')
        raw_text_ = _v(result.get('raw_text') or '')

        # ── For HCA: validate qqi_number is alphanumeric (contains letters) ──
        # If Gemini returned digits-only, reset and let cron retry next run
        import re as _re2
        if is_hca and reg_num:
            has_letter = bool(_re2.search(r'[A-Za-z]', reg_num))
            if not has_letter:
                # Digits-only — invalid QQI number, reset so cron retries
                col.update_one(
                    {"_id": staff['_id']},
                    {"$set": {
                        "qqi_number":            "",
                        "qualification_fetched": False,
                        "qualification_note":    f"digits-only QQI rejected: {reg_num} — will retry",
                        "qualification_doc_url": doc_url,
                    }}
                )
                return jsonify({
                    "success":         False,
                    "email":           email,
                    "staff_name":      full_name,
                    "user_type":       user_type,
                    "doc_found":       True,
                    "qqi_number":      "",
                    "rejected":        reg_num,
                    "reason":          "QQI number must be alphanumeric (contain letters). Digits-only rejected.",
                    "remaining_count": remaining_total,
                    "message":         f"Rejected digits-only QQI '{reg_num}' for {full_name} — will retry.",
                })
            reg_num = reg_num.strip()

        recon_note = _v(result.get('reconstruction_note') or '')
        _mark_done({
            save_field:               reg_num,
            "qualification_doc_type": qual_doc.get('document_type_name', ''),
            "qualification_doc_url":  doc_url,
            "qualification_raw":      raw_text_,
            "qualification_data":     result,
            "qualification_confidence": _v(result.get('confidence') or ''),
            "qualification_reconstruction": recon_note,
            "qualification_note":     f"extracted from {qual_doc.get('document_type_name','')}",
        })

        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "user_type":       user_type,
            "doc_found":       True,
            save_field:        reg_num,
            "doc_type":        qual_doc.get('document_type_name', ''),
            "readable":        bool(reg_num),
            "remaining_count": max(0, remaining_total - 1),
            "message": (
                f"{save_field} {'extracted: ' + reg_num if reg_num else 'not readable'} "
                f"for {full_name} — {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"qualification_note": "Gemini JSON parse error"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"qualification_note": f"error: {e}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Nurse NMBI list ───────────────────────────────────────────

@admin_bp.route('/live-staffs/export/nmbi-xlsx')
@admin_required
def live_staff_export_nmbi_xlsx():
    """Export Nurse staff with NMBI numbers to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {"$or": [
                {"user_type": {"$regex": "nurse", "$options": "i"}},
                {"user_type": {"$regex": "nursing", "$options": "i"}},
            ]},
            {"section_1_personal_details": 1, "email": 1,
             "user_type": 1, "nmbi_number": 1, "qualification_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        wb, ws = _build_qual_xlsx(docs, 'Nurses — NMBI', 'nmbi_number', 'NMBI Number')

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="nmbi_numbers_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Export: HCA QQI list ──────────────────────────────────────────────

@admin_bp.route('/live-staffs/export/qqi-xlsx')
@admin_required
def live_staff_export_qqi_xlsx():
    """Export Healthcare Assistant staff with QQI numbers to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {"$or": [
                {"user_type": {"$regex": "healthcare assistant", "$options": "i"}},
                {"user_type": {"$regex": "\bhca\b", "$options": "i"}},
            ]},
            {"section_1_personal_details": 1, "email": 1,
             "user_type": 1, "qqi_number": 1, "qualification_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        wb, ws = _build_qual_xlsx(docs, 'HCA — QQI Level 5', 'qqi_number', 'QQI Number')

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="qqi_numbers_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


def _build_qual_xlsx(docs, sheet_title, number_field, number_label):
    """Shared builder for NMBI/QQI Excel exports."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
    ALT  = 'EFF6FF'; WARN  = 'FFFBEB'

    h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
    h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
    h_align = Alignment(horizontal='center', vertical='center')
    b_font  = Font(name='Arial', size=10)
    l_align = Alignment(horizontal='left', vertical='center')
    c_align = Alignment(horizontal='center', vertical='center')
    thin    = Side(style='thin', color='CCCCCC')
    border  = Border(left=thin, right=thin, top=thin, bottom=thin)
    green_b = Border(left=thin, right=thin, top=thin,
                     bottom=Side(style='medium', color=GREEN))

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_title[:31]

    headers    = ['Sno', 'Name', 'Email', number_label, 'Status']
    col_widths = [5, 32, 40, 25, 18]

    for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
        cell = ws.cell(row=1, column=ci, value=hdr)
        cell.font      = h_font
        cell.fill      = h_fill
        cell.alignment = h_align
        cell.border    = green_b
        ws.column_dimensions[cell.column_letter].width = width
    ws.row_dimensions[1].height = 24
    ws.freeze_panes = 'A2'

    for ri, doc in enumerate(docs, start=2):
        s1     = doc.get('section_1_personal_details') or {}
        name   = _v(s1.get('full_name') or '')
        email  = _v(doc.get('email') or '')
        number = _v(doc.get(number_field) or '')
        fetched= doc.get('qualification_fetched', False)

        if number:
            status = 'Found'
            row_fill = None
        elif fetched:
            status = 'Not Found'
            row_fill = PatternFill('solid', start_color='FFDDDD', end_color='FFDDDD')
        else:
            status = 'Not Checked'
            row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)

        alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT) if ri % 2 == 0 and not row_fill else None

        for ci, (val, align) in enumerate(
            [(ri-1, c_align), (name, l_align), (email, l_align), (number, c_align), (status, c_align)],
            start=1
        ):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.font      = b_font
            cell.alignment = align
            cell.border    = border
            cell.fill      = row_fill or alt_fill or PatternFill()

        ws.row_dimensions[ri].height = 17

    ws.cell(row=len(docs)+2, column=1,
            value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

    return wb, ws



# ── Cron: Extract CPR/BLS certificate details one staff at a time ─────

@admin_bp.route('/live-staffs/cron/sync-cpr-certificate', methods=['GET', 'POST'])
def live_staff_cron_sync_cpr_certificate():
    """
    Cron job — processes ONE staff member per call.

    Finds "Cpr/Bls" document from XN Portal, sends to Gemini AI to extract:
      - Certificate name
      - Staff name as printed on certificate
      - Expiry date

    Saves to live_staffs:
      - cpr_certificate_name
      - cpr_staff_name
      - cpr_expiry_date
      - cpr_fetched = True

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url   = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key    = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country= os.environ.get('XN_APP_COUNTRY', '')
    gemini_key = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"cpr_fetched": {"$exists": False}},
            {"cpr_fetched": False},
            {"cpr_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff CPR/BLS certificates already extracted.",
            "remaining_count": 0,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or
                   s1.get('email_address') or '')

    def _mark_done(fields):
        fields["cpr_fetched"]    = True
        fields["cpr_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"cpr_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         f"Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Call XN Portal API ────────────────────────────────────────────
    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"cpr_note": f"API error: {e}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"cpr_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"cpr_note": "no documents returned"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "doc_found":       False,
            "message":         f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Find CPR/BLS document ─────────────────────────────────────────
    cpr_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in ('cpr/bls', 'cpr', 'bls', 'basic life support',
                                        'cardiopulmonary')) and d.get('url'):
            cpr_doc = d
            break

    if not cpr_doc:
        _mark_done({"cpr_note": "no CPR/BLS document found"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "doc_found":       False,
            "message":         f"No CPR/BLS document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (cpr_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"cpr_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "doc_found":       True,
            "skipped":         True,
            "reason":          "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message":         f"Skipped {full_name} ({email}) — CPR doc has no URL",
        })

    # ── Download and extract with Gemini ──────────────────────────────
    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"cpr_note": "document URL returned 404 — skipped"})
            return jsonify({
                "success":         True,
                "email":           email,
                "staff_name":      full_name,
                "doc_found":       True,
                "skipped":         True,
                "reason":          "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message":         f"Skipped {full_name} ({email}) — CPR doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this CPR/BLS certificate:
1. Certificate name (e.g. "Basic Life Support", "CPR/AED Certificate", "BLS for Healthcare Providers")
2. Staff name as printed on the certificate
3. Expiry date (the date the certificate expires or renewal is due)

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue/completion date if visible>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg','jpg')) else                     'png'  if 'png' in content_type else 'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "cpr_certificate_name": cert_name,
            "cpr_staff_name":       cert_staff,
            "cpr_expiry_date":      expiry_date,
            "cpr_issue_date":       issue_date,
            "cpr_issuing_body":     issuing_body,
            "cpr_doc_url":          doc_url,
            "cpr_doc_type":         cpr_doc.get('document_type_name', ''),
            "cpr_note":             "extracted successfully",
        })

        return jsonify({
            "success":              True,
            "email":                email,
            "staff_name":           full_name,
            "doc_found":            True,
            "certificate_name":     cert_name,
            "staff_name_on_cert":   cert_staff,
            "expiry_date":          expiry_date,
            "issuing_body":         issuing_body,
            "remaining_count":      max(0, remaining_total - 1),
            "message": (
                f"CPR/BLS cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"cpr_note": "Gemini JSON parse error"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"cpr_note": f"error: {e}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: CPR/BLS certificates to Excel ────────────────────────────

@admin_bp.route('/live-staffs/export/cpr-xlsx')
@admin_required
def live_staff_export_cpr_xlsx():
    """Export CPR/BLS certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "cpr_certificate_name": 1, "cpr_staff_name": 1,
             "cpr_expiry_date": 1, "cpr_issue_date": 1,
             "cpr_issuing_body": 1, "cpr_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY  = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT   = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left', vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'CPR-BLS Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 30, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('cpr_certificate_name') or '')
            cert_s   = _v(doc.get('cpr_staff_name') or '')
            expiry   = _v(doc.get('cpr_expiry_date') or '')
            issue    = _v(doc.get('cpr_issue_date') or '')
            issuer   = _v(doc.get('cpr_issuing_body') or '')
            fetched  = doc.get('cpr_fetched', False)

            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align,l_align,l_align,l_align,l_align,c_align,c_align,l_align,c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font      = b_font
                cell.alignment = align
                cell.border    = border
                cell.fill      = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="cpr_certificates_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



@admin_bp.route('/live-staffs/export/missing-nmbi-xlsx')
@admin_required
def live_staff_export_missing_nmbi_xlsx():
    """Export Nurses with missing NMBI registration number."""
    try:
        import io as _io
        docs = list(_staffs_col().find(
            {"$and": [
                {"$or": [
                    {"user_type": {"$regex": "nurse", "$options": "i"}},
                    {"user_type": {"$regex": "nursing", "$options": "i"}},
                ]},
                {"$or": [
                    {"nmbi_number": {"$exists": False}},
                    {"nmbi_number": None},
                    {"nmbi_number": ""},
                ]},
            ]},
            {"section_1_personal_details": 1, "email": 1,
             "user_type": 1, "nmbi_number": 1, "qualification_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        wb, ws = _build_qual_xlsx(docs, 'Missing NMBI Numbers', 'nmbi_number', 'NMBI Number')
        buf = _io.BytesIO()
        wb.save(buf)
        return Response(buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="missing_nmbi_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/export/missing-qqi-xlsx')
@admin_required
def live_staff_export_missing_qqi_xlsx():
    """Export Healthcare Assistants with missing QQI number."""
    try:
        import io as _io
        docs = list(_staffs_col().find(
            {"$and": [
                {"$or": [
                    {"user_type": {"$regex": "healthcare assistant", "$options": "i"}},
                    {"user_type": {"$regex": "\bhca\b", "$options": "i"}},
                ]},
                {"$or": [
                    {"qqi_number": {"$exists": False}},
                    {"qqi_number": None},
                    {"qqi_number": ""},
                ]},
            ]},
            {"section_1_personal_details": 1, "email": 1,
             "user_type": 1, "qqi_number": 1, "qualification_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        wb, ws = _build_qual_xlsx(docs, 'Missing QQI Numbers', 'qqi_number', 'QQI Number')
        buf = _io.BytesIO()
        wb.save(buf)
        return Response(buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="missing_qqi_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



@admin_bp.route('/live-staffs/export/passport-xlsx')
@admin_required
def live_staff_export_passport_xlsx():
    """Export staff with saved passport IDs to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {"passport_fetched": True},
            {"section_1_personal_details": 1, "email": 1,
             "passport_id": 1, "passport_data": 1, "passport_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY  = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT   = 'EFF6FF'; RED   = 'FFDDDD'; WARN  = 'FFF3CD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left', vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Passport IDs'

        headers    = ['Sno', 'Name', 'Email', 'Passport ID',
                      'Nationality', 'DOB', 'Expiry Date', 'Status']
        col_widths = [5, 30, 36, 18, 18, 14, 14, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:H{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1          = doc.get('section_1_personal_details') or {}
            name        = _v(s1.get('full_name') or '')
            email       = _v(doc.get('email') or '')
            passport_id = _v(doc.get('passport_id') or '')
            pdata       = doc.get('passport_data') or {}
            nationality = _v(pdata.get('nationality') or '')
            dob         = _v(pdata.get('date_of_birth') or '')
            expiry      = _v(pdata.get('expiry_date') or '')

            if passport_id:
                status   = 'Found'
                row_fill = None
            else:
                status   = 'Not Readable'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, passport_id,
                        nationality, dob, expiry, status]
            aligns   = [c_align, l_align, l_align, c_align,
                        l_align, c_align, c_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="passport_ids_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



@admin_bp.route('/live-staffs/cron/reset-hca-qualification', methods=['GET', 'POST'])
@admin_required
def live_staff_reset_hca_qualification():
    """
    Reset qualification_fetched for ALL HCA staff so sync-qualification
    cron re-runs and re-extracts their QQI numbers.

    Also resets any HCA with a digits-only qqi_number (invalid).

    GET /admin/live-staffs/cron/reset-hca-qualification
    """
    import re as _re2
    col = _staffs_col()

    hca_query = {"$or": [
        {"user_type": {"$regex": "healthcare assistant", "$options": "i"}},
        {"user_type": {"$regex": r"\bhca\b", "$options": "i"}},
    ]}

    # Find HCA with digits-only or missing qqi_number
    reset_query = {"$and": [
        hca_query,
        {"$or": [
            {"qualification_fetched": True},
            {"qqi_number": {"$exists": False}},
            {"qqi_number": None},
            {"qqi_number": ""},
        ]},
    ]}

    all_hca  = list(col.find(hca_query, {"qqi_number": 1}))
    to_reset = []
    for doc in all_hca:
        qqi = _v(doc.get('qqi_number') or '')
        # Reset if empty OR digits-only (no letters)
        if not qqi or not _re2.search(r'[A-Za-z]', qqi):
            to_reset.append(doc['_id'])

    if not to_reset:
        return jsonify({
            "success": True,
            "reset_count": 0,
            "message": "No HCA staff need resetting — all have valid alphanumeric QQI numbers.",
        })

    result = col.update_many(
        {"_id": {"$in": to_reset}},
        {"$set": {
            "qualification_fetched": False,
            "qqi_number":            "",
            "qualification_note":    "reset for re-extraction",
        }}
    )

    return jsonify({
        "success":     True,
        "reset_count": result.modified_count,
        "total_hca":   len(all_hca),
        "message":     f"Reset {result.modified_count} HCA staff — cron will re-extract their QQI numbers.",
    })



# ── Cron: Extract Infection Prevention & Control Certificate ──────────

@admin_bp.route('/live-staffs/cron/sync-ipc-certificate', methods=['GET', 'POST'])
def live_staff_cron_sync_ipc_certificate():
    """
    Cron job — processes ONE staff member per call.

    Finds "Infection Prevention Control Certificate" document from XN Portal,
    sends to Gemini AI to extract:
      - Certificate name
      - Staff name as printed on certificate
      - Expiry date / completion date

    Saves to live_staffs:
      - ipc_certificate_name
      - ipc_staff_name
      - ipc_expiry_date
      - ipc_issue_date
      - ipc_issuing_body
      - ipc_fetched = True

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"ipc_fetched": {"$exists": False}},
            {"ipc_fetched": False},
            {"ipc_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff IPC certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["ipc_fetched"]    = True
        fields["ipc_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"ipc_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Call XN Portal API ────────────────────────────────────────────
    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"ipc_note": f"API error: {e}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"ipc_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"ipc_note": "no documents returned"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "doc_found":       False,
            "message":         f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Find IPC document ─────────────────────────────────────────────
    ipc_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'infection prevention control certificate',
            'infection prevention and control',
            'infection prevention & control',
            'infection control certificate',
            'infection prevention',
            'ipc certificate',
            'ipc',
        )) and d.get('url'):
            ipc_doc = d
            break

    if not ipc_doc:
        _mark_done({"ipc_note": "no IPC document found"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "doc_found":       False,
            "message":         f"No IPC certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (ipc_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"ipc_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "doc_found":       True,
            "skipped":         True,
            "reason":          "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message":         f"Skipped {full_name} ({email}) — IPC doc has no URL",
        })

    # ── Download and extract with Gemini ──────────────────────────────
    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"ipc_note": "document URL returned 404 — skipped",
                        "ipc_doc_404": True})
            return jsonify({
                "success":         True,
                "email":           email,
                "staff_name":      full_name,
                "doc_found":       True,
                "skipped":         True,
                "reason":          "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message":         f"Skipped {full_name} ({email}) — IPC doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Infection Prevention and Control (IPC) certificate:
1. Certificate name (e.g. "Infection Prevention and Control", "IPC Training Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "ipc_certificate_name": cert_name,
            "ipc_staff_name":       cert_staff,
            "ipc_expiry_date":      expiry_date,
            "ipc_issue_date":       issue_date,
            "ipc_issuing_body":     issuing_body,
            "ipc_doc_url":          doc_url,
            "ipc_doc_type":         ipc_doc.get('document_type_name', ''),
            "ipc_note":             "extracted successfully",
        })

        return jsonify({
            "success":              True,
            "email":                email,
            "staff_name":           full_name,
            "doc_found":            True,
            "certificate_name":     cert_name,
            "staff_name_on_cert":   cert_staff,
            "expiry_date":          expiry_date,
            "issue_date":           issue_date,
            "issuing_body":         issuing_body,
            "remaining_count":      max(0, remaining_total - 1),
            "message": (
                f"IPC cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"ipc_note": "Gemini JSON parse error"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"ipc_note": f"error: {e}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "error":           str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: IPC certificates to Excel ────────────────────────────────

@admin_bp.route('/live-staffs/export/ipc-xlsx')
@admin_required
def live_staff_export_ipc_xlsx():
    """Export IPC certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "ipc_certificate_name": 1, "ipc_staff_name": 1,
             "ipc_expiry_date": 1, "ipc_issue_date": 1,
             "ipc_issuing_body": 1, "ipc_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left', vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'IPC Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 35, 28, 16, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('ipc_certificate_name') or '')
            cert_s   = _v(doc.get('ipc_staff_name') or '')
            expiry   = _v(doc.get('ipc_expiry_date') or '')
            issue    = _v(doc.get('ipc_issue_date') or '')
            issuer   = _v(doc.get('ipc_issuing_body') or '')
            fetched  = doc.get('ipc_fetched', False)

            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align,l_align,l_align,l_align,l_align,c_align,c_align,l_align,c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="ipc_certificates_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Hand Hygiene Certificate ────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-hand-hygiene', methods=['GET', 'POST'])
def live_staff_cron_sync_hand_hygiene():
    """
    Cron job — processes ONE staff member per call.
    Finds "Hand Hygiene" document, extracts certificate details via Gemini.
    Saves: hh_certificate_name, hh_staff_name, hh_expiry_date,
           hh_issue_date, hh_issuing_body, hh_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"hh_fetched": {"$exists": False}},
            {"hh_fetched": False},
            {"hh_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Hand Hygiene certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["hh_fetched"]    = True
        fields["hh_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"hh_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"hh_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"hh_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"hh_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    hh_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'hand hygiene', 'hand washing', 'hand hygiene certificate',
            'handhygiene', 'hand-hygiene',
        )) and d.get('url'):
            hh_doc = d
            break

    if not hh_doc:
        _mark_done({"hh_note": "no Hand Hygiene document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Hand Hygiene certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (hh_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"hh_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Hand Hygiene doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"hh_note": "document URL 404 — skipped", "hh_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Hand Hygiene doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Hand Hygiene certificate:
1. Certificate name (e.g. "Hand Hygiene", "Hand Washing Training Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg','jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "hh_certificate_name": cert_name,
            "hh_staff_name":       cert_staff,
            "hh_expiry_date":      expiry_date,
            "hh_issue_date":       issue_date,
            "hh_issuing_body":     issuing_body,
            "hh_doc_url":          doc_url,
            "hh_doc_type":         hh_doc.get('document_type_name', ''),
            "hh_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Hand Hygiene cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"hh_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"hh_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Hand Hygiene certificates to Excel ────────────────────────

@admin_bp.route('/live-staffs/export/hand-hygiene-xlsx')
@admin_required
def live_staff_export_hand_hygiene_xlsx():
    """Export Hand Hygiene certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "hh_certificate_name": 1, "hh_staff_name": 1,
             "hh_expiry_date": 1, "hh_issue_date": 1,
             "hh_issuing_body": 1, "hh_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left', vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Hand Hygiene Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('hh_certificate_name') or '')
            cert_s   = _v(doc.get('hh_staff_name') or '')
            expiry   = _v(doc.get('hh_expiry_date') or '')
            issue    = _v(doc.get('hh_issue_date') or '')
            issuer   = _v(doc.get('hh_issuing_body') or '')
            fetched  = doc.get('hh_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align,l_align,l_align,l_align,l_align,c_align,c_align,l_align,c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="hand_hygiene_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Children First Certificate ──────────────────────────

@admin_bp.route('/live-staffs/cron/sync-children-first', methods=['GET', 'POST'])
def live_staff_cron_sync_children_first():
    """
    Cron job — processes ONE staff member per call.
    Finds "Children First" document, extracts certificate details via Gemini.
    Saves: cf_certificate_name, cf_staff_name, cf_expiry_date,
           cf_issue_date, cf_issuing_body, cf_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"cf_fetched": {"$exists": False}},
            {"cf_fetched": False},
            {"cf_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Children First certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["cf_fetched"]    = True
        fields["cf_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"cf_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"cf_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"cf_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"cf_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    cf_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'children first', 'childrenfirst', 'children first certificate',
            'children first training', 'children first e-learning',
        )) and d.get('url'):
            cf_doc = d
            break

    if not cf_doc:
        _mark_done({"cf_note": "no Children First document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Children First certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (cf_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"cf_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Children First doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"cf_note": "document URL 404 — skipped", "cf_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Children First doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Children First certificate or training record:
1. Certificate name (e.g. "Children First", "Children First e-Learning", "Safeguarding Children")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg','jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "cf_certificate_name": cert_name,
            "cf_staff_name":       cert_staff,
            "cf_expiry_date":      expiry_date,
            "cf_issue_date":       issue_date,
            "cf_issuing_body":     issuing_body,
            "cf_doc_url":          doc_url,
            "cf_doc_type":         cf_doc.get('document_type_name', ''),
            "cf_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Children First cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"cf_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"cf_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Children First certificates to Excel ──────────────────────

@admin_bp.route('/live-staffs/export/children-first-xlsx')
@admin_required
def live_staff_export_children_first_xlsx():
    """Export Children First certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "cf_certificate_name": 1, "cf_staff_name": 1,
             "cf_expiry_date": 1, "cf_issue_date": 1,
             "cf_issuing_body": 1, "cf_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left', vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Children First Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('cf_certificate_name') or '')
            cert_s   = _v(doc.get('cf_staff_name') or '')
            expiry   = _v(doc.get('cf_expiry_date') or '')
            issue    = _v(doc.get('cf_issue_date') or '')
            issuer   = _v(doc.get('cf_issuing_body') or '')
            fetched  = doc.get('cf_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align,l_align,l_align,l_align,l_align,c_align,c_align,l_align,c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="children_first_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Push passport number to DOC API one staff at a time ─────────

@admin_bp.route('/live-staffs/cron/push-passport-number', methods=['GET', 'POST'])
def live_staff_cron_push_passport_number():
    """
    Cron job — processes ONE staff member per call.

    Finds staff where:
      - passport_id is set (already extracted by sync-passport cron)
      - staff_id (XN Portal staff ID) exists
      - passport_number_pushed is not True (not yet pushed to DOC API)

    Calls:
      POST {DOC_BASE_URL}/api/staff/passport-number-update
      Headers: Api-Key, X-App-Country
      Form:    staff_id, passport_number

    Sets passport_number_pushed = True on success.
    Stores passport_push_email in live_staffs for reference.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    import requests as _req

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    doc_base_url = os.environ.get('DOC_BASE_URL', '').rstrip('/')
    doc_api_key  = os.environ.get('PASSPORT_API_KEY', '') or os.environ.get('DOC_API_KEY', '')
    app_country  = os.environ.get('XN_APP_COUNTRY', 'ie')

    if not doc_base_url:
        return jsonify({"success": False, "error": "DOC_BASE_URL not set"}), 500
    if not doc_api_key:
        return jsonify({"success": False, "error": "PASSPORT_API_KEY not set in .env"}), 500

    col = _staffs_col()

    # Only process staff that:
    # 1. Have a passport_id extracted
    # 2. Have a staff_id (XN Portal)
    # 3. Have not been pushed yet
    pending_query = {
        "$and": [
            {"passport_id": {"$exists": True}},
            {"passport_id": {"$ne": None}},
            {"passport_id": {"$ne": ""}},
            {"$or": [
                {"staff_id": {"$exists": True}},
                {"xn_staff_id": {"$exists": True}},
            ]},
            {"$or": [
                {"passport_number_pushed": {"$exists": False}},
                {"passport_number_pushed": False},
                {"passport_number_pushed": None},
            ]},
        ]
    }

    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All passport numbers already pushed — nothing to do.",
            "remaining_count": 0,
        })

    s1          = staff.get('section_1_personal_details') or {}
    full_name   = _v(s1.get('full_name') or '')
    email       = _v(staff.get('email') or s1.get('email_address') or '')
    passport_id = _v(staff.get('passport_id') or '')
    xn_staff_id = _v(staff.get('staff_id') or staff.get('xn_staff_id') or '')

    def _mark_done(fields):
        fields["passport_push_attempted_at"] = datetime.utcnow()
        fields["passport_push_email"]        = email
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not xn_staff_id:
        _mark_done({
            "passport_number_pushed": False,
            "passport_push_note":     "skipped — no staff_id found",
        })
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "skipped":         True,
            "reason":          "No staff_id / xn_staff_id found",
            "remaining_count": max(0, remaining_total - 1),
            "message":         f"Skipped {full_name} ({email}) — no staff_id",
        })

    if not passport_id:
        _mark_done({
            "passport_number_pushed": False,
            "passport_push_note":     "skipped — passport_id is empty",
        })
        return jsonify({
            "success":         True,
            "email":           email,
            "staff_name":      full_name,
            "skipped":         True,
            "reason":          "passport_id is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message":         f"Skipped {full_name} ({email}) — passport_id empty",
        })

    # ── Call DOC API ──────────────────────────────────────────────────
    url     = f"{doc_base_url}/api/staff/passport-number-update"
    headers = {
        "Api-Key":       doc_api_key,
        "X-App-Country": app_country,
    }
    form_data = {
        "staff_id":        xn_staff_id,
        "passport_number": passport_id,
    }

    try:
        resp = _req.post(url, headers=headers, data=form_data, timeout=30)

        if resp.status_code in (200, 201):
            try:
                resp_data = resp.json()
            except Exception:
                resp_data = {"raw": resp.text[:200]}

            _mark_done({
                "passport_number_pushed":      True,
                "passport_push_note":          "pushed successfully",
                "passport_push_response":      resp_data,
            })

            return jsonify({
                "success":         True,
                "email":           email,
                "staff_name":      full_name,
                "staff_id":        xn_staff_id,
                "passport_number": passport_id,
                "pushed":          True,
                "api_response":    resp_data,
                "remaining_count": max(0, remaining_total - 1),
                "message": (
                    f"Passport number pushed for {full_name} ({email}) "
                    f"[{passport_id}] — {max(0, remaining_total - 1)} remaining."
                ),
            })

        else:
            # Non-200 — log full response
            try:
                err_json = resp.json()
            except Exception:
                err_json = None
            err_text = resp.text[:500]

            # 422 = invalid staff_id — skip permanently, do not retry
            if resp.status_code == 422:
                _mark_done({
                    "passport_number_pushed": True,   # mark done so cron skips next time
                    "passport_push_note":     f"422 invalid staff_id — skipped permanently",
                    "passport_push_response": err_json or err_text,
                })
                return jsonify({
                    "success":         True,
                    "email":           email,
                    "staff_name":      full_name,
                    "staff_id":        xn_staff_id,
                    "passport_number": passport_id,
                    "pushed":          False,
                    "skipped":         True,
                    "reason":          "422 — invalid staff_id, skipped permanently",
                    "http_status":     422,
                    "api_error":       err_text,
                    "remaining_count": max(0, remaining_total - 1),
                    "message":         f"Skipped {full_name} ({email}) — invalid staff_id (422)",
                })

            _mark_done({
                "passport_number_pushed":  False,
                "passport_push_note":      f"API returned {resp.status_code}: {err_text}",
                "passport_push_response":  err_json or err_text,
            })
            return jsonify({
                "success":         False,
                "email":           email,
                "staff_name":      full_name,
                "staff_id":        xn_staff_id,
                "passport_number": passport_id,
                "pushed":          False,
                "http_status":     resp.status_code,
                "api_error":       err_text,
                "api_error_json":  err_json,
                "remaining_count": max(0, remaining_total - 1),
                "message":         f"API error {resp.status_code} for {full_name} ({email}): {err_text}",
            })

    except Exception as e:
        _mark_done({
            "passport_number_pushed": False,
            "passport_push_note":     f"request error: {e}",
        })
        return jsonify({
            "success":         False,
            "email":           email,
            "staff_name":      full_name,
            "error":           str(e),
            "remaining_count": max(0, remaining_total - 1),
        })



@admin_bp.route('/live-staffs/export/safeguarding-xlsx')
@admin_required
def live_staff_export_safeguarding_xlsx():
    """Export Safeguarding Adults At Risk certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "sg_certificate_name": 1, "sg_staff_name": 1,
             "sg_expiry_date": 1, "sg_issue_date": 1,
             "sg_issuing_body": 1, "sg_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Safeguarding Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 35, 28, 16, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y', '%m/%Y', '%Y-%m-%d', '%d-%m-%Y', '%B %Y', '%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('sg_certificate_name') or '')
            cert_s   = _v(doc.get('sg_staff_name') or '')
            expiry   = _v(doc.get('sg_expiry_date') or '')
            issue    = _v(doc.get('sg_issue_date') or '')
            issuer   = _v(doc.get('sg_issuing_body') or '')
            fetched  = doc.get('sg_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT) \
                       if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="safeguarding_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/cron/sync-safeguarding', methods=['GET', 'POST'])
def live_staff_cron_sync_safeguarding():
    """
    Cron job — processes ONE staff member per call.
    Finds "Safeguarding Adults At Risk" document, extracts details via Gemini.
    Sets sg_fetched = True on every outcome (success, skip, error).
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"sg_fetched": {"$exists": False}},
            {"sg_fetched": False},
            {"sg_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Safeguarding certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["sg_fetched"]    = True
        fields["sg_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"sg_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Call XN Portal API ────────────────────────────────────────────
    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"sg_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"sg_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else \
                (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"sg_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    # ── Find Safeguarding document ────────────────────────────────────
    sg_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'safeguarding adults at risk',
            'safeguarding adults',
            'safeguarding at risk',
            'safeguarding certificate',
            'safeguarding training',
            'safeguarding',
        )) and d.get('url'):
            sg_doc = d
            break

    if not sg_doc:
        _mark_done({"sg_note": "no Safeguarding document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Safeguarding Adults At Risk certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (sg_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"sg_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Safeguarding doc has no URL",
        })

    # ── Download ──────────────────────────────────────────────────────
    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"sg_note": "document URL 404 — skipped", "sg_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Safeguarding doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Safeguarding Adults At Risk certificate:
1. Certificate name (e.g. "Safeguarding Adults At Risk", "Safeguarding Training Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else \
                    'png'  if 'png'  in content_type else \
                    'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "sg_certificate_name": cert_name,
            "sg_staff_name":       cert_staff,
            "sg_expiry_date":      expiry_date,
            "sg_issue_date":       issue_date,
            "sg_issuing_body":     issuing_body,
            "sg_doc_url":          doc_url,
            "sg_doc_type":         sg_doc.get('document_type_name', ''),
            "sg_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Safeguarding cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"sg_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"sg_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export ────────────────────────────────────────────────────────────


# ── Cron: Extract Garda Vetting Document details ─────────────────────

@admin_bp.route('/live-staffs/cron/sync-garda-vetting', methods=['GET', 'POST'])
def live_staff_cron_sync_garda_vetting():
    """
    Cron job — processes ONE staff member per call.
    Finds "Garda Vetting Document", extracts details via Gemini AI.
    Saves: garda_cert_name, garda_staff_name, garda_issue_date,
           garda_reference_number, garda_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"garda_fetched": {"$exists": False}},
            {"garda_fetched": False},
            {"garda_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Garda Vetting documents already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["garda_fetched"]    = True
        fields["garda_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"garda_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"garda_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"garda_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"garda_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    garda_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'garda vetting document',
            'garda vetting',
            'garda',
            'vetting document',
            'national vetting bureau',
            'nvb',
        )) and d.get('url'):
            garda_doc = d
            break

    if not garda_doc:
        _mark_done({"garda_note": "no Garda Vetting document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Garda Vetting document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (garda_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"garda_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Garda doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"garda_note": "document URL 404 — skipped",
                        "garda_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Garda doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor specialising in Irish Garda Vetting documents.

Extract the following details from this Garda Vetting disclosure document:
1. Document / certificate name (e.g. "Garda Vetting Disclosure", "National Vetting Bureau Disclosure")
2. Staff name as printed on the document
3. Issue date / date of disclosure
4. Reference number or vetting reference (if shown)
5. Issuing body (e.g. "National Vetting Bureau", "An Garda Siochana")

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact document title as printed>",
  "staff_name_on_cert": "<name as printed on document>",
  "issue_date": "<date of issue or disclosure as printed>",
  "reference_number": "<vetting reference number if visible>",
  "issuing_body": "<organization that issued the document>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result     = _cjson.loads(raw_out)
        cert_name  = _v(result.get('certificate_name') or '')
        cert_staff = _v(result.get('staff_name_on_cert') or '')
        issue_date = _v(result.get('issue_date') or '')
        ref_num    = _v(result.get('reference_number') or '')
        issuer     = _v(result.get('issuing_body') or '')

        _mark_done({
            "garda_cert_name":      cert_name,
            "garda_staff_name":     cert_staff,
            "garda_issue_date":     issue_date,
            "garda_reference":      ref_num,
            "garda_issuing_body":   issuer,
            "garda_doc_url":        doc_url,
            "garda_doc_type":       garda_doc.get('document_type_name', ''),
            "garda_note":           "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "issue_date":         issue_date,
            "reference_number":   ref_num,
            "issuing_body":       issuer,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Garda Vetting extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"garda_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"garda_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Garda Vetting to Excel ───────────────────────────────────

@admin_bp.route('/live-staffs/export/garda-xlsx')
@admin_required
def live_staff_export_garda_xlsx():
    """Export Garda Vetting document details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "garda_cert_name": 1, "garda_staff_name": 1,
             "garda_issue_date": 1, "garda_reference": 1,
             "garda_issuing_body": 1, "garda_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Garda Vetting'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Issue Date', 'Reference No', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 20, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('garda_cert_name') or '')
            cert_s   = _v(doc.get('garda_staff_name') or '')
            issue    = _v(doc.get('garda_issue_date') or '')
            ref_n    = _v(doc.get('garda_reference') or '')
            issuer   = _v(doc.get('garda_issuing_body') or '')
            fetched  = doc.get('garda_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Doc Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, issue, ref_n, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="garda_vetting_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract CPI/MAPA/PMAV Certificate ──────────────────────────

@admin_bp.route('/live-staffs/cron/sync-cpi-mapa', methods=['GET', 'POST'])
def live_staff_cron_sync_cpi_mapa():
    """
    Cron job — processes ONE staff member per call.
    Finds "Cpi/ Mapa/Pmav" document, extracts details via Gemini AI.
    Saves: cpi_certificate_name, cpi_staff_name, cpi_expiry_date,
           cpi_issue_date, cpi_issuing_body, cpi_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"cpi_fetched": {"$exists": False}},
            {"cpi_fetched": False},
            {"cpi_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff CPI/MAPA/PMAV certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["cpi_fetched"]    = True
        fields["cpi_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"cpi_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"cpi_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"cpi_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"cpi_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    cpi_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'cpi/ mapa/pmav', 'cpi/mapa/pmav', 'cpi mapa pmav',
            'cpi', 'mapa', 'pmav',
            'crisis prevention', 'management of actual',
            'prevention management',
        )) and d.get('url'):
            cpi_doc = d
            break

    if not cpi_doc:
        _mark_done({"cpi_note": "no CPI/MAPA/PMAV document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No CPI/MAPA/PMAV certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (cpi_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"cpi_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — CPI/MAPA/PMAV doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"cpi_note": "document URL 404 — skipped",
                        "cpi_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — CPI/MAPA/PMAV doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this CPI / MAPA / PMAV training certificate:
1. Certificate name (e.g. "CPI Non-Violent Crisis Intervention", "MAPA Foundation", "PMAV Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "cpi_certificate_name": cert_name,
            "cpi_staff_name":       cert_staff,
            "cpi_expiry_date":      expiry_date,
            "cpi_issue_date":       issue_date,
            "cpi_issuing_body":     issuing_body,
            "cpi_doc_url":          doc_url,
            "cpi_doc_type":         cpi_doc.get('document_type_name', ''),
            "cpi_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"CPI/MAPA/PMAV cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"cpi_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"cpi_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: CPI/MAPA/PMAV certificates to Excel ───────────────────────

@admin_bp.route('/live-staffs/export/cpi-xlsx')
@admin_required
def live_staff_export_cpi_xlsx():
    """Export CPI/MAPA/PMAV certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "cpi_certificate_name": 1, "cpi_staff_name": 1,
             "cpi_expiry_date": 1, "cpi_issue_date": 1,
             "cpi_issuing_body": 1, "cpi_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'CPI-MAPA-PMAV Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('cpi_certificate_name') or '')
            cert_s   = _v(doc.get('cpi_staff_name') or '')
            expiry   = _v(doc.get('cpi_expiry_date') or '')
            issue    = _v(doc.get('cpi_issue_date') or '')
            issuer   = _v(doc.get('cpi_issuing_body') or '')
            fetched  = doc.get('cpi_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="cpi_mapa_pmav_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Employment Contract Signed details ─────────────────

@admin_bp.route('/live-staffs/cron/sync-employment-contract', methods=['GET', 'POST'])
def live_staff_cron_sync_employment_contract():
    """
    Cron job — processes ONE staff member per call.
    Finds "Employment Contract Signed" document, extracts details via Gemini.
    Saves: ec_contract_name, ec_staff_name, ec_signed_date,
           ec_employer_name, ec_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"ec_fetched": {"$exists": False}},
            {"ec_fetched": False},
            {"ec_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Employment Contracts already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["ec_fetched"]    = True
        fields["ec_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"ec_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"ec_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"ec_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"ec_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    ec_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'employment contract signed',
            'employment contract',
            'signed contract',
            'contract signed',
            'contract of employment',
        )) and d.get('url'):
            ec_doc = d
            break

    if not ec_doc:
        _mark_done({"ec_note": "no Employment Contract document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Employment Contract found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (ec_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"ec_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Employment Contract has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"ec_note": "document URL 404 — skipped",
                        "ec_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Employment Contract URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this Employment Contract:
1. Contract / document name (e.g. "Employment Contract", "Contract of Employment", "Staff Agreement")
2. Employee / staff name as printed on the contract
3. Date the contract was signed (signature date)
4. Employer name (the company or organisation offering the contract)

Return ONLY a JSON object — no markdown, no explanation:
{
  "contract_name": "<exact document title as printed>",
  "staff_name_on_doc": "<employee name as printed>",
  "signed_date": "<date the contract was signed as printed>",
  "employer_name": "<employer or company name>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result        = _cjson.loads(raw_out)
        contract_name = _v(result.get('contract_name') or '')
        staff_on_doc  = _v(result.get('staff_name_on_doc') or '')
        signed_date   = _v(result.get('signed_date') or '')
        employer_name = _v(result.get('employer_name') or '')

        _mark_done({
            "ec_contract_name": contract_name,
            "ec_staff_name":    staff_on_doc,
            "ec_signed_date":   signed_date,
            "ec_employer_name": employer_name,
            "ec_doc_url":       doc_url,
            "ec_doc_type":      ec_doc.get('document_type_name', ''),
            "ec_note":          "extracted successfully",
        })

        return jsonify({
            "success":          True,
            "email":            email,
            "staff_name":       full_name,
            "doc_found":        True,
            "contract_name":    contract_name,
            "staff_name_on_doc": staff_on_doc,
            "signed_date":      signed_date,
            "employer_name":    employer_name,
            "remaining_count":  max(0, remaining_total - 1),
            "message": (
                f"Employment Contract extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"ec_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"ec_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Employment Contracts to Excel ─────────────────────────────

@admin_bp.route('/live-staffs/export/employment-contract-xlsx')
@admin_required
def live_staff_export_employment_contract_xlsx():
    """Export Employment Contract details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "ec_contract_name": 1, "ec_staff_name": 1,
             "ec_signed_date": 1, "ec_employer_name": 1, "ec_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Employment Contracts'

        headers    = ['Sno', 'Staff Name', 'Email', 'Contract Name',
                      'Name on Contract', 'Signed Date', 'Employer', 'Status']
        col_widths = [5, 28, 36, 30, 28, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:H{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1        = doc.get('section_1_personal_details') or {}
            name      = _v(s1.get('full_name') or '')
            email     = _v(doc.get('email') or '')
            cont_n    = _v(doc.get('ec_contract_name') or '')
            cont_s    = _v(doc.get('ec_staff_name') or '')
            signed    = _v(doc.get('ec_signed_date') or '')
            employer  = _v(doc.get('ec_employer_name') or '')
            fetched   = doc.get('ec_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cont_n:
                status   = 'No Contract Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cont_n, cont_s, signed, employer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="employment_contracts_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Open Disclosure Certificate ─────────────────────────

@admin_bp.route('/live-staffs/cron/sync-open-disclosure', methods=['GET', 'POST'])
def live_staff_cron_sync_open_disclosure_new():
    """
    Cron job — processes ONE staff member per call.
    Finds "Open Disclosure" document, extracts details via Gemini AI.
    Saves: od_certificate_name, od_staff_name, od_expiry_date,
           od_issue_date, od_issuing_body, od_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"od_fetched": {"$exists": False}},
            {"od_fetched": False},
            {"od_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Open Disclosure certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["od_fetched"]    = True
        fields["od_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"od_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"od_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"od_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"od_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    od_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'open disclosure',
            'the open disclosure',
            'open disclosure certificate',
            'open disclosure training',
        )) and d.get('url'):
            od_doc = d
            break

    if not od_doc:
        _mark_done({"od_note": "no Open Disclosure document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Open Disclosure certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (od_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"od_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Open Disclosure doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"od_note": "document URL 404 — skipped", "od_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Open Disclosure doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Open Disclosure certificate or training record:
1. Certificate name (e.g. "Open Disclosure", "The Open Disclosure", "Open Disclosure Training")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "od_certificate_name": cert_name,
            "od_staff_name":       cert_staff,
            "od_expiry_date":      expiry_date,
            "od_issue_date":       issue_date,
            "od_issuing_body":     issuing_body,
            "od_doc_url":          doc_url,
            "od_doc_type":         od_doc.get('document_type_name', ''),
            "od_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Open Disclosure cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"od_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"od_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Open Disclosure certificates to Excel ─────────────────────

@admin_bp.route('/live-staffs/export/open-disclosure-xlsx')
@admin_required
def live_staff_export_open_disclosure_xlsx_new():
    """Export Open Disclosure certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "od_certificate_name": 1, "od_staff_name": 1,
             "od_expiry_date": 1, "od_issue_date": 1,
             "od_issuing_body": 1, "od_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Open Disclosure Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y',
                        '%B %Y','%b %Y','%d %B, %Y','%d %b, %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip().rstrip(','), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('od_certificate_name') or '')
            cert_s   = _v(doc.get('od_staff_name') or '')
            expiry   = _v(doc.get('od_expiry_date') or '')
            issue    = _v(doc.get('od_issue_date') or '')
            issuer   = _v(doc.get('od_issuing_body') or '')
            fetched  = doc.get('od_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="open_disclosure_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Reviewers list — rotated alternately per staff ────────────────────
_PCC_REVIEWERS = [
    'Letty Mathew',
    'Valencia Da Silva',
    'Ann Maria',
    'Audrey Maguire',
    'Liberata Gama',
]
_PCC_COMPLIANCE_OFFICER = 'Betsy Daniel'


def _build_pcc_docx(doc, reviewer_index=0):
    """
    Generate the PCC Self-Declaration & Risk Assessment DOCX for a staff member.
    Pure Python — uses python-docx only, no Node.js required.
    Returns bytes of the generated DOCX.
    """
    import io as _io
    from datetime import datetime as _dt, timedelta as _td
    from docx import Document as _Doc
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    s1          = doc.get('section_1_personal_details') or {}
    full_name   = _v(s1.get('full_name') or '')
    dob         = _v(s1.get('date_of_birth') or '')
    nationality = _v(s1.get('nationality') or '')
    role        = _v(doc.get('user_type') or '')
    first_shift = _v(doc.get('first_shift_date') or '')

    # Employment history for Section 3
    s5 = doc.get('section_5_employment_history') or {}
    emp_entries = [e for e in (s5.get('entries') or [])
                   if e.get('employer') or e.get('position')]

    # Signature image bytes from GCS (if available)
    sig_bytes = None
    sig_blob  = _v(doc.get('signature_gcs_blob') or '')
    if sig_blob:
        try:
            sig_bytes = _gcs_download(sig_blob)
        except Exception:
            sig_bytes = None

    # Reviewer rotation
    reviewer = _PCC_REVIEWERS[reviewer_index % len(_PCC_REVIEWERS)]

    # Parse first_shift date
    first_shift_formatted = ''
    date_reviewed         = ''
    if first_shift:
        for fmt in ('%d/%m/%Y','%Y-%m-%d','%d-%m-%Y','%d %B %Y','%d %b %Y'):
            try:
                d_shift = _dt.strptime(first_shift.strip(), fmt)
                first_shift_formatted = d_shift.strftime('%d %B %Y')
                # Date Reviewed = first shift + 1 day
                date_reviewed = (d_shift + _td(days=1)).strftime('%d %B %Y')
                break
            except Exception:
                continue

    # ── Helpers ──────────────────────────────────────────────────────
    NAVY  = RGBColor(0x1B, 0x3A, 0x6B)
    GREEN = RGBColor(0x2E, 0x9E, 0x44)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    GRAY  = RGBColor(0x64, 0x74, 0x8B)

    def _set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _set_cell_border(cell, color='CCCCCC', sz=4):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    str(sz))
            b.set(qn('w:color'), color)
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def _para_border_bottom(para, color='1B3A6B', sz=8):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    str(sz))
        bot.set(qn('w:color'), color)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _add_run(para, text, bold=False, size=10, color=None, italic=False):
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Arial'
        if color:
            run.font.color.rgb = color
        return run

    document = _Doc()

    # Page margins
    for section in document.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Default font
    document.styles['Normal'].font.name = 'Arial'
    document.styles['Normal'].font.size = Pt(10)

    def sp(before=4, after=4):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        return p

    def section_heading(text):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
        _para_border_bottom(p, '1B3A6B', 8)
        run = p.add_run('  ' + text + '  ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.font.color.rgb = WHITE
        # Shading on run
        rPr  = run._r.get_or_add_rPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1B3A6B')
        rPr.append(shd)
        return p

    def label_value(label, value, bold_label=True):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _add_run(p, label + '  ', bold=bold_label, size=10)
        _add_run(p, value if value else '_' * 35, bold=False, size=10)
        return p

    def body_text(text, size=9.5):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _add_run(p, text, size=size)
        return p

    def numbered_item(num, text):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.8)
        _add_run(p, f'{num}.  ', bold=True, size=9.5)
        _add_run(p, text, size=9.5)
        return p

    def checkbox_item(text, checked=False):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _add_run(p, ('☑' if checked else '☐') + '  ', size=11)
        _add_run(p, text, size=9.5)
        return p

    def blank_line(label):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        _add_run(p, label + '  ', bold=True, size=9.5)
        _add_run(p, '_' * 40, size=9.5, color=GRAY)
        return p

    # ── Header ────────────────────────────────────────────────────────
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_border_bottom(p, '2E9E44', 6)
    _add_run(p, 'XPRESS HEALTH', bold=True, size=18, color=NAVY)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, 'SELF-DECLARATION & INTERNATIONAL RISK ASSESSMENT FORM', bold=True, size=11, color=GRAY)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, 'International Police Clearance Certificate (PCC) — Residential History Declaration',
             italic=True, size=9, color=GRAY)
    sp(4, 8)

    # ── Section 1 ─────────────────────────────────────────────────────
    section_heading('SECTION 1 — EMPLOYEE DETAILS')
    sp(2)
    tbl = document.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'
    for row in tbl.rows:
        for cell in row.cells:
            for border_el in cell._tc.iter(qn('w:tcBorders')):
                border_el.getparent().remove(border_el)
    tbl.cell(0,0).text = ''; tbl.cell(0,1).text = ''
    tbl.cell(1,0).text = ''; tbl.cell(1,1).text = ''
    for p_, label_, val_ in [
        (tbl.cell(0,0).paragraphs[0], 'Employee Name:', full_name),
        (tbl.cell(0,1).paragraphs[0], 'Nationality:', nationality),
        (tbl.cell(1,0).paragraphs[0], 'Position / Role:', role),
        (tbl.cell(1,1).paragraphs[0], '', ''),
    ]:
        if label_:
            _add_run(p_, label_ + '  ', bold=True, size=10)
            _add_run(p_, val_ if val_ else '_' * 28, size=10)
    sp(4)

    # ── Section 2 ─────────────────────────────────────────────────────
    section_heading('SECTION 2 — PURPOSE OF THIS DECLARATION')
    body_text('As part of the recruitment and compliance process, an International Police Clearance Certificate (PCC) is required from all countries where you have resided. As you are currently unable to provide the required certificate(s), this Self-Declaration & Risk Assessment Form must be completed.')
    body_text('Completion of this form does not exempt you from the requirement to obtain and submit the PCC. You remain obligated to provide the certificate(s) as soon as reasonably practicable.')
    sp(2)

    # ── Section 3 — extract from extracted_cv via Gemini ────────────
    section_heading('SECTION 3 — INTERNATIONAL RESIDENTIAL HISTORY')
    body_text('Please list ALL countries (other than your current country of residence) where you have lived for six (6) months or more since the age of 18. Include the reason for your stay.')
    sp(2)

    # Parse employment history from extracted_cv using Gemini
    hist_rows = []
    extracted_cv_ = _v(doc.get('extracted_cv') or '')
    has_cv_ = (extracted_cv_ and not extracted_cv_.startswith('[')
               and extracted_cv_ != 'No doc found')

    IRELAND_TERMS = {
        'ireland', 'irish', 'republic of ireland', 'eire', 'éire',
        'ie', 'northern ireland', 'dublin', 'cork', 'galway', 'limerick',
        'waterford', 'kilkenny', 'wexford', 'wicklow', 'kildare', 'meath',
        'louth', 'monaghan', 'cavan', 'donegal', 'sligo', 'leitrim', 'roscommon',
        'mayo', 'galway', 'clare', 'tipperary', 'kilkenny', 'carlow', 'laois',
        'offaly', 'westmeath', 'longford', 'louth', 'kerry', 'cork', 'hse',
    }

    def _is_ireland(country, city_region=''):
        text = (country + ' ' + city_region).lower()
        return any(t in text for t in IRELAND_TERMS)

    if has_cv_:
        try:

            from google import genai as _genai2
            _gemini_key2 = __import__('os').environ.get('GEMINI_API_KEY', '')
            if _gemini_key2:
                _prompt2 = f"""You are a CV analyser. Extract all employment history entries from the CV text below.

For each job role extract:
- country: the country where the job was based
- city_region: the city, region, county or employer location
- from_date: start date (MM/YYYY format if possible)
- to_date: end date (MM/YYYY format) or "Present" if current
- reason: the job title / role description

Return ONLY a JSON array — no markdown:
[
  {{"country": "...", "city_region": "...", "from_date": "...", "to_date": "...", "reason": "..."}}
]

If country is not mentioned, infer from city/employer name if possible, otherwise leave blank.

CV TEXT:
{extracted_cv_[:6000]}
"""
                _client2 = _genai2.Client(api_key=_gemini_key2)
                _resp2   = _client2.models.generate_content(
                    model='gemini-2.5-flash', contents=_prompt2
                )
                _raw2 = (_resp2.text or '').strip()
                _raw2 = _re2.sub(r'^```(?:json)?\s*', '', _raw2, flags=_re2.MULTILINE)
                _raw2 = _re2.sub(r'```\s*$', '', _raw2, flags=_re2.MULTILINE).strip()
                _parsed2 = _cjson2.loads(_raw2)
                for _e in _parsed2[:20]:
                    country_    = _v(_e.get('country') or '')
                    city_region = _v(_e.get('city_region') or '')
                    # Skip Ireland entries
                    if _is_ireland(country_, city_region):
                        continue
                    hist_rows.append((
                        country_,
                        city_region,
                        _v(_e.get('from_date') or ''),
                        _v(_e.get('to_date') or 'Present'),
                        _v(_e.get('reason') or ''),
                    ))
        except Exception:
            pass

    # Fallback to section_5 entries if Gemini failed — also filter Ireland
    if not hist_rows:
        for e in emp_entries[:8]:
            country_    = _v(e.get('country') or '')
            city_region = _v(e.get('employer') or '')
            if _is_ireland(country_, city_region):
                continue
            hist_rows.append((
                country_,
                city_region,
                _v(e.get('from') or ''),
                _v(e.get('to') or 'Present'),
                _v(e.get('position') or ''),
            ))

    # Always show at least 4 blank rows (even if all entries were Ireland)
    while len(hist_rows) < 4:
        hist_rows.append(('', '', '', '', ''))

    # 4 columns — City/Region removed
    cols_   = ['Country', 'From (MM/YYYY)', 'To (MM/YYYY)', 'Reason for Stay']
    # Strip city_region (index 1) from each row
    hist_rows_4 = [(r[0], r[2], r[3], r[4]) for r in hist_rows]
    htbl  = document.add_table(rows=len(hist_rows_4) + 1, cols=4)
    htbl.style = 'Table Grid'
    for ci, hdr in enumerate(cols_):
        c  = htbl.cell(0, ci)
        _set_cell_bg(c, '1B3A6B')
        _set_cell_border(c, '1B3A6B')
        p_ = c.paragraphs[0]
        p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_, hdr, bold=True, size=9, color=WHITE)
    for ri, row_vals in enumerate(hist_rows_4, start=1):
        for ci, val in enumerate(row_vals):
            c = htbl.cell(ri, ci)
            _set_cell_border(c, 'CCCCCC')
            _add_run(c.paragraphs[0], val, size=9)
    sp(6)

    # ── Section 4 ─────────────────────────────────────────────────────
    section_heading('SECTION 4 — SELF-DECLARATION')
    body_text('I declare that:')
    numbered_item(1, 'I have accurately listed all countries in which I have resided for six (6) months or more since the age of 18.')
    numbered_item(2, 'I understand that I have not yet provided the required International Police Clearance Certificate(s) for the country/countries listed above.')
    numbered_item(3, 'I confirm that I have NEVER been convicted of a criminal offence in any country, nor am I currently the subject of any criminal investigation, prosecution, or pending criminal proceedings, EXCEPT as disclosed below.')
    sp(2)
    p = document.add_paragraph()
    _add_run(p, 'Disclosure of Criminal History (if applicable — leave blank if none):', italic=True, size=9, color=GRAY)
    for _ in range(3):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _para_border_bottom(p, 'CCCCCC', 4)
        _add_run(p, '', size=9)
    sp(4)

    # ── Section 5 ─────────────────────────────────────────────────────
    section_heading('SECTION 5 — RISK ACKNOWLEDGEMENT')
    body_text('I understand and acknowledge that:')
    numbered_item(1, 'Failure to provide the required International Police Clearance Certificate(s) may affect my compliance status and/or eligibility for certain assignments or postings.')
    numbered_item(2, 'My employment and/or continued engagement may be subject to additional compliance reviews, conditions, or restrictions until the PCC is received.')
    numbered_item(3, 'If any information provided in this declaration is found to be false, misleading, or incomplete, it may result in disciplinary action, withdrawal of employment offer, termination of engagement, or notification to relevant regulatory or statutory authorities.')
    numbered_item(4, 'I remain responsible for pursuing the required Police Clearance Certificate(s) and providing them to Xpress Health as soon as they become available.')
    sp(4)

    # ── Section 6 ─────────────────────────────────────────────────────
    section_heading('SECTION 6 — CONSENT')
    body_text('I authorise Xpress Health to:')
    numbered_item(1, 'Verify the information provided in this declaration where required.')
    numbered_item(2, 'Request additional information or supporting documentation relevant to my residential history or compliance status.')
    numbered_item(3, 'Conduct further background checks as permitted under applicable laws and regulations.')
    numbered_item(4, 'Retain this form as part of my personnel and compliance records.')
    sp(4)

    # ── Section 7 ─────────────────────────────────────────────────────
    section_heading('SECTION 7 — EMPLOYEE DECLARATION & SIGNATURE')
    body_text('I declare that the information provided in this form is true, complete, and accurate to the best of my knowledge and belief. I understand that knowingly providing false or misleading information may result in disciplinary action, including termination of employment or engagement.')
    sp(4)
    stbl = document.add_table(rows=2, cols=2)
    stbl.style = 'Table Grid'
    for row in stbl.rows:
        for cell in row.cells:
            for b in cell._tc.iter(qn('w:tcBorders')):
                b.getparent().remove(b)

    # Employee Signature cell — embed image if available
    sig_cell = stbl.cell(0, 0)
    _add_run(sig_cell.paragraphs[0], 'Employee Signature  ', bold=True, size=9.5)
    if sig_bytes:
        try:
            from docx.shared import Inches as _Inches
            import io as _sig_io
            sig_p = sig_cell.add_paragraph()
            sig_run = sig_p.add_run()
            sig_run.add_picture(_sig_io.BytesIO(sig_bytes), width=_Inches(1.5))
        except Exception:
            _add_run(sig_cell.paragraphs[0], '_' * 32, size=9.5, color=GRAY)
    else:
        _add_run(sig_cell.paragraphs[0], '_' * 32, size=9.5, color=GRAY)

    # Date cell = first shift date
    _add_run(stbl.cell(0,1).paragraphs[0], 'Date  ', bold=True, size=9.5)
    _add_run(stbl.cell(0,1).paragraphs[0], first_shift_formatted or '_' * 32, size=9.5)

    # Employee Full Name (Print) — use full_name from Section 1
    _add_run(stbl.cell(1,0).paragraphs[0], 'Employee Full Name (Print)  ', bold=True, size=9.5)
    _add_run(stbl.cell(1,0).paragraphs[0], full_name, size=9.5)
    sp(6)

    body_text('Compliance Decision:')
    checkbox_item('Acceptable — pending PCC submission', checked=True)
    checkbox_item('Further Information Required')
    checkbox_item('Escalated for Risk Review')
    checkbox_item('Not Accepted')
    sp(6)

    # ── For Office Use Only — just above footer text ──────────────────
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    rr = p.add_run('  FOR OFFICE USE ONLY  ')
    rr.bold = True; rr.font.size = Pt(12); rr.font.name = 'Arial'
    rr.font.color.rgb = WHITE
    rPr = rr._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1B3A6B')
    rPr.append(shd)
    sp(2)

    otbl = document.add_table(rows=4, cols=2)
    otbl.style = 'Table Grid'
    # Approval On = Date Reviewed + 1 day
    approval_on = ''
    if date_reviewed:
        try:
            _d = _dt.strptime(date_reviewed, '%d %B %Y')
            approval_on = (_d + _td(days=1)).strftime('%d %B %Y')
        except Exception:
            approval_on = date_reviewed

    office_rows = [
        ('Reviewed By:', reviewer),
        ('Date Reviewed:', date_reviewed),
        ('Approved By (Compliance Officer):', _PCC_COMPLIANCE_OFFICER),
        ('Approval On:', approval_on),
    ]
    for ri, (label_, val_) in enumerate(office_rows):
        lc = otbl.cell(ri, 0)
        vc = otbl.cell(ri, 1)
        _set_cell_border(lc, 'CCCCCC')
        _set_cell_border(vc, 'CCCCCC')
        if ri % 2 == 0:
            _set_cell_bg(lc, 'EFF6FF')
        _add_run(lc.paragraphs[0], label_, bold=True, size=9.5)
        _add_run(vc.paragraphs[0], val_, size=9.5)
    sp(6)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, 'Xpress Health — Confidential & Compliance Document  |  Not for distribution',
             italic=True, size=8, color=GRAY)

    buf = _io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _ai_pcc_col():
    from flask import current_app
    return current_app.db.live_staff_ai_pcc


# ── PCC Generate ──────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/ai-pcc/generate', methods=['POST'])
@admin_required
def live_staff_ai_pcc_generate():
    """Generate PCC Self-Declaration form for a staff member."""
    data     = request.get_json(silent=True) or {}
    staff_id = (data.get('staff_id') or '').strip()
    if not staff_id:
        return jsonify({"success": False, "error": "Missing staff_id"}), 400

    try:
        doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not doc:
            return jsonify({"success": False, "error": "Staff not found"}), 404

        s1        = doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff')
        emp_code  = _v(doc.get('employee_code') or '')

        # Rotate reviewer based on total docs count
        total = _ai_pcc_col().count_documents({})
        reviewer_index = total % len(_PCC_REVIEWERS)

        docx_bytes = _build_pcc_docx(doc, reviewer_index=reviewer_index)
        safe_name  = full_name.replace(' ', '_').replace('/', '_')
        filename   = f"PCC_{safe_name}.docx"
        gcs_blob   = f"pcc/{filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        col      = _ai_pcc_col()
        existing = col.find_one({"staff_id": staff_id})
        rec = {
            "staff_id":      staff_id,
            "staff_name":    full_name,
            "employee_code": emp_code,
            "filename":      filename,
            "gcs_blob":      gcs_blob,
            "reviewer":      _PCC_REVIEWERS[reviewer_index % len(_PCC_REVIEWERS)],
            "generated_at":  datetime.utcnow(),
        }
        if existing:
            col.update_one({"_id": existing["_id"]}, {"$set": rec})
            rec_id = str(existing["_id"])
        else:
            rec_id = str(col.insert_one(rec).inserted_id)

        download_url = _gcs_signed_url(gcs_blob) or ''
        return jsonify({
            "success":      True,
            "pcc_id":       rec_id,
            "staff_name":   full_name,
            "filename":     filename,
            "gcs_blob":     gcs_blob,
            "download_url": download_url,
            "reviewer":     _PCC_REVIEWERS[reviewer_index % len(_PCC_REVIEWERS)],
            "generated_at": datetime.utcnow().isoformat(),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── PCC Download ──────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/ai-pcc/download/<pcc_id>')
@admin_required
def live_staff_ai_pcc_download(pcc_id):
    """Download saved PCC DOCX from GCS."""
    try:
        rec = _ai_pcc_col().find_one({"_id": ObjectId(pcc_id)})
        if not rec or not rec.get('gcs_blob'):
            return jsonify({"success": False, "error": "PCC not found"}), 404
        docx_bytes = _gcs_download(rec['gcs_blob'])
        return Response(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f'attachment; filename="{rec["filename"]}"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── PCC Saved check ───────────────────────────────────────────────────

@admin_bp.route('/live-staffs/ai-pcc/saved/<staff_id>')
@admin_required
def live_staff_ai_pcc_saved(staff_id):
    """Check if a saved PCC exists for this staff member."""
    rec = _ai_pcc_col().find_one({"staff_id": staff_id})
    if rec:
        return jsonify({
            "success":    True,
            "found":      True,
            "pcc_id":     str(rec["_id"]),
            "filename":   rec.get("filename", ""),
            "reviewer":   rec.get("reviewer", ""),
            "generated_at": rec["generated_at"].strftime("%d %b %Y %H:%M") if rec.get("generated_at") else "",
        })
    return jsonify({"success": True, "found": False})


# ── PCC Upload (replace) ──────────────────────────────────────────────

@admin_bp.route('/live-staffs/ai-pcc/upload/<staff_id>', methods=['POST'])
@admin_required
def live_staff_ai_pcc_upload(staff_id):
    """Upload an edited PCC DOCX to replace saved version in GCS."""
    f = request.files.get('file')
    if not f:
        return jsonify({"success": False, "error": "No file uploaded"}), 400
    try:
        doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not doc:
            return jsonify({"success": False, "error": "Staff not found"}), 404
        s1        = doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff')
        safe_name = full_name.replace(' ', '_').replace('/', '_')
        filename  = f"PCC_{safe_name}.docx"
        gcs_blob  = f"pcc/{filename}"
        docx_bytes = f.read()
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        _ai_pcc_col().update_one(
            {"staff_id": staff_id},
            {"$set": {"gcs_blob": gcs_blob, "filename": filename,
                      "generated_at": datetime.utcnow()}},
            upsert=True
        )
        return jsonify({"success": True, "gcs_blob": gcs_blob, "filename": filename})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract GDPR Certificate ───────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-gdpr', methods=['GET', 'POST'])
def live_staff_cron_sync_gdpr():
    """
    Cron job — processes ONE staff member per call.
    Finds "Gdpr" document, extracts details via Gemini AI.
    Saves: gdpr_certificate_name, gdpr_staff_name, gdpr_expiry_date,
           gdpr_issue_date, gdpr_issuing_body, gdpr_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"gdpr_fetched": {"$exists": False}},
            {"gdpr_fetched": False},
            {"gdpr_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff GDPR certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["gdpr_fetched"]    = True
        fields["gdpr_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"gdpr_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"gdpr_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"gdpr_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"gdpr_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    gdpr_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'gdpr', 'general data protection', 'data protection',
            'gdpr certificate', 'gdpr training',
        )) and d.get('url'):
            gdpr_doc = d
            break

    if not gdpr_doc:
        _mark_done({"gdpr_note": "no GDPR document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No GDPR certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (gdpr_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"gdpr_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — GDPR doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"gdpr_note": "document URL 404 — skipped", "gdpr_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — GDPR doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this GDPR / Data Protection certificate or training record:
1. Certificate name (e.g. "GDPR Awareness", "General Data Protection Regulation Training", "Data Protection Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "gdpr_certificate_name": cert_name,
            "gdpr_staff_name":       cert_staff,
            "gdpr_expiry_date":      expiry_date,
            "gdpr_issue_date":       issue_date,
            "gdpr_issuing_body":     issuing_body,
            "gdpr_doc_url":          doc_url,
            "gdpr_doc_type":         gdpr_doc.get('document_type_name', ''),
            "gdpr_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"GDPR cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"gdpr_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"gdpr_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: GDPR certificates to Excel ───────────────────────────────

@admin_bp.route('/live-staffs/export/gdpr-xlsx')
@admin_required
def live_staff_export_gdpr_xlsx():
    """Export GDPR certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "gdpr_certificate_name": 1, "gdpr_staff_name": 1,
             "gdpr_expiry_date": 1, "gdpr_issue_date": 1,
             "gdpr_issuing_body": 1, "gdpr_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'GDPR Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('gdpr_certificate_name') or '')
            cert_s   = _v(doc.get('gdpr_staff_name') or '')
            expiry   = _v(doc.get('gdpr_expiry_date') or '')
            issue    = _v(doc.get('gdpr_issue_date') or '')
            issuer   = _v(doc.get('gdpr_issuing_body') or '')
            fetched  = doc.get('gdpr_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="gdpr_certificates_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Dignity At Work Certificate ─────────────────────────

@admin_bp.route('/live-staffs/cron/sync-dignity-at-work', methods=['GET', 'POST'])
def live_staff_cron_sync_dignity_at_work():
    """
    Cron job — processes ONE staff member per call.
    Finds "Dignity At Work" document, extracts details via Gemini AI.
    Saves: daw_certificate_name, daw_staff_name, daw_expiry_date,
           daw_issue_date, daw_issuing_body, daw_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"daw_fetched": {"$exists": False}},
            {"daw_fetched": False},
            {"daw_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Dignity At Work certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["daw_fetched"]    = True
        fields["daw_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"daw_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"daw_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"daw_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"daw_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    daw_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'dignity at work',
            'dignity & work',
            'dignity and work',
            'dignity at work certificate',
            'dignity at work training',
        )) and d.get('url'):
            daw_doc = d
            break

    if not daw_doc:
        _mark_done({"daw_note": "no Dignity At Work document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Dignity At Work certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (daw_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"daw_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Dignity At Work doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"daw_note": "document URL 404 — skipped", "daw_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Dignity At Work doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Dignity At Work certificate or training record:
1. Certificate name (e.g. "Dignity At Work", "Dignity & Respect at Work", "Dignity At Work Training Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "daw_certificate_name": cert_name,
            "daw_staff_name":       cert_staff,
            "daw_expiry_date":      expiry_date,
            "daw_issue_date":       issue_date,
            "daw_issuing_body":     issuing_body,
            "daw_doc_url":          doc_url,
            "daw_doc_type":         daw_doc.get('document_type_name', ''),
            "daw_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Dignity At Work cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"daw_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"daw_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Dignity At Work certificates to Excel ─────────────────────

@admin_bp.route('/live-staffs/export/dignity-at-work-xlsx')
@admin_required
def live_staff_export_dignity_at_work_xlsx():
    """Export Dignity At Work certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "daw_certificate_name": 1, "daw_staff_name": 1,
             "daw_expiry_date": 1, "daw_issue_date": 1,
             "daw_issuing_body": 1, "daw_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Dignity At Work Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('daw_certificate_name') or '')
            cert_s   = _v(doc.get('daw_staff_name') or '')
            expiry   = _v(doc.get('daw_expiry_date') or '')
            issue    = _v(doc.get('daw_issue_date') or '')
            issuer   = _v(doc.get('daw_issuing_body') or '')
            fetched  = doc.get('daw_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="dignity_at_work_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract HACCP/Food Safety Certificate ───────────────────────

@admin_bp.route('/live-staffs/cron/sync-haccp', methods=['GET', 'POST'])
def live_staff_cron_sync_haccp():
    """
    Cron job — processes ONE staff member per call.
    Finds "Haccp/Food Safety" document, extracts details via Gemini AI.
    Saves: haccp_certificate_name, haccp_staff_name, haccp_expiry_date,
           haccp_issue_date, haccp_issuing_body, haccp_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"haccp_fetched": {"$exists": False}},
            {"haccp_fetched": False},
            {"haccp_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff HACCP/Food Safety certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["haccp_fetched"]    = True
        fields["haccp_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"haccp_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"haccp_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"haccp_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"haccp_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    haccp_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'haccp/food safety', 'haccp', 'food safety',
            'food hygiene', 'food safety certificate',
            'haccp certificate', 'food safety training',
        )) and d.get('url'):
            haccp_doc = d
            break

    if not haccp_doc:
        _mark_done({"haccp_note": "no HACCP/Food Safety document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No HACCP/Food Safety certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (haccp_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"haccp_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — HACCP doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"haccp_note": "document URL 404 — skipped", "haccp_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — HACCP doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this HACCP / Food Safety certificate or training record:
1. Certificate name (e.g. "HACCP", "Food Safety Certificate", "Food Hygiene & Safety", "HACCP/Food Safety Training")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "haccp_certificate_name": cert_name,
            "haccp_staff_name":       cert_staff,
            "haccp_expiry_date":      expiry_date,
            "haccp_issue_date":       issue_date,
            "haccp_issuing_body":     issuing_body,
            "haccp_doc_url":          doc_url,
            "haccp_doc_type":         haccp_doc.get('document_type_name', ''),
            "haccp_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"HACCP/Food Safety cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"haccp_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"haccp_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: HACCP/Food Safety certificates to Excel ───────────────────

@admin_bp.route('/live-staffs/export/haccp-xlsx')
@admin_required
def live_staff_export_haccp_xlsx():
    """Export HACCP/Food Safety certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "haccp_certificate_name": 1, "haccp_staff_name": 1,
             "haccp_expiry_date": 1, "haccp_issue_date": 1,
             "haccp_issuing_body": 1, "haccp_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'HACCP-Food Safety Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('haccp_certificate_name') or '')
            cert_s   = _v(doc.get('haccp_staff_name') or '')
            expiry   = _v(doc.get('haccp_expiry_date') or '')
            issue    = _v(doc.get('haccp_issue_date') or '')
            issuer   = _v(doc.get('haccp_issuing_body') or '')
            fetched  = doc.get('haccp_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="haccp_food_safety_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Cyber Security Certificate ─────────────────────────

@admin_bp.route('/live-staffs/cron/sync-cyber-security', methods=['GET', 'POST'])
def live_staff_cron_sync_cyber_security():
    """
    Cron job — processes ONE staff member per call.
    Finds "Cyber Security" document, extracts details via Gemini AI.
    Saves: cs_certificate_name, cs_staff_name, cs_expiry_date,
           cs_issue_date, cs_issuing_body, cs_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"cs_fetched": {"$exists": False}},
            {"cs_fetched": False},
            {"cs_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Cyber Security certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["cs_fetched"]    = True
        fields["cs_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"cs_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"cs_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"cs_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"cs_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    cs_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'cyber security', 'cybersecurity',
            'cyber security certificate', 'cyber security training',
            'cyber awareness', 'information security',
        )) and d.get('url'):
            cs_doc = d
            break

    if not cs_doc:
        _mark_done({"cs_note": "no Cyber Security document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Cyber Security certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (cs_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"cs_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Cyber Security doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"cs_note": "document URL 404 — skipped", "cs_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Cyber Security doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Cyber Security certificate or training record:
1. Certificate name (e.g. "Cyber Security Awareness", "Cybersecurity Training Certificate", "Information Security Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "cs_certificate_name": cert_name,
            "cs_staff_name":       cert_staff,
            "cs_expiry_date":      expiry_date,
            "cs_issue_date":       issue_date,
            "cs_issuing_body":     issuing_body,
            "cs_doc_url":          doc_url,
            "cs_doc_type":         cs_doc.get('document_type_name', ''),
            "cs_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Cyber Security cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"cs_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"cs_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Cyber Security certificates to Excel ──────────────────────

@admin_bp.route('/live-staffs/export/cyber-security-xlsx')
@admin_required
def live_staff_export_cyber_security_xlsx():
    """Export Cyber Security certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "cs_certificate_name": 1, "cs_staff_name": 1,
             "cs_expiry_date": 1, "cs_issue_date": 1,
             "cs_issuing_body": 1, "cs_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Cyber Security Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('cs_certificate_name') or '')
            cert_s   = _v(doc.get('cs_staff_name') or '')
            expiry   = _v(doc.get('cs_expiry_date') or '')
            issue    = _v(doc.get('cs_issue_date') or '')
            issuer   = _v(doc.get('cs_issuing_body') or '')
            fetched  = doc.get('cs_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="cyber_security_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Health Declaration Form ────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-health-declaration', methods=['GET', 'POST'])
def live_staff_cron_sync_health_declaration():
    """
    Cron job — processes ONE staff member per call.
    Finds "Health Declaration Form" document, extracts details via Gemini AI.
    Saves: hdf_certificate_name, hdf_staff_name, hdf_signed_date,
           hdf_issuing_body, hdf_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"hdf_fetched": {"$exists": False}},
            {"hdf_fetched": False},
            {"hdf_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Health Declaration Forms already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["hdf_fetched"]    = True
        fields["hdf_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"hdf_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"hdf_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"hdf_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"hdf_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    hdf_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'health declaration form',
            'health declaration',
            'medical declaration',
            'occupational health declaration',
            'health questionnaire',
        )) and d.get('url'):
            hdf_doc = d
            break

    if not hdf_doc:
        _mark_done({"hdf_note": "no Health Declaration Form found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Health Declaration Form found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (hdf_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"hdf_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Health Declaration doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"hdf_note": "document URL 404 — skipped", "hdf_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Health Declaration doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this Health Declaration Form:
1. Document / form name (e.g. "Health Declaration Form", "Occupational Health Declaration", "Medical Declaration")
2. Staff / employee name as printed on the form
3. Date the form was signed or completed
4. Issuing body or organisation (if shown)

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact form title as printed>",
  "staff_name_on_cert": "<name as printed on form>",
  "signed_date": "<date the form was signed or completed>",
  "issuing_body": "<organization or company name if shown>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        signed_date  = _v(result.get('signed_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "hdf_certificate_name": cert_name,
            "hdf_staff_name":       cert_staff,
            "hdf_signed_date":      signed_date,
            "hdf_issuing_body":     issuing_body,
            "hdf_doc_url":          doc_url,
            "hdf_doc_type":         hdf_doc.get('document_type_name', ''),
            "hdf_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "signed_date":        signed_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Health Declaration Form extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"hdf_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"hdf_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Health Declaration Forms to Excel ─────────────────────────

@admin_bp.route('/live-staffs/export/health-declaration-xlsx')
@admin_required
def live_staff_export_health_declaration_xlsx():
    """Export Health Declaration Form details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "hdf_certificate_name": 1, "hdf_staff_name": 1,
             "hdf_signed_date": 1, "hdf_issuing_body": 1, "hdf_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Health Declaration Forms'

        headers    = ['Sno', 'Staff Name', 'Email', 'Form Name',
                      'Name on Form', 'Signed Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:H{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('hdf_certificate_name') or '')
            cert_s   = _v(doc.get('hdf_staff_name') or '')
            signed   = _v(doc.get('hdf_signed_date') or '')
            issuer   = _v(doc.get('hdf_issuing_body') or '')
            fetched  = doc.get('hdf_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Form Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, signed, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="health_declaration_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Generate PCC Self-Declaration form one staff at a time ──────

@admin_bp.route('/live-staffs/cron/generate-pcc', methods=['GET', 'POST'])
def live_staff_cron_generate_pcc():
    """
    Cron job — generates PCC Self-Declaration form for ONE staff member per call.

    Finds staff where pcc_generated is not True.
    Generates DOCX, uploads to GCS, saves to live_staff_ai_pcc collection.
    Reviewer rotates through _PCC_REVIEWERS alternately.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    col     = _staffs_col()
    pcc_col = _ai_pcc_col()

    pending_query = {
        "$or": [
            {"pcc_generated": {"$exists": False}},
            {"pcc_generated": False},
            {"pcc_generated": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff PCC forms already generated.",
            "remaining_count": 0,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')
    emp_code  = _v(staff.get('employee_code') or '')

    def _mark_done(fields):
        fields["pcc_generated"]    = True
        fields["pcc_generated_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    try:
        # Rotate reviewer based on total generated so far
        total_generated    = pcc_col.count_documents({})
        reviewer_index     = total_generated % len(_PCC_REVIEWERS)

        docx_bytes = _build_pcc_docx(staff, reviewer_index=reviewer_index)

        safe_name  = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        filename   = f"PCC_{safe_name}.docx"
        gcs_blob   = f"pcc/{filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        existing = pcc_col.find_one({"staff_id": staff_id})
        rec = {
            "staff_id":      staff_id,
            "staff_name":    full_name,
            "employee_code": emp_code,
            "filename":      filename,
            "gcs_blob":      gcs_blob,
            "reviewer":      _PCC_REVIEWERS[reviewer_index % len(_PCC_REVIEWERS)],
            "generated_at":  datetime.utcnow(),
        }
        if existing:
            pcc_col.update_one({"_id": existing["_id"]}, {"$set": rec})
            rec_id = str(existing["_id"])
        else:
            rec_id = str(pcc_col.insert_one(rec).inserted_id)

        download_url = _gcs_signed_url(gcs_blob) or ''
        _mark_done({
            "pcc_gcs_blob":    gcs_blob,
            "pcc_filename":    filename,
            "pcc_download_url": download_url,
        })

        return jsonify({
            "success":         True,
            "staff_name":      full_name,
            "email":           email,
            "pcc_id":          rec_id,
            "filename":        filename,
            "gcs_blob":        gcs_blob,
            "download_url":    download_url,
            "reviewer":        _PCC_REVIEWERS[reviewer_index % len(_PCC_REVIEWERS)],
            "remaining_count": max(0, remaining_total - 1),
            "message": (
                f"PCC generated for {full_name} — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except Exception as e:
        _mark_done({"pcc_note": f"error: {e}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "staff_name":      full_name,
            "error":           str(e),
            "remaining_count": max(0, remaining_total - 1),
        })



# ── Export: Profile Check (matching Profile_check.xlsx format) ────────

@admin_bp.route('/live-staffs/export/profile-check-xlsx')
@admin_required
def live_staff_export_profile_check_xlsx():
    """
    Export staff certificate profile check to Excel.
    Sheets: Passport, CPR, IPC, Handhygiene, Safeguard, Garda, Children
    Matches the structure of Profile_check.xlsx.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import (Font, PatternFill, Alignment,
                                     Border, Side, numbers)
        import io as _io
        from datetime import date as _date
        today = _date.today()

        NAVY   = '1B3A6B'; GREEN  = '2E9E44'; WHITE  = 'FFFFFF'
        ALT    = 'EFF6FF'; RED    = 'FFDDDD'; WARN   = 'FFF3CD'
        MISMATCH_RED = 'FFE0E0'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        b_font  = Font(name='Arial', size=9)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        def _name_match(staff_name, cert_name):
            """Check if staff name matches cert name (simple fuzzy)."""
            if not staff_name or not cert_name:
                return True
            s = staff_name.strip().lower()
            c = cert_name.strip().lower()
            if s == c:
                return True
            # Check if all parts of cert name appear in staff name or vice versa
            s_parts = set(s.split())
            c_parts = set(c.split())
            return len(s_parts & c_parts) >= min(2, len(c_parts))

        def _fmt_date(val):
            if not val:
                return ''
            if hasattr(val, 'strftime'):
                return val.strftime('%d/%m/%Y')
            return str(val)

        def _is_expired(expiry):
            if not expiry:
                return None
            if hasattr(expiry, 'date'):
                return expiry.date() < today
            if hasattr(expiry, 'year'):
                return expiry < today
            # string
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y',
                        '%d %B, %Y','%d %b, %Y'):
                try:
                    from datetime import datetime as _dt2
                    d = _dt2.strptime(str(expiry).strip().rstrip(','), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        def _write_sheet(ws, headers, rows, col_widths):
            for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
                cell = ws.cell(row=1, column=ci, value=hdr)
                cell.font      = h_font
                cell.fill      = h_fill
                cell.alignment = h_align
                cell.border    = green_b
                ws.column_dimensions[
                    cell.column_letter].width = width
            ws.row_dimensions[1].height = 28
            ws.freeze_panes = 'A2'
            ws.auto_filter.ref = (
                f'A1:{ws.cell(row=1, column=len(headers)).column_letter}1'
            )
            for ri, row_data in enumerate(rows, start=2):
                is_mismatch = row_data.get('_mismatch', False)
                is_expired  = row_data.get('_expired', False)
                row_fill = (PatternFill('solid', start_color=RED,  end_color=RED)
                            if is_expired else
                            PatternFill('solid', start_color=WARN, end_color=WARN)
                            if is_mismatch else
                            PatternFill('solid', start_color=ALT,  end_color=ALT)
                            if ri % 2 == 0 else None)
                for ci, key in enumerate(row_data.get('_keys', []), start=1):
                    val  = row_data.get(key, '')
                    cell = ws.cell(row=ri, column=ci, value=val)
                    cell.font      = b_font
                    cell.alignment = c_align if ci == 1 else l_align
                    cell.border    = border
                    if row_fill:
                        cell.fill = row_fill
                ws.row_dimensions[ri].height = 16

        docs = list(_staffs_col().find({}, {
            'section_1_personal_details': 1, 'email': 1,
            'passport_id': 1, 'passport_data': 1, 'passport_fetched': 1,
            'cpr_certificate_name':1,'cpr_staff_name':1,'cpr_expiry_date':1,
            'cpr_issue_date':1,'cpr_issuing_body':1,'cpr_fetched':1,
            'ipc_certificate_name':1,'ipc_staff_name':1,'ipc_expiry_date':1,
            'ipc_issue_date':1,'ipc_issuing_body':1,'ipc_fetched':1,
            'hh_certificate_name':1,'hh_staff_name':1,'hh_expiry_date':1,
            'hh_issue_date':1,'hh_issuing_body':1,'hh_fetched':1,
            'sg_certificate_name':1,'sg_staff_name':1,'sg_expiry_date':1,
            'sg_issue_date':1,'sg_issuing_body':1,'sg_fetched':1,
            'garda_cert_name':1,'garda_staff_name':1,'garda_issue_date':1,
            'garda_reference':1,'garda_issuing_body':1,'garda_fetched':1,
            'cf_certificate_name':1,'cf_staff_name':1,'cf_expiry_date':1,
            'cf_issue_date':1,'cf_issuing_body':1,'cf_fetched':1,
        }))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        wb = Workbook()
        wb.remove(wb.active)  # remove default sheet

        # ── Helper to build a row dict ────────────────────────────────
        def _row(sno, name, email, cert_n, cert_s, expiry_raw, issue_raw,
                 issuer, status, extra=None):
            mismatch     = not _name_match(name, cert_s)
            mismatch_lbl = 'Mismatch' if mismatch else 'Match'
            expired      = _is_expired(expiry_raw) is True
            d = {
                'sno':      sno,
                'name':     name,
                'email':    email,
                'cert_n':   cert_n,
                'cert_s':   cert_s,
                'expiry':   _fmt_date(expiry_raw),
                'issue':    _fmt_date(issue_raw),
                'issuer':   issuer,
                'status':   status if status else ('Found' if cert_n else 'No Cert Found'),
                'mismatch': mismatch_lbl,
                '_mismatch': mismatch,
                '_expired':  expired,
            }
            if extra:
                d.update(extra)
            return d

        # ── Passport ──────────────────────────────────────────────────
        ws_p = wb.create_sheet('Passport')
        hdrs_p = ['Sno','Staff Name','Email','Passport ID','Nationality',
                  'Date of Birth','Expiry Date','Country','Status','Name Match']
        wids_p = [5,28,36,16,18,14,14,16,14,12]
        rows_p = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            pid   = _v(doc.get('passport_id') or '')
            pd_   = doc.get('passport_data') or {}
            nat   = _v(pd_.get('nationality') or '')
            dob   = _v(pd_.get('date_of_birth') or '')
            exp   = _v(pd_.get('expiry_date') or '')
            ctry  = _v(pd_.get('country') or '')
            st    = 'Found' if pid else ('Not Checked' if not doc.get('passport_fetched') else 'No ID Found')
            rows_p.append({
                '_keys': ['sno','name','email','pid','nat','dob','exp','ctry','st','match'],
                'sno': i, 'name': name, 'email': email, 'pid': pid,
                'nat': nat, 'dob': dob, 'exp': exp, 'ctry': ctry, 'st': st,
                'match': 'Match', '_mismatch': False, '_expired': False,
            })
        _write_sheet(ws_p, hdrs_p, rows_p, wids_p)

        # ── CPR ───────────────────────────────────────────────────────
        ws_c = wb.create_sheet('CPR')
        hdrs_c = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                  'Expiry Date','Issue Date','Issuing Body','Status','Staff Name Mismatch']
        wids_c = [5,28,36,32,28,14,14,28,14,14]
        rows_c = []
        for i, doc in enumerate(docs, start=1):
            s1   = doc.get('section_1_personal_details') or {}
            name = _v(s1.get('full_name') or '')
            email= _v(doc.get('email') or '')
            r    = _row(i,name,email,
                        _v(doc.get('cpr_certificate_name') or ''),
                        _v(doc.get('cpr_staff_name') or ''),
                        doc.get('cpr_expiry_date'),
                        doc.get('cpr_issue_date'),
                        _v(doc.get('cpr_issuing_body') or ''),
                        None if doc.get('cpr_fetched') else 'Not Checked')
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch']
            rows_c.append(r)
        _write_sheet(ws_c, hdrs_c, rows_c, wids_c)

        # ── IPC ───────────────────────────────────────────────────────
        ws_i = wb.create_sheet('IPC')
        hdrs_i = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                  'Expiry Date','Issue Date','Issuing Body','Status',
                  'Staff Name Mismatch','Reason for Mismatch']
        wids_i = [5,28,36,35,28,14,14,28,14,14,28]
        rows_i = []
        for i, doc in enumerate(docs, start=1):
            s1   = doc.get('section_1_personal_details') or {}
            name = _v(s1.get('full_name') or '')
            email= _v(doc.get('email') or '')
            cert_s = _v(doc.get('ipc_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            reason = (f'mismatch - {cert_s}' if mismatch and cert_s else '')
            r = _row(i,name,email,
                     _v(doc.get('ipc_certificate_name') or ''),
                     cert_s,
                     doc.get('ipc_expiry_date'),
                     doc.get('ipc_issue_date'),
                     _v(doc.get('ipc_issuing_body') or ''),
                     None if doc.get('ipc_fetched') else 'Not Checked')
            r['reason'] = reason
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch','reason']
            rows_i.append(r)
        _write_sheet(ws_i, hdrs_i, rows_i, wids_i)

        # ── Hand Hygiene ──────────────────────────────────────────────
        ws_h = wb.create_sheet('Handhygiene')
        hdrs_h = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                  'Expiry Date','Issue Date','Issuing Body','Status',
                  'Staff Name Mismatch','Reason for Mismatch']
        wids_h = [5,28,36,32,28,14,14,28,14,14,28]
        rows_h = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            cert_s= _v(doc.get('hh_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            reason = (f'mismatch - {cert_s}' if mismatch and cert_s else '')
            r = _row(i,name,email,
                     _v(doc.get('hh_certificate_name') or ''),
                     cert_s,
                     doc.get('hh_expiry_date'),
                     doc.get('hh_issue_date'),
                     _v(doc.get('hh_issuing_body') or ''),
                     None if doc.get('hh_fetched') else 'Not Checked')
            r['reason'] = reason
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch','reason']
            rows_h.append(r)
        _write_sheet(ws_h, hdrs_h, rows_h, wids_h)

        # ── Safeguard ─────────────────────────────────────────────────
        ws_s = wb.create_sheet('Safeguard')
        hdrs_s = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                  'Expiry Date','Issue Date','Issuing Body','Status',
                  'Staff Name Mismatch','Reason for Mismatch']
        wids_s = [5,28,36,38,28,14,14,28,14,14,28]
        rows_sg = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            cert_s= _v(doc.get('sg_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            reason = (f'mismatch - {cert_s}' if mismatch and cert_s else '')
            r = _row(i,name,email,
                     _v(doc.get('sg_certificate_name') or ''),
                     cert_s,
                     doc.get('sg_expiry_date'),
                     doc.get('sg_issue_date'),
                     _v(doc.get('sg_issuing_body') or ''),
                     None if doc.get('sg_fetched') else 'Not Checked')
            r['reason'] = reason
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch','reason']
            rows_sg.append(r)
        _write_sheet(ws_s, hdrs_s, rows_sg, wids_s)

        # ── Garda ─────────────────────────────────────────────────────
        ws_g = wb.create_sheet('Garda')
        hdrs_g = ['Sno','Staff Name','Email','Document Name','Name on Doc',
                  'Issue Date','Reference No','Issuing Body','Status',
                  'Staff Name Mismatch','Reason for Mismatch']
        wids_g = [5,28,36,32,28,14,24,28,14,14,28]
        rows_g = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            cert_s= _v(doc.get('garda_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            reason = (f'mismatch - {cert_s}' if mismatch and cert_s else '')
            cert_n= _v(doc.get('garda_cert_name') or '')
            r = {
                '_keys': ['sno','name','email','cert_n','cert_s',
                          'issue','ref_n','issuer','status','mismatch','reason'],
                'sno': i, 'name': name, 'email': email,
                'cert_n': cert_n, 'cert_s': cert_s,
                'issue': _fmt_date(doc.get('garda_issue_date')),
                'ref_n': _v(doc.get('garda_reference') or ''),
                'issuer': _v(doc.get('garda_issuing_body') or ''),
                'status': 'Found' if cert_n else ('Not Checked' if not doc.get('garda_fetched') else 'No Doc Found'),
                'mismatch': 'Mismatch' if mismatch else 'Match',
                'reason': reason,
                '_mismatch': mismatch, '_expired': False,
            }
            rows_g.append(r)
        _write_sheet(ws_g, hdrs_g, rows_g, wids_g)

        # ── Children First ────────────────────────────────────────────
        ws_ch = wb.create_sheet('Children')
        hdrs_ch = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                   'Expiry Date','Issue Date','Issuing Body','Status',
                   'Staff Name Mismatch','Expiry Date Exceeded']
        wids_ch = [5,28,36,38,28,14,14,28,14,14,16]
        rows_ch = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            cert_s= _v(doc.get('cf_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            exp_raw  = doc.get('cf_expiry_date')
            expired  = _is_expired(exp_raw) is True
            r = _row(i,name,email,
                     _v(doc.get('cf_certificate_name') or ''),
                     cert_s,
                     exp_raw,
                     doc.get('cf_issue_date'),
                     _v(doc.get('cf_issuing_body') or ''),
                     None if doc.get('cf_fetched') else 'Not Checked')
            r['expired_lbl'] = 'Expired' if expired else ''
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch','expired_lbl']
            rows_ch.append(r)
        _write_sheet(ws_ch, hdrs_ch, rows_ch, wids_ch)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="Profile_check_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Code Of Conduct Certificate ─────────────────────────

@admin_bp.route('/live-staffs/cron/sync-code-of-conduct', methods=['GET', 'POST'])
def live_staff_cron_sync_code_of_conduct():
    """
    Cron job — processes ONE staff member per call.
    Finds "Code Of Conduct" document, extracts details via Gemini AI.
    Saves: coc_certificate_name, coc_staff_name, coc_expiry_date,
           coc_issue_date, coc_issuing_body, coc_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"coc_fetched": {"$exists": False}},
            {"coc_fetched": False},
            {"coc_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Code Of Conduct certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["coc_fetched"]    = True
        fields["coc_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"coc_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"coc_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"coc_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"coc_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    coc_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'code of conduct', 'codeofconduct',
            'code of conduct certificate', 'code of conduct training',
            'staff code of conduct',
        )) and d.get('url'):
            coc_doc = d
            break

    if not coc_doc:
        _mark_done({"coc_note": "no Code Of Conduct document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Code Of Conduct certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (coc_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"coc_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Code Of Conduct doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"coc_note": "document URL 404 — skipped", "coc_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Code Of Conduct doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Code Of Conduct certificate or signed document:
1. Certificate / document name (e.g. "Code Of Conduct", "Staff Code of Conduct", "Code of Conduct Agreement")
2. Staff name as printed on the document
3. Expiry date or renewal date (if shown)
4. Issue / completion / signing date
5. Issuing body or organisation

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact document title as printed>",
  "staff_name_on_cert": "<name as printed on document>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue, signing or completion date as printed>",
  "issuing_body": "<organization that issued the document>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "coc_certificate_name": cert_name,
            "coc_staff_name":       cert_staff,
            "coc_expiry_date":      expiry_date,
            "coc_issue_date":       issue_date,
            "coc_issuing_body":     issuing_body,
            "coc_doc_url":          doc_url,
            "coc_doc_type":         coc_doc.get('document_type_name', ''),
            "coc_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Code Of Conduct extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"coc_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"coc_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Code Of Conduct to Excel ─────────────────────────────────

@admin_bp.route('/live-staffs/export/code-of-conduct-xlsx')
@admin_required
def live_staff_export_code_of_conduct_xlsx():
    """Export Code Of Conduct certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "coc_certificate_name": 1, "coc_staff_name": 1,
             "coc_expiry_date": 1, "coc_issue_date": 1,
             "coc_issuing_body": 1, "coc_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Code Of Conduct'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('coc_certificate_name') or '')
            cert_s   = _v(doc.get('coc_staff_name') or '')
            expiry   = _v(doc.get('coc_expiry_date') or '')
            issue    = _v(doc.get('coc_issue_date') or '')
            issuer   = _v(doc.get('coc_issuing_body') or '')
            fetched  = doc.get('coc_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="code_of_conduct_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract References Document ────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-references', methods=['GET', 'POST'])
def live_staff_cron_sync_references():
    """
    Cron job — processes ONE staff member per call.
    Finds "References" document, extracts details via Gemini AI.
    Saves: ref_document_name, ref_staff_name, ref_signed_date,
           ref_issuing_body, ref_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"ref_fetched": {"$exists": False}},
            {"ref_fetched": False},
            {"ref_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff References already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["ref_fetched"]    = True
        fields["ref_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"ref_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"ref_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"ref_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"ref_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    ref_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'references', 'reference letter', 'reference document',
            'employment reference', 'character reference', 'professional reference',
        )) and d.get('url'):
            ref_doc = d
            break

    if not ref_doc:
        _mark_done({"ref_note": "no References document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No References document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (ref_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"ref_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — References doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"ref_note": "document URL 404 — skipped", "ref_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — References doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this References document or reference letter:
1. Document name (e.g. "Reference Letter", "Employment Reference", "Character Reference", "References")
2. Name of the person the reference is for (the staff member / applicant)
3. Date the document was signed or issued
4. Name of the referee or issuing organisation (the person or company providing the reference)

Return ONLY a JSON object — no markdown, no explanation:
{
  "document_name": "<exact document title or type>",
  "staff_name_on_doc": "<name of the applicant / staff member the reference is for>",
  "signed_date": "<date the reference was signed or issued>",
  "referee_name": "<name of the person or organisation providing the reference>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        doc_name     = _v(result.get('document_name') or '')
        staff_on_doc = _v(result.get('staff_name_on_doc') or '')
        signed_date  = _v(result.get('signed_date') or '')
        referee_name = _v(result.get('referee_name') or '')

        _mark_done({
            "ref_document_name": doc_name,
            "ref_staff_name":    staff_on_doc,
            "ref_signed_date":   signed_date,
            "ref_referee_name":  referee_name,
            "ref_doc_url":       doc_url,
            "ref_doc_type":      ref_doc.get('document_type_name', ''),
            "ref_note":          "extracted successfully",
        })

        return jsonify({
            "success":          True,
            "email":            email,
            "staff_name":       full_name,
            "doc_found":        True,
            "document_name":    doc_name,
            "staff_name_on_doc": staff_on_doc,
            "signed_date":      signed_date,
            "referee_name":     referee_name,
            "remaining_count":  max(0, remaining_total - 1),
            "message": (
                f"References extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"ref_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"ref_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: References to Excel ───────────────────────────────────────

@admin_bp.route('/live-staffs/export/references-xlsx')
@admin_required
def live_staff_export_references_xlsx():
    """Export References document details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "ref_document_name": 1, "ref_staff_name": 1,
             "ref_signed_date": 1, "ref_referee_name": 1, "ref_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'References'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Signed Date', 'Referee Name', 'Status']
        col_widths = [5, 28, 36, 30, 28, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:H{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            doc_n    = _v(doc.get('ref_document_name') or '')
            doc_s    = _v(doc.get('ref_staff_name') or '')
            signed   = _v(doc.get('ref_signed_date') or '')
            referee  = _v(doc.get('ref_referee_name') or '')
            fetched  = doc.get('ref_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not doc_n:
                status   = 'No Doc Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, doc_n, doc_s, signed, referee, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="references_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Police Clearance Certificate ────────────────────────

@admin_bp.route('/live-staffs/cron/sync-police-clearance', methods=['GET', 'POST'])
def live_staff_cron_sync_police_clearance():
    """
    Cron job — processes ONE staff member per call.
    Finds "Police Clearance Certificate ( From Country Of Birth )" document.
    Saves: pcc2_certificate_name, pcc2_staff_name, pcc2_expiry_date,
           pcc2_issue_date, pcc2_issuing_body, pcc2_country, pcc2_fetched = True
    Note: pcc2_ prefix used to avoid conflict with PCC Self-Declaration (pcc_).
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"pcc2_fetched": {"$exists": False}},
            {"pcc2_fetched": False},
            {"pcc2_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Police Clearance Certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["pcc2_fetched"]    = True
        fields["pcc2_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"pcc2_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"pcc2_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"pcc2_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"pcc2_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    pcc2_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'police clearance certificate',
            'police clearance',
            'police certificate',
            'garda clearance',
            'criminal record certificate',
            'certificate of good conduct',
            'country of birth',
        )) and d.get('url'):
            pcc2_doc = d
            break

    if not pcc2_doc:
        _mark_done({"pcc2_note": "no Police Clearance Certificate found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Police Clearance Certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (pcc2_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"pcc2_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Police Clearance doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"pcc2_note": "document URL 404 — skipped", "pcc2_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Police Clearance doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Police Clearance Certificate (from country of birth):
1. Certificate name (e.g. "Police Clearance Certificate", "Certificate of Good Conduct", "Criminal Record Certificate")
2. Staff / applicant name as printed on the certificate
3. Country the certificate is issued for (country of birth)
4. Expiry date (if shown)
5. Issue / date of issue
6. Issuing body or authority (e.g. Police authority name, Ministry of Interior)

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<applicant name as printed>",
  "country": "<country this PCC was issued for>",
  "expiry_date": "<expiry date if shown>",
  "issue_date": "<date of issue as printed>",
  "issuing_body": "<issuing authority or organisation>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        country      = _v(result.get('country') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "pcc2_certificate_name": cert_name,
            "pcc2_staff_name":       cert_staff,
            "pcc2_country":          country,
            "pcc2_expiry_date":      expiry_date,
            "pcc2_issue_date":       issue_date,
            "pcc2_issuing_body":     issuing_body,
            "pcc2_doc_url":          doc_url,
            "pcc2_doc_type":         pcc2_doc.get('document_type_name', ''),
            "pcc2_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "country":            country,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Police Clearance Certificate extracted for {full_name} "
                f"(country: {country or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"pcc2_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"pcc2_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Police Clearance Certificates to Excel ────────────────────

@admin_bp.route('/live-staffs/export/police-clearance-xlsx')
@admin_required
def live_staff_export_police_clearance_xlsx():
    """Export Police Clearance Certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "pcc2_certificate_name": 1, "pcc2_staff_name": 1,
             "pcc2_country": 1, "pcc2_expiry_date": 1,
             "pcc2_issue_date": 1, "pcc2_issuing_body": 1, "pcc2_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Police Clearance Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Country', 'Expiry Date',
                      'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 18, 14, 14, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:J{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('pcc2_certificate_name') or '')
            cert_s   = _v(doc.get('pcc2_staff_name') or '')
            country  = _v(doc.get('pcc2_country') or '')
            expiry   = _v(doc.get('pcc2_expiry_date') or '')
            issue    = _v(doc.get('pcc2_issue_date') or '')
            issuer   = _v(doc.get('pcc2_issuing_body') or '')
            fetched  = doc.get('pcc2_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s,
                        country, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        l_align, c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="police_clearance_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Fire Safety Certificate ────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-fire-safety', methods=['GET', 'POST'])
def live_staff_cron_sync_fire_safety():
    """
    Cron job — processes ONE staff member per call.
    Finds "Fire Safety" document, extracts details via Gemini AI.
    Saves: fs_certificate_name, fs_staff_name, fs_expiry_date,
           fs_issue_date, fs_issuing_body, fs_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"fs_fetched": {"$exists": False}},
            {"fs_fetched": False},
            {"fs_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Fire Safety certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["fs_fetched"]    = True
        fields["fs_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"fs_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"fs_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"fs_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"fs_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    fs_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'fire safety', 'fire safety certificate',
            'fire safety training', 'fire prevention',
            'fire awareness', 'fire warden',
        )) and d.get('url'):
            fs_doc = d
            break

    if not fs_doc:
        _mark_done({"fs_note": "no Fire Safety document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Fire Safety certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (fs_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"fs_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Fire Safety doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"fs_note": "document URL 404 — skipped", "fs_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Fire Safety doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Fire Safety certificate or training record:
1. Certificate name (e.g. "Fire Safety", "Fire Safety Awareness", "Fire Warden Training Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "fs_certificate_name": cert_name,
            "fs_staff_name":       cert_staff,
            "fs_expiry_date":      expiry_date,
            "fs_issue_date":       issue_date,
            "fs_issuing_body":     issuing_body,
            "fs_doc_url":          doc_url,
            "fs_doc_type":         fs_doc.get('document_type_name', ''),
            "fs_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Fire Safety cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"fs_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"fs_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Fire Safety certificates to Excel ─────────────────────────

@admin_bp.route('/live-staffs/export/fire-safety-xlsx')
@admin_required
def live_staff_export_fire_safety_xlsx():
    """Export Fire Safety certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "fs_certificate_name": 1, "fs_staff_name": 1,
             "fs_expiry_date": 1, "fs_issue_date": 1,
             "fs_issuing_body": 1, "fs_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Fire Safety Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('fs_certificate_name') or '')
            cert_s   = _v(doc.get('fs_staff_name') or '')
            expiry   = _v(doc.get('fs_expiry_date') or '')
            issue    = _v(doc.get('fs_issue_date') or '')
            issuer   = _v(doc.get('fs_issuing_body') or '')
            fetched  = doc.get('fs_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="fire_safety_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract QQI Level 5 Or Equivalent Certificate ──────────────

@admin_bp.route('/live-staffs/cron/sync-qqi-level5', methods=['GET', 'POST'])
def live_staff_cron_sync_qqi_level5():
    """
    Cron job — processes ONE staff member per call.
    Finds "Qqi Level 5 Or Equivalent" document, extracts details via Gemini AI.
    Saves: qqi5_certificate_name, qqi5_staff_name, qqi5_expiry_date,
           qqi5_issue_date, qqi5_issuing_body, qqi5_award_level, qqi5_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"qqi5_fetched": {"$exists": False}},
            {"qqi5_fetched": False},
            {"qqi5_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff QQI Level 5 certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["qqi5_fetched"]    = True
        fields["qqi5_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"qqi5_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"qqi5_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"qqi5_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"qqi5_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    qqi5_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'qqi level 5', 'qqi level5', 'qqi 5',
            'level 5', 'level5',
            'qqi level 5 or equivalent',
            'fetac level 5',
            'care skills', 'healthcare support',
        )) and d.get('url'):
            qqi5_doc = d
            break

    if not qqi5_doc:
        _mark_done({"qqi5_note": "no QQI Level 5 document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No QQI Level 5 certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (qqi5_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"qqi5_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — QQI Level 5 doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"qqi5_note": "document URL 404 — skipped", "qqi5_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — QQI Level 5 doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this QQI Level 5 (or equivalent) qualification certificate:
1. Certificate / award name (e.g. "QQI Level 5 Healthcare Support", "FETAC Level 5 Care Skills", "NVQ Level 3 Health and Social Care")
2. Staff / student name as printed on the certificate
3. Award level (e.g. "Level 5", "Level 3", "QQI Level 5")
4. Date of issue or award date
5. Expiry date (if shown — most QQI certs do not expire)
6. Issuing body (e.g. "QQI", "FETAC", "City & Guilds", "BTEC")

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "award_level": "<award level, e.g. Level 5>",
  "issue_date": "<date of issue or award as printed>",
  "expiry_date": "<expiry date if shown, otherwise null>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        award_level  = _v(result.get('award_level') or '')
        issue_date   = _v(result.get('issue_date') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "qqi5_certificate_name": cert_name,
            "qqi5_staff_name":       cert_staff,
            "qqi5_award_level":      award_level,
            "qqi5_issue_date":       issue_date,
            "qqi5_expiry_date":      expiry_date,
            "qqi5_issuing_body":     issuing_body,
            "qqi5_doc_url":          doc_url,
            "qqi5_doc_type":         qqi5_doc.get('document_type_name', ''),
            "qqi5_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "award_level":        award_level,
            "issue_date":         issue_date,
            "expiry_date":        expiry_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"QQI Level 5 cert extracted for {full_name} "
                f"({award_level or 'Level 5'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"qqi5_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"qqi5_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: QQI Level 5 certificates to Excel ─────────────────────────

@admin_bp.route('/live-staffs/export/qqi-level5-xlsx')
@admin_required
def live_staff_export_qqi_level5_xlsx():
    """Export QQI Level 5 Or Equivalent certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "qqi5_certificate_name": 1, "qqi5_staff_name": 1,
             "qqi5_award_level": 1, "qqi5_issue_date": 1,
             "qqi5_expiry_date": 1, "qqi5_issuing_body": 1, "qqi5_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'QQI Level 5'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Award Level', 'Issue Date',
                      'Expiry Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 38, 28, 12, 14, 14, 22, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:J{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('qqi5_certificate_name') or '')
            cert_s   = _v(doc.get('qqi5_staff_name') or '')
            level    = _v(doc.get('qqi5_award_level') or '')
            issue    = _v(doc.get('qqi5_issue_date') or '')
            expiry   = _v(doc.get('qqi5_expiry_date') or '')
            issuer   = _v(doc.get('qqi5_issuing_body') or '')
            fetched  = doc.get('qqi5_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s,
                        level, issue, expiry, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="qqi_level5_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract HSE Clearance Pass ─────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-hse-clearance', methods=['GET', 'POST'])
def live_staff_cron_sync_hse_clearance():
    """
    Cron job — processes ONE staff member per call.
    Finds "Hse Clearance Pass" document, extracts details via Gemini AI.
    Saves: hcp_certificate_name, hcp_staff_name, hcp_issue_date,
           hcp_expiry_date, hcp_issuing_body, hcp_pass_number, hcp_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"hcp_fetched": {"$exists": False}},
            {"hcp_fetched": False},
            {"hcp_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff HSE Clearance Passes already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["hcp_fetched"]    = True
        fields["hcp_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"hcp_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"hcp_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"hcp_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"hcp_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    hcp_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'hse clearance pass', 'hse clearance', 'hse pass',
            'health service executive clearance',
            'hse occupational health', 'hse pre-employment',
            'hse pre employment',
        )) and d.get('url'):
            hcp_doc = d
            break

    if not hcp_doc:
        _mark_done({"hcp_note": "no HSE Clearance Pass found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No HSE Clearance Pass found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (hcp_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"hcp_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — HSE Clearance doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"hcp_note": "document URL 404 — skipped", "hcp_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — HSE Clearance doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this HSE Clearance Pass or occupational health clearance document:
1. Document / certificate name (e.g. "HSE Clearance Pass", "Occupational Health Clearance", "Pre-Employment Health Assessment")
2. Staff / employee name as printed on the document
3. Pass / reference number (if shown)
4. Issue date or clearance date
5. Expiry date (if shown)
6. Issuing body (e.g. "HSE", "Health Service Executive", "Occupational Health Dept")

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact document title as printed>",
  "staff_name_on_cert": "<name as printed on document>",
  "pass_number": "<pass or reference number if visible>",
  "issue_date": "<date of issue or clearance as printed>",
  "expiry_date": "<expiry date if shown>",
  "issuing_body": "<organization that issued the clearance>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        pass_number  = _v(result.get('pass_number') or '')
        issue_date   = _v(result.get('issue_date') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "hcp_certificate_name": cert_name,
            "hcp_staff_name":       cert_staff,
            "hcp_pass_number":      pass_number,
            "hcp_issue_date":       issue_date,
            "hcp_expiry_date":      expiry_date,
            "hcp_issuing_body":     issuing_body,
            "hcp_doc_url":          doc_url,
            "hcp_doc_type":         hcp_doc.get('document_type_name', ''),
            "hcp_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "pass_number":        pass_number,
            "issue_date":         issue_date,
            "expiry_date":        expiry_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"HSE Clearance Pass extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"hcp_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"hcp_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: HSE Clearance Pass to Excel ──────────────────────────────

@admin_bp.route('/live-staffs/export/hse-clearance-xlsx')
@admin_required
def live_staff_export_hse_clearance_xlsx():
    """Export HSE Clearance Pass details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "hcp_certificate_name": 1, "hcp_staff_name": 1,
             "hcp_pass_number": 1, "hcp_issue_date": 1,
             "hcp_expiry_date": 1, "hcp_issuing_body": 1, "hcp_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'HSE Clearance Pass'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Pass Number', 'Issue Date',
                      'Expiry Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 18, 14, 14, 24, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:J{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('hcp_certificate_name') or '')
            cert_s   = _v(doc.get('hcp_staff_name') or '')
            pass_n   = _v(doc.get('hcp_pass_number') or '')
            issue    = _v(doc.get('hcp_issue_date') or '')
            expiry   = _v(doc.get('hcp_expiry_date') or '')
            issuer   = _v(doc.get('hcp_issuing_body') or '')
            fetched  = doc.get('hcp_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Doc Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s,
                        pass_n, issue, expiry, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="hse_clearance_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Driving Licence ────────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-driving-licence', methods=['GET', 'POST'])
def live_staff_cron_sync_driving_licence():
    """
    Cron job — processes ONE staff member per call.
    Finds "Driving Licence" document, extracts details via Gemini AI.
    Saves: dl_staff_name, dl_licence_number, dl_expiry_date,
           dl_issue_date, dl_country, dl_categories, dl_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"dl_fetched": {"$exists": False}},
            {"dl_fetched": False},
            {"dl_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Driving Licences already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["dl_fetched"]    = True
        fields["dl_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"dl_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"dl_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"dl_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"dl_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    dl_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'driving licence', 'driving license',
            'driver licence', 'driver license',
            "driver's licence", "driver's license",
        )) and d.get('url'):
            dl_doc = d
            break

    if not dl_doc:
        _mark_done({"dl_note": "no Driving Licence found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Driving Licence found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (dl_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"dl_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Driving Licence has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"dl_note": "document URL 404 — skipped", "dl_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Driving Licence URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this Driving Licence:
1. Full name of the licence holder as printed
2. Licence number (field 5 on EU licences)
3. Date of birth (field 3)
4. Issue date (field 4a)
5. Expiry date (field 4b)
6. Issuing country or authority (field 4c or issuing body)
7. Licence categories (field 9/10/11 or entitlement codes, e.g. "B", "AM, B", "C1, D1")

Return ONLY a JSON object — no markdown, no explanation:
{
  "staff_name_on_doc": "<full name as printed on licence>",
  "licence_number": "<licence number>",
  "date_of_birth": "<date of birth as printed>",
  "issue_date": "<date of issue as printed>",
  "expiry_date": "<expiry date as printed>",
  "issuing_country": "<issuing country or authority>",
  "categories": "<licence categories or entitlement codes>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result         = _cjson.loads(raw_out)
        staff_on_doc   = _v(result.get('staff_name_on_doc') or '')
        licence_number = _v(result.get('licence_number') or '')
        dob            = _v(result.get('date_of_birth') or '')
        issue_date     = _v(result.get('issue_date') or '')
        expiry_date    = _v(result.get('expiry_date') or '')
        country        = _v(result.get('issuing_country') or '')
        categories     = _v(result.get('categories') or '')

        _mark_done({
            "dl_staff_name":     staff_on_doc,
            "dl_licence_number": licence_number,
            "dl_dob":            dob,
            "dl_issue_date":     issue_date,
            "dl_expiry_date":    expiry_date,
            "dl_country":        country,
            "dl_categories":     categories,
            "dl_doc_url":        doc_url,
            "dl_doc_type":       dl_doc.get('document_type_name', ''),
            "dl_note":           "extracted successfully",
        })

        return jsonify({
            "success":          True,
            "email":            email,
            "staff_name":       full_name,
            "doc_found":        True,
            "staff_name_on_doc": staff_on_doc,
            "licence_number":   licence_number,
            "expiry_date":      expiry_date,
            "issue_date":       issue_date,
            "country":          country,
            "categories":       categories,
            "remaining_count":  max(0, remaining_total - 1),
            "message": (
                f"Driving Licence extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"dl_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"dl_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Driving Licences to Excel ────────────────────────────────

@admin_bp.route('/live-staffs/export/driving-licence-xlsx')
@admin_required
def live_staff_export_driving_licence_xlsx():
    """Export Driving Licence details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "dl_staff_name": 1, "dl_licence_number": 1, "dl_dob": 1,
             "dl_issue_date": 1, "dl_expiry_date": 1,
             "dl_country": 1, "dl_categories": 1, "dl_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Driving Licences'

        headers    = ['Sno', 'Staff Name', 'Email', 'Name on Licence',
                      'Licence Number', 'Date of Birth', 'Issue Date',
                      'Expiry Date', 'Country', 'Categories', 'Status']
        col_widths = [5, 28, 36, 28, 18, 14, 14, 14, 16, 16, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:K{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1        = doc.get('section_1_personal_details') or {}
            name      = _v(s1.get('full_name') or '')
            email     = _v(doc.get('email') or '')
            dl_name   = _v(doc.get('dl_staff_name') or '')
            dl_num    = _v(doc.get('dl_licence_number') or '')
            dob       = _v(doc.get('dl_dob') or '')
            issue     = _v(doc.get('dl_issue_date') or '')
            expiry    = _v(doc.get('dl_expiry_date') or '')
            country   = _v(doc.get('dl_country') or '')
            cats      = _v(doc.get('dl_categories') or '')
            fetched   = doc.get('dl_fetched', False)
            expired   = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not dl_num:
                status   = 'No Licence Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, dl_name, dl_num,
                        dob, issue, expiry, country, cats, status]
            aligns   = [c_align, l_align, l_align, l_align, c_align,
                        c_align, c_align, c_align, l_align, c_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="driving_licences_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Point Scale Document ───────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-point-scale', methods=['GET', 'POST'])
def live_staff_cron_sync_point_scale():
    """
    Cron job — processes ONE staff member per call.
    Finds "Point Scale Document" document, extracts details via Gemini AI.
    Saves: psd_document_name, psd_staff_name, psd_signed_date,
           psd_point_scale, psd_issuing_body, psd_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"psd_fetched": {"$exists": False}},
            {"psd_fetched": False},
            {"psd_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Point Scale Documents already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["psd_fetched"]    = True
        fields["psd_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"psd_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"psd_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"psd_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"psd_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    psd_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'point scale document', 'point scale',
            'points scale document', 'points scale',
            'salary scale', 'pay scale', 'incremental scale',
        )) and d.get('url'):
            psd_doc = d
            break

    if not psd_doc:
        _mark_done({"psd_note": "no Point Scale Document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Point Scale Document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (psd_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"psd_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Point Scale doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"psd_note": "document URL 404 — skipped", "psd_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Point Scale doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this Point Scale Document (salary/pay scale document):
1. Document name (e.g. "Point Scale Document", "Salary Scale", "Pay Scale Agreement")
2. Staff / employee name as printed on the document
3. Point scale or salary scale value (e.g. "Point 1", "Scale 5", "€32,000")
4. Date the document was signed or issued
5. Issuing body or employer name

Return ONLY a JSON object — no markdown, no explanation:
{
  "document_name": "<exact document title as printed>",
  "staff_name_on_doc": "<name as printed on document>",
  "point_scale": "<point scale, salary scale or pay value as shown>",
  "signed_date": "<date signed or issued as printed>",
  "issuing_body": "<employer or issuing organisation name>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        doc_name     = _v(result.get('document_name') or '')
        staff_on_doc = _v(result.get('staff_name_on_doc') or '')
        point_scale  = _v(result.get('point_scale') or '')
        signed_date  = _v(result.get('signed_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "psd_document_name": doc_name,
            "psd_staff_name":    staff_on_doc,
            "psd_point_scale":   point_scale,
            "psd_signed_date":   signed_date,
            "psd_issuing_body":  issuing_body,
            "psd_doc_url":       doc_url,
            "psd_doc_type":      psd_doc.get('document_type_name', ''),
            "psd_note":          "extracted successfully",
        })

        return jsonify({
            "success":          True,
            "email":            email,
            "staff_name":       full_name,
            "doc_found":        True,
            "document_name":    doc_name,
            "staff_name_on_doc": staff_on_doc,
            "point_scale":      point_scale,
            "signed_date":      signed_date,
            "issuing_body":     issuing_body,
            "remaining_count":  max(0, remaining_total - 1),
            "message": (
                f"Point Scale Document extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"psd_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"psd_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Point Scale Documents to Excel ───────────────────────────

@admin_bp.route('/live-staffs/export/point-scale-xlsx')
@admin_required
def live_staff_export_point_scale_xlsx():
    """Export Point Scale Document details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "psd_document_name": 1, "psd_staff_name": 1,
             "psd_point_scale": 1, "psd_signed_date": 1,
             "psd_issuing_body": 1, "psd_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Point Scale Documents'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Point Scale', 'Signed Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 30, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            doc_n    = _v(doc.get('psd_document_name') or '')
            doc_s    = _v(doc.get('psd_staff_name') or '')
            scale    = _v(doc.get('psd_point_scale') or '')
            signed   = _v(doc.get('psd_signed_date') or '')
            issuer   = _v(doc.get('psd_issuing_body') or '')
            fetched  = doc.get('psd_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not doc_n:
                status   = 'No Doc Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, doc_n, doc_s, scale, signed, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="point_scale_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Application Form ───────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-application-form', methods=['GET', 'POST'])
def live_staff_cron_sync_application_form():
    """
    Cron job — processes ONE staff member per call.
    Finds "Application Form" document, extracts details via Gemini AI.
    Saves: af_document_name, af_staff_name, af_signed_date,
           af_position_applied, af_issuing_body, af_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"af_fetched": {"$exists": False}},
            {"af_fetched": False},
            {"af_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Application Forms already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["af_fetched"]    = True
        fields["af_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"af_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"af_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"af_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"af_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    af_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'application form', 'job application',
            'employment application', 'application for employment',
        )) and d.get('url'):
            af_doc = d
            break

    if not af_doc:
        _mark_done({"af_note": "no Application Form found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Application Form found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (af_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"af_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Application Form has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"af_note": "document URL 404 — skipped", "af_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Application Form URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this Application Form (job / employment application):
1. Document name (e.g. "Application Form", "Job Application Form", "Application for Employment")
2. Applicant / staff name as printed on the form
3. Position or role applied for (if shown)
4. Date the form was signed or submitted
5. Organisation or employer the application is addressed to

Return ONLY a JSON object — no markdown, no explanation:
{
  "document_name": "<exact document title as printed>",
  "staff_name_on_doc": "<applicant name as printed on form>",
  "position_applied": "<position or role applied for>",
  "signed_date": "<date signed or submitted as printed>",
  "issuing_body": "<organisation or employer the form is addressed to>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result           = _cjson.loads(raw_out)
        doc_name         = _v(result.get('document_name') or '')
        staff_on_doc     = _v(result.get('staff_name_on_doc') or '')
        position_applied = _v(result.get('position_applied') or '')
        signed_date      = _v(result.get('signed_date') or '')
        issuing_body     = _v(result.get('issuing_body') or '')

        _mark_done({
            "af_document_name":    doc_name,
            "af_staff_name":       staff_on_doc,
            "af_position_applied": position_applied,
            "af_signed_date":      signed_date,
            "af_issuing_body":     issuing_body,
            "af_doc_url":          doc_url,
            "af_doc_type":         af_doc.get('document_type_name', ''),
            "af_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "document_name":      doc_name,
            "staff_name_on_doc":  staff_on_doc,
            "position_applied":   position_applied,
            "signed_date":        signed_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Application Form extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"af_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"af_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Application Forms to Excel ───────────────────────────────

@admin_bp.route('/live-staffs/export/application-form-xlsx')
@admin_required
def live_staff_export_application_form_xlsx():
    """Export Application Form details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "af_document_name": 1, "af_staff_name": 1,
             "af_position_applied": 1, "af_signed_date": 1,
             "af_issuing_body": 1, "af_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Application Forms'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Form', 'Position Applied', 'Signed Date',
                      'Organisation', 'Status']
        col_widths = [5, 28, 36, 28, 28, 28, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            doc_n    = _v(doc.get('af_document_name') or '')
            doc_s    = _v(doc.get('af_staff_name') or '')
            pos      = _v(doc.get('af_position_applied') or '')
            signed   = _v(doc.get('af_signed_date') or '')
            org      = _v(doc.get('af_issuing_body') or '')
            fetched  = doc.get('af_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not doc_n:
                status   = 'No Form Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, doc_n, doc_s, pos, signed, org, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        l_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="application_forms_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Interview Notes ────────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-interview-notes', methods=['GET', 'POST'])
def live_staff_cron_sync_interview_notes():
    """
    Cron job — processes ONE staff member per call.
    Finds "Interview Notes" document, extracts details via Gemini AI.
    Saves: in_document_name, in_staff_name, in_interview_date,
           in_interviewer, in_outcome, in_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"in_fetched": {"$exists": False}},
            {"in_fetched": False},
            {"in_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Interview Notes already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["in_fetched"]    = True
        fields["in_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"in_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"in_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"in_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"in_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    in_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'interview notes', 'interview note',
            'interview record', 'interview assessment',
            'interview feedback', 'interview evaluation',
        )) and d.get('url'):
            in_doc = d
            break

    if not in_doc:
        _mark_done({"in_note": "no Interview Notes document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Interview Notes found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (in_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"in_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Interview Notes has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"in_note": "document URL 404 — skipped", "in_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Interview Notes URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this Interview Notes document:
1. Document name (e.g. "Interview Notes", "Interview Assessment", "Interview Record")
2. Candidate / staff name as printed on the document
3. Date of the interview
4. Interviewer name(s) (if shown)
5. Interview outcome or recommendation (e.g. "Recommended", "Successful", "Not Recommended", "Pending")

Return ONLY a JSON object — no markdown, no explanation:
{
  "document_name": "<exact document title as printed>",
  "staff_name_on_doc": "<candidate name as printed>",
  "interview_date": "<date of interview as printed>",
  "interviewer": "<interviewer name(s) if visible>",
  "outcome": "<interview outcome or recommendation if visible>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result         = _cjson.loads(raw_out)
        doc_name       = _v(result.get('document_name') or '')
        staff_on_doc   = _v(result.get('staff_name_on_doc') or '')
        interview_date = _v(result.get('interview_date') or '')
        interviewer    = _v(result.get('interviewer') or '')
        outcome        = _v(result.get('outcome') or '')

        _mark_done({
            "in_document_name":  doc_name,
            "in_staff_name":     staff_on_doc,
            "in_interview_date": interview_date,
            "in_interviewer":    interviewer,
            "in_outcome":        outcome,
            "in_doc_url":        doc_url,
            "in_doc_type":       in_doc.get('document_type_name', ''),
            "in_note":           "extracted successfully",
        })

        return jsonify({
            "success":          True,
            "email":            email,
            "staff_name":       full_name,
            "doc_found":        True,
            "document_name":    doc_name,
            "staff_name_on_doc": staff_on_doc,
            "interview_date":   interview_date,
            "interviewer":      interviewer,
            "outcome":          outcome,
            "remaining_count":  max(0, remaining_total - 1),
            "message": (
                f"Interview Notes extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"in_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"in_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Interview Notes to Excel ─────────────────────────────────

@admin_bp.route('/live-staffs/export/interview-notes-xlsx')
@admin_required
def live_staff_export_interview_notes_xlsx():
    """Export Interview Notes details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "in_document_name": 1, "in_staff_name": 1,
             "in_interview_date": 1, "in_interviewer": 1,
             "in_outcome": 1, "in_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Interview Notes'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Interview Date', 'Interviewer',
                      'Outcome', 'Status']
        col_widths = [5, 28, 36, 28, 28, 16, 28, 20, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            doc_n    = _v(doc.get('in_document_name') or '')
            doc_s    = _v(doc.get('in_staff_name') or '')
            int_date = _v(doc.get('in_interview_date') or '')
            intrvwr  = _v(doc.get('in_interviewer') or '')
            outcome  = _v(doc.get('in_outcome') or '')
            fetched  = doc.get('in_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not doc_n:
                status   = 'No Notes Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, doc_n, doc_s,
                        int_date, intrvwr, outcome, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, l_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="interview_notes_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Visa Document ──────────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-visa', methods=['GET', 'POST'])
def live_staff_cron_sync_visa():
    """
    Cron job — processes ONE staff member per call.
    Finds "Visa" document, extracts details via Gemini AI.
    Saves: visa_document_name, visa_staff_name, visa_type,
           visa_issue_date, visa_expiry_date, visa_country, visa_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"visa_fetched": {"$exists": False}},
            {"visa_fetched": False},
            {"visa_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Visa documents already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["visa_fetched"]    = True
        fields["visa_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"visa_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"visa_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"visa_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"visa_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    visa_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'visa', 'work visa', 'work permit',
            'residence permit', 'stamp 4', 'stamp 1',
            'gnib card', 'irish residence permit', 'irp',
        )) and d.get('url'):
            visa_doc = d
            break

    if not visa_doc:
        _mark_done({"visa_note": "no Visa document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Visa document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (visa_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"visa_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Visa doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"visa_note": "document URL 404 — skipped", "visa_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Visa doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this Visa or residence permit document:
1. Document name (e.g. "Visa", "Work Permit", "Irish Residence Permit", "Stamp 4", "GNIB Card")
2. Holder name as printed on the document
3. Visa / permit type or category (e.g. "Work Permit", "Stamp 4", "Type D", "IRP")
4. Issue date
5. Expiry date
6. Issuing country or authority

Return ONLY a JSON object — no markdown, no explanation:
{
  "document_name": "<exact document title as printed>",
  "staff_name_on_doc": "<holder name as printed>",
  "visa_type": "<visa or permit type / category>",
  "issue_date": "<issue date as printed>",
  "expiry_date": "<expiry date as printed>",
  "issuing_country": "<issuing country or authority>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result          = _cjson.loads(raw_out)
        doc_name        = _v(result.get('document_name') or '')
        staff_on_doc    = _v(result.get('staff_name_on_doc') or '')
        visa_type       = _v(result.get('visa_type') or '')
        issue_date      = _v(result.get('issue_date') or '')
        expiry_date     = _v(result.get('expiry_date') or '')
        issuing_country = _v(result.get('issuing_country') or '')

        _mark_done({
            "visa_document_name":  doc_name,
            "visa_staff_name":     staff_on_doc,
            "visa_type":           visa_type,
            "visa_issue_date":     issue_date,
            "visa_expiry_date":    expiry_date,
            "visa_issuing_country": issuing_country,
            "visa_doc_url":        doc_url,
            "visa_doc_type":       visa_doc.get('document_type_name', ''),
            "visa_note":           "extracted successfully",
        })

        return jsonify({
            "success":           True,
            "email":             email,
            "staff_name":        full_name,
            "doc_found":         True,
            "document_name":     doc_name,
            "staff_name_on_doc": staff_on_doc,
            "visa_type":         visa_type,
            "issue_date":        issue_date,
            "expiry_date":       expiry_date,
            "issuing_country":   issuing_country,
            "remaining_count":   max(0, remaining_total - 1),
            "message": (
                f"Visa extracted for {full_name} ({email}) "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"visa_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"visa_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Visa documents to Excel ──────────────────────────────────

@admin_bp.route('/live-staffs/export/visa-xlsx')
@admin_required
def live_staff_export_visa_xlsx():
    """Export Visa document details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "visa_document_name": 1, "visa_staff_name": 1,
             "visa_type": 1, "visa_issue_date": 1,
             "visa_expiry_date": 1, "visa_issuing_country": 1, "visa_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Visa Documents'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Visa Type', 'Issue Date',
                      'Expiry Date', 'Issuing Country', 'Status']
        col_widths = [5, 28, 36, 26, 28, 18, 14, 14, 20, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:J{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1        = doc.get('section_1_personal_details') or {}
            name      = _v(s1.get('full_name') or '')
            email     = _v(doc.get('email') or '')
            doc_n     = _v(doc.get('visa_document_name') or '')
            doc_s     = _v(doc.get('visa_staff_name') or '')
            vtype     = _v(doc.get('visa_type') or '')
            issue     = _v(doc.get('visa_issue_date') or '')
            expiry    = _v(doc.get('visa_expiry_date') or '')
            country   = _v(doc.get('visa_issuing_country') or '')
            fetched   = doc.get('visa_fetched', False)
            expired   = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not doc_n:
                status   = 'No Visa Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, doc_n, doc_s,
                        vtype, issue, expiry, country, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        l_align, c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="visa_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Address Proof Document ─────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-address-proof', methods=['GET', 'POST'])
def live_staff_cron_sync_address_proof():
    """
    Cron job — processes ONE staff member per call.
    Finds "address_proof" / "Address proof" document, extracts details via Gemini AI.
    Saves: ap_document_name, ap_staff_name, ap_address, ap_issue_date,
           ap_issuing_body, ap_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"ap_fetched": {"$exists": False}},
            {"ap_fetched": False},
            {"ap_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Address Proof documents already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["ap_fetched"]    = True
        fields["ap_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"ap_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"ap_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"ap_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"ap_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    ap_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'address proof', 'address_proof', 'proof of address',
            'utility bill', 'bank statement', 'council tax',
            'proof of residence', 'residence proof',
        )) and d.get('url'):
            ap_doc = d
            break

    if not ap_doc:
        _mark_done({"ap_note": "no Address Proof document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Address Proof found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (ap_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"ap_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Address Proof has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"ap_note": "document URL 404 — skipped", "ap_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Address Proof URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this address proof document (e.g. utility bill, bank statement, council tax letter):
1. Document type / name (e.g. "Utility Bill", "Bank Statement", "Proof of Address", "Council Tax Letter")
2. Name of the person the document is addressed to
3. Full address as printed on the document
4. Date of the document (issue date or statement date)
5. Issuing organisation (e.g. "AIB Bank", "Electric Ireland", "Dublin City Council")

Return ONLY a JSON object — no markdown, no explanation:
{
  "document_name": "<document type or title>",
  "staff_name_on_doc": "<name as addressed on document>",
  "address": "<full address as printed>",
  "issue_date": "<date of document as printed>",
  "issuing_body": "<organisation that issued the document>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        doc_name     = _v(result.get('document_name') or '')
        staff_on_doc = _v(result.get('staff_name_on_doc') or '')
        address      = _v(result.get('address') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "ap_document_name": doc_name,
            "ap_staff_name":    staff_on_doc,
            "ap_address":       address,
            "ap_issue_date":    issue_date,
            "ap_issuing_body":  issuing_body,
            "ap_doc_url":       doc_url,
            "ap_doc_type":      ap_doc.get('document_type_name', ''),
            "ap_note":          "extracted successfully",
        })

        return jsonify({
            "success":           True,
            "email":             email,
            "staff_name":        full_name,
            "doc_found":         True,
            "document_name":     doc_name,
            "staff_name_on_doc": staff_on_doc,
            "address":           address,
            "issue_date":        issue_date,
            "issuing_body":      issuing_body,
            "remaining_count":   max(0, remaining_total - 1),
            "message": (
                f"Address Proof extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"ap_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"ap_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Address Proof to Excel ───────────────────────────────────

@admin_bp.route('/live-staffs/export/address-proof-xlsx')
@admin_required
def live_staff_export_address_proof_xlsx():
    """Export Address Proof document details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "ap_document_name": 1, "ap_staff_name": 1,
             "ap_address": 1, "ap_issue_date": 1,
             "ap_issuing_body": 1, "ap_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Address Proof'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Address', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 24, 28, 40, 14, 24, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            doc_n    = _v(doc.get('ap_document_name') or '')
            doc_s    = _v(doc.get('ap_staff_name') or '')
            address  = _v(doc.get('ap_address') or '')
            issue    = _v(doc.get('ap_issue_date') or '')
            issuer   = _v(doc.get('ap_issuing_body') or '')
            fetched  = doc.get('ap_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not doc_n:
                status   = 'No Doc Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, doc_n, doc_s, address, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        l_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                if ci == 6:  # Address — allow wrap
                    cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="address_proof_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract PPE Document ────────────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-ppe', methods=['GET', 'POST'])
def live_staff_cron_sync_ppe():
    """
    Cron job — processes ONE staff member per call.
    Finds "Ppe" document, extracts details via Gemini AI.
    Saves: ppe_document_name, ppe_staff_name, ppe_signed_date,
           ppe_issuing_body, ppe_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"ppe_fetched": {"$exists": False}},
            {"ppe_fetched": False},
            {"ppe_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff PPE documents already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["ppe_fetched"]    = True
        fields["ppe_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"ppe_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"ppe_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"ppe_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"ppe_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    ppe_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'ppe', 'personal protective equipment',
            'ppe acknowledgement', 'ppe form', 'ppe policy',
            'ppe agreement', 'ppe checklist',
        )) and d.get('url'):
            ppe_doc = d
            break

    if not ppe_doc:
        _mark_done({"ppe_note": "no PPE document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No PPE document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (ppe_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"ppe_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — PPE doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"ppe_note": "document URL 404 — skipped", "ppe_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — PPE doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this PPE (Personal Protective Equipment) document, acknowledgement or policy form:
1. Document name (e.g. "PPE Acknowledgement", "Personal Protective Equipment Form", "PPE Policy Agreement")
2. Staff / employee name as printed on the document
3. Date the document was signed or issued
4. Issuing body or employer name

Return ONLY a JSON object — no markdown, no explanation:
{
  "document_name": "<exact document title as printed>",
  "staff_name_on_doc": "<name as printed on document>",
  "signed_date": "<date signed or issued as printed>",
  "issuing_body": "<employer or issuing organisation name>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        doc_name     = _v(result.get('document_name') or '')
        staff_on_doc = _v(result.get('staff_name_on_doc') or '')
        signed_date  = _v(result.get('signed_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "ppe_document_name": doc_name,
            "ppe_staff_name":    staff_on_doc,
            "ppe_signed_date":   signed_date,
            "ppe_issuing_body":  issuing_body,
            "ppe_doc_url":       doc_url,
            "ppe_doc_type":      ppe_doc.get('document_type_name', ''),
            "ppe_note":          "extracted successfully",
        })

        return jsonify({
            "success":           True,
            "email":             email,
            "staff_name":        full_name,
            "doc_found":         True,
            "document_name":     doc_name,
            "staff_name_on_doc": staff_on_doc,
            "signed_date":       signed_date,
            "issuing_body":      issuing_body,
            "remaining_count":   max(0, remaining_total - 1),
            "message": (
                f"PPE document extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"ppe_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"ppe_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: PPE documents to Excel ───────────────────────────────────

@admin_bp.route('/live-staffs/export/ppe-xlsx')
@admin_required
def live_staff_export_ppe_xlsx():
    """Export PPE document details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "ppe_document_name": 1, "ppe_staff_name": 1,
             "ppe_signed_date": 1, "ppe_issuing_body": 1, "ppe_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'PPE Documents'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Signed Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 30, 28, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:H{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            doc_n    = _v(doc.get('ppe_document_name') or '')
            doc_s    = _v(doc.get('ppe_staff_name') or '')
            signed   = _v(doc.get('ppe_signed_date') or '')
            issuer   = _v(doc.get('ppe_issuing_body') or '')
            fetched  = doc.get('ppe_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not doc_n:
                status   = 'No Doc Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, doc_n, doc_s, signed, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="ppe_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Manual And People Handling Document ─────────────────

@admin_bp.route('/live-staffs/cron/sync-manual-handling', methods=['GET', 'POST'])
def live_staff_cron_sync_manual_handling():
    """
    Cron job — processes ONE staff member per call.
    Finds "Manual And People Handling Documents" document.
    Saves: mph_certificate_name, mph_staff_name, mph_expiry_date,
           mph_issue_date, mph_issuing_body, mph_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"mph_fetched": {"$exists": False}},
            {"mph_fetched": False},
            {"mph_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Manual Handling documents already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["mph_fetched"]    = True
        fields["mph_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"mph_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"mph_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"mph_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"mph_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    mph_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'manual and people handling',
            'manual handling',
            'people handling',
            'manual handling certificate',
            'manual handling training',
            'patient handling',
            'moving and handling',
        )) and d.get('url'):
            mph_doc = d
            break

    if not mph_doc:
        _mark_done({"mph_note": "no Manual Handling document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Manual Handling document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (mph_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"mph_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Manual Handling doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"mph_note": "document URL 404 — skipped", "mph_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Manual Handling doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Manual and People Handling training certificate or document:
1. Certificate name (e.g. "Manual Handling", "People Handling", "Manual & People Handling Training Certificate", "Moving and Handling")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "mph_certificate_name": cert_name,
            "mph_staff_name":       cert_staff,
            "mph_expiry_date":      expiry_date,
            "mph_issue_date":       issue_date,
            "mph_issuing_body":     issuing_body,
            "mph_doc_url":          doc_url,
            "mph_doc_type":         mph_doc.get('document_type_name', ''),
            "mph_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Manual Handling cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"mph_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"mph_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Manual Handling certificates to Excel ─────────────────────

@admin_bp.route('/live-staffs/export/manual-handling-xlsx')
@admin_required
def live_staff_export_manual_handling_xlsx():
    """Export Manual and People Handling certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "mph_certificate_name": 1, "mph_staff_name": 1,
             "mph_expiry_date": 1, "mph_issue_date": 1,
             "mph_issuing_body": 1, "mph_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Manual Handling'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 36, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('mph_certificate_name') or '')
            cert_s   = _v(doc.get('mph_staff_name') or '')
            expiry   = _v(doc.get('mph_expiry_date') or '')
            issue    = _v(doc.get('mph_issue_date') or '')
            issuer   = _v(doc.get('mph_issuing_body') or '')
            fetched  = doc.get('mph_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="manual_handling_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract NMBI Qualification Document ─────────────────────────

@admin_bp.route('/live-staffs/cron/sync-nmbi-qualification', methods=['GET', 'POST'])
def live_staff_cron_sync_nmbi_qualification():
    """
    Cron job — processes ONE staff member per call.
    Finds "Nmbi Qualification" document, extracts details via Gemini AI.
    Saves: nmbi_q_certificate_name, nmbi_q_staff_name, nmbi_q_reg_number,
           nmbi_q_expiry_date, nmbi_q_issue_date, nmbi_q_issuing_body, nmbi_q_fetched = True
    Note: nmbi_q_ prefix to avoid conflict with nmbi_number field from sync-qualification cron.
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"nmbi_q_fetched": {"$exists": False}},
            {"nmbi_q_fetched": False},
            {"nmbi_q_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff NMBI Qualification documents already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["nmbi_q_fetched"]    = True
        fields["nmbi_q_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"nmbi_q_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"nmbi_q_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"nmbi_q_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"nmbi_q_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    nmbi_q_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'nmbi qualification', 'nmbi', 'nursing and midwifery board',
            'nursing midwifery board', 'nmbi registration',
            'nmbi certificate', 'nmbi pin',
        )) and d.get('url'):
            nmbi_q_doc = d
            break

    if not nmbi_q_doc:
        _mark_done({"nmbi_q_note": "no NMBI Qualification document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No NMBI Qualification document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (nmbi_q_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"nmbi_q_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — NMBI doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"nmbi_q_note": "document URL 404 — skipped", "nmbi_q_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — NMBI doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor specialising in Irish nursing qualifications.

Extract the following details from this NMBI (Nursing and Midwifery Board of Ireland) qualification certificate or registration document:
1. Certificate / document name (e.g. "NMBI Registration Certificate", "NMBI Qualification", "Certificate of Registration")
2. Nurse / midwife name as printed on the document
3. NMBI PIN or registration number
4. Expiry date or renewal date (if shown)
5. Issue / registration date
6. Issuing body (should be "NMBI" or "Nursing and Midwifery Board of Ireland")

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact document title as printed>",
  "staff_name_on_cert": "<name as printed on document>",
  "registration_number": "<NMBI PIN or registration number>",
  "expiry_date": "<expiry or renewal date as printed>",
  "issue_date": "<issue or registration date as printed>",
  "issuing_body": "<issuing organisation>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result      = _cjson.loads(raw_out)
        cert_name   = _v(result.get('certificate_name') or '')
        cert_staff  = _v(result.get('staff_name_on_cert') or '')
        reg_number  = _v(result.get('registration_number') or '')
        expiry_date = _v(result.get('expiry_date') or '')
        issue_date  = _v(result.get('issue_date') or '')
        issuing_body= _v(result.get('issuing_body') or '')

        _mark_done({
            "nmbi_q_certificate_name": cert_name,
            "nmbi_q_staff_name":       cert_staff,
            "nmbi_q_reg_number":       reg_number,
            "nmbi_q_expiry_date":      expiry_date,
            "nmbi_q_issue_date":       issue_date,
            "nmbi_q_issuing_body":     issuing_body,
            "nmbi_q_doc_url":          doc_url,
            "nmbi_q_doc_type":         nmbi_q_doc.get('document_type_name', ''),
            "nmbi_q_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "registration_number": reg_number,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"NMBI Qualification extracted for {full_name} "
                f"(PIN: {reg_number or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"nmbi_q_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"nmbi_q_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: NMBI Qualification to Excel ──────────────────────────────

@admin_bp.route('/live-staffs/export/nmbi-qualification-xlsx')
@admin_required
def live_staff_export_nmbi_qualification_xlsx():
    """Export NMBI Qualification document details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "nmbi_q_certificate_name": 1, "nmbi_q_staff_name": 1,
             "nmbi_q_reg_number": 1, "nmbi_q_expiry_date": 1,
             "nmbi_q_issue_date": 1, "nmbi_q_issuing_body": 1, "nmbi_q_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'NMBI Qualification'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'NMBI PIN / Reg No', 'Expiry Date',
                      'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 18, 14, 14, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:J{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('nmbi_q_certificate_name') or '')
            cert_s   = _v(doc.get('nmbi_q_staff_name') or '')
            reg_n    = _v(doc.get('nmbi_q_reg_number') or '')
            expiry   = _v(doc.get('nmbi_q_expiry_date') or '')
            issue    = _v(doc.get('nmbi_q_issue_date') or '')
            issuer   = _v(doc.get('nmbi_q_issuing_body') or '')
            fetched  = doc.get('nmbi_q_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Doc Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s,
                        reg_n, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="nmbi_qualification_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Medication Management Certificate ───────────────────

@admin_bp.route('/live-staffs/cron/sync-medication-management', methods=['GET', 'POST'])
def live_staff_cron_sync_medication_management():
    """
    Cron job — processes ONE staff member per call.
    Finds "Medication Management" document, extracts details via Gemini AI.
    Saves: mm_certificate_name, mm_staff_name, mm_expiry_date,
           mm_issue_date, mm_issuing_body, mm_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"mm_fetched": {"$exists": False}},
            {"mm_fetched": False},
            {"mm_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Medication Management certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["mm_fetched"]    = True
        fields["mm_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"mm_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"mm_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"mm_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"mm_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    mm_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'medication management', 'medication administration',
            'medicines management', 'drug administration',
            'safe administration of medication',
            'medication management certificate',
            'medication management training',
        )) and d.get('url'):
            mm_doc = d
            break

    if not mm_doc:
        _mark_done({"mm_note": "no Medication Management document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Medication Management certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (mm_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"mm_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Medication Management doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"mm_note": "document URL 404 — skipped", "mm_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Medication Management doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Medication Management or Medication Administration training certificate:
1. Certificate name (e.g. "Medication Management", "Safe Administration of Medication", "Medicines Management Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "mm_certificate_name": cert_name,
            "mm_staff_name":       cert_staff,
            "mm_expiry_date":      expiry_date,
            "mm_issue_date":       issue_date,
            "mm_issuing_body":     issuing_body,
            "mm_doc_url":          doc_url,
            "mm_doc_type":         mm_doc.get('document_type_name', ''),
            "mm_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Medication Management cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"mm_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"mm_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Medication Management certificates to Excel ───────────────

@admin_bp.route('/live-staffs/export/medication-management-xlsx')
@admin_required
def live_staff_export_medication_management_xlsx():
    """Export Medication Management certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "mm_certificate_name": 1, "mm_staff_name": 1,
             "mm_expiry_date": 1, "mm_issue_date": 1,
             "mm_issuing_body": 1, "mm_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Medication Management'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 36, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('mm_certificate_name') or '')
            cert_s   = _v(doc.get('mm_staff_name') or '')
            expiry   = _v(doc.get('mm_expiry_date') or '')
            issue    = _v(doc.get('mm_issue_date') or '')
            issuer   = _v(doc.get('mm_issuing_body') or '')
            fetched  = doc.get('mm_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="medication_management_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


def _export_json(items):
    serialized = _serialize(items)
    payload    = json.dumps({"records": serialized}, indent=2, ensure_ascii=False)
    return Response(
        payload,
        mimetype='application/json',
        headers={"Content-Disposition": f'attachment; filename="live_staffs_{_now_slug()}.json"'}
    )


def _export_csv(items):
    flat_rows = [_flatten(_serialize(doc)) for doc in items]
    all_keys  = _ordered_keys(flat_rows)

    buf    = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=all_keys, extrasaction='ignore')
    writer.writeheader()
    for row in flat_rows:
        writer.writerow({k: row.get(k, '') for k in all_keys})

    return Response(
        buf.getvalue(),
        mimetype='text/csv',
        headers={"Content-Disposition": f'attachment; filename="live_staffs_{_now_slug()}.csv"'}
    )


# ── Internal builders ─────────────────────────────────────────────────

def _build_doc(data):
    """Build a MongoDB document from flat form POST data."""
    return {
        "recruitment_id": data.get('recruitment_id'),
        "email":          (data.get('email') or '').strip().lower(),
        "employee_code":  (data.get('employee_code') or '').strip(),
        "user_type":      (data.get('user_type') or '').strip(),
        "status":         (data.get('status') or 'active').strip(),

        "section_1_personal_details": {
            "full_name":      data.get('full_name', ''),
            "previous_names": data.get('previous_names', ''),
            "date_of_birth":  data.get('date_of_birth', ''),
            "address":        data.get('address', ''),
            "eircode_postcode": data.get('eircode_postcode', ''),
            "mobile_number":  data.get('mobile_number', ''),
            "email_address":  (data.get('email') or '').strip().lower(),
            "pps_number":     data.get('pps_number', ''),
            "nationality":    data.get('nationality', ''),
            "work_permit_visa_status": {
                "permission_to_work": data.get('permission_to_work', ''),
                "visa_type":          data.get('visa_type', ''),
            },
            "nmbi_pin_number": data.get('nmbi_pin_number', ''),
        },

        "section_2_identity_verification": {
            "passport_number":        data.get('passport_number', ''),
            "expiry_date":            data.get('passport_expiry', ''),
            "driving_licence_number": data.get('driving_licence_number', ''),
            "documents_submitted": {
                "passport":          bool(data.get('doc_passport')),
                "birth_certificate": bool(data.get('doc_birth_cert')),
                "driving_licence":   bool(data.get('doc_driving')),
                "proof_of_address":  bool(data.get('doc_address')),
            },
            "verified_by":       data.get('verified_by', ''),
            "verification_date": data.get('verification_date', ''),
        },

        "section_3_professional_registration": {
            "registration_number_pin":  data.get('registration_number_pin', ''),
            "divisions_registered_in":  data.get('divisions_registered_in', []),
            "registration_expiry_date": data.get('registration_expiry_date', ''),
            "nmbi_active_declaration":  bool(data.get('nmbi_active_declaration')),
        },

        "section_8_garda_vetting_police_clearance": {
            "garda_vetting_submitted":    bool(data.get('garda_vetting_submitted')),
            "police_clearance_submitted": bool(data.get('police_clearance_submitted')),
        },

        "section_9_occupational_health": {
            "occupational_health_screening": bool(data.get('occupational_health_screening')),
            "immunisation_records_provided": bool(data.get('immunisation_records_provided')),
            "fit_for_nursing_duties":        bool(data.get('fit_for_nursing_duties')),
            "covid_19_vaccine":              data.get('covid_19_vaccine', ''),
            "tuberculosis_vaccine":          data.get('tuberculosis_vaccine', ''),
            "hepatitis_antibody":            data.get('hepatitis_antibody', ''),
            "mmr_vaccine":                   data.get('mmr_vaccine', ''),
        },

        "section_10_mandatory_training": {
            "manual_handling":              data.get('manual_handling', ''),
            "cpr_bls":                      data.get('cpr_bls', ''),
            "fire_safety":                  data.get('fire_safety', ''),
            "infection_prevention_control": data.get('infection_prevention_control', ''),
            "hand_hygiene":                 data.get('hand_hygiene', ''),
            "safeguarding":                 data.get('safeguarding', ''),
            "children_first":               data.get('children_first', ''),
            "cyber_security":               data.get('cyber_security', ''),
            "dignity_at_work":              data.get('dignity_at_work', ''),
            "open_disclosure":              data.get('open_disclosure', ''),
            "mapa_pmav":                    data.get('mapa_pmav', ''),
        },
    }


def _map_import_record(rec):
    """Map a full JSON record into a clean MongoDB doc."""
    s1  = rec.get('section_1_personal_details', {})
    s2  = rec.get('section_2_identity_verification', {})
    s3  = rec.get('section_3_professional_registration', {})
    s4  = rec.get('section_4_qualifications', {})
    s5  = rec.get('section_5_employment_history', {})
    s6  = rec.get('section_6_employment_gaps', [])
    s7  = rec.get('section_7_references', {})
    s8  = rec.get('section_8_garda_vetting_police_clearance', {})
    s9  = rec.get('section_9_occupational_health', {})
    s10 = rec.get('section_10_mandatory_training', {})
    s11 = rec.get('section_11_criminal_convictions_declaration', {})
    s12 = rec.get('section_12_declaration', {})

    return {
        "recruitment_id": rec.get('recruitment_id'),
        "email":          (rec.get('email') or '').strip().lower(),
        "employee_code":  rec.get('employee_code', ''),
        "user_type":      rec.get('user_type', ''),
        "status":         rec.get('status', 'found'),

        "section_1_personal_details":                  s1,
        "section_2_identity_verification":             s2,
        "section_3_professional_registration":         s3,
        "section_4_qualifications":                    s4,
        "section_5_employment_history":                s5,
        "section_6_employment_gaps":                   s6,
        "section_7_references":                        s7,
        "section_8_garda_vetting_police_clearance":    s8,
        "section_9_occupational_health":               s9,
        "section_10_mandatory_training":               s10,
        "section_11_criminal_convictions_declaration": s11,
        "section_12_declaration":                      s12,
    }


# ── Flatten helpers for CSV export ───────────────────────────────────

def _flatten(doc, prefix='', result=None):
    if result is None:
        result = {}
    for k, v in doc.items():
        key = f"{prefix}{k}" if prefix else k
        if isinstance(v, dict):
            _flatten(v, key + '.', result)
        elif isinstance(v, list):
            result[key] = '; '.join(
                json.dumps(i) if isinstance(i, dict) else str(i)
                for i in v
            )
        else:
            result[key] = v if v is not None else ''
    return result


_PREFERRED_KEY_ORDER = [
    '_id', 'recruitment_id', 'employee_code', 'email', 'user_type', 'status', 'created_at'
]


def _ordered_keys(rows):
    seen    = set()
    ordered = []
    for k in _PREFERRED_KEY_ORDER:
        for row in rows:
            if k in row and k not in seen:
                ordered.append(k)
                seen.add(k)
                break
    for row in rows:
        for k in row:
            if k not in seen:
                ordered.append(k)
                seen.add(k)
    return ordered


def _now_slug():
    return datetime.utcnow().strftime('%Y%m%d_%H%M%S')
