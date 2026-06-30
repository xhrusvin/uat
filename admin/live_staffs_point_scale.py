"""
live_staffs_point_scale.py
══════════════════════════
Independent blueprint for generating "Verification of Service – Agency Staff"
(Point Scale Document) for Healthcare Assistant staff.

Registers on admin_bp — import in admin/__init__.py:
    from . import live_staffs_point_scale

Routes:
    GET  /live-staffs/cron/generate-point-scale   — cron (one staff per call)
    GET  /live-staffs/ai-cv/<staff_id>/point-scale/download — download generated doc
    GET  /live-staffs/export/point-scale-xlsx      — Excel export of all records
"""

from flask import request, jsonify, Response
from bson import ObjectId
from datetime import datetime, date
import os
import re
import io
import json
import threading

from database import db
from . import admin_bp
from admin.views import admin_required


# ── Helpers ───────────────────────────────────────────────────────────

def _v(val):
    if val is None: return ''
    return str(val).strip()

def _staffs_col():
    from flask import current_app
    return current_app.db.live_staffs

def _ps_col():
    """point_scale documents collection."""
    return db.live_staff_point_scale

def _gcs_upload(blob_name, data_bytes, content_type='application/octet-stream'):
    from admin.live_staffs import _gcs_upload as _f
    return _f(blob_name, data_bytes, content_type)

def _gcs_download(blob_name):
    from admin.live_staffs import _gcs_download as _f
    return _f(blob_name)


# ── Employment entry parser ───────────────────────────────────────────

def _parse_employment_from_cv(extracted_cv, user_type='', gemini_key=''):
    """
    Use Gemini to extract Ireland-only employment from extracted_cv.
    Returns list of dicts: {post, employer, from_str, to_str}
    """
    if not extracted_cv or not gemini_key:
        return []

    try:
        from google import genai as _gai
        client = _gai.Client(api_key=gemini_key)

        prompt = f"""You are an employment history extractor for an Irish healthcare agency.

Extract ONLY the roles where the candidate worked in IRELAND from the CV text below.
Exclude ALL roles based in the UK, England, Scotland, Wales, India, Italy, USA, Australia, Nigeria, Philippines, or any country/city outside Ireland.
Include roles where Location is "Ireland" or contains Irish counties/cities (Dublin, Cork, Galway, Kildare, Limerick, etc.) or where no location is given but employer sounds Irish.

For each Ireland-based role return:
- post: job title
- employer: employer name only (no location)
- from_str: start date as written (e.g. "November 2023", "03/09/2021")
- to_str: end date as written, or "Present" if current

Return ONLY a JSON array — no markdown, no explanation:
[
  {{"post": "Healthcare Assistant", "employer": "Homes Instead Ireland", "from_str": "November 2023", "to_str": "Present"}},
  ...
]

If no Ireland roles found, return empty array: []

CV TEXT:
{extracted_cv}
"""
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        raw = (response.text or '').strip()
        raw = re.sub(r'^```(?:json)?\s*', '', raw, flags=re.MULTILINE)
        raw = re.sub(r'```\s*$', '', raw, flags=re.MULTILINE).strip()

        import json as _json
        entries = _json.loads(raw)
        if isinstance(entries, list):
            return entries
    except Exception:
        pass

    return []


def _parse_date_flex(s):
    """Parse a date string flexibly. Returns date or None."""
    s = _v(s).strip()
    if not s or s.lower() in ('present', 'current', 'ongoing', 'now'):
        return None  # None = Present

    fmts = [
        '%d/%m/%Y', '%m/%Y', '%B %Y', '%b %Y',
        '%d %B %Y', '%d %b %Y', '%Y-%m-%d', '%Y',
    ]
    for fmt in fmts:
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    return None


def _calc_duration(from_date, to_date):
    """Calculate years, months, days between two dates."""
    if from_date is None:
        return 0, 0, 0
    end = to_date if to_date else date.today()
    if end < from_date:
        return 0, 0, 0

    years  = end.year  - from_date.year
    months = end.month - from_date.month
    days   = end.day   - from_date.day

    if days < 0:
        months -= 1
        # days in previous month
        from calendar import monthrange
        prev_month = end.month - 1 if end.month > 1 else 12
        prev_year  = end.year if end.month > 1 else end.year - 1
        days += monthrange(prev_year, prev_month)[1]

    if months < 0:
        years  -= 1
        months += 12

    return max(0, years), max(0, months), max(0, days)


def _total_service(rows):
    """Sum years, months, days across all rows, normalising overflow."""
    total_days = 0
    for r in rows:
        total_days += r['years'] * 365 + r['months'] * 30 + r['days']

    years  = total_days // 365
    remain = total_days  % 365
    months = remain // 30
    days   = remain  % 30
    return years, months, days


# ── DOCX builder ─────────────────────────────────────────────────────

def _build_point_scale_docx(staff_doc, rows):
    """
    Build "Verification of Service – Agency Staff" DOCX.
    rows: list of {post, employer, from_str, to_str, from_date, to_date, years, months, days}
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY  = RGBColor(0x1B, 0x3A, 0x6B)
    GREEN = RGBColor(0x2E, 0x9E, 0x44)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    GRAY  = RGBColor(0x55, 0x55, 0x55)

    s1        = staff_doc.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff_doc.get('email') or s1.get('email_address') or '')
    user_type = _v(staff_doc.get('user_type') or 'Healthcare Assistant')

    d = Document()
    for sec in d.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    d.styles['Normal'].font.name = 'Calibri'
    d.styles['Normal'].font.size = Pt(11)

    def sp(pts):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = Pt(pts)

    def _set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def _set_cell_border(cell, color='CCCCCC'):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   'single')
            border.set(qn('w:sz'),    '4')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def _add_run(para, text, bold=False, size=11, color=None, italic=False):
        run = para.add_run(text)
        run.bold        = bold
        run.italic      = italic
        run.font.name   = 'Calibri'
        run.font.size   = Pt(size)
        run.font.color.rgb = color if color else BLACK
        return run

    # ── Logo ─────────────────────────────────────────────────────────
    # Logo — try static/image/logo.png (as served by /static/image/logo.png)
    logo_loaded = False
    for logo_candidate in [
        'static/image/logo.png',
        'static/img/logo.png',
        os.path.join(os.path.dirname(__file__), '..', 'static', 'image', 'logo.png'),
        os.path.join(os.path.dirname(__file__), '..', 'static', 'img',   'logo.png'),
    ]:
        if os.path.exists(logo_candidate):
            try:
                p_logo = d.add_paragraph()
                p_logo.paragraph_format.space_before = Pt(0)
                p_logo.paragraph_format.space_after  = Pt(4)
                p_logo.add_run().add_picture(logo_candidate, width=Inches(1.8))
                logo_loaded = True
                break
            except Exception:
                continue
    if not logo_loaded:
        p_logo = d.add_paragraph()
        r_logo = p_logo.add_run('XPRESS HEALTH')
        r_logo.bold = True; r_logo.font.size = Pt(14)
        r_logo.font.color.rgb = NAVY; r_logo.font.name = 'Calibri'

    sp(8)

    # ── Main Title ────────────────────────────────────────────────────
    p_title = d.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after  = Pt(4)
    r_title = _add_run(p_title, 'Verification of Service \u2013 Agency Staff',
                       bold=False, size=26, color=NAVY)
    # Underline via border
    from docx.oxml.ns import qn as _qn
    pPr  = p_title._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(_qn('w:val'),   'single')
    bot.set(_qn('w:sz'),    '6')
    bot.set(_qn('w:color'), '1B3A6B')
    pBdr.append(bot)
    pPr.append(pBdr)

    sp(14)

    # ── Staff Details section ─────────────────────────────────────────
    p_sd = d.add_paragraph()
    _add_run(p_sd, 'Staff Details', bold=True, size=12, color=GREEN)

    def field_line(label, value):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _add_run(p, f'{label}: ', bold=False, size=11, color=BLACK)
        _add_run(p, value, bold=False, size=11, color=BLACK)

    field_line('Full Name',       full_name)
    field_line('Role / Designation', user_type)
    # Contact Email removed from document

    sp(12)

    # ── Service Details section ───────────────────────────────────────
    p_srv = d.add_paragraph()
    _add_run(p_srv, 'Service Details', bold=True, size=12, color=GREEN)

    sp(4)

    # ── Table ─────────────────────────────────────────────────────────
    headers = ['Post', 'HSE Location\n/ Employer', 'From Date', 'To Date',
               'Years', 'Months', 'Days', 'Verified']
    col_widths_cm = [3.0, 4.0, 3.0, 3.0, 1.4, 1.6, 1.4, 1.8]

    tbl = d.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'

    # Header row
    hdr_row = tbl.rows[0]
    for ci, (hdr, w) in enumerate(zip(headers, col_widths_cm)):
        cell = hdr_row.cells[ci]
        cell.width = Cm(w)
        _set_cell_bg(cell, '1B3A6B')
        _set_cell_border(cell, '1B3A6B')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = _add_run(p, hdr, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))

    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        row_data = [
            row.get('post', user_type),
            row.get('employer', ''),
            row.get('from_str', ''),
            row.get('to_str', '').capitalize() if row.get('to_str','').lower() in ('present','current','ongoing') else row.get('to_str',''),
            str(row.get('years', 0)),
            str(row.get('months', 0)),
            str(row.get('days', 0)),
            'Yes',
        ]
        bg = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, (val, w) in enumerate(zip(row_data, col_widths_cm)):
            cell = tr.cells[ci]
            cell.width = Cm(w)
            _set_cell_bg(cell, bg)
            _set_cell_border(cell, 'CCCCCC')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            _add_run(p, val, size=10)

    sp(12)

    # ── Declaration ───────────────────────────────────────────────────
    p_decl = d.add_paragraph()
    p_decl.paragraph_format.space_before = Pt(4)
    p_decl.paragraph_format.space_after  = Pt(16)
    _add_run(p_decl,
             'I hereby confirm that the above details are accurate and verified as per agency records.',
             size=11)

    # ── Signatory ─────────────────────────────────────────────────────
    field_line('Authorised Signatory', 'Betsy Daniel')
    field_line('Designation',          'Recruitment Head')

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ── Extract employment rows from staff doc ────────────────────────────

def _is_ireland_employer(employer_str, location_str=''):
    """
    Return True if the employer/location is in Ireland or location unknown.
    Excludes only when an explicit non-Ireland country/city is found.
    """
    text = (employer_str + ' ' + location_str).lower()

    # Explicit non-Ireland locations — exclude these
    non_ireland = [
        'india', 'indian', 'bangalore', 'mumbai', 'delhi', 'chennai',
        'hyderabad', 'kolkata', 'pune', 'kerala', 'karnataka',
        'uk', 'united kingdom', 'england', 'scotland', 'wales',
        'london', 'manchester', 'birmingham', 'leeds', 'bristol',
        'southampton', 'dorset', 'bournemouth', 'poole', 'christchurch',
        'guernsey', 'jersey', 'isle of man',
        'usa', 'united states', 'america', 'australia', 'canada',
        'nigeria', 'kenya', 'ghana', 'zimbabwe', 'south africa',
        'philippines', 'pakistan', 'bangladesh', 'sri lanka',
        'malaysia', 'singapore', 'uae', 'dubai', 'saudi',
        'netherlands', ' nl', 'sint oedenrode', 'poland', 'polska',
        'germany', 'france', 'spain', 'italy', 'portugal',
    ]
    if any(kw in text for kw in non_ireland):
        return False

    # If location explicitly says Ireland or Irish place — include
    # If no location info at all — default INCLUDE (assume Ireland for HCA staff)
    return True


def _parse_employment_regex(extracted_cv, user_type=''):
    """Regex-based fallback parser — handles both multi-line and label formats."""
    if not extracted_cv:
        return []

    lines = extracted_cv.split('\n')
    exp_lines = []; in_exp = False
    exp_stop    = {'education','qualifications','training','certifications','key skills','skills','references','additional information','profile','summary'}
    exp_headers = {'professional experience','work experience','employment history','employment','experience','career history','positions held'}
    for line in lines:
        l = line.strip(); ll = l.lower()
        if any(ll.startswith(h) for h in exp_headers): in_exp = True; continue
        if in_exp and l and any(ll.startswith(h) for h in exp_stop): break
        if in_exp: exp_lines.append(l)
    if not exp_lines:
        exp_lines = [l.strip() for l in lines if l.strip()]

    entries = []

    # Format B: Job Title: / Employer: / Location: / Dates:
    label_re = re.compile(r'^(?:job\s+title|title)\s*[:\-]\s*(.+)', re.I)
    emp_re   = re.compile(r'^employer\s*[:\-]\s*(.+)', re.I)
    loc_re   = re.compile(r'^location\s*[:\-]\s*(.+)', re.I)
    date_re2 = re.compile(r'^dates?\s*[:\-]\s*(.+)', re.I)
    date_split = re.compile(r'(.+?)\s*(?:to|-|–|—)\s*(present|current|ongoing|\w+\s+\d{4}|\d{4})', re.I)

    cur = {}
    for l in exp_lines:
        ml=label_re.match(l); me=emp_re.match(l); mlo=loc_re.match(l); md=date_re2.match(l)
        if ml:
            if cur.get('employer') and cur.get('dates'): entries.append(cur); cur={}
            cur['post'] = ml.group(1).strip()
        elif me:  cur['employer']  = me.group(1).strip()
        elif mlo: cur['location']  = mlo.group(1).strip()
        elif md:  cur['dates']     = md.group(1).strip()
    if cur.get('employer') and cur.get('dates'): entries.append(cur)

    if entries:
        result = []
        for e in entries:
            if not _is_ireland_employer(e.get('employer',''), e.get('location','')):
                continue
            ds = date_split.search(e.get('dates',''))
            result.append({
                'post':     e.get('post') or user_type,
                'employer': e.get('employer',''),
                'from_str': ds.group(1).strip() if ds else e.get('dates',''),
                'to_str':   ds.group(2).strip()  if ds else 'Present',
            })
        if result: return result

    # Format A: multi-line Title / Employer / Dates
    date_re = re.compile(
        r'^(\d{1,2}/\d{1,2}/\d{4}|\w+\s+\d{4}|\d{4})\s*(?:to|-|–|—)\s*'
        r'(\d{1,2}/\d{1,2}/\d{4}|\w+\s+\d{4}|\d{4}|present|current|ongoing)', re.I)
    result = []; i = 0
    while i < len(exp_lines):
        dm = date_re.match(exp_lines[i].strip())
        if dm:
            post_text=''; employer_text=''
            for back in range(1,4):
                prev = exp_lines[i-back].strip() if i-back>=0 else ''
                if not prev or date_re.match(prev): break
                if not employer_text: employer_text=prev
                elif not post_text:   post_text=prev
                else: break
            if _is_ireland_employer(employer_text):
                result.append({'post': post_text or user_type, 'employer': employer_text,
                                'from_str': dm.group(1).strip(), 'to_str': dm.group(2).strip()})
        i += 1
    return result


def _fetch_hse_cv_text(email, gemini_key=''):
    """
    Fetch 'Hse Cv' document from XN Portal for the given email,
    extract text using Gemini vision, and return raw text.
    """
    import requests as _req

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')

    if not base_url or not email:
        return ''

    hdrs = {"Api-Key": api_key, "X-App-Country": app_country,
            "Content-Type": "application/json", "Accept": "application/json"}

    # Get document list
    try:
        r = _req.post(f"{base_url}/ai/recruitments/user-document-list",
                      json={"email": email}, headers=hdrs, timeout=30)
        if r.status_code == 405:
            r = _req.get(f"{base_url}/ai/recruitments/user-document-list",
                         params={"email": email}, headers=hdrs, timeout=30)
        r.raise_for_status()
        api_data = r.json().get('data')
        docs = api_data if isinstance(api_data, list) else \
               (api_data.get('documents') or [] if isinstance(api_data, dict) else [])
    except Exception:
        return ''

    # Find Hse Cv document
    hse_cv_url = None
    for d in docs:
        name = (d.get('document_type_name') or '').strip().lower()
        if name in ('hse cv', 'hse_cv', 'hsecv') and d.get('url'):
            hse_cv_url = d['url']
            break

    if not hse_cv_url:
        return ''

    # Download document
    try:
        dl_hdrs = {k: v for k, v in hdrs.items() if k != 'Content-Type'}
        dl = _req.get(hse_cv_url, headers=dl_hdrs, timeout=60)
        dl.raise_for_status()
        raw_bytes    = dl.content
        content_type = dl.headers.get('Content-Type', '').lower()
        url_lower    = hse_cv_url.lower().split('?')[0]
    except Exception:
        return ''

    # Extract text
    raw_text = ''

    # PDF text extraction
    if 'pdf' in content_type or url_lower.endswith('.pdf'):
        try:
            import pdfplumber, io as _io
            with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                raw_text = '\n'.join(p.extract_text() or '' for p in pdf.pages).strip()
        except Exception:
            pass

    # DOCX text extraction
    if not raw_text and ('wordprocessingml' in content_type or
                         url_lower.endswith('.docx') or url_lower.endswith('.doc')):
        try:
            from docx import Document as _DDoc
            import io as _io
            ddoc = _DDoc(_io.BytesIO(raw_bytes))
            raw_text = '\n'.join(p.text for p in ddoc.paragraphs).strip()
        except Exception:
            pass

    # Gemini vision fallback for scanned PDFs
    if not raw_text and gemini_key:
        try:
            import base64 as _b64
            from google import genai as _gai
            client = _gai.Client(api_key=gemini_key)
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": [
                    {"inline_data": {"mime_type": "application/pdf",
                                     "data": _b64.b64encode(raw_bytes).decode()}},
                    {"text": "Extract all text from this CV document. Return only the plain text."}
                ]}]
            )
            raw_text = (response.text or '').strip()
        except Exception:
            pass

    return raw_text


def _get_employment_rows(staff_doc, gemini_key=None):
    """
    Get Ireland-only employment rows from HSE CV (document_type_name: 'Hse Cv')
    fetched from XN Portal. No other data source used.
    """
    if gemini_key is None:
        gemini_key = os.environ.get('GEMINI_API_KEY', '')

    user_type = _v(staff_doc.get('user_type') or 'Healthcare Assistant')
    s1        = staff_doc.get('section_1_personal_details') or {}
    email     = _v(staff_doc.get('email') or s1.get('email_address') or '')

    if not email:
        return []

    # Fetch HSE CV text from XN Portal
    cv_text = _fetch_hse_cv_text(email, gemini_key)
    if not cv_text:
        return []

    # Extract Ireland-only employment — Gemini first, regex fallback
    parsed = []
    if gemini_key:
        parsed = _parse_employment_from_cv(cv_text, user_type, gemini_key)
    if not parsed:
        parsed = _parse_employment_regex(cv_text, user_type)

    rows = []
    for e in parsed:
        from_d  = _parse_date_flex(_v(e.get('from_str') or ''))
        to_d    = _parse_date_flex(_v(e.get('to_str') or ''))
        y, m, dys = _calc_duration(from_d, to_d)
        rows.append({
            'post':     _v(e.get('post') or user_type),
            'employer': _v(e.get('employer') or ''),
            'from_str': from_d.strftime('%d/%m/%Y') if from_d else _v(e.get('from_str') or ''),
            'to_str':   to_d.strftime('%d/%m/%Y') if to_d else 'Present',
            'years':    y, 'months': m, 'days': dys,
        })

    return rows


# ── Cron: Generate Point Scale document ──────────────────────────────

@admin_bp.route('/live-staffs/cron/generate-point-scale', methods=['GET', 'POST'])
def live_staff_cron_generate_point_scale():
    """
    Cron — generates Verification of Service document for ONE HCA staff per call.
    Only processes staff where user_type contains 'Healthcare Assistant'.

    Auth: cron_key query param or X-Cron-Key header.
    """
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    ps_col     = _ps_col()
    staffs_col = _staffs_col()

    # Find staff already processed
    existing_ids = set(
        str(r['staff_id'])
        for r in ps_col.find({}, {"staff_id": 1})
        if r.get('staff_id')
    )

    # Find next HCA staff without a point scale document
    pending_query = {
        "$and": [
            {"user_type": {"$regex": "Healthcare Assistant", "$options": "i"}},
            {"_id": {"$nin": [ObjectId(i) for i in existing_ids if len(i) == 24]}},
        ]
    }
    remaining_total = staffs_col.count_documents(pending_query)
    staff = staffs_col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All Healthcare Assistant staff point scale documents generated.",
            "remaining_count": 0,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    # Mark as processing immediately to avoid duplicate runs
    ps_col.insert_one({
        "staff_id":    staff_id,
        "staff_name":  full_name,
        "status":      "processing",
        "created_at":  datetime.utcnow(),
    })

    gemini_key = os.environ.get('GEMINI_API_KEY', '')

    def _do_generate():
        try:
            rows = _get_employment_rows(staff, gemini_key=gemini_key)

            if not rows:
                ps_col.update_one(
                    {"staff_id": staff_id},
                    {"$set": {
                        "status":  "no_employment_data",
                        "note":    "No Hse Cv document found in XN Portal or no Ireland employment extracted",
                        "updated_at": datetime.utcnow(),
                    }}
                )
                return

            docx_bytes = _build_point_scale_docx(staff, rows)
            safe_name  = (full_name or 'staff').replace(' ', '_').replace('/', '_')
            filename   = f"PointScale_{safe_name}.docx"
            gcs_blob   = f"point_scale/{filename}"

            _gcs_upload(
                gcs_blob, docx_bytes,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            tot_y, tot_m, tot_d = _total_service(rows)

            ps_col.update_one(
                {"staff_id": staff_id},
                {"$set": {
                    "staff_name":    full_name,
                    "email":         email,
                    "user_type":     _v(staff.get('user_type')),
                    "filename":      filename,
                    "gcs_blob":      gcs_blob,
                    "rows":          rows,
                    "total_years":   tot_y,
                    "total_months":  tot_m,
                    "total_days":    tot_d,
                    "status":        "generated",
                    "generated_at":  datetime.utcnow(),
                }}
            )

        except Exception as e:
            ps_col.update_one(
                {"staff_id": staff_id},
                {"$set": {
                    "status": "error",
                    "error":  str(e),
                    "updated_at": datetime.utcnow(),
                }}
            )

    threading.Thread(target=_do_generate, daemon=True).start()

    return jsonify({
        "success":         True,
        "staff_id":        staff_id,
        "staff_name":      full_name,
        "email":           email,
        "remaining_count": max(0, remaining_total - 1),
        "message":         f"Point scale generation started for {full_name} — {max(0, remaining_total - 1)} remaining.",
    })


# ── Download generated document ───────────────────────────────────────

@admin_bp.route('/live-staffs/api/vos-saved/<staff_id>')
@admin_required
def live_staff_api_vos_saved(staff_id):
    """Check if a Verification of Service document exists for this staff member."""
    rec = _ps_col().find_one({"staff_id": staff_id, "status": "generated"})
    if rec:
        tot_y = rec.get('total_years', 0)
        tot_m = rec.get('total_months', 0)
        tot_d = rec.get('total_days', 0)
        total_str = f"{tot_y} Year{'s' if tot_y!=1 else ''}, {tot_m} Month{'s' if tot_m!=1 else ''}, {tot_d} Day{'s' if tot_d!=1 else ''}"
        return jsonify({
            "success":      True,
            "found":        True,
            "gcs_blob":     rec.get("gcs_blob", ""),
            "generated_at": rec["generated_at"].strftime("%d %b %Y %H:%M")
                            if rec.get("generated_at") else "",
            "total":        total_str,
        })
    return jsonify({"success": True, "found": False})


@admin_bp.route('/live-staffs/vos/download/<staff_id>')
@admin_required
def live_staff_vos_download(staff_id):
    """Download the generated Verification of Service document.
    Accessible via admin session (browser) or ?cron_key= (local script).
    """
    # Also allow cron_key for local script downloads
    _ = request.args.get('cron_key')  # accepted but not required — @admin_required handles browser
    try:
        rec = _ps_col().find_one({"staff_id": staff_id})
        if not rec:
            # Try generating on-demand
            staff_doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
            if not staff_doc:
                return jsonify({"success": False, "error": "Staff not found"}), 404
            gemini_key = os.environ.get('GEMINI_API_KEY', '')
            rows = _get_employment_rows(staff_doc, gemini_key=gemini_key)
            if not rows:
                return jsonify({"success": False, "error": "No employment data"}), 404
            docx_bytes = _build_point_scale_docx(staff_doc, rows)
            s1         = staff_doc.get('section_1_personal_details') or {}
            full_name  = _v(s1.get('full_name') or 'staff')
            filename   = f"PointScale_{full_name.replace(' ','_')}.docx"
        else:
            if rec.get('status') != 'generated' or not rec.get('gcs_blob'):
                return jsonify({"success": False,
                                "error": f"Document not ready (status: {rec.get('status')})"}), 404
            docx_bytes = _gcs_download(rec['gcs_blob'])
            filename   = rec.get('filename', 'PointScale.docx')

        return Response(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Regenerate on-demand for a single staff ───────────────────────────

@admin_bp.route('/live-staffs/vos/generate/<staff_id>', methods=['POST'])
@admin_required
def live_staff_vos_generate(staff_id):
    """Regenerate point scale document for a specific staff member."""
    try:
        staff_doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not staff_doc:
            return jsonify({"success": False, "error": "Staff not found"}), 404

        gemini_key = os.environ.get('GEMINI_API_KEY', '')
        rows = _get_employment_rows(staff_doc, gemini_key=gemini_key)
        if not rows:
            return jsonify({"success": False, "error": "No employment data found"}), 400

        s1        = staff_doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff')
        email     = _v(staff_doc.get('email') or s1.get('email_address') or '')

        docx_bytes = _build_point_scale_docx(staff_doc, rows)
        safe_name  = full_name.replace(' ', '_').replace('/', '_')
        filename   = f"PointScale_{safe_name}.docx"
        gcs_blob   = f"point_scale/{filename}"

        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        tot_y, tot_m, tot_d = _total_service(rows)

        _ps_col().update_one(
            {"staff_id": staff_id},
            {"$set": {
                "staff_name":   full_name,
                "email":        email,
                "user_type":    _v(staff_doc.get('user_type')),
                "filename":     filename,
                "gcs_blob":     gcs_blob,
                "rows":         rows,
                "total_years":  tot_y,
                "total_months": tot_m,
                "total_days":   tot_d,
                "status":       "generated",
                "generated_at": datetime.utcnow(),
            }},
            upsert=True
        )

        return jsonify({
            "success":     True,
            "staff_name":  full_name,
            "filename":    filename,
            "total":       f"{tot_y} Years, {tot_m} Months, {tot_d} Days",
            "row_count":   len(rows),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Excel export ──────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/vos/upload/<staff_id>', methods=['POST'])
@admin_required
def live_staff_vos_upload(staff_id):
    """Replace the saved VOS document with an uploaded .docx file."""
    f = request.files.get('file')
    if not f:
        return jsonify({"success": False, "error": "No file uploaded"}), 400
    if not f.filename.lower().endswith('.docx'):
        return jsonify({"success": False, "error": "Only .docx files accepted"}), 400
    try:
        staff_doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not staff_doc:
            return jsonify({"success": False, "error": "Staff not found"}), 404

        s1        = staff_doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff')
        safe_name = full_name.replace(' ', '_').replace('/', '_')
        filename  = f"PointScale_{safe_name}.docx"
        gcs_blob  = f"point_scale/{filename}"

        docx_bytes = f.read()
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        _ps_col().update_one(
            {"staff_id": staff_id},
            {"$set": {
                "gcs_blob":      gcs_blob,
                "filename":      filename,
                "status":        "generated",
                "generated_at":  datetime.utcnow(),
                "uploaded_by":   "admin",
            }},
            upsert=True
        )
        return jsonify({"success": True, "gcs_blob": gcs_blob, "filename": filename})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500




# ── API: Get all generated VOS point scale documents ─────────────────


@admin_bp.route('/live-staffs/api/vos-docs', methods=['GET'])
def live_staff_api_vos_docs():
    """One staff per call. Auth: cron_key."""
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = request.args.get('cron_key') or request.headers.get('X-Cron-Key', '')
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    col = _ps_col()

    # Ensure index exists for fast lookup (no-op if already exists)
    try:
        col.create_index([("status", 1), ("downloaded", 1)], background=True)
    except Exception:
        pass

    doc = col.find_one_and_update(
        {"status": "generated", "downloaded": {"$ne": True}},
        {"$set": {"downloaded": True}},
        projection={"staff_id": 1, "staff_name": 1, "email": 1},
    )

    if not doc:
        return jsonify({"success": True, "done": True, "message": "All documents downloaded"})

    return jsonify({
        "success":      True,
        "done":         False,
        "staff_name":   doc.get('staff_name', ''),
        "email":        doc.get('email', ''),
        "download_url": f"{request.host_url.rstrip('/')}/admin/live-staffs/vos/cron-download/{doc.get('staff_id', '')}?cron_key={os.environ.get('CRON_SECRET', '')}",
    })



@admin_bp.route('/live-staffs/vos/cron-download/<staff_id>')
def live_staff_vos_cron_download(staff_id):
    """No-auth download for local script — cron_key required."""
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = request.args.get('cron_key') or request.headers.get('X-Cron-Key', '')
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401
    try:
        rec = _ps_col().find_one({"staff_id": staff_id, "status": "generated"})
        if not rec or not rec.get('gcs_blob'):
            return jsonify({"success": False, "error": "Document not found"}), 404
        from admin.live_staffs import _gcs_download as _dl
        docx_bytes = _dl(rec['gcs_blob'])
        filename   = rec.get('filename', 'document.docx')
        return Response(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
@admin_bp.route('/live-staffs/export/vos-missing-xlsx')
@admin_required
def live_staff_export_vos_missing_xlsx():
    """Export Healthcare Assistant staff who do NOT have a generated Point Scale document."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        ps_col     = _ps_col()
        staffs_col = _staffs_col()

        hca_staff = list(staffs_col.find(
            {"user_type": {"$regex": "Healthcare Assistant", "$options": "i"}},
            {"section_1_personal_details": 1, "email": 1, "user_type": 1}
        ))

        ps_records = {r['staff_id']: r for r in ps_col.find({}) if r.get('staff_id')}

        missing = []
        for s in hca_staff:
            sid = str(s['_id'])
            rec = ps_records.get(sid)
            if rec and rec.get('status') == 'generated':
                continue  # already has document — skip
            s1 = s.get('section_1_personal_details') or {}
            reason = 'Not started'
            detail = ''
            if rec:
                status_map = {
                    'processing':          'Processing',
                    'no_employment_data':  'No employment data found',
                    'error':               'Error during generation',
                }
                reason = status_map.get(rec.get('status'), rec.get('status', 'Unknown'))
                detail = rec.get('note') or rec.get('error') or ''
            missing.append({
                "staff_name": _v(s1.get('full_name') or ''),
                "email":      _v(s.get('email') or s1.get('email_address') or ''),
                "user_type":  _v(s.get('user_type') or ''),
                "reason":     reason,
                "detail":     detail,
            })

        missing.sort(key=lambda d: d['staff_name'].lower())

        NAVY  = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT   = 'EFF6FF'; WARN  = 'FFF3CD'

        h_font  = Font(name='Calibri', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Calibri', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        g_bot   = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))
        warn_fill = PatternFill('solid', start_color=WARN, end_color=WARN)

        wb = Workbook()
        ws = wb.active
        ws.title = 'Missing Point Scale Docs'

        headers    = ['Sno', 'Staff Name', 'Email', 'User Type', 'Reason', 'Detail']
        col_widths = [5, 30, 36, 24, 22, 50]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = g_bot
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:F{len(missing)+1}'

        for ri, m in enumerate(missing, start=2):
            row_vals = [ri - 1, m['staff_name'], m['email'], m['user_type'], m['reason'], m['detail']]
            aligns   = [c_align, l_align, l_align, l_align, l_align, l_align]
            alt_fill = warn_fill if ri % 2 == 0 else PatternFill()
            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border; cell.fill = alt_fill
            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(missing)+2, column=1,
                value=f'Total missing: {len(missing)}').font = Font(name='Calibri', bold=True, size=9)

        buf = io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="point_scale_missing_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/live-staffs/export/vos-xlsx')
@admin_required
def live_staff_export_vos_xlsx():
    """Export all generated point scale documents to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        docs = list(_ps_col().find({}))
        docs.sort(key=lambda d: _v(d.get('staff_name') or '').lower())

        NAVY  = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT   = 'EFF6FF'; RED   = 'FFDDDD'; WARN  = 'FFF3CD'

        h_font  = Font(name='Calibri', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Calibri', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        g_bot   = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Point Scale Documents'

        headers    = ['Sno', 'Staff Name', 'Email', 'User Type',
                      'Total Years', 'Total Months', 'Total Days',
                      'Total Service', 'Status', 'Generated At']
        col_widths = [5, 30, 36, 24, 12, 13, 10, 28, 14, 20]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = g_bot
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:J{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            tot_y = doc.get('total_years',  '')
            tot_m = doc.get('total_months', '')
            tot_d = doc.get('total_days',   '')
            status = _v(doc.get('status') or '')
            gen_at = doc.get('generated_at')
            gen_str = gen_at.strftime('%d %b %Y %H:%M') if gen_at else ''

            total_str = (f"{tot_y} Yr{'s' if tot_y!=1 else ''}, "
                         f"{tot_m} Mo, {tot_d} D") if tot_y != '' else ''

            if status == 'generated':
                row_fill = None
            elif status == 'error':
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif status == 'no_employment_data':
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            else:
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT) \
                       if ri % 2 == 0 and not row_fill else None

            row_vals = [
                ri-1,
                _v(doc.get('staff_name') or ''),
                _v(doc.get('email') or ''),
                _v(doc.get('user_type') or ''),
                tot_y, tot_m, tot_d, total_str, status, gen_str,
            ]
            aligns = [c_align, l_align, l_align, l_align,
                      c_align, c_align, c_align, l_align, c_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Calibri', bold=True, size=9)

        buf = io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="point_scale_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
