"""
live_staffs_vos_nurse.py
═════════════════════════
Independent blueprint — Verification of Service – Agency Staff
for NURSE staff (user_type contains 'Nurse' / 'RGN' / 'RNM' etc.)

Design matches Point_Scale_Document_1.pdf exactly:
  - Logo top-right
  - Title centred bold navy
  - Agency Name / HSE Site fields
  - Staff Details in 2-col table (bold label | value)
  - Service Details table (8 cols)
  - Total Eligible Service line
  - Declaration text
  - Signatory block (page 2 style)

Data source: document_type_name "Hse Cv" from XN Portal
             → Gemini extracts Ireland-only nursing roles
             → regex fallback

Routes:
    GET  /live-staffs/cron/generate-vos-nurse      — cron (one staff per call)
    POST /live-staffs/vos-nurse/generate/<id>      — on-demand
    GET  /live-staffs/vos-nurse/download/<id>      — download DOCX
    POST /live-staffs/vos-nurse/upload/<id>        — replace with edited DOCX
    GET  /live-staffs/api/vos-nurse-saved/<id>     — check if saved
    GET  /live-staffs/export/vos-nurse-xlsx        — Excel export
"""

from flask import request, jsonify, Response
from bson import ObjectId
from datetime import datetime, date
import os, re, io, json, threading

from database import db
from . import admin_bp
from admin.views import admin_required


# ── Helpers ───────────────────────────────────────────────────────────

def _v(val):
    if val is None: return ''
    return str(val).strip()

def _staffs_col():
    from flask import current_app
    return current_app.db.live_staffs

def _vos_nurse_col():
    return db.live_staff_vos_nurse

def _gcs_upload(blob_name, data_bytes, content_type='application/octet-stream'):
    from admin.live_staffs import _gcs_upload as _f
    return _f(blob_name, data_bytes, content_type)

def _gcs_download(blob_name):
    from admin.live_staffs import _gcs_download as _f
    return _f(blob_name)

def _is_nurse(user_type):
    ut = (user_type or '').lower()
    return any(k in ut for k in ('nurse','rgn','rnm','midwife','nchd','rn '))

# ── Date / Duration helpers ───────────────────────────────────────────

def _parse_date_flex(s):
    s = _v(s)
    if not s or s.lower() in ('present','current','ongoing','now'): return None
    # Handle "Present (June 2026)" → treat as present
    if 'present' in s.lower(): return None
    fmts = ['%d/%m/%Y','%m/%Y','%B %Y','%b %Y','%d %B %Y','%d %b %Y','%Y-%m-%d','%Y']
    for fmt in fmts:
        try: return datetime.strptime(s.strip(), fmt).date()
        except ValueError: pass
    return None

def _calc_duration(from_date, to_date):
    if from_date is None: return 0, 0, 0
    end = to_date if to_date else date.today()
    if end < from_date: return 0, 0, 0
    years  = end.year  - from_date.year
    months = end.month - from_date.month
    days   = end.day   - from_date.day
    if days < 0:
        months -= 1
        from calendar import monthrange
        prev_month = end.month - 1 if end.month > 1 else 12
        prev_year  = end.year if end.month > 1 else end.year - 1
        days += monthrange(prev_year, prev_month)[1]
    if months < 0: years -= 1; months += 12
    return max(0,years), max(0,months), max(0,days)

def _total_service(rows):
    total_days = sum(r['years']*365 + r['months']*30 + r['days'] for r in rows)
    years  = total_days // 365
    remain = total_days  % 365
    months = remain // 30
    days   = remain  % 30
    return years, months, days

# ── Non-Ireland exclusion ─────────────────────────────────────────────

def _is_non_ireland(employer_str, location_str=''):
    text = (employer_str + ' ' + location_str).lower()
    non_ie = [
        'australia','australian','new zealand','uk','united kingdom',
        'england','scotland','wales','london','manchester','birmingham',
        'india','bangalore','mumbai','delhi','chennai','hyderabad',
        'usa','united states','canada','nigeria','kenya','ghana',
        'philippines','pakistan','saudi','uae','dubai','singapore',
        'germany','france','spain','italy','netherlands','poland',
    ]
    return any(k in text for k in non_ie)

# ── Fetch HSE CV from XN Portal ───────────────────────────────────────

def _fetch_hse_cv_text(email, gemini_key=''):
    import requests as _req
    base_url    = os.environ.get('LIVE_STAFF_URL','').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY','')
    app_country = os.environ.get('XN_APP_COUNTRY','')
    if not base_url or not email: return ''
    hdrs = {"Api-Key": api_key, "X-App-Country": app_country,
            "Content-Type": "application/json", "Accept": "application/json"}
    try:
        r = _req.post(f"{base_url}/ai/recruitments/user-document-list",
                      json={"email": email}, headers=hdrs, timeout=30)
        if r.status_code == 405:
            r = _req.get(f"{base_url}/ai/recruitments/user-document-list",
                         params={"email": email}, headers=hdrs, timeout=30)
        r.raise_for_status()
        api_data = r.json().get('data')
        docs = api_data if isinstance(api_data, list) else \
               (api_data.get('documents') or [] if isinstance(api_data, dict) else [])
    except Exception: return ''

    hse_cv_url = next(
        (d['url'] for d in docs
         if (d.get('document_type_name') or '').strip().lower() in ('hse cv','hse_cv','hsecv')
         and d.get('url')), None
    )
    if not hse_cv_url: return ''

    try:
        dl_hdrs = {k: v for k, v in hdrs.items() if k != 'Content-Type'}
        dl = _req.get(hse_cv_url, headers=dl_hdrs, timeout=60)
        dl.raise_for_status()
        raw_bytes = dl.content
        ct = dl.headers.get('Content-Type','').lower()
        ul = hse_cv_url.lower().split('?')[0]
    except Exception: return ''

    raw_text = ''
    if 'pdf' in ct or ul.endswith('.pdf'):
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                raw_text = '\n'.join(p.extract_text() or '' for p in pdf.pages).strip()
        except Exception: pass
    if not raw_text and ('wordprocessingml' in ct or ul.endswith('.docx')):
        try:
            from docx import Document as _DDoc
            ddoc = _DDoc(io.BytesIO(raw_bytes))
            raw_text = '\n'.join(p.text for p in ddoc.paragraphs).strip()
        except Exception: pass
    if not raw_text and gemini_key:
        try:
            import base64 as _b64
            from google import genai as _gai
            client = _gai.Client(api_key=gemini_key)
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[{"parts": [
                    {"inline_data": {"mime_type": "application/pdf",
                                     "data": _b64.b64encode(raw_bytes).decode()}},
                    {"text": "Extract all text from this CV document. Return only the plain text."}
                ]}]
            )
            raw_text = (response.text or '').strip()
        except Exception: pass
    return raw_text

# ── Employment extraction ─────────────────────────────────────────────

def _extract_with_gemini(cv_text, user_type, gemini_key):
    if not cv_text or not gemini_key: return []
    try:
        from google import genai as _gai
        client = _gai.Client(api_key=gemini_key)
        prompt = f"""You are an employment history extractor for an Irish healthcare agency.

Extract ONLY the roles where the candidate worked in IRELAND from the CV text below.
Exclude ALL roles outside Ireland (UK, England, Australia, India, Italy, USA, Canada, etc.).
Include roles where location mentions Ireland, Irish counties/cities, or HSE sites.

For each Ireland role return:
- post: job title (e.g. "Agency Registered Nurse")
- employer: HSE location/employer name only
- from_str: start date as written
- to_str: end date as written or "Present"

Return ONLY a JSON array — no markdown, no explanation:
[{{"post": "...", "employer": "...", "from_str": "...", "to_str": "..."}}]
If no Ireland roles found return: []

CV TEXT:
{cv_text}
"""
        response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
        raw = re.sub(r'^```(?:json)?\s*', '', (response.text or '').strip(), flags=re.MULTILINE)
        raw = re.sub(r'```\s*$', '', raw, flags=re.MULTILINE).strip()
        result = json.loads(raw)
        return result if isinstance(result, list) else []
    except Exception: return []

def _extract_with_regex(cv_text, user_type):
    if not cv_text: return []
    lines = cv_text.split('\n')
    exp_lines = []; in_exp = False
    exp_stop    = {'education','qualifications','training','certifications','key skills','skills','references','additional information','profile','summary'}
    exp_headers = {'professional experience','work experience','employment history','employment','experience','career history','positions held'}
    for line in lines:
        l = line.strip(); ll = l.lower()
        if any(ll.startswith(h) for h in exp_headers): in_exp = True; continue
        if in_exp and l and any(ll.startswith(h) for h in exp_stop): break
        if in_exp: exp_lines.append(l)
    if not exp_lines: exp_lines = [l.strip() for l in lines if l.strip()]

    entries = []
    label_re = re.compile(r'^(?:job\s+title|title)\s*[:\-]\s*(.+)', re.I)
    emp_re   = re.compile(r'^employer\s*[:\-]\s*(.+)', re.I)
    loc_re   = re.compile(r'^location\s*[:\-]\s*(.+)', re.I)
    date_re2 = re.compile(r'^dates?\s*[:\-]\s*(.+)', re.I)
    date_split = re.compile(r'(.+?)\s*(?:to|-|–|—)\s*(present|current|ongoing|\w+\s+\d{4}|\d{4})', re.I)
    cur = {}
    for l in exp_lines:
        ml=label_re.match(l); me=emp_re.match(l); mlo=loc_re.match(l); md=date_re2.match(l)
        if ml:
            if cur.get('employer') and cur.get('dates'): entries.append(cur); cur={}
            cur['post'] = ml.group(1).strip()
        elif me:  cur['employer'] = me.group(1).strip()
        elif mlo: cur['location'] = mlo.group(1).strip()
        elif md:  cur['dates']    = md.group(1).strip()
    if cur.get('employer') and cur.get('dates'): entries.append(cur)

    if entries:
        result = []
        for e in entries:
            if _is_non_ireland(e.get('employer',''), e.get('location','')): continue
            ds = date_split.search(e.get('dates',''))
            result.append({'post': e.get('post') or user_type, 'employer': e.get('employer',''),
                           'from_str': ds.group(1).strip() if ds else e.get('dates',''),
                           'to_str':   ds.group(2).strip() if ds else 'Present'})
        if result: return result

    date_re = re.compile(
        r'^(\d{1,2}/\d{1,2}/\d{4}|\w+\s+\d{4}|\d{4})\s*(?:to|-|–|—)\s*'
        r'(\d{1,2}/\d{1,2}/\d{4}|\w+\s+\d{4}|\d{4}|present|current|ongoing)', re.I)
    result = []; i = 0
    while i < len(exp_lines):
        dm = date_re.match(exp_lines[i].strip())
        if dm:
            post_t=''; emp_t=''
            for back in range(1,4):
                prev = exp_lines[i-back].strip() if i-back>=0 else ''
                if not prev or date_re.match(prev): break
                if not emp_t: emp_t=prev
                elif not post_t: post_t=prev
                else: break
            if not _is_non_ireland(emp_t):
                result.append({'post': post_t or user_type, 'employer': emp_t,
                                'from_str': dm.group(1).strip(), 'to_str': dm.group(2).strip()})
        i += 1
    return result

def _get_employment_rows(staff_doc, gemini_key=None):
    if gemini_key is None: gemini_key = os.environ.get('GEMINI_API_KEY','')
    user_type = _v(staff_doc.get('user_type') or 'Registered Nurse')
    s1  = staff_doc.get('section_1_personal_details') or {}
    email = _v(staff_doc.get('email') or s1.get('email_address') or '')
    if not email: return []

    cv_text = _fetch_hse_cv_text(email, gemini_key)
    if not cv_text: return []

    parsed = _extract_with_gemini(cv_text, user_type, gemini_key)
    if not parsed: parsed = _extract_with_regex(cv_text, user_type)

    rows = []
    for e in parsed:
        from_d  = _parse_date_flex(_v(e.get('from_str') or ''))
        to_d    = _parse_date_flex(_v(e.get('to_str') or ''))
        y, m, dys = _calc_duration(from_d, to_d)
        rows.append({
            'post':     _v(e.get('post') or user_type),
            'employer': _v(e.get('employer') or ''),
            'from_str': from_d.strftime('%d/%m/%Y') if from_d else _v(e.get('from_str') or ''),
            'to_str':   to_d.strftime('%d/%m/%Y') if to_d else 'Present',
            'years': y, 'months': m, 'days': dys,
        })
    return rows

# ── DOCX builder ─────────────────────────────────────────────────────

def _build_vos_nurse_docx(staff_doc, rows):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY  = RGBColor(0x1B, 0x3A, 0x6B)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    s1        = staff_doc.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff_doc.get('email') or s1.get('email_address') or '')
    user_type = _v(staff_doc.get('user_type') or 'Registered Nurse')
    pps       = _v(s1.get('pps_number') or '')

    d = Document()
    for sec in d.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)
    d.styles['Normal'].font.name = 'Calibri'
    d.styles['Normal'].font.size = Pt(11)

    def _run(para, text, bold=False, size=11, color=None, italic=False):
        r = para.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.name = 'Calibri'; r.font.size = Pt(size)
        r.font.color.rgb = color if color else BLACK
        return r

    def _set_cell_bg(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

    def _set_tbl_border(cell, color='CCCCCC'):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
            b.set(qn('w:color'), color); tcBorders.append(b)
        tcPr.append(tcBorders)

    def sp(pts=6):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = Pt(pts)

    # ── Logo top-right ────────────────────────────────────────────────
    logo_added = False
    for logo_path in [
        'static/image/logo.png',
        'static/img/logo.png',
        os.path.join(os.path.dirname(__file__), '..', 'static', 'image', 'logo.png'),
    ]:
        if os.path.exists(logo_path):
            try:
                p_logo = d.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_logo.paragraph_format.space_before = Pt(0)
                p_logo.paragraph_format.space_after  = Pt(4)
                p_logo.add_run().add_picture(logo_path, width=Inches(1.6))
                logo_added = True
                break
            except Exception: pass
    if not logo_added:
        p_logo = d.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p_logo, 'XPRESS HEALTH', bold=True, size=13, color=NAVY)

    sp(6)

    # ── Title ─────────────────────────────────────────────────────────
    p_title = d.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after  = Pt(10)
    r_title = _run(p_title, 'Verification of Service \u2013 Agency Staff', bold=True, size=18, color=NAVY)

    # ── Agency / HSE Site fields ──────────────────────────────────────
    def _agency_line(label, value='', blank=False):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        _run(p, label + ':  ', bold=True, size=11)
        if blank:
            _run(p, '________________________________', size=11)
        else:
            _run(p, value, size=11)

    sp(10)

    # ── Staff Details table (2 col) ────────────────────────────────────
    p_sd = d.add_paragraph()
    p_sd.paragraph_format.space_before = Pt(4)
    p_sd.paragraph_format.space_after  = Pt(6)
    _run(p_sd, 'Staff Details', bold=True, size=12, color=NAVY)

    staff_rows = [
        ('Full Name',         full_name),
        ('Role / Designation', user_type),
    ]
    staff_tbl = d.add_table(rows=len(staff_rows), cols=2)
    staff_tbl.style = 'Table Grid'
    staff_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (label, value) in enumerate(staff_rows):
        row = staff_tbl.rows[ri]
        # Col widths
        row.cells[0].width = Cm(6.0)
        row.cells[1].width = Cm(11.0)
        _set_tbl_border(row.cells[0], '888888')
        _set_tbl_border(row.cells[1], '888888')
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after  = Pt(4)
        _run(p0, label, bold=True, size=11)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after  = Pt(4)
        _run(p1, value, size=11)

    sp(12)

    # ── Service Details table ─────────────────────────────────────────
    p_srv = d.add_paragraph()
    p_srv.paragraph_format.space_before = Pt(4)
    p_srv.paragraph_format.space_after  = Pt(6)
    _run(p_srv, 'Service Details', bold=True, size=12, color=NAVY)

    headers    = ['Post', 'HSE\nLocation', 'From\nDate', 'To Date', 'Years', 'Months', 'Days', 'Verified']
    col_widths = [Cm(2.8), Cm(4.2), Cm(2.6), Cm(2.6), Cm(1.3), Cm(1.5), Cm(1.3), Cm(1.7)]

    svc_tbl = d.add_table(rows=1 + len(rows), cols=8)
    svc_tbl.style = 'Table Grid'

    # Header row
    hrow = svc_tbl.rows[0]
    for ci, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hrow.cells[ci]
        cell.width = w
        _set_cell_bg(cell, '1B3A6B')
        _set_tbl_border(cell, '1B3A6B')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, hdr, bold=True, size=10, color=WHITE)

    # Data rows
    for ri, row in enumerate(rows):
        to_display = row.get('to_str','')
        if to_display.lower() in ('present','') or not _parse_date_flex(to_display):
            today = date.today()
            to_display = f"Present\n(June {today.year})"

        row_data = [
            row.get('post', user_type),
            row.get('employer',''),
            row.get('from_str',''),
            to_display,
            str(row.get('years',0)),
            str(row.get('months',0)),
            str(row.get('days',0)),
            'yes',
        ]
        bg = 'FFFFFF'
        tr = svc_tbl.rows[ri + 1]
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = tr.cells[ci]
            cell.width = w
            _set_cell_bg(cell, bg)
            _set_tbl_border(cell, '888888')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _run(p, val, size=10)

    sp(10)

    # ── Total Eligible Service ────────────────────────────────────────
    tot_y, tot_m, tot_d = _total_service(rows)
    p_tot = d.add_paragraph()
    p_tot.paragraph_format.space_before = Pt(4)
    p_tot.paragraph_format.space_after  = Pt(10)
    _run(p_tot,
         f'Total Eligible Service for Incremental Credit: '
         f'{tot_y} Year{"s" if tot_y!=1 else ""}, '
         f'{tot_m} Month{"s" if tot_m!=1 else ""}, '
         f'{tot_d} Day{"s" if tot_d!=1 else ""}',
         size=11)

    # ── Declaration ────────────────────────────────────────────────────
    p_decl = d.add_paragraph()
    p_decl.paragraph_format.space_before = Pt(4)
    p_decl.paragraph_format.space_after  = Pt(20)
    _run(p_decl, 'I hereby confirm that the above details are accurate and verified as per agency records.', size=11)

    # ── Signatory block ───────────────────────────────────────────────
    def _sig_line(label, value='', blank=False):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, label + ': ', bold=False, size=11)
        if blank:
            _run(p, '______________________________', size=11)
        else:
            _run(p, value, size=11)

    _sig_line('Authorised Signatory', 'Betsy D')
    _sig_line('Designation', 'Head of Recruitment')
    _sig_line('Agency Stamp & Signature', blank=True)
    _sig_line('Date', blank=True)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()

# ── Cron ──────────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/cron/generate-vos-nurse', methods=['GET','POST'])
def live_staff_cron_generate_vos_nurse():
    """
    Cron — generate VOS document for ONE Nurse staff per call.
    Auth: cron_key query param or X-Cron-Key header.
    """
    cron_secret = os.environ.get('CRON_SECRET','')
    if cron_secret:
        provided = (request.args.get('cron_key') or request.headers.get('X-Cron-Key',''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    col      = _vos_nurse_col()
    staffs   = _staffs_col()

    existing_ids = set(
        str(r['staff_id']) for r in col.find({}, {"staff_id":1}) if r.get('staff_id')
    )

    # Nurse user types
    nurse_query = {
        "$and": [
            {"user_type": {"$regex": "nurse|rgn|rnm|midwife|nchd", "$options": "i"}},
            {"_id": {"$nin": [ObjectId(i) for i in existing_ids if len(i)==24]}},
        ]
    }
    remaining = staffs.count_documents(nurse_query)
    staff     = staffs.find_one(nurse_query)

    if not staff:
        return jsonify({"success": True, "message": "All nurse VOS documents generated.", "remaining_count": 0})

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    col.insert_one({"staff_id": staff_id, "staff_name": full_name, "status": "processing", "created_at": datetime.utcnow()})

    gemini_key = os.environ.get('GEMINI_API_KEY','')

    def _do_generate():
        try:
            rows = _get_employment_rows(staff, gemini_key=gemini_key)
            if not rows:
                col.update_one({"staff_id": staff_id}, {"$set": {
                    "status": "no_employment_data",
                    "note":   "No Hse Cv found in XN Portal or no Ireland nursing roles extracted",
                    "updated_at": datetime.utcnow(),
                }}); return

            docx_bytes = _build_vos_nurse_docx(staff, rows)
            safe_name  = (full_name or 'staff').replace(' ','_').replace('/','_')
            filename   = f"VOS_Nurse_{safe_name}.docx"
            gcs_blob   = f"vos_nurse/{filename}"
            _gcs_upload(gcs_blob, docx_bytes,
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            tot_y, tot_m, tot_d = _total_service(rows)
            col.update_one({"staff_id": staff_id}, {"$set": {
                "staff_name":   full_name, "email": email,
                "user_type":    _v(staff.get('user_type')),
                "filename":     filename,  "gcs_blob": gcs_blob,
                "rows":         rows,
                "total_years":  tot_y, "total_months": tot_m, "total_days": tot_d,
                "status":       "generated", "generated_at": datetime.utcnow(),
            }})
        except Exception as e:
            col.update_one({"staff_id": staff_id}, {"$set": {
                "status": "error", "error": str(e), "updated_at": datetime.utcnow(),
            }})

    threading.Thread(target=_do_generate, daemon=True).start()
    return jsonify({
        "success": True, "staff_id": staff_id, "staff_name": full_name, "email": email,
        "remaining_count": max(0, remaining-1),
        "message": f"VOS generation started for {full_name} — {max(0,remaining-1)} remaining.",
    })

# ── API: check saved ──────────────────────────────────────────────────

@admin_bp.route('/live-staffs/api/vos-nurse-saved/<staff_id>')
@admin_required
def live_staff_api_vos_nurse_saved(staff_id):
    rec = _vos_nurse_col().find_one({"staff_id": staff_id, "status": "generated"})
    if rec:
        tot_y = rec.get('total_years',0); tot_m = rec.get('total_months',0); tot_d = rec.get('total_days',0)
        return jsonify({
            "success": True, "found": True,
            "generated_at": rec["generated_at"].strftime("%d %b %Y %H:%M") if rec.get("generated_at") else "",
            "total": f"{tot_y} Year{'s' if tot_y!=1 else ''}, {tot_m} Month{'s' if tot_m!=1 else ''}, {tot_d} Day{'s' if tot_d!=1 else ''}",
        })
    return jsonify({"success": True, "found": False})

# ── Download ──────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/vos-nurse/download/<staff_id>')
@admin_required
def live_staff_vos_nurse_download(staff_id):
    try:
        rec = _vos_nurse_col().find_one({"staff_id": staff_id})
        if not rec or rec.get('status') != 'generated' or not rec.get('gcs_blob'):
            # On-demand generate
            staff_doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
            if not staff_doc: return jsonify({"error": "Staff not found"}), 404
            rows = _get_employment_rows(staff_doc)
            if not rows: return jsonify({"error": "No employment data"}), 404
            docx_bytes = _build_vos_nurse_docx(staff_doc, rows)
            s1 = staff_doc.get('section_1_personal_details') or {}
            filename = f"VOS_Nurse_{_v(s1.get('full_name') or 'staff').replace(' ','_')}.docx"
        else:
            docx_bytes = _gcs_download(rec['gcs_blob'])
            filename   = rec.get('filename','VOS_Nurse.docx')
        return Response(docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f'attachment; filename="{filename}"'})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ── On-demand generate ────────────────────────────────────────────────

@admin_bp.route('/live-staffs/vos-nurse/generate/<staff_id>', methods=['POST'])
@admin_required
def live_staff_vos_nurse_generate(staff_id):
    try:
        staff_doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not staff_doc: return jsonify({"success": False, "error": "Staff not found"}), 404
        gemini_key = os.environ.get('GEMINI_API_KEY','')
        rows = _get_employment_rows(staff_doc, gemini_key=gemini_key)
        if not rows: return jsonify({"success": False, "error": "No Ireland nursing roles found in HSE CV"}), 400

        s1        = staff_doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff')
        email     = _v(staff_doc.get('email') or s1.get('email_address') or '')
        docx_bytes= _build_vos_nurse_docx(staff_doc, rows)
        safe_name = full_name.replace(' ','_').replace('/','_')
        filename  = f"VOS_Nurse_{safe_name}.docx"
        gcs_blob  = f"vos_nurse/{filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        tot_y, tot_m, tot_d = _total_service(rows)
        _vos_nurse_col().update_one(
            {"staff_id": staff_id},
            {"$set": {"staff_name": full_name, "email": email,
                      "user_type": _v(staff_doc.get('user_type')),
                      "filename": filename, "gcs_blob": gcs_blob, "rows": rows,
                      "total_years": tot_y, "total_months": tot_m, "total_days": tot_d,
                      "status": "generated", "generated_at": datetime.utcnow()}},
            upsert=True
        )
        return jsonify({"success": True, "staff_name": full_name, "filename": filename,
                        "total": f"{tot_y} Years, {tot_m} Months, {tot_d} Days", "row_count": len(rows)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# ── Upload ────────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/vos-nurse/upload/<staff_id>', methods=['POST'])
@admin_required
def live_staff_vos_nurse_upload(staff_id):
    f = request.files.get('file')
    if not f: return jsonify({"success": False, "error": "No file uploaded"}), 400
    if not f.filename.lower().endswith('.docx'): return jsonify({"success": False, "error": "Only .docx accepted"}), 400
    try:
        staff_doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not staff_doc: return jsonify({"success": False, "error": "Staff not found"}), 404
        s1        = staff_doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff').replace(' ','_').replace('/','_')
        filename  = f"VOS_Nurse_{full_name}.docx"
        gcs_blob  = f"vos_nurse/{filename}"
        docx_bytes= f.read()
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        _vos_nurse_col().update_one({"staff_id": staff_id},
            {"$set": {"gcs_blob": gcs_blob, "filename": filename,
                      "status": "generated", "generated_at": datetime.utcnow(), "uploaded_by": "admin"}},
            upsert=True)
        return jsonify({"success": True, "gcs_blob": gcs_blob, "filename": filename})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# ── Excel export ──────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/export/vos-nurse-xlsx')
@admin_required
def live_staff_export_vos_nurse_xlsx():
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        docs = sorted(list(_vos_nurse_col().find({})), key=lambda d: _v(d.get('staff_name') or '').lower())
        NAVY='1B3A6B'; WHITE='FFFFFF'; ALT='EFF6FF'; RED='FFDDDD'; WARN='FFF3CD'
        h_font=Font(name='Calibri',bold=True,color=WHITE,size=10)
        h_fill=PatternFill('solid',start_color=NAVY,end_color=NAVY)
        h_align=Alignment(horizontal='center',vertical='center')
        b_font=Font(name='Calibri',size=10)
        thin=Side(style='thin',color='CCCCCC')
        border=Border(left=thin,right=thin,top=thin,bottom=thin)
        wb = Workbook(); ws = wb.active; ws.title = 'VOS Nurse Documents'
        headers=['Sno','Staff Name','Email','User Type','Total Years','Total Months','Total Days','Total Service','Status','Generated At']
        col_widths=[5,30,36,24,12,13,10,28,14,20]
        for ci,(hdr,w) in enumerate(zip(headers,col_widths),1):
            cell=ws.cell(row=1,column=ci,value=hdr)
            cell.font=h_font; cell.fill=h_fill; cell.alignment=h_align; cell.border=border
            ws.column_dimensions[cell.column_letter].width=w
        ws.freeze_panes='A2'
        for ri,doc in enumerate(docs,2):
            tot_y=doc.get('total_years',''); tot_m=doc.get('total_months',''); tot_d=doc.get('total_days','')
            status=_v(doc.get('status') or '')
            gen_at=doc.get('generated_at')
            gen_str=gen_at.strftime('%d %b %Y %H:%M') if gen_at else ''
            total_str=f"{tot_y} Yr, {tot_m} Mo, {tot_d} D" if tot_y!='' else ''
            fill=PatternFill('solid',start_color=RED,end_color=RED) if status=='error' else \
                 PatternFill('solid',start_color=WARN,end_color=WARN) if status!='generated' else \
                 (PatternFill('solid',start_color=ALT,end_color=ALT) if ri%2==0 else PatternFill())
            row_vals=[ri-1,_v(doc.get('staff_name') or ''),_v(doc.get('email') or ''),_v(doc.get('user_type') or ''),
                      tot_y,tot_m,tot_d,total_str,status,gen_str]
            for ci,val in enumerate(row_vals,1):
                cell=ws.cell(row=ri,column=ci,value=val)
                cell.font=b_font; cell.border=border; cell.fill=fill
                cell.alignment=Alignment(horizontal='center' if ci not in (2,3,4,8) else 'left',vertical='center')
        buf=io.BytesIO(); wb.save(buf)
        return Response(buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":f'attachment; filename="vos_nurse_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
