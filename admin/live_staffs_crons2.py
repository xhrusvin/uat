from flask import request, jsonify, Response
from bson import ObjectId
from datetime import datetime
import json
import json as _cjson
import json as _cjson2
import base64
import csv
import io
import re as _re
import re
import os
import threading

from database import db
from . import admin_bp
from admin.views import admin_required

# ── Helpers — lazy wrappers to avoid circular imports ─────────────────

def _v(val):
    if val is None: return ''
    return str(val).strip()

def _staffs_col():
    from flask import current_app
    return current_app.db.live_staffs

def _gcs_upload(blob_name, data_bytes, content_type='application/octet-stream'):
    from admin.live_staffs import _gcs_upload as _f
    return _f(blob_name, data_bytes, content_type)

def _gcs_download(blob_name):
    from admin.live_staffs import _gcs_download as _f
    return _f(blob_name)

def _gcs_signed_url(blob_name, expiry_minutes=60):
    from admin.live_staffs import _gcs_signed_url as _f
    return _f(blob_name, expiry_minutes)

def _ai_pcc_col():
    from flask import current_app
    return current_app.db.live_staff_ai_pcc

def _ai_cvs_col():
    from flask import current_app
    return current_app.db.live_staff_ai_cvs

def _ai_interviews_col():
    from flask import current_app
    return current_app.db.live_staff_ai_interviews

def _ai_appforms_col():
    from flask import current_app
    return current_app.db.live_staff_ai_appforms

def _build_pcc_docx(doc, reviewer_index=0):
    from admin.live_staffs import _build_pcc_docx as _f
    return _f(doc, reviewer_index)

def _build_ai_cv_docx(doc, cv_text):
    from admin.live_staffs import _build_ai_cv_docx as _f
    return _f(doc, cv_text)

def _build_ai_interview_docx(doc, interview_text):
    from admin.live_staffs import _build_ai_interview_docx as _f
    return _f(doc, interview_text)

def _build_ai_appform_docx(doc, appform_text):
    from admin.live_staffs import _build_ai_appform_docx as _f
    return _f(doc, appform_text)

def _build_appform_docx(doc):
    from admin.live_staffs import _build_appform_docx as _f
    return _f(doc)

def _build_interview_docx(doc, interview_text):
    from admin.live_staffs import _build_interview_docx as _f
    return _f(doc, interview_text)

def _build_doc(template, context):
    from admin.live_staffs import _build_doc as _f
    return _f(template, context)

def _build_qual_xlsx(doc):
    from admin.live_staffs import _build_qual_xlsx as _f
    return _f(doc)

def _extract_text_from_url(url, headers=None):
    from admin.live_staffs_crons import _extract_text_from_url as _etf
    return _etf(url, headers)

def _get_pcc_reviewers():
    from admin.live_staffs import _PCC_REVIEWERS
    return _PCC_REVIEWERS

def _get_compliance_officer():
    from admin.live_staffs import _PCC_COMPLIANCE_OFFICER
    return _PCC_COMPLIANCE_OFFICER

def _push_hse_document_background(staff_id_str, doc_type_key,
                                   gcs_blob, filename,
                                   user_type=None):
    from admin.live_staffs import _push_hse_document_background as _f
    return _f(staff_id_str, doc_type_key, gcs_blob, filename, user_type)


@admin_bp.route('/live-staffs/cron/sync-garda-vetting', methods=['GET', 'POST'])
def live_staff_cron_sync_garda_vetting():
    """
    Cron job — processes ONE staff member per call.
    Finds "Garda Vetting Document", extracts details via Gemini AI.
    Saves: garda_cert_name, garda_staff_name, garda_issue_date,
           garda_reference_number, garda_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"garda_fetched": {"$exists": False}},
            {"garda_fetched": False},
            {"garda_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Garda Vetting documents already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["garda_fetched"]    = True
        fields["garda_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"garda_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"garda_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"garda_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"garda_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    garda_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'garda vetting document',
            'garda vetting',
            'garda',
            'vetting document',
            'national vetting bureau',
            'nvb',
        )) and d.get('url'):
            garda_doc = d
            break

    if not garda_doc:
        _mark_done({"garda_note": "no Garda Vetting document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Garda Vetting document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (garda_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"garda_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Garda doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"garda_note": "document URL 404 — skipped",
                        "garda_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Garda doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor specialising in Irish Garda Vetting documents.

Extract the following details from this Garda Vetting disclosure document:
1. Document / certificate name (e.g. "Garda Vetting Disclosure", "National Vetting Bureau Disclosure")
2. Staff name as printed on the document
3. Issue date / date of disclosure
4. Reference number or vetting reference (if shown)
5. Issuing body (e.g. "National Vetting Bureau", "An Garda Siochana")

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact document title as printed>",
  "staff_name_on_cert": "<name as printed on document>",
  "issue_date": "<date of issue or disclosure as printed>",
  "reference_number": "<vetting reference number if visible>",
  "issuing_body": "<organization that issued the document>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result     = _cjson.loads(raw_out)
        cert_name  = _v(result.get('certificate_name') or '')
        cert_staff = _v(result.get('staff_name_on_cert') or '')
        issue_date = _v(result.get('issue_date') or '')
        ref_num    = _v(result.get('reference_number') or '')
        issuer     = _v(result.get('issuing_body') or '')

        _mark_done({
            "garda_cert_name":      cert_name,
            "garda_staff_name":     cert_staff,
            "garda_issue_date":     issue_date,
            "garda_reference":      ref_num,
            "garda_issuing_body":   issuer,
            "garda_doc_url":        doc_url,
            "garda_doc_type":       garda_doc.get('document_type_name', ''),
            "garda_note":           "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "issue_date":         issue_date,
            "reference_number":   ref_num,
            "issuing_body":       issuer,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Garda Vetting extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"garda_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"garda_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Garda Vetting to Excel ───────────────────────────────────

@admin_bp.route('/live-staffs/export/garda-xlsx')
@admin_required
def live_staff_export_garda_xlsx():
    """Export Garda Vetting document details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "garda_cert_name": 1, "garda_staff_name": 1,
             "garda_issue_date": 1, "garda_reference": 1,
             "garda_issuing_body": 1, "garda_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Garda Vetting'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Issue Date', 'Reference No', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 20, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('garda_cert_name') or '')
            cert_s   = _v(doc.get('garda_staff_name') or '')
            issue    = _v(doc.get('garda_issue_date') or '')
            ref_n    = _v(doc.get('garda_reference') or '')
            issuer   = _v(doc.get('garda_issuing_body') or '')
            fetched  = doc.get('garda_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Doc Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, issue, ref_n, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="garda_vetting_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract CPI/MAPA/PMAV Certificate ──────────────────────────

@admin_bp.route('/live-staffs/cron/sync-cpi-mapa', methods=['GET', 'POST'])
def live_staff_cron_sync_cpi_mapa():
    """
    Cron job — processes ONE staff member per call.
    Finds "Cpi/ Mapa/Pmav" document, extracts details via Gemini AI.
    Saves: cpi_certificate_name, cpi_staff_name, cpi_expiry_date,
           cpi_issue_date, cpi_issuing_body, cpi_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"cpi_fetched": {"$exists": False}},
            {"cpi_fetched": False},
            {"cpi_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff CPI/MAPA/PMAV certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["cpi_fetched"]    = True
        fields["cpi_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"cpi_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"cpi_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"cpi_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"cpi_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    cpi_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'cpi/ mapa/pmav', 'cpi/mapa/pmav', 'cpi mapa pmav',
            'cpi', 'mapa', 'pmav',
            'crisis prevention', 'management of actual',
            'prevention management',
        )) and d.get('url'):
            cpi_doc = d
            break

    if not cpi_doc:
        _mark_done({"cpi_note": "no CPI/MAPA/PMAV document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No CPI/MAPA/PMAV certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (cpi_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"cpi_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — CPI/MAPA/PMAV doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"cpi_note": "document URL 404 — skipped",
                        "cpi_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — CPI/MAPA/PMAV doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this CPI / MAPA / PMAV training certificate:
1. Certificate name (e.g. "CPI Non-Violent Crisis Intervention", "MAPA Foundation", "PMAV Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "cpi_certificate_name": cert_name,
            "cpi_staff_name":       cert_staff,
            "cpi_expiry_date":      expiry_date,
            "cpi_issue_date":       issue_date,
            "cpi_issuing_body":     issuing_body,
            "cpi_doc_url":          doc_url,
            "cpi_doc_type":         cpi_doc.get('document_type_name', ''),
            "cpi_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"CPI/MAPA/PMAV cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"cpi_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"cpi_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: CPI/MAPA/PMAV certificates to Excel ───────────────────────

@admin_bp.route('/live-staffs/export/cpi-xlsx')
@admin_required
def live_staff_export_cpi_xlsx():
    """Export CPI/MAPA/PMAV certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "cpi_certificate_name": 1, "cpi_staff_name": 1,
             "cpi_expiry_date": 1, "cpi_issue_date": 1,
             "cpi_issuing_body": 1, "cpi_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'CPI-MAPA-PMAV Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('cpi_certificate_name') or '')
            cert_s   = _v(doc.get('cpi_staff_name') or '')
            expiry   = _v(doc.get('cpi_expiry_date') or '')
            issue    = _v(doc.get('cpi_issue_date') or '')
            issuer   = _v(doc.get('cpi_issuing_body') or '')
            fetched  = doc.get('cpi_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="cpi_mapa_pmav_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Employment Contract Signed details ─────────────────

@admin_bp.route('/live-staffs/cron/sync-employment-contract', methods=['GET', 'POST'])
def live_staff_cron_sync_employment_contract():
    """
    Cron job — processes ONE staff member per call.
    Finds "Employment Contract Signed" document, extracts details via Gemini.
    Saves: ec_contract_name, ec_staff_name, ec_signed_date,
           ec_employer_name, ec_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"ec_fetched": {"$exists": False}},
            {"ec_fetched": False},
            {"ec_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Employment Contracts already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["ec_fetched"]    = True
        fields["ec_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"ec_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"ec_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"ec_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"ec_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    ec_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'employment contract signed',
            'employment contract',
            'signed contract',
            'contract signed',
            'contract of employment',
        )) and d.get('url'):
            ec_doc = d
            break

    if not ec_doc:
        _mark_done({"ec_note": "no Employment Contract document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Employment Contract found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (ec_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"ec_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Employment Contract has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"ec_note": "document URL 404 — skipped",
                        "ec_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Employment Contract URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this Employment Contract:
1. Contract / document name (e.g. "Employment Contract", "Contract of Employment", "Staff Agreement")
2. Employee / staff name as printed on the contract
3. Date the contract was signed (signature date)
4. Employer name (the company or organisation offering the contract)

Return ONLY a JSON object — no markdown, no explanation:
{
  "contract_name": "<exact document title as printed>",
  "staff_name_on_doc": "<employee name as printed>",
  "signed_date": "<date the contract was signed as printed>",
  "employer_name": "<employer or company name>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result        = _cjson.loads(raw_out)
        contract_name = _v(result.get('contract_name') or '')
        staff_on_doc  = _v(result.get('staff_name_on_doc') or '')
        signed_date   = _v(result.get('signed_date') or '')
        employer_name = _v(result.get('employer_name') or '')

        _mark_done({
            "ec_contract_name": contract_name,
            "ec_staff_name":    staff_on_doc,
            "ec_signed_date":   signed_date,
            "ec_employer_name": employer_name,
            "ec_doc_url":       doc_url,
            "ec_doc_type":      ec_doc.get('document_type_name', ''),
            "ec_note":          "extracted successfully",
        })

        return jsonify({
            "success":          True,
            "email":            email,
            "staff_name":       full_name,
            "doc_found":        True,
            "contract_name":    contract_name,
            "staff_name_on_doc": staff_on_doc,
            "signed_date":      signed_date,
            "employer_name":    employer_name,
            "remaining_count":  max(0, remaining_total - 1),
            "message": (
                f"Employment Contract extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"ec_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"ec_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Employment Contracts to Excel ─────────────────────────────

@admin_bp.route('/live-staffs/export/employment-contract-xlsx')
@admin_required
def live_staff_export_employment_contract_xlsx():
    """Export Employment Contract details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "ec_contract_name": 1, "ec_staff_name": 1,
             "ec_signed_date": 1, "ec_employer_name": 1, "ec_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Employment Contracts'

        headers    = ['Sno', 'Staff Name', 'Email', 'Contract Name',
                      'Name on Contract', 'Signed Date', 'Employer', 'Status']
        col_widths = [5, 28, 36, 30, 28, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:H{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1        = doc.get('section_1_personal_details') or {}
            name      = _v(s1.get('full_name') or '')
            email     = _v(doc.get('email') or '')
            cont_n    = _v(doc.get('ec_contract_name') or '')
            cont_s    = _v(doc.get('ec_staff_name') or '')
            signed    = _v(doc.get('ec_signed_date') or '')
            employer  = _v(doc.get('ec_employer_name') or '')
            fetched   = doc.get('ec_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cont_n:
                status   = 'No Contract Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cont_n, cont_s, signed, employer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="employment_contracts_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Open Disclosure Certificate ─────────────────────────

@admin_bp.route('/live-staffs/cron/sync-open-disclosure', methods=['GET', 'POST'])
def live_staff_cron_sync_open_disclosure_new():
    """
    Cron job — processes ONE staff member per call.
    Finds "Open Disclosure" document, extracts details via Gemini AI.
    Saves: od_certificate_name, od_staff_name, od_expiry_date,
           od_issue_date, od_issuing_body, od_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"od_fetched": {"$exists": False}},
            {"od_fetched": False},
            {"od_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Open Disclosure certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["od_fetched"]    = True
        fields["od_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"od_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"od_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"od_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"od_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    od_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'open disclosure',
            'the open disclosure',
            'open disclosure certificate',
            'open disclosure training',
        )) and d.get('url'):
            od_doc = d
            break

    if not od_doc:
        _mark_done({"od_note": "no Open Disclosure document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Open Disclosure certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (od_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"od_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Open Disclosure doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"od_note": "document URL 404 — skipped", "od_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Open Disclosure doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Open Disclosure certificate or training record:
1. Certificate name (e.g. "Open Disclosure", "The Open Disclosure", "Open Disclosure Training")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "od_certificate_name": cert_name,
            "od_staff_name":       cert_staff,
            "od_expiry_date":      expiry_date,
            "od_issue_date":       issue_date,
            "od_issuing_body":     issuing_body,
            "od_doc_url":          doc_url,
            "od_doc_type":         od_doc.get('document_type_name', ''),
            "od_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Open Disclosure cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"od_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"od_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Open Disclosure certificates to Excel ─────────────────────

@admin_bp.route('/live-staffs/export/open-disclosure-xlsx')
@admin_required
def live_staff_export_open_disclosure_xlsx_new():
    """Export Open Disclosure certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "od_certificate_name": 1, "od_staff_name": 1,
             "od_expiry_date": 1, "od_issue_date": 1,
             "od_issuing_body": 1, "od_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Open Disclosure Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y',
                        '%B %Y','%b %Y','%d %B, %Y','%d %b, %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip().rstrip(','), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('od_certificate_name') or '')
            cert_s   = _v(doc.get('od_staff_name') or '')
            expiry   = _v(doc.get('od_expiry_date') or '')
            issue    = _v(doc.get('od_issue_date') or '')
            issuer   = _v(doc.get('od_issuing_body') or '')
            fetched  = doc.get('od_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="open_disclosure_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Reviewers list — rotated alternately per staff ────────────────────
_PCC_REVIEWERS = [
    'Letty Mathew',
    'Valencia Da Silva',
    'Ann Maria',
    'Audrey Maguire',
    'Liberata Gama',
]
_PCC_COMPLIANCE_OFFICER = 'Betsy Daniel'


def _build_pcc_docx(doc, reviewer_index=0):
    """
    Generate the PCC Self-Declaration & Risk Assessment DOCX for a staff member.
    Pure Python — uses python-docx only, no Node.js required.
    Returns bytes of the generated DOCX.
    """
    import io as _io
    from datetime import datetime as _dt, timedelta as _td
    from docx import Document as _Doc
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    s1          = doc.get('section_1_personal_details') or {}
    full_name   = _v(s1.get('full_name') or '')
    dob         = _v(s1.get('date_of_birth') or '')
    nationality = _v(s1.get('nationality') or '')
    role        = _v(doc.get('user_type') or '')
    first_shift = _v(doc.get('first_shift_date') or '')

    # Employment history for Section 3
    s5 = doc.get('section_5_employment_history') or {}
    emp_entries = [e for e in (s5.get('entries') or [])
                   if e.get('employer') or e.get('position')]

    # Signature image bytes from GCS (if available)
    sig_bytes = None
    sig_blob  = _v(doc.get('signature_gcs_blob') or '')
    if sig_blob:
        try:
            sig_bytes = _gcs_download(sig_blob)
        except Exception:
            sig_bytes = None

    # Reviewer rotation
    reviewer = _get_pcc_reviewers()[reviewer_index % len(_get_pcc_reviewers())]

    # Parse first_shift date
    first_shift_formatted = ''
    date_reviewed         = ''
    if first_shift:
        for fmt in ('%d/%m/%Y','%Y-%m-%d','%d-%m-%Y','%d %B %Y','%d %b %Y'):
            try:
                d_shift = _dt.strptime(first_shift.strip(), fmt)
                first_shift_formatted = d_shift.strftime('%d %B %Y')
                # Date Reviewed = first shift + 1 day
                date_reviewed = (d_shift + _td(days=1)).strftime('%d %B %Y')
                break
            except Exception:
                continue

    # ── Helpers ──────────────────────────────────────────────────────
    NAVY  = RGBColor(0x1B, 0x3A, 0x6B)
    GREEN = RGBColor(0x2E, 0x9E, 0x44)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    GRAY  = RGBColor(0x64, 0x74, 0x8B)

    def _set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _set_cell_border(cell, color='CCCCCC', sz=4):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    str(sz))
            b.set(qn('w:color'), color)
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def _para_border_bottom(para, color='1B3A6B', sz=8):
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    str(sz))
        bot.set(qn('w:color'), color)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _add_run(para, text, bold=False, size=10, color=None, italic=False):
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Arial'
        if color:
            run.font.color.rgb = color
        return run

    document = _Doc()

    # Page margins
    for section in document.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Default font
    document.styles['Normal'].font.name = 'Arial'
    document.styles['Normal'].font.size = Pt(10)

    def sp(before=4, after=4):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        return p

    def section_heading(text):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
        _para_border_bottom(p, '1B3A6B', 8)
        run = p.add_run('  ' + text + '  ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.font.color.rgb = WHITE
        # Shading on run
        rPr  = run._r.get_or_add_rPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1B3A6B')
        rPr.append(shd)
        return p

    def label_value(label, value, bold_label=True):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _add_run(p, label + '  ', bold=bold_label, size=10)
        _add_run(p, value if value else '_' * 35, bold=False, size=10)
        return p

    def body_text(text, size=9.5):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _add_run(p, text, size=size)
        return p

    def numbered_item(num, text):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.8)
        _add_run(p, f'{num}.  ', bold=True, size=9.5)
        _add_run(p, text, size=9.5)
        return p

    def checkbox_item(text, checked=False):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _add_run(p, ('☑' if checked else '☐') + '  ', size=11)
        _add_run(p, text, size=9.5)
        return p

    def blank_line(label):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        _add_run(p, label + '  ', bold=True, size=9.5)
        _add_run(p, '_' * 40, size=9.5, color=GRAY)
        return p

    # ── Header ────────────────────────────────────────────────────────
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_border_bottom(p, '2E9E44', 6)
    _add_run(p, 'XPRESS HEALTH', bold=True, size=18, color=NAVY)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, 'SELF-DECLARATION & INTERNATIONAL RISK ASSESSMENT FORM', bold=True, size=11, color=GRAY)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, 'International Police Clearance Certificate (PCC) — Residential History Declaration',
             italic=True, size=9, color=GRAY)
    sp(4, 8)

    # ── Section 1 ─────────────────────────────────────────────────────
    section_heading('SECTION 1 — EMPLOYEE DETAILS')
    sp(2)
    tbl = document.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'
    for row in tbl.rows:
        for cell in row.cells:
            for border_el in cell._tc.iter(qn('w:tcBorders')):
                border_el.getparent().remove(border_el)
    tbl.cell(0,0).text = ''; tbl.cell(0,1).text = ''
    tbl.cell(1,0).text = ''; tbl.cell(1,1).text = ''
    for p_, label_, val_ in [
        (tbl.cell(0,0).paragraphs[0], 'Employee Name:', full_name),
        (tbl.cell(0,1).paragraphs[0], 'Nationality:', nationality),
        (tbl.cell(1,0).paragraphs[0], 'Position / Role:', role),
        (tbl.cell(1,1).paragraphs[0], '', ''),
    ]:
        if label_:
            _add_run(p_, label_ + '  ', bold=True, size=10)
            _add_run(p_, val_ if val_ else '_' * 28, size=10)
    sp(4)

    # ── Section 2 ─────────────────────────────────────────────────────
    section_heading('SECTION 2 — PURPOSE OF THIS DECLARATION')
    body_text('As part of the recruitment and compliance process, an International Police Clearance Certificate (PCC) is required from all countries where you have resided. As you are currently unable to provide the required certificate(s), this Self-Declaration & Risk Assessment Form must be completed.')
    body_text('Completion of this form does not exempt you from the requirement to obtain and submit the PCC. You remain obligated to provide the certificate(s) as soon as reasonably practicable.')
    sp(2)

    # ── Section 3 — extract from extracted_cv via Gemini ────────────
    section_heading('SECTION 3 — INTERNATIONAL RESIDENTIAL HISTORY')
    body_text('Please list ALL countries (other than your current country of residence) where you have lived for six (6) months or more since the age of 18. Include the reason for your stay.')
    sp(2)

    # Parse employment history from extracted_cv using Gemini
    hist_rows = []
    extracted_cv_ = _v(doc.get('extracted_cv') or '')
    has_cv_ = (extracted_cv_ and not extracted_cv_.startswith('[')
               and extracted_cv_ != 'No doc found')

    IRELAND_TERMS = {
        'ireland', 'irish', 'republic of ireland', 'eire', 'éire',
        'ie', 'northern ireland', 'dublin', 'cork', 'galway', 'limerick',
        'waterford', 'kilkenny', 'wexford', 'wicklow', 'kildare', 'meath',
        'louth', 'monaghan', 'cavan', 'donegal', 'sligo', 'leitrim', 'roscommon',
        'mayo', 'galway', 'clare', 'tipperary', 'kilkenny', 'carlow', 'laois',
        'offaly', 'westmeath', 'longford', 'louth', 'kerry', 'cork', 'hse',
    }

    def _is_ireland(country, city_region=''):
        text = (country + ' ' + city_region).lower()
        return any(t in text for t in IRELAND_TERMS)

    if has_cv_:
        try:

            from google import genai as _genai2
            _gemini_key2 = __import__('os').environ.get('GEMINI_API_KEY', '')
            if _gemini_key2:
                _prompt2 = f"""You are a CV analyser. Extract all employment history entries from the CV text below.

For each job role extract:
- country: the country where the job was based
- city_region: the city, region, county or employer location
- from_date: start date (MM/YYYY format if possible)
- to_date: end date (MM/YYYY format) or "Present" if current
- reason: the job title / role description

Return ONLY a JSON array — no markdown:
[
  {{"country": "...", "city_region": "...", "from_date": "...", "to_date": "...", "reason": "..."}}
]

If country is not mentioned, infer from city/employer name if possible, otherwise leave blank.

CV TEXT:
{extracted_cv_[:6000]}
"""
                _client2 = _genai2.Client(api_key=_gemini_key2)
                _resp2   = _client2.models.generate_content(
                    model='gemini-2.5-flash', contents=_prompt2
                )
                _raw2 = (_resp2.text or '').strip()
                _raw2 = _re2.sub(r'^```(?:json)?\s*', '', _raw2, flags=_re2.MULTILINE)
                _raw2 = _re2.sub(r'```\s*$', '', _raw2, flags=_re2.MULTILINE).strip()
                _parsed2 = _cjson2.loads(_raw2)
                for _e in _parsed2[:20]:
                    country_    = _v(_e.get('country') or '')
                    city_region = _v(_e.get('city_region') or '')
                    # Skip Ireland entries
                    if _is_ireland(country_, city_region):
                        continue
                    hist_rows.append((
                        country_,
                        city_region,
                        _v(_e.get('from_date') or ''),
                        _v(_e.get('to_date') or 'Present'),
                        _v(_e.get('reason') or ''),
                    ))
        except Exception:
            pass

    # Fallback to section_5 entries if Gemini failed — also filter Ireland
    if not hist_rows:
        for e in emp_entries[:8]:
            country_    = _v(e.get('country') or '')
            city_region = _v(e.get('employer') or '')
            if _is_ireland(country_, city_region):
                continue
            hist_rows.append((
                country_,
                city_region,
                _v(e.get('from') or ''),
                _v(e.get('to') or 'Present'),
                _v(e.get('position') or ''),
            ))

    # Always show at least 4 blank rows (even if all entries were Ireland)
    while len(hist_rows) < 4:
        hist_rows.append(('', '', '', '', ''))

    # 4 columns — City/Region removed
    cols_   = ['Country', 'From (MM/YYYY)', 'To (MM/YYYY)', 'Reason for Stay']
    # Strip city_region (index 1) from each row
    hist_rows_4 = [(r[0], r[2], r[3], r[4]) for r in hist_rows]
    htbl  = document.add_table(rows=len(hist_rows_4) + 1, cols=4)
    htbl.style = 'Table Grid'
    for ci, hdr in enumerate(cols_):
        c  = htbl.cell(0, ci)
        _set_cell_bg(c, '1B3A6B')
        _set_cell_border(c, '1B3A6B')
        p_ = c.paragraphs[0]
        p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_, hdr, bold=True, size=9, color=WHITE)
    for ri, row_vals in enumerate(hist_rows_4, start=1):
        for ci, val in enumerate(row_vals):
            c = htbl.cell(ri, ci)
            _set_cell_border(c, 'CCCCCC')
            _add_run(c.paragraphs[0], val, size=9)
    sp(6)

    # ── Section 4 ─────────────────────────────────────────────────────
    section_heading('SECTION 4 — SELF-DECLARATION')
    body_text('I declare that:')
    numbered_item(1, 'I have accurately listed all countries in which I have resided for six (6) months or more since the age of 18.')
    numbered_item(2, 'I understand that I have not yet provided the required International Police Clearance Certificate(s) for the country/countries listed above.')
    numbered_item(3, 'I confirm that I have NEVER been convicted of a criminal offence in any country, nor am I currently the subject of any criminal investigation, prosecution, or pending criminal proceedings, EXCEPT as disclosed below.')
    sp(2)
    p = document.add_paragraph()
    _add_run(p, 'Disclosure of Criminal History (if applicable — leave blank if none):', italic=True, size=9, color=GRAY)
    for _ in range(3):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _para_border_bottom(p, 'CCCCCC', 4)
        _add_run(p, '', size=9)
    sp(4)

    # ── Section 5 ─────────────────────────────────────────────────────
    section_heading('SECTION 5 — RISK ACKNOWLEDGEMENT')
    body_text('I understand and acknowledge that:')
    numbered_item(1, 'Failure to provide the required International Police Clearance Certificate(s) may affect my compliance status and/or eligibility for certain assignments or postings.')
    numbered_item(2, 'My employment and/or continued engagement may be subject to additional compliance reviews, conditions, or restrictions until the PCC is received.')
    numbered_item(3, 'If any information provided in this declaration is found to be false, misleading, or incomplete, it may result in disciplinary action, withdrawal of employment offer, termination of engagement, or notification to relevant regulatory or statutory authorities.')
    numbered_item(4, 'I remain responsible for pursuing the required Police Clearance Certificate(s) and providing them to Xpress Health as soon as they become available.')
    sp(4)

    # ── Section 6 ─────────────────────────────────────────────────────
    section_heading('SECTION 6 — CONSENT')
    body_text('I authorise Xpress Health to:')
    numbered_item(1, 'Verify the information provided in this declaration where required.')
    numbered_item(2, 'Request additional information or supporting documentation relevant to my residential history or compliance status.')
    numbered_item(3, 'Conduct further background checks as permitted under applicable laws and regulations.')
    numbered_item(4, 'Retain this form as part of my personnel and compliance records.')
    sp(4)

    # ── Section 7 ─────────────────────────────────────────────────────
    section_heading('SECTION 7 — EMPLOYEE DECLARATION & SIGNATURE')
    body_text('I declare that the information provided in this form is true, complete, and accurate to the best of my knowledge and belief. I understand that knowingly providing false or misleading information may result in disciplinary action, including termination of employment or engagement.')
    sp(4)
    stbl = document.add_table(rows=2, cols=2)
    stbl.style = 'Table Grid'
    for row in stbl.rows:
        for cell in row.cells:
            for b in cell._tc.iter(qn('w:tcBorders')):
                b.getparent().remove(b)

    # Employee Signature cell — embed image if available
    sig_cell = stbl.cell(0, 0)
    _add_run(sig_cell.paragraphs[0], 'Employee Signature  ', bold=True, size=9.5)
    if sig_bytes:
        try:
            from docx.shared import Inches as _Inches
            import io as _sig_io
            sig_p = sig_cell.add_paragraph()
            sig_run = sig_p.add_run()
            sig_run.add_picture(_sig_io.BytesIO(sig_bytes), width=_Inches(1.5))
        except Exception:
            _add_run(sig_cell.paragraphs[0], '_' * 32, size=9.5, color=GRAY)
    else:
        _add_run(sig_cell.paragraphs[0], '_' * 32, size=9.5, color=GRAY)

    # Date cell = first shift date
    _add_run(stbl.cell(0,1).paragraphs[0], 'Date  ', bold=True, size=9.5)
    _add_run(stbl.cell(0,1).paragraphs[0], first_shift_formatted or '_' * 32, size=9.5)

    # Employee Full Name (Print) — use full_name from Section 1
    _add_run(stbl.cell(1,0).paragraphs[0], 'Employee Full Name (Print)  ', bold=True, size=9.5)
    _add_run(stbl.cell(1,0).paragraphs[0], full_name, size=9.5)
    sp(6)

    # ── For Office Use Only ── 1st ────────────────────────────────────
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    rr = p.add_run('  FOR OFFICE USE ONLY  ')
    rr.bold = True; rr.font.size = Pt(12); rr.font.name = 'Arial'
    rr.font.color.rgb = WHITE
    rPr = rr._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1B3A6B')
    rPr.append(shd)
    sp(4)

    # ── Compliance Decision ── 2nd ────────────────────────────────────
    body_text('Compliance Decision:')
    checkbox_item('Acceptable — pending PCC submission', checked=True)
    checkbox_item('Further Information Required')
    checkbox_item('Escalated for Risk Review')
    checkbox_item('Not Accepted')
    sp(6)

    # ── Office Use Table ── 3rd (last before footer) ──────────────────
    # Approval On = Date Reviewed + 1 day
    approval_on = ''
    if date_reviewed:
        try:
            _d = _dt.strptime(date_reviewed, '%d %B %Y')
            approval_on = (_d + _td(days=1)).strftime('%d %B %Y')
        except Exception:
            approval_on = date_reviewed

    otbl = document.add_table(rows=4, cols=2)
    otbl.style = 'Table Grid'
    office_rows = [
        ('Reviewed By:', reviewer),
        ('Date Reviewed:', date_reviewed),
        ('Approved By (Compliance Officer):', _get_compliance_officer()),
        ('Approval On:', approval_on),
    ]
    for ri, (label_, val_) in enumerate(office_rows):
        lc = otbl.cell(ri, 0)
        vc = otbl.cell(ri, 1)
        _set_cell_border(lc, 'CCCCCC')
        _set_cell_border(vc, 'CCCCCC')
        if ri % 2 == 0:
            _set_cell_bg(lc, 'EFF6FF')
        _add_run(lc.paragraphs[0], label_, bold=True, size=9.5)
        _add_run(vc.paragraphs[0], val_, size=9.5)
    sp(6)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, 'Xpress Health — Confidential & Compliance Document  |  Not for distribution',
             italic=True, size=8, color=GRAY)

    buf = _io.BytesIO()
    document.save(buf)
    return buf.getvalue()



# ── PCC Generate ──────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/ai-pcc/generate', methods=['POST'])
@admin_required
def live_staff_ai_pcc_generate():
    """Generate PCC Self-Declaration form for a staff member."""
    data     = request.get_json(silent=True) or {}
    staff_id = (data.get('staff_id') or '').strip()
    if not staff_id:
        return jsonify({"success": False, "error": "Missing staff_id"}), 400

    try:
        doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not doc:
            return jsonify({"success": False, "error": "Staff not found"}), 404

        s1        = doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff')
        emp_code  = _v(doc.get('employee_code') or '')

        # Rotate reviewer based on total docs count
        total = _ai_pcc_col().count_documents({})
        reviewer_index = total % len(_get_pcc_reviewers())

        docx_bytes = _build_pcc_docx(doc, reviewer_index=reviewer_index)
        safe_name  = full_name.replace(' ', '_').replace('/', '_')
        filename   = f"PCC_{safe_name}.docx"
        gcs_blob   = f"pcc/{filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        col      = _ai_pcc_col()
        existing = col.find_one({"staff_id": staff_id})
        rec = {
            "staff_id":      staff_id,
            "staff_name":    full_name,
            "employee_code": emp_code,
            "filename":      filename,
            "gcs_blob":      gcs_blob,
            "reviewer":      _get_pcc_reviewers()[reviewer_index % len(_get_pcc_reviewers())],
            "generated_at":  datetime.utcnow(),
        }
        if existing:
            col.update_one({"_id": existing["_id"]}, {"$set": rec})
            rec_id = str(existing["_id"])
        else:
            rec_id = str(col.insert_one(rec).inserted_id)

        download_url = _gcs_signed_url(gcs_blob) or ''
        return jsonify({
            "success":      True,
            "pcc_id":       rec_id,
            "staff_name":   full_name,
            "filename":     filename,
            "gcs_blob":     gcs_blob,
            "download_url": download_url,
            "reviewer":     _get_pcc_reviewers()[reviewer_index % len(_get_pcc_reviewers())],
            "generated_at": datetime.utcnow().isoformat(),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── PCC Download ──────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/ai-pcc/download/<pcc_id>')
@admin_required
def live_staff_ai_pcc_download(pcc_id):
    """Download saved PCC DOCX from GCS."""
    try:
        rec = _ai_pcc_col().find_one({"_id": ObjectId(pcc_id)})
        if not rec or not rec.get('gcs_blob'):
            return jsonify({"success": False, "error": "PCC not found"}), 404
        docx_bytes = _gcs_download(rec['gcs_blob'])
        return Response(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f'attachment; filename="{rec["filename"]}"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── PCC Saved check ───────────────────────────────────────────────────

@admin_bp.route('/live-staffs/ai-pcc/saved/<staff_id>')
@admin_required
def live_staff_ai_pcc_saved(staff_id):
    """Check if a saved PCC exists for this staff member."""
    rec = _ai_pcc_col().find_one({"staff_id": staff_id})
    if rec:
        return jsonify({
            "success":    True,
            "found":      True,
            "pcc_id":     str(rec["_id"]),
            "filename":   rec.get("filename", ""),
            "reviewer":   rec.get("reviewer", ""),
            "generated_at": rec["generated_at"].strftime("%d %b %Y %H:%M") if rec.get("generated_at") else "",
        })
    return jsonify({"success": True, "found": False})


# ── PCC Upload (replace) ──────────────────────────────────────────────

@admin_bp.route('/live-staffs/ai-pcc/upload/<staff_id>', methods=['POST'])
@admin_required
def live_staff_ai_pcc_upload(staff_id):
    """Upload an edited PCC DOCX to replace saved version in GCS."""
    f = request.files.get('file')
    if not f:
        return jsonify({"success": False, "error": "No file uploaded"}), 400
    try:
        doc = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not doc:
            return jsonify({"success": False, "error": "Staff not found"}), 404
        s1        = doc.get('section_1_personal_details') or {}
        full_name = _v(s1.get('full_name') or 'staff')
        safe_name = full_name.replace(' ', '_').replace('/', '_')
        filename  = f"PCC_{safe_name}.docx"
        gcs_blob  = f"pcc/{filename}"
        docx_bytes = f.read()
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        _ai_pcc_col().update_one(
            {"staff_id": staff_id},
            {"$set": {"gcs_blob": gcs_blob, "filename": filename,
                      "generated_at": datetime.utcnow()}},
            upsert=True
        )
        return jsonify({"success": True, "gcs_blob": gcs_blob, "filename": filename})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract GDPR Certificate ───────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-gdpr', methods=['GET', 'POST'])
def live_staff_cron_sync_gdpr():
    """
    Cron job — processes ONE staff member per call.
    Finds "Gdpr" document, extracts details via Gemini AI.
    Saves: gdpr_certificate_name, gdpr_staff_name, gdpr_expiry_date,
           gdpr_issue_date, gdpr_issuing_body, gdpr_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"gdpr_fetched": {"$exists": False}},
            {"gdpr_fetched": False},
            {"gdpr_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff GDPR certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["gdpr_fetched"]    = True
        fields["gdpr_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"gdpr_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"gdpr_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"gdpr_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"gdpr_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    gdpr_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'gdpr', 'general data protection', 'data protection',
            'gdpr certificate', 'gdpr training',
        )) and d.get('url'):
            gdpr_doc = d
            break

    if not gdpr_doc:
        _mark_done({"gdpr_note": "no GDPR document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No GDPR certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (gdpr_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"gdpr_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — GDPR doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"gdpr_note": "document URL 404 — skipped", "gdpr_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — GDPR doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this GDPR / Data Protection certificate or training record:
1. Certificate name (e.g. "GDPR Awareness", "General Data Protection Regulation Training", "Data Protection Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "gdpr_certificate_name": cert_name,
            "gdpr_staff_name":       cert_staff,
            "gdpr_expiry_date":      expiry_date,
            "gdpr_issue_date":       issue_date,
            "gdpr_issuing_body":     issuing_body,
            "gdpr_doc_url":          doc_url,
            "gdpr_doc_type":         gdpr_doc.get('document_type_name', ''),
            "gdpr_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"GDPR cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"gdpr_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"gdpr_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: GDPR certificates to Excel ───────────────────────────────

@admin_bp.route('/live-staffs/export/gdpr-xlsx')
@admin_required
def live_staff_export_gdpr_xlsx():
    """Export GDPR certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "gdpr_certificate_name": 1, "gdpr_staff_name": 1,
             "gdpr_expiry_date": 1, "gdpr_issue_date": 1,
             "gdpr_issuing_body": 1, "gdpr_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'GDPR Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('gdpr_certificate_name') or '')
            cert_s   = _v(doc.get('gdpr_staff_name') or '')
            expiry   = _v(doc.get('gdpr_expiry_date') or '')
            issue    = _v(doc.get('gdpr_issue_date') or '')
            issuer   = _v(doc.get('gdpr_issuing_body') or '')
            fetched  = doc.get('gdpr_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="gdpr_certificates_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Dignity At Work Certificate ─────────────────────────

@admin_bp.route('/live-staffs/cron/sync-dignity-at-work', methods=['GET', 'POST'])
def live_staff_cron_sync_dignity_at_work():
    """
    Cron job — processes ONE staff member per call.
    Finds "Dignity At Work" document, extracts details via Gemini AI.
    Saves: daw_certificate_name, daw_staff_name, daw_expiry_date,
           daw_issue_date, daw_issuing_body, daw_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"daw_fetched": {"$exists": False}},
            {"daw_fetched": False},
            {"daw_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Dignity At Work certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["daw_fetched"]    = True
        fields["daw_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"daw_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"daw_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"daw_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"daw_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    daw_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'dignity at work',
            'dignity & work',
            'dignity and work',
            'dignity at work certificate',
            'dignity at work training',
        )) and d.get('url'):
            daw_doc = d
            break

    if not daw_doc:
        _mark_done({"daw_note": "no Dignity At Work document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Dignity At Work certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (daw_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"daw_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Dignity At Work doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"daw_note": "document URL 404 — skipped", "daw_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Dignity At Work doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Dignity At Work certificate or training record:
1. Certificate name (e.g. "Dignity At Work", "Dignity & Respect at Work", "Dignity At Work Training Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "daw_certificate_name": cert_name,
            "daw_staff_name":       cert_staff,
            "daw_expiry_date":      expiry_date,
            "daw_issue_date":       issue_date,
            "daw_issuing_body":     issuing_body,
            "daw_doc_url":          doc_url,
            "daw_doc_type":         daw_doc.get('document_type_name', ''),
            "daw_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Dignity At Work cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"daw_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"daw_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Dignity At Work certificates to Excel ─────────────────────

@admin_bp.route('/live-staffs/export/dignity-at-work-xlsx')
@admin_required
def live_staff_export_dignity_at_work_xlsx():
    """Export Dignity At Work certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "daw_certificate_name": 1, "daw_staff_name": 1,
             "daw_expiry_date": 1, "daw_issue_date": 1,
             "daw_issuing_body": 1, "daw_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Dignity At Work Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('daw_certificate_name') or '')
            cert_s   = _v(doc.get('daw_staff_name') or '')
            expiry   = _v(doc.get('daw_expiry_date') or '')
            issue    = _v(doc.get('daw_issue_date') or '')
            issuer   = _v(doc.get('daw_issuing_body') or '')
            fetched  = doc.get('daw_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="dignity_at_work_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract HACCP/Food Safety Certificate ───────────────────────

@admin_bp.route('/live-staffs/cron/sync-haccp', methods=['GET', 'POST'])
def live_staff_cron_sync_haccp():
    """
    Cron job — processes ONE staff member per call.
    Finds "Haccp/Food Safety" document, extracts details via Gemini AI.
    Saves: haccp_certificate_name, haccp_staff_name, haccp_expiry_date,
           haccp_issue_date, haccp_issuing_body, haccp_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"haccp_fetched": {"$exists": False}},
            {"haccp_fetched": False},
            {"haccp_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff HACCP/Food Safety certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["haccp_fetched"]    = True
        fields["haccp_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"haccp_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"haccp_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"haccp_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"haccp_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    haccp_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'haccp/food safety', 'haccp', 'food safety',
            'food hygiene', 'food safety certificate',
            'haccp certificate', 'food safety training',
        )) and d.get('url'):
            haccp_doc = d
            break

    if not haccp_doc:
        _mark_done({"haccp_note": "no HACCP/Food Safety document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No HACCP/Food Safety certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (haccp_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"haccp_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — HACCP doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"haccp_note": "document URL 404 — skipped", "haccp_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — HACCP doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this HACCP / Food Safety certificate or training record:
1. Certificate name (e.g. "HACCP", "Food Safety Certificate", "Food Hygiene & Safety", "HACCP/Food Safety Training")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "haccp_certificate_name": cert_name,
            "haccp_staff_name":       cert_staff,
            "haccp_expiry_date":      expiry_date,
            "haccp_issue_date":       issue_date,
            "haccp_issuing_body":     issuing_body,
            "haccp_doc_url":          doc_url,
            "haccp_doc_type":         haccp_doc.get('document_type_name', ''),
            "haccp_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"HACCP/Food Safety cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"haccp_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"haccp_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: HACCP/Food Safety certificates to Excel ───────────────────

@admin_bp.route('/live-staffs/export/haccp-xlsx')
@admin_required
def live_staff_export_haccp_xlsx():
    """Export HACCP/Food Safety certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "haccp_certificate_name": 1, "haccp_staff_name": 1,
             "haccp_expiry_date": 1, "haccp_issue_date": 1,
             "haccp_issuing_body": 1, "haccp_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'HACCP-Food Safety Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('haccp_certificate_name') or '')
            cert_s   = _v(doc.get('haccp_staff_name') or '')
            expiry   = _v(doc.get('haccp_expiry_date') or '')
            issue    = _v(doc.get('haccp_issue_date') or '')
            issuer   = _v(doc.get('haccp_issuing_body') or '')
            fetched  = doc.get('haccp_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="haccp_food_safety_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Cyber Security Certificate ─────────────────────────

@admin_bp.route('/live-staffs/cron/sync-cyber-security', methods=['GET', 'POST'])
def live_staff_cron_sync_cyber_security():
    """
    Cron job — processes ONE staff member per call.
    Finds "Cyber Security" document, extracts details via Gemini AI.
    Saves: cs_certificate_name, cs_staff_name, cs_expiry_date,
           cs_issue_date, cs_issuing_body, cs_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"cs_fetched": {"$exists": False}},
            {"cs_fetched": False},
            {"cs_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Cyber Security certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["cs_fetched"]    = True
        fields["cs_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"cs_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"cs_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"cs_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"cs_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    cs_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'cyber security', 'cybersecurity',
            'cyber security certificate', 'cyber security training',
            'cyber awareness', 'information security',
        )) and d.get('url'):
            cs_doc = d
            break

    if not cs_doc:
        _mark_done({"cs_note": "no Cyber Security document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Cyber Security certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (cs_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"cs_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Cyber Security doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"cs_note": "document URL 404 — skipped", "cs_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Cyber Security doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Cyber Security certificate or training record:
1. Certificate name (e.g. "Cyber Security Awareness", "Cybersecurity Training Certificate", "Information Security Certificate")
2. Staff name as printed on the certificate
3. Expiry date or renewal date (if shown)
4. Issue / completion date
5. Issuing body or training provider

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<name as printed on certificate>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue or completion date as printed>",
  "issuing_body": "<organization that issued the certificate>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "cs_certificate_name": cert_name,
            "cs_staff_name":       cert_staff,
            "cs_expiry_date":      expiry_date,
            "cs_issue_date":       issue_date,
            "cs_issuing_body":     issuing_body,
            "cs_doc_url":          doc_url,
            "cs_doc_type":         cs_doc.get('document_type_name', ''),
            "cs_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Cyber Security cert extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"cs_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"cs_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Cyber Security certificates to Excel ──────────────────────

@admin_bp.route('/live-staffs/export/cyber-security-xlsx')
@admin_required
def live_staff_export_cyber_security_xlsx():
    """Export Cyber Security certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "cs_certificate_name": 1, "cs_staff_name": 1,
             "cs_expiry_date": 1, "cs_issue_date": 1,
             "cs_issuing_body": 1, "cs_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Cyber Security Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('cs_certificate_name') or '')
            cert_s   = _v(doc.get('cs_staff_name') or '')
            expiry   = _v(doc.get('cs_expiry_date') or '')
            issue    = _v(doc.get('cs_issue_date') or '')
            issuer   = _v(doc.get('cs_issuing_body') or '')
            fetched  = doc.get('cs_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="cyber_security_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Health Declaration Form ────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-health-declaration', methods=['GET', 'POST'])
def live_staff_cron_sync_health_declaration():
    """
    Cron job — processes ONE staff member per call.
    Finds "Health Declaration Form" document, extracts details via Gemini AI.
    Saves: hdf_certificate_name, hdf_staff_name, hdf_signed_date,
           hdf_issuing_body, hdf_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"hdf_fetched": {"$exists": False}},
            {"hdf_fetched": False},
            {"hdf_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Health Declaration Forms already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["hdf_fetched"]    = True
        fields["hdf_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"hdf_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"hdf_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"hdf_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"hdf_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    hdf_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'health declaration form',
            'health declaration',
            'medical declaration',
            'occupational health declaration',
            'health questionnaire',
        )) and d.get('url'):
            hdf_doc = d
            break

    if not hdf_doc:
        _mark_done({"hdf_note": "no Health Declaration Form found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Health Declaration Form found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (hdf_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"hdf_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Health Declaration doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"hdf_note": "document URL 404 — skipped", "hdf_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Health Declaration doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this Health Declaration Form:
1. Document / form name (e.g. "Health Declaration Form", "Occupational Health Declaration", "Medical Declaration")
2. Staff / employee name as printed on the form
3. Date the form was signed or completed
4. Issuing body or organisation (if shown)

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact form title as printed>",
  "staff_name_on_cert": "<name as printed on form>",
  "signed_date": "<date the form was signed or completed>",
  "issuing_body": "<organization or company name if shown>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        signed_date  = _v(result.get('signed_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "hdf_certificate_name": cert_name,
            "hdf_staff_name":       cert_staff,
            "hdf_signed_date":      signed_date,
            "hdf_issuing_body":     issuing_body,
            "hdf_doc_url":          doc_url,
            "hdf_doc_type":         hdf_doc.get('document_type_name', ''),
            "hdf_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "signed_date":        signed_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Health Declaration Form extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"hdf_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"hdf_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Health Declaration Forms to Excel ─────────────────────────

@admin_bp.route('/live-staffs/export/health-declaration-xlsx')
@admin_required
def live_staff_export_health_declaration_xlsx():
    """Export Health Declaration Form details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "hdf_certificate_name": 1, "hdf_staff_name": 1,
             "hdf_signed_date": 1, "hdf_issuing_body": 1, "hdf_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Health Declaration Forms'

        headers    = ['Sno', 'Staff Name', 'Email', 'Form Name',
                      'Name on Form', 'Signed Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:H{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('hdf_certificate_name') or '')
            cert_s   = _v(doc.get('hdf_staff_name') or '')
            signed   = _v(doc.get('hdf_signed_date') or '')
            issuer   = _v(doc.get('hdf_issuing_body') or '')
            fetched  = doc.get('hdf_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Form Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, signed, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="health_declaration_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Generate PCC Self-Declaration form one staff at a time ──────

@admin_bp.route('/live-staffs/cron/generate-pcc', methods=['GET', 'POST'])
def live_staff_cron_generate_pcc():
    """
    Cron job — generates PCC Self-Declaration form for ONE staff member per call.

    Finds staff where pcc_generated is not True.
    Generates DOCX, uploads to GCS, saves to live_staff_ai_pcc collection.
    Reviewer rotates through _PCC_REVIEWERS alternately.

    Protect with ?cron_key=<CRON_SECRET> env var.
    """
    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    col     = _staffs_col()
    pcc_col = _ai_pcc_col()

    pending_query = {
        "$or": [
            {"pcc_generated": {"$exists": False}},
            {"pcc_generated": False},
            {"pcc_generated": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff PCC forms already generated.",
            "remaining_count": 0,
        })

    staff_id  = str(staff['_id'])
    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')
    emp_code  = _v(staff.get('employee_code') or '')

    def _mark_done(fields):
        fields["pcc_generated"]    = True
        fields["pcc_generated_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    try:
        # Rotate reviewer based on total generated so far
        total_generated    = pcc_col.count_documents({})
        reviewer_index     = total_generated % len(_get_pcc_reviewers())

        docx_bytes = _build_pcc_docx(staff, reviewer_index=reviewer_index)

        safe_name  = (full_name or 'staff').replace(' ', '_').replace('/', '_')
        filename   = f"PCC_{safe_name}.docx"
        gcs_blob   = f"pcc/{filename}"
        _gcs_upload(gcs_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        existing = pcc_col.find_one({"staff_id": staff_id})
        rec = {
            "staff_id":      staff_id,
            "staff_name":    full_name,
            "employee_code": emp_code,
            "filename":      filename,
            "gcs_blob":      gcs_blob,
            "reviewer":      _get_pcc_reviewers()[reviewer_index % len(_get_pcc_reviewers())],
            "generated_at":  datetime.utcnow(),
        }
        if existing:
            pcc_col.update_one({"_id": existing["_id"]}, {"$set": rec})
            rec_id = str(existing["_id"])
        else:
            rec_id = str(pcc_col.insert_one(rec).inserted_id)

        download_url = _gcs_signed_url(gcs_blob) or ''
        _mark_done({
            "pcc_gcs_blob":    gcs_blob,
            "pcc_filename":    filename,
            "pcc_download_url": download_url,
        })

        return jsonify({
            "success":         True,
            "staff_name":      full_name,
            "email":           email,
            "pcc_id":          rec_id,
            "filename":        filename,
            "gcs_blob":        gcs_blob,
            "download_url":    download_url,
            "reviewer":        _get_pcc_reviewers()[reviewer_index % len(_get_pcc_reviewers())],
            "remaining_count": max(0, remaining_total - 1),
            "message": (
                f"PCC generated for {full_name} — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except Exception as e:
        _mark_done({"pcc_note": f"error: {e}"})
        return jsonify({
            "success":         False,
            "email":           email,
            "staff_name":      full_name,
            "error":           str(e),
            "remaining_count": max(0, remaining_total - 1),
        })



# ── Export: Profile Check (matching Profile_check.xlsx format) ────────

@admin_bp.route('/live-staffs/export/profile-check-xlsx')
@admin_required
def live_staff_export_profile_check_xlsx():
    """
    Export staff certificate profile check to Excel.
    Sheets: Passport, CPR, IPC, Handhygiene, Safeguard, Garda, Children
    Matches the structure of Profile_check.xlsx.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import (Font, PatternFill, Alignment,
                                     Border, Side, numbers)
        import io as _io
        from datetime import date as _date
        today = _date.today()

        NAVY   = '1B3A6B'; GREEN  = '2E9E44'; WHITE  = 'FFFFFF'
        ALT    = 'EFF6FF'; RED    = 'FFDDDD'; WARN   = 'FFF3CD'
        MISMATCH_RED = 'FFE0E0'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        b_font  = Font(name='Arial', size=9)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        def _name_match(staff_name, cert_name):
            """Check if staff name matches cert name (simple fuzzy)."""
            if not staff_name or not cert_name:
                return True
            s = staff_name.strip().lower()
            c = cert_name.strip().lower()
            if s == c:
                return True
            # Check if all parts of cert name appear in staff name or vice versa
            s_parts = set(s.split())
            c_parts = set(c.split())
            return len(s_parts & c_parts) >= min(2, len(c_parts))

        def _fmt_date(val):
            if not val:
                return ''
            if hasattr(val, 'strftime'):
                return val.strftime('%d/%m/%Y')
            return str(val)

        def _is_expired(expiry):
            if not expiry:
                return None
            if hasattr(expiry, 'date'):
                return expiry.date() < today
            if hasattr(expiry, 'year'):
                return expiry < today
            # string
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y',
                        '%d %B, %Y','%d %b, %Y'):
                try:
                    from datetime import datetime as _dt2
                    d = _dt2.strptime(str(expiry).strip().rstrip(','), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        def _write_sheet(ws, headers, rows, col_widths):
            for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
                cell = ws.cell(row=1, column=ci, value=hdr)
                cell.font      = h_font
                cell.fill      = h_fill
                cell.alignment = h_align
                cell.border    = green_b
                ws.column_dimensions[
                    cell.column_letter].width = width
            ws.row_dimensions[1].height = 28
            ws.freeze_panes = 'A2'
            ws.auto_filter.ref = (
                f'A1:{ws.cell(row=1, column=len(headers)).column_letter}1'
            )
            for ri, row_data in enumerate(rows, start=2):
                is_mismatch = row_data.get('_mismatch', False)
                is_expired  = row_data.get('_expired', False)
                row_fill = (PatternFill('solid', start_color=RED,  end_color=RED)
                            if is_expired else
                            PatternFill('solid', start_color=WARN, end_color=WARN)
                            if is_mismatch else
                            PatternFill('solid', start_color=ALT,  end_color=ALT)
                            if ri % 2 == 0 else None)
                for ci, key in enumerate(row_data.get('_keys', []), start=1):
                    val  = row_data.get(key, '')
                    cell = ws.cell(row=ri, column=ci, value=val)
                    cell.font      = b_font
                    cell.alignment = c_align if ci == 1 else l_align
                    cell.border    = border
                    if row_fill:
                        cell.fill = row_fill
                ws.row_dimensions[ri].height = 16

        docs = list(_staffs_col().find({}, {
            'section_1_personal_details': 1, 'email': 1,
            'passport_id': 1, 'passport_data': 1, 'passport_fetched': 1,
            'cpr_certificate_name':1,'cpr_staff_name':1,'cpr_expiry_date':1,
            'cpr_issue_date':1,'cpr_issuing_body':1,'cpr_fetched':1,
            'ipc_certificate_name':1,'ipc_staff_name':1,'ipc_expiry_date':1,
            'ipc_issue_date':1,'ipc_issuing_body':1,'ipc_fetched':1,
            'hh_certificate_name':1,'hh_staff_name':1,'hh_expiry_date':1,
            'hh_issue_date':1,'hh_issuing_body':1,'hh_fetched':1,
            'sg_certificate_name':1,'sg_staff_name':1,'sg_expiry_date':1,
            'sg_issue_date':1,'sg_issuing_body':1,'sg_fetched':1,
            'garda_cert_name':1,'garda_staff_name':1,'garda_issue_date':1,
            'garda_reference':1,'garda_issuing_body':1,'garda_fetched':1,
            'cf_certificate_name':1,'cf_staff_name':1,'cf_expiry_date':1,
            'cf_issue_date':1,'cf_issuing_body':1,'cf_fetched':1,
        }))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        wb = Workbook()
        wb.remove(wb.active)  # remove default sheet

        # ── Helper to build a row dict ────────────────────────────────
        def _row(sno, name, email, cert_n, cert_s, expiry_raw, issue_raw,
                 issuer, status, extra=None):
            mismatch     = not _name_match(name, cert_s)
            mismatch_lbl = 'Mismatch' if mismatch else 'Match'
            expired      = _is_expired(expiry_raw) is True
            d = {
                'sno':      sno,
                'name':     name,
                'email':    email,
                'cert_n':   cert_n,
                'cert_s':   cert_s,
                'expiry':   _fmt_date(expiry_raw),
                'issue':    _fmt_date(issue_raw),
                'issuer':   issuer,
                'status':   status if status else ('Found' if cert_n else 'No Cert Found'),
                'mismatch': mismatch_lbl,
                '_mismatch': mismatch,
                '_expired':  expired,
            }
            if extra:
                d.update(extra)
            return d

        # ── Passport ──────────────────────────────────────────────────
        ws_p = wb.create_sheet('Passport')
        hdrs_p = ['Sno','Staff Name','Email','Passport ID','Nationality',
                  'Date of Birth','Expiry Date','Country','Status','Name Match']
        wids_p = [5,28,36,16,18,14,14,16,14,12]
        rows_p = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            pid   = _v(doc.get('passport_id') or '')
            pd_   = doc.get('passport_data') or {}
            nat   = _v(pd_.get('nationality') or '')
            dob   = _v(pd_.get('date_of_birth') or '')
            exp   = _v(pd_.get('expiry_date') or '')
            ctry  = _v(pd_.get('country') or '')
            st    = 'Found' if pid else ('Not Checked' if not doc.get('passport_fetched') else 'No ID Found')
            rows_p.append({
                '_keys': ['sno','name','email','pid','nat','dob','exp','ctry','st','match'],
                'sno': i, 'name': name, 'email': email, 'pid': pid,
                'nat': nat, 'dob': dob, 'exp': exp, 'ctry': ctry, 'st': st,
                'match': 'Match', '_mismatch': False, '_expired': False,
            })
        _write_sheet(ws_p, hdrs_p, rows_p, wids_p)

        # ── CPR ───────────────────────────────────────────────────────
        ws_c = wb.create_sheet('CPR')
        hdrs_c = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                  'Expiry Date','Issue Date','Issuing Body','Status','Staff Name Mismatch']
        wids_c = [5,28,36,32,28,14,14,28,14,14]
        rows_c = []
        for i, doc in enumerate(docs, start=1):
            s1   = doc.get('section_1_personal_details') or {}
            name = _v(s1.get('full_name') or '')
            email= _v(doc.get('email') or '')
            r    = _row(i,name,email,
                        _v(doc.get('cpr_certificate_name') or ''),
                        _v(doc.get('cpr_staff_name') or ''),
                        doc.get('cpr_expiry_date'),
                        doc.get('cpr_issue_date'),
                        _v(doc.get('cpr_issuing_body') or ''),
                        None if doc.get('cpr_fetched') else 'Not Checked')
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch']
            rows_c.append(r)
        _write_sheet(ws_c, hdrs_c, rows_c, wids_c)

        # ── IPC ───────────────────────────────────────────────────────
        ws_i = wb.create_sheet('IPC')
        hdrs_i = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                  'Expiry Date','Issue Date','Issuing Body','Status',
                  'Staff Name Mismatch','Reason for Mismatch']
        wids_i = [5,28,36,35,28,14,14,28,14,14,28]
        rows_i = []
        for i, doc in enumerate(docs, start=1):
            s1   = doc.get('section_1_personal_details') or {}
            name = _v(s1.get('full_name') or '')
            email= _v(doc.get('email') or '')
            cert_s = _v(doc.get('ipc_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            reason = (f'mismatch - {cert_s}' if mismatch and cert_s else '')
            r = _row(i,name,email,
                     _v(doc.get('ipc_certificate_name') or ''),
                     cert_s,
                     doc.get('ipc_expiry_date'),
                     doc.get('ipc_issue_date'),
                     _v(doc.get('ipc_issuing_body') or ''),
                     None if doc.get('ipc_fetched') else 'Not Checked')
            r['reason'] = reason
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch','reason']
            rows_i.append(r)
        _write_sheet(ws_i, hdrs_i, rows_i, wids_i)

        # ── Hand Hygiene ──────────────────────────────────────────────
        ws_h = wb.create_sheet('Handhygiene')
        hdrs_h = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                  'Expiry Date','Issue Date','Issuing Body','Status',
                  'Staff Name Mismatch','Reason for Mismatch']
        wids_h = [5,28,36,32,28,14,14,28,14,14,28]
        rows_h = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            cert_s= _v(doc.get('hh_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            reason = (f'mismatch - {cert_s}' if mismatch and cert_s else '')
            r = _row(i,name,email,
                     _v(doc.get('hh_certificate_name') or ''),
                     cert_s,
                     doc.get('hh_expiry_date'),
                     doc.get('hh_issue_date'),
                     _v(doc.get('hh_issuing_body') or ''),
                     None if doc.get('hh_fetched') else 'Not Checked')
            r['reason'] = reason
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch','reason']
            rows_h.append(r)
        _write_sheet(ws_h, hdrs_h, rows_h, wids_h)

        # ── Safeguard ─────────────────────────────────────────────────
        ws_s = wb.create_sheet('Safeguard')
        hdrs_s = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                  'Expiry Date','Issue Date','Issuing Body','Status',
                  'Staff Name Mismatch','Reason for Mismatch']
        wids_s = [5,28,36,38,28,14,14,28,14,14,28]
        rows_sg = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            cert_s= _v(doc.get('sg_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            reason = (f'mismatch - {cert_s}' if mismatch and cert_s else '')
            r = _row(i,name,email,
                     _v(doc.get('sg_certificate_name') or ''),
                     cert_s,
                     doc.get('sg_expiry_date'),
                     doc.get('sg_issue_date'),
                     _v(doc.get('sg_issuing_body') or ''),
                     None if doc.get('sg_fetched') else 'Not Checked')
            r['reason'] = reason
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch','reason']
            rows_sg.append(r)
        _write_sheet(ws_s, hdrs_s, rows_sg, wids_s)

        # ── Garda ─────────────────────────────────────────────────────
        ws_g = wb.create_sheet('Garda')
        hdrs_g = ['Sno','Staff Name','Email','Document Name','Name on Doc',
                  'Issue Date','Reference No','Issuing Body','Status',
                  'Staff Name Mismatch','Reason for Mismatch']
        wids_g = [5,28,36,32,28,14,24,28,14,14,28]
        rows_g = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            cert_s= _v(doc.get('garda_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            reason = (f'mismatch - {cert_s}' if mismatch and cert_s else '')
            cert_n= _v(doc.get('garda_cert_name') or '')
            r = {
                '_keys': ['sno','name','email','cert_n','cert_s',
                          'issue','ref_n','issuer','status','mismatch','reason'],
                'sno': i, 'name': name, 'email': email,
                'cert_n': cert_n, 'cert_s': cert_s,
                'issue': _fmt_date(doc.get('garda_issue_date')),
                'ref_n': _v(doc.get('garda_reference') or ''),
                'issuer': _v(doc.get('garda_issuing_body') or ''),
                'status': 'Found' if cert_n else ('Not Checked' if not doc.get('garda_fetched') else 'No Doc Found'),
                'mismatch': 'Mismatch' if mismatch else 'Match',
                'reason': reason,
                '_mismatch': mismatch, '_expired': False,
            }
            rows_g.append(r)
        _write_sheet(ws_g, hdrs_g, rows_g, wids_g)

        # ── Children First ────────────────────────────────────────────
        ws_ch = wb.create_sheet('Children')
        hdrs_ch = ['Sno','Staff Name','Email','Certificate Name','Name on Cert',
                   'Expiry Date','Issue Date','Issuing Body','Status',
                   'Staff Name Mismatch','Expiry Date Exceeded']
        wids_ch = [5,28,36,38,28,14,14,28,14,14,16]
        rows_ch = []
        for i, doc in enumerate(docs, start=1):
            s1    = doc.get('section_1_personal_details') or {}
            name  = _v(s1.get('full_name') or '')
            email = _v(doc.get('email') or '')
            cert_s= _v(doc.get('cf_staff_name') or '')
            mismatch = not _name_match(name, cert_s)
            exp_raw  = doc.get('cf_expiry_date')
            expired  = _is_expired(exp_raw) is True
            r = _row(i,name,email,
                     _v(doc.get('cf_certificate_name') or ''),
                     cert_s,
                     exp_raw,
                     doc.get('cf_issue_date'),
                     _v(doc.get('cf_issuing_body') or ''),
                     None if doc.get('cf_fetched') else 'Not Checked')
            r['expired_lbl'] = 'Expired' if expired else ''
            r['_keys'] = ['sno','name','email','cert_n','cert_s',
                          'expiry','issue','issuer','status','mismatch','expired_lbl']
            rows_ch.append(r)
        _write_sheet(ws_ch, hdrs_ch, rows_ch, wids_ch)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="Profile_check_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Code Of Conduct Certificate ─────────────────────────

@admin_bp.route('/live-staffs/cron/sync-code-of-conduct', methods=['GET', 'POST'])
def live_staff_cron_sync_code_of_conduct():
    """
    Cron job — processes ONE staff member per call.
    Finds "Code Of Conduct" document, extracts details via Gemini AI.
    Saves: coc_certificate_name, coc_staff_name, coc_expiry_date,
           coc_issue_date, coc_issuing_body, coc_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"coc_fetched": {"$exists": False}},
            {"coc_fetched": False},
            {"coc_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Code Of Conduct certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["coc_fetched"]    = True
        fields["coc_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"coc_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"coc_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"coc_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"coc_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    coc_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'code of conduct', 'codeofconduct',
            'code of conduct certificate', 'code of conduct training',
            'staff code of conduct',
        )) and d.get('url'):
            coc_doc = d
            break

    if not coc_doc:
        _mark_done({"coc_note": "no Code Of Conduct document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Code Of Conduct certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (coc_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"coc_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Code Of Conduct doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"coc_note": "document URL 404 — skipped", "coc_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Code Of Conduct doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Code Of Conduct certificate or signed document:
1. Certificate / document name (e.g. "Code Of Conduct", "Staff Code of Conduct", "Code of Conduct Agreement")
2. Staff name as printed on the document
3. Expiry date or renewal date (if shown)
4. Issue / completion / signing date
5. Issuing body or organisation

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact document title as printed>",
  "staff_name_on_cert": "<name as printed on document>",
  "expiry_date": "<expiry or renewal date as printed, e.g. 01/06/2027 or June 2027>",
  "issue_date": "<issue, signing or completion date as printed>",
  "issuing_body": "<organization that issued the document>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "coc_certificate_name": cert_name,
            "coc_staff_name":       cert_staff,
            "coc_expiry_date":      expiry_date,
            "coc_issue_date":       issue_date,
            "coc_issuing_body":     issuing_body,
            "coc_doc_url":          doc_url,
            "coc_doc_type":         coc_doc.get('document_type_name', ''),
            "coc_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Code Of Conduct extracted for {full_name} "
                f"(expires: {expiry_date or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"coc_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"coc_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Code Of Conduct to Excel ─────────────────────────────────

@admin_bp.route('/live-staffs/export/code-of-conduct-xlsx')
@admin_required
def live_staff_export_code_of_conduct_xlsx():
    """Export Code Of Conduct certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "coc_certificate_name": 1, "coc_staff_name": 1,
             "coc_expiry_date": 1, "coc_issue_date": 1,
             "coc_issuing_body": 1, "coc_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Code Of Conduct'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Expiry Date', 'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 16, 16, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:I{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('coc_certificate_name') or '')
            cert_s   = _v(doc.get('coc_staff_name') or '')
            expiry   = _v(doc.get('coc_expiry_date') or '')
            issue    = _v(doc.get('coc_issue_date') or '')
            issuer   = _v(doc.get('coc_issuing_body') or '')
            fetched  = doc.get('coc_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="code_of_conduct_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract References Document ────────────────────────────────

@admin_bp.route('/live-staffs/cron/sync-references', methods=['GET', 'POST'])
def live_staff_cron_sync_references():
    """
    Cron job — processes ONE staff member per call.
    Finds "References" document, extracts details via Gemini AI.
    Saves: ref_document_name, ref_staff_name, ref_signed_date,
           ref_issuing_body, ref_fetched = True
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"ref_fetched": {"$exists": False}},
            {"ref_fetched": False},
            {"ref_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff References already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["ref_fetched"]    = True
        fields["ref_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"ref_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"ref_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"ref_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"ref_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    ref_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'references', 'reference letter', 'reference document',
            'employment reference', 'character reference', 'professional reference',
        )) and d.get('url'):
            ref_doc = d
            break

    if not ref_doc:
        _mark_done({"ref_note": "no References document found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No References document found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (ref_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"ref_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — References doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"ref_note": "document URL 404 — skipped", "ref_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — References doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a document data extractor.

Extract the following details from this References document or reference letter:
1. Document name (e.g. "Reference Letter", "Employment Reference", "Character Reference", "References")
2. Name of the person the reference is for (the staff member / applicant)
3. Date the document was signed or issued
4. Name of the referee or issuing organisation (the person or company providing the reference)

Return ONLY a JSON object — no markdown, no explanation:
{
  "document_name": "<exact document title or type>",
  "staff_name_on_doc": "<name of the applicant / staff member the reference is for>",
  "signed_date": "<date the reference was signed or issued>",
  "referee_name": "<name of the person or organisation providing the reference>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nDOCUMENT TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        doc_name     = _v(result.get('document_name') or '')
        staff_on_doc = _v(result.get('staff_name_on_doc') or '')
        signed_date  = _v(result.get('signed_date') or '')
        referee_name = _v(result.get('referee_name') or '')

        _mark_done({
            "ref_document_name": doc_name,
            "ref_staff_name":    staff_on_doc,
            "ref_signed_date":   signed_date,
            "ref_referee_name":  referee_name,
            "ref_doc_url":       doc_url,
            "ref_doc_type":      ref_doc.get('document_type_name', ''),
            "ref_note":          "extracted successfully",
        })

        return jsonify({
            "success":          True,
            "email":            email,
            "staff_name":       full_name,
            "doc_found":        True,
            "document_name":    doc_name,
            "staff_name_on_doc": staff_on_doc,
            "signed_date":      signed_date,
            "referee_name":     referee_name,
            "remaining_count":  max(0, remaining_total - 1),
            "message": (
                f"References extracted for {full_name} ({email}) "
                f"— {max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"ref_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"ref_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: References to Excel ───────────────────────────────────────

@admin_bp.route('/live-staffs/export/references-xlsx')
@admin_required
def live_staff_export_references_xlsx():
    """Export References document details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "ref_document_name": 1, "ref_staff_name": 1,
             "ref_signed_date": 1, "ref_referee_name": 1, "ref_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'References'

        headers    = ['Sno', 'Staff Name', 'Email', 'Document Name',
                      'Name on Doc', 'Signed Date', 'Referee Name', 'Status']
        col_widths = [5, 28, 36, 30, 28, 16, 30, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:H{len(docs)+1}'

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            doc_n    = _v(doc.get('ref_document_name') or '')
            doc_s    = _v(doc.get('ref_staff_name') or '')
            signed   = _v(doc.get('ref_signed_date') or '')
            referee  = _v(doc.get('ref_referee_name') or '')
            fetched  = doc.get('ref_fetched', False)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not doc_n:
                status   = 'No Doc Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, doc_n, doc_s, signed, referee, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="references_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Police Clearance Certificate ────────────────────────

@admin_bp.route('/live-staffs/cron/sync-police-clearance', methods=['GET', 'POST'])
def live_staff_cron_sync_police_clearance():
    """
    Cron job — processes ONE staff member per call.
    Finds "Police Clearance Certificate ( From Country Of Birth )" document.
    Saves: pcc2_certificate_name, pcc2_staff_name, pcc2_expiry_date,
           pcc2_issue_date, pcc2_issuing_body, pcc2_country, pcc2_fetched = True
    Note: pcc2_ prefix used to avoid conflict with PCC Self-Declaration (pcc_).
    """
    import requests as _req
    from google import genai as google_genai

    cron_secret = os.environ.get('CRON_SECRET', '')
    if cron_secret:
        provided = (request.args.get('cron_key') or
                    request.headers.get('X-Cron-Key', ''))
        if provided != cron_secret:
            return jsonify({"success": False, "error": "Unauthorised"}), 401

    base_url    = os.environ.get('LIVE_STAFF_URL', '').rstrip('/')
    api_key     = os.environ.get('XN_PORTAL_API_KEY', '')
    app_country = os.environ.get('XN_APP_COUNTRY', '')
    gemini_key  = os.environ.get('GEMINI_API_KEY', '')

    if not base_url:
        return jsonify({"success": False, "error": "LIVE_STAFF_URL not set"}), 500
    if not gemini_key:
        return jsonify({"success": False, "error": "GEMINI_API_KEY not set"}), 500

    col = _staffs_col()

    pending_query = {
        "$or": [
            {"pcc2_fetched": {"$exists": False}},
            {"pcc2_fetched": False},
            {"pcc2_fetched": None},
        ]
    }
    remaining_total = col.count_documents(pending_query)
    staff           = col.find_one(pending_query)

    if not staff:
        return jsonify({
            "success":         True,
            "message":         "All staff Police Clearance Certificates already extracted.",
            "remaining_count": 0,
        })

    s1        = staff.get('section_1_personal_details') or {}
    full_name = _v(s1.get('full_name') or '')
    email     = _v(staff.get('email') or s1.get('email_address') or '')

    def _mark_done(fields):
        fields["pcc2_fetched"]    = True
        fields["pcc2_fetched_at"] = datetime.utcnow()
        col.update_one({"_id": staff['_id']}, {"$set": fields})

    if not email:
        _mark_done({"pcc2_note": "skipped — no email"})
        return jsonify({
            "success":         True,
            "message":         "Skipped — no email",
            "remaining_count": max(0, remaining_total - 1),
        })

    endpoint    = f"{base_url}/ai/recruitments/user-document-list"
    api_headers = {
        "Api-Key":       api_key,
        "X-App-Country": app_country,
        "Content-Type":  "application/json",
        "Accept":        "application/json",
    }

    try:
        resp = _req.post(endpoint, json={"email": email},
                         headers=api_headers, timeout=30)
        if resp.status_code == 405:
            resp = _req.get(endpoint, params={"email": email},
                            headers=api_headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _mark_done({"pcc2_note": f"API error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": f"API error: {e}",
            "remaining_count": max(0, remaining_total - 1),
        })

    if not data.get('success'):
        _mark_done({"pcc2_note": f"API error: {data.get('message')}"})
        return jsonify({
            "success": False, "email": email,
            "error": data.get('message', 'API error'),
            "remaining_count": max(0, remaining_total - 1),
        })

    api_data  = data.get('data')
    documents = api_data if isinstance(api_data, list) else                 (api_data.get('documents') or [] if isinstance(api_data, dict) else [])

    if not documents:
        _mark_done({"pcc2_note": "no documents returned"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No documents returned for {email}",
            "remaining_count": max(0, remaining_total - 1),
        })

    pcc2_doc = None
    for d in documents:
        doc_name = (d.get('document_type_name') or '').strip().lower()
        if any(t in doc_name for t in (
            'police clearance certificate',
            'police clearance',
            'police certificate',
            'garda clearance',
            'criminal record certificate',
            'certificate of good conduct',
            'country of birth',
        )) and d.get('url'):
            pcc2_doc = d
            break

    if not pcc2_doc:
        _mark_done({"pcc2_note": "no Police Clearance Certificate found"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": False,
            "message": f"No Police Clearance Certificate found for {full_name}",
            "remaining_count": max(0, remaining_total - 1),
        })

    doc_url = (pcc2_doc.get('url') or '').strip()

    if not doc_url:
        _mark_done({"pcc2_note": "document found but URL is empty — skipped"})
        return jsonify({
            "success": True, "email": email, "staff_name": full_name,
            "doc_found": True, "skipped": True,
            "reason": "Document URL is empty",
            "remaining_count": max(0, remaining_total - 1),
            "message": f"Skipped {full_name} ({email}) — Police Clearance doc has no URL",
        })

    try:
        dl_headers = {k: v for k, v in api_headers.items() if k != 'Content-Type'}
        dl_resp    = _req.get(doc_url, headers=dl_headers, timeout=60)

        if dl_resp.status_code == 404:
            _mark_done({"pcc2_note": "document URL 404 — skipped", "pcc2_doc_404": True})
            return jsonify({
                "success": True, "email": email, "staff_name": full_name,
                "doc_found": True, "skipped": True,
                "reason": "Document URL returned 404",
                "remaining_count": max(0, remaining_total - 1),
                "message": f"Skipped {full_name} ({email}) — Police Clearance doc URL 404",
            })

        dl_resp.raise_for_status()
        raw_bytes    = dl_resp.content
        content_type = dl_resp.headers.get('Content-Type', '').lower()

        client = google_genai.Client(api_key=gemini_key)

        prompt_text = """You are a certificate data extractor.

Extract the following details from this Police Clearance Certificate (from country of birth):
1. Certificate name (e.g. "Police Clearance Certificate", "Certificate of Good Conduct", "Criminal Record Certificate")
2. Staff / applicant name as printed on the certificate
3. Country the certificate is issued for (country of birth)
4. Expiry date (if shown)
5. Issue / date of issue
6. Issuing body or authority (e.g. Police authority name, Ministry of Interior)

Return ONLY a JSON object — no markdown, no explanation:
{
  "certificate_name": "<exact certificate title as printed>",
  "staff_name_on_cert": "<applicant name as printed>",
  "country": "<country this PCC was issued for>",
  "expiry_date": "<expiry date if shown>",
  "issue_date": "<date of issue as printed>",
  "issuing_body": "<issuing authority or organisation>"
}

If a field is not visible, set it to null.
"""

        is_image = any(t in content_type for t in ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf   = 'pdf' in content_type or doc_url.lower().split('?')[0].endswith('.pdf')

        if is_image:
            ext   = 'jpeg' if any(t in content_type for t in ('jpeg', 'jpg')) else                     'png'  if 'png'  in content_type else                     'webp' if 'webp' in content_type else 'jpeg'
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": prompt_text}
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}]
            )
        else:
            try:
                import io as _io, pdfplumber
                with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                    raw_text = chr(10).join(p.extract_text() or '' for p in pdf.pages).strip()
            except Exception:
                raw_text = raw_bytes.decode('utf-8', errors='replace').strip()
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt_text + "\n\nCERTIFICATE TEXT:\n" + raw_text[:5000]
            )

        raw_out = (response.text or '').strip()
        raw_out = _re.sub(r'^```(?:json)?\s*', '', raw_out, flags=_re.MULTILINE)
        raw_out = _re.sub(r'```\s*$', '', raw_out, flags=_re.MULTILINE).strip()

        result       = _cjson.loads(raw_out)
        cert_name    = _v(result.get('certificate_name') or '')
        cert_staff   = _v(result.get('staff_name_on_cert') or '')
        country      = _v(result.get('country') or '')
        expiry_date  = _v(result.get('expiry_date') or '')
        issue_date   = _v(result.get('issue_date') or '')
        issuing_body = _v(result.get('issuing_body') or '')

        _mark_done({
            "pcc2_certificate_name": cert_name,
            "pcc2_staff_name":       cert_staff,
            "pcc2_country":          country,
            "pcc2_expiry_date":      expiry_date,
            "pcc2_issue_date":       issue_date,
            "pcc2_issuing_body":     issuing_body,
            "pcc2_doc_url":          doc_url,
            "pcc2_doc_type":         pcc2_doc.get('document_type_name', ''),
            "pcc2_note":             "extracted successfully",
        })

        return jsonify({
            "success":            True,
            "email":              email,
            "staff_name":         full_name,
            "doc_found":          True,
            "certificate_name":   cert_name,
            "staff_name_on_cert": cert_staff,
            "country":            country,
            "expiry_date":        expiry_date,
            "issue_date":         issue_date,
            "issuing_body":       issuing_body,
            "remaining_count":    max(0, remaining_total - 1),
            "message": (
                f"Police Clearance Certificate extracted for {full_name} "
                f"(country: {country or 'unknown'}) — "
                f"{max(0, remaining_total - 1)} remaining."
            ),
        })

    except _cjson.JSONDecodeError:
        _mark_done({"pcc2_note": "Gemini JSON parse error"})
        return jsonify({
            "success": False, "email": email,
            "error": "Gemini returned non-JSON",
            "remaining_count": max(0, remaining_total - 1),
        })
    except Exception as e:
        _mark_done({"pcc2_note": f"error: {e}"})
        return jsonify({
            "success": False, "email": email,
            "error": str(e),
            "remaining_count": max(0, remaining_total - 1),
        })


# ── Export: Police Clearance Certificates to Excel ────────────────────

@admin_bp.route('/live-staffs/export/police-clearance-xlsx')
@admin_required
def live_staff_export_police_clearance_xlsx():
    """Export Police Clearance Certificate details to Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        import io as _io

        docs = list(_staffs_col().find(
            {},
            {"section_1_personal_details": 1, "email": 1,
             "pcc2_certificate_name": 1, "pcc2_staff_name": 1,
             "pcc2_country": 1, "pcc2_expiry_date": 1,
             "pcc2_issue_date": 1, "pcc2_issuing_body": 1, "pcc2_fetched": 1}
        ))
        docs.sort(key=lambda d: _v(
            (d.get('section_1_personal_details') or {}).get('full_name') or ''
        ).lower())

        NAVY = '1B3A6B'; GREEN = '2E9E44'; WHITE = 'FFFFFF'
        ALT  = 'EFF6FF'; WARN  = 'FFF3CD'; RED   = 'FFDDDD'

        h_font  = Font(name='Arial', bold=True, color=WHITE, size=10)
        h_fill  = PatternFill('solid', start_color=NAVY, end_color=NAVY)
        h_align = Alignment(horizontal='center', vertical='center')
        b_font  = Font(name='Arial', size=10)
        l_align = Alignment(horizontal='left',   vertical='center')
        c_align = Alignment(horizontal='center', vertical='center')
        thin    = Side(style='thin', color='CCCCCC')
        border  = Border(left=thin, right=thin, top=thin, bottom=thin)
        green_b = Border(left=thin, right=thin, top=thin,
                         bottom=Side(style='medium', color=GREEN))

        wb = Workbook()
        ws = wb.active
        ws.title = 'Police Clearance Certificates'

        headers    = ['Sno', 'Staff Name', 'Email', 'Certificate Name',
                      'Name on Cert', 'Country', 'Expiry Date',
                      'Issue Date', 'Issuing Body', 'Status']
        col_widths = [5, 28, 36, 32, 28, 18, 14, 14, 28, 14]

        for ci, (hdr, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=ci, value=hdr)
            cell.font = h_font; cell.fill = h_fill
            cell.alignment = h_align; cell.border = green_b
            ws.column_dimensions[cell.column_letter].width = width
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f'A1:J{len(docs)+1}'

        from datetime import date as _date
        today = _date.today()

        def _is_expired(expiry_str):
            if not expiry_str:
                return None
            for fmt in ('%d/%m/%Y','%m/%Y','%Y-%m-%d','%d-%m-%Y','%B %Y','%b %Y'):
                try:
                    from datetime import datetime as _dt
                    d = _dt.strptime(expiry_str.strip(), fmt).date()
                    return d < today
                except Exception:
                    continue
            return None

        for ri, doc in enumerate(docs, start=2):
            s1       = doc.get('section_1_personal_details') or {}
            name     = _v(s1.get('full_name') or '')
            email    = _v(doc.get('email') or '')
            cert_n   = _v(doc.get('pcc2_certificate_name') or '')
            cert_s   = _v(doc.get('pcc2_staff_name') or '')
            country  = _v(doc.get('pcc2_country') or '')
            expiry   = _v(doc.get('pcc2_expiry_date') or '')
            issue    = _v(doc.get('pcc2_issue_date') or '')
            issuer   = _v(doc.get('pcc2_issuing_body') or '')
            fetched  = doc.get('pcc2_fetched', False)
            expired  = _is_expired(expiry)

            if not fetched:
                status   = 'Not Checked'
                row_fill = PatternFill('solid', start_color=WARN, end_color=WARN)
            elif not cert_n:
                status   = 'No Cert Found'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is True:
                status   = 'EXPIRED'
                row_fill = PatternFill('solid', start_color=RED, end_color=RED)
            elif expired is False:
                status   = 'Valid'
                row_fill = None
            else:
                status   = 'Found'
                row_fill = None

            alt_fill = PatternFill('solid', start_color=ALT, end_color=ALT)                        if ri % 2 == 0 and not row_fill else None

            row_vals = [ri-1, name, email, cert_n, cert_s,
                        country, expiry, issue, issuer, status]
            aligns   = [c_align, l_align, l_align, l_align, l_align,
                        l_align, c_align, c_align, l_align, c_align]

            for ci, (val, align) in enumerate(zip(row_vals, aligns), start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = b_font; cell.alignment = align
                cell.border = border
                cell.fill = row_fill or alt_fill or PatternFill()

            ws.row_dimensions[ri].height = 17

        ws.cell(row=len(docs)+2, column=1,
                value=f'Total: {len(docs)}').font = Font(name='Arial', bold=True, size=9)

        buf = _io.BytesIO()
        wb.save(buf)
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition":
                     f'attachment; filename="police_clearance_{datetime.utcnow().strftime("%Y%m%d")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# ── Cron: Extract Fire Safety Certificate ────────────────────────────

