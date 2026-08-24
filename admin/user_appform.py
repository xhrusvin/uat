"""
Application Form PDF generation + preview.

Pipeline
--------
  1. Load the user from MongoDB collection `users` (by mongo _id OR xn_user_id).
  2. POST {LIVE_STAFF_URL}/ai/recruitments/detail  {"_id": <xn_user_id>}
       -> authoritative profile + signature_url
  3. Download the signature image.
  4. POST {LIVE_STAFF_URL}/ai/recruitments/user-document-list
       -> document list
  5. If a document with document_type_name == "Application" has a url:
       download it, extract structured fields via Gemini, and use those to
       enrich / override the form.
  6. If there is NO Application form, look for Driving Licence / Passport /
       Birth Certificate / Proof of Address and tick whichever are present.
  7. Merge everything (mongo <- detail API <- Gemini) and render the
       Application Form PDF with reportlab.

Endpoints
---------
  GET/POST /admin/appform/preview/<user_id>    inline PDF (browser preview)
  GET/POST /admin/appform/download/<user_id>   PDF as attachment
  GET      /admin/appform/data/<user_id>       JSON of the merged context (debug)
  POST     /admin/appform/upload/<user_id>     push the PDF to the HSE document API

Env (same as live_staffs_crons.py)
----------------------------------
  LIVE_STAFF_URL, XN_PORTAL_API_KEY, XN_APP_COUNTRY, GEMINI_API_KEY

Optional: APPFORM_LOGO_PATH overrides the embedded Xpress Health logo.

Upload env (see admin/hse_document_upload.py)
---------------------------------------------
  HSE_UPLOAD_URL       base url of the admin API (falls back to LIVE_STAFF_URL)
  HSE_UPLOAD_API_KEY   Api-Key header (falls back to XN_PORTAL_API_KEY)
"""

from flask import request, jsonify, Response
from bson import ObjectId
from bson.errors import InvalidId
from datetime import datetime
import base64
import io as _io
import json as _cjson
import os
import re as _re

from database import db
from . import admin_bp
from admin.views import admin_required
from admin.hse_document_upload import upload_hse_document, HSE_APPLICATION_FORM

# ──────────────────────────────────────────────────────────────────────
# Small helpers
# ──────────────────────────────────────────────────────────────────────


def _v(val):
    """Normalise anything to a clean string."""
    if val is None:
        return ''
    return str(val).strip()


def _users_col():
    return db.users


def _api_headers(json_body=True):
    h = {
        "Api-Key":       os.environ.get('XN_PORTAL_API_KEY', ''),
        "X-App-Country": os.environ.get('XN_APP_COUNTRY', ''),
        "Accept":        "application/json",
    }
    if json_body:
        h["Content-Type"] = "application/json"
    return h


def _base_url():
    return os.environ.get('LIVE_STAFF_URL', '').rstrip('/')


def _xn_post(path, payload, timeout=30):
    """
    POST to the XN Portal, falling back to GET on 405 (same pattern the
    existing crons use). Returns the parsed `data` node.
    Raises RuntimeError with a readable message on failure.
    """
    import requests as _req

    base = _base_url()
    if not base:
        raise RuntimeError("LIVE_STAFF_URL not set in environment")

    url = f"{base}{path}"
    try:
        resp = _req.post(url, json=payload, headers=_api_headers(), timeout=timeout)
        if resp.status_code == 405:
            resp = _req.get(url, params=payload,
                            headers=_api_headers(False), timeout=timeout)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        raise RuntimeError(f"{path} call failed: {e}")

    if not data.get('success'):
        raise RuntimeError(f"{path} returned success=false: "
                           f"{data.get('message', 'unknown error')}")

    return data.get('data')


def _download(url, timeout=60):
    """Download a URL. Returns (bytes, content_type) or (None, '')."""
    import requests as _req
    try:
        r = _req.get(url, timeout=timeout)
        if r.status_code == 404:
            return None, ''
        r.raise_for_status()
        return r.content, (r.headers.get('Content-Type') or '').lower()
    except Exception:
        return None, ''


def _yn(val):
    """0/1/'yes'/'no'/True/False -> 'Yes' / 'No' / ''."""
    if val is None or val == '':
        return ''
    if isinstance(val, bool):
        return 'Yes' if val else 'No'
    s = _v(val).lower()
    if s in ('1', 'yes', 'true', 'y'):
        return 'Yes'
    if s in ('0', 'no', 'false', 'n'):
        return 'No'
    return _v(val)


def _truthy(val):
    return _yn(val) == 'Yes'


def _fmt_date(val):
    """Best-effort date formatting to dd/mm/yyyy."""
    if not val:
        return ''
    if isinstance(val, datetime):
        return val.strftime('%d/%m/%Y')
    s = _v(val)
    for fmt in ('%Y-%m-%d', '%Y-%m-%dT%H:%M:%S', '%d/%m/%Y', '%d-%m-%Y'):
        try:
            return datetime.strptime(s[:len(fmt) + 2].strip(), fmt).strftime('%d/%m/%Y')
        except Exception:
            continue
    return s


def _exp(years, months):
    """(4, 9) -> '4 years 9 months'."""
    y, m = years or 0, months or 0
    parts = []
    if y:
        parts.append(f"{y} year{'s' if y != 1 else ''}")
    if m:
        parts.append(f"{m} month{'s' if m != 1 else ''}")
    return ' '.join(parts) or ''


def _first(*vals):
    """First non-empty value."""
    for v in vals:
        if v not in (None, '', [], {}):
            return v
    return ''


# ──────────────────────────────────────────────────────────────────────
# Lookup resolution — same convention as user_cv.py / user_point_scale.py
# ──────────────────────────────────────────────────────────────────────


def _oid(val):
    try:
        return ObjectId(str(val))
    except Exception:
        return None


def _lookup_name(collection_name, oid_str):
    """
    Soft lookup for the *_id reference fields on `users`
    (county_id, country_id, gender_id, visa_type_id ...).
    Returns '' on any failure so it never breaks form generation.
    """
    oid = _oid(oid_str)
    if oid is None:
        return ''
    try:
        doc = db[collection_name].find_one(
            {"_id": oid}, {"name": 1, "title": 1, "label": 1}
        )
        if not doc:
            return ''
        return _v(doc.get('name') or doc.get('title') or doc.get('label') or '')
    except Exception:
        return ''


# ──────────────────────────────────────────────────────────────────────
# Gemini extraction of an existing Application form
# ──────────────────────────────────────────────────────────────────────

_GEMINI_PROMPT = """You are a data extractor for the Xpress Health Application Form.

Read the attached completed application form and extract each field exactly
as written on it.

Return ONLY a JSON object - no markdown fences, no explanation - in this shape:
{
  "full_name": null,
  "email": null,
  "role": null,
  "address": null,
  "eir_code": null,
  "mobile_number": null,
  "work_permit_visa_status": null,
  "pps_number": null,
  "id_documents_provided": [],
  "is_nurse_nmbi": null,
  "is_hca_qqi": null,
  "total_years_of_experience": null,
  "references": [
    {"name": null, "position": null, "organisation": null,
     "telephone": null, "email": null}
  ],
  "vaccinated": null
}

Rules:
- Any field not filled in on the form must be null. Never guess or invent.
- "work_permit_visa_status" is the full string as printed, e.g. "Yes; Stamp 4".
- "id_documents_provided" lists only the TICKED boxes under Identity
  Verification, using exactly these spellings: "Passport",
  "Birth Certificate", "Driving Licence", "Proof of Address".
  An unticked or empty box must NOT appear in the list.
- "is_nurse_nmbi" and "is_hca_qqi" are true/false based on which box is
  ticked in Section 3.
- "references" holds up to 3 entries in the order they appear on the form.
  Omit reference blocks that are entirely blank.
- "vaccinated" is "Yes" or "No" as answered on the form.
- "total_years_of_experience" is the string as printed, e.g. "3 years, 0 months".
"""


def _extract_appform_via_gemini(raw_bytes, content_type, source_url=''):
    """
    Send the downloaded Application form to Gemini and return a dict of
    extracted fields. Returns ({}, error_string) on any failure so the
    caller can carry on with mongo/API data alone.
    """
    gemini_key = os.environ.get('GEMINI_API_KEY', '')
    if not gemini_key:
        return {}, 'GEMINI_API_KEY not set'
    if not raw_bytes:
        return {}, 'no bytes to extract'

    try:
        from google import genai as google_genai
        client = google_genai.Client(api_key=gemini_key)

        url_path = (source_url or '').split('?')[0].lower()
        is_image = any(t in content_type for t in
                       ('image/', 'jpeg', 'jpg', 'png', 'webp'))
        is_pdf = 'pdf' in content_type or url_path.endswith('.pdf')

        if is_image:
            ext = ('jpeg' if any(t in content_type for t in ('jpeg', 'jpg'))
                   else 'png' if 'png' in content_type
                   else 'webp' if 'webp' in content_type else 'jpeg')
            parts = [
                {"inline_data": {"mime_type": f"image/{ext}",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": _GEMINI_PROMPT},
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}])

        elif is_pdf:
            parts = [
                {"inline_data": {"mime_type": "application/pdf",
                                 "data": base64.b64encode(raw_bytes).decode()}},
                {"text": _GEMINI_PROMPT},
            ]
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=[{"parts": parts}])

        else:
            # DOCX / unknown -> pre-extract text
            raw_text = ''
            try:
                import docx as _docx
                d = _docx.Document(_io.BytesIO(raw_bytes))
                raw_text = '\n'.join(p.text for p in d.paragraphs)
                for tbl in d.tables:
                    for row in tbl.rows:
                        raw_text += '\n' + ' | '.join(c.text for c in row.cells)
            except Exception:
                try:
                    import pdfplumber
                    with pdfplumber.open(_io.BytesIO(raw_bytes)) as pdf:
                        raw_text = '\n'.join(p.extract_text() or ''
                                             for p in pdf.pages)
                except Exception:
                    raw_text = raw_bytes.decode('utf-8', errors='replace')

            raw_text = raw_text.strip()
            if not raw_text:
                return {}, 'could not read document text'

            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=_GEMINI_PROMPT + "\n\nAPPLICATION FORM TEXT:\n" + raw_text[:20000])

        out = (response.text or '').strip()
        out = _re.sub(r'^```(?:json)?\s*', '', out, flags=_re.MULTILINE)
        out = _re.sub(r'```\s*$', '', out, flags=_re.MULTILINE).strip()

        parsed = _cjson.loads(out)
        if not isinstance(parsed, dict):
            return {}, 'Gemini returned non-object JSON'
        # Strip nulls so merge logic can use simple truthiness
        return {k: v for k, v in parsed.items() if v not in (None, '', [], {})}, ''

    except _cjson.JSONDecodeError:
        return {}, 'Gemini returned non-JSON'
    except Exception as e:
        return {}, f'Gemini error: {e}'


# ──────────────────────────────────────────────────────────────────────
# Gather everything needed for one form
# ──────────────────────────────────────────────────────────────────────

_ID_DOC_TYPES = ('Driving Licence', 'Passport', 'Birth Certificate', 'Proof of Address')

# Matching aliases -> canonical label
_ID_DOC_ALIASES = {
    'Driving Licence':   ('driving licence', 'driving license', 'drivers licence',
                          'driver licence', 'driver license', 'drivers license'),
    'Passport':          ('passport',),
    'Birth Certificate': ('birth certificate', 'birth cert'),
    'Proof of Address':  ('proof of address', 'address proof', 'utility bill'),
}


def _find_user(user_id):
    """Look up a users record by mongo _id, xn_user_id, or email."""
    col = _users_col()
    key = _v(user_id)

    if len(key) == 24:
        try:
            rec = col.find_one({"_id": ObjectId(key)})
            if rec:
                return rec
        except (InvalidId, TypeError):
            pass

    return col.find_one({"$or": [
        {"xn_user_id": key},
        {"email": key},
    ]})


def _gather_appform_context(user_id):
    """
    Build the merged context dict used to render the PDF.
    Returns (ctx, error_string). ctx is None when the user cannot be found.
    """
    user = _find_user(user_id)
    if not user:
        return None, f"No user found in 'users' for id/email: {user_id}"

    xn_id = _v(user.get('xn_user_id'))
    email = _v(user.get('email'))

    meta = {
        "mongo_id":        str(user.get('_id')),
        "xn_user_id":      xn_id,
        "detail_api_ok":   False,
        "detail_error":    '',
        "documents_ok":    False,
        "documents_error": '',
        "appform_found":   False,
        "appform_url":     '',
        "gemini_ok":       False,
        "gemini_error":    '',
        "signature_ok":    False,
        "id_docs_source":  '',
    }

    # ── 2. Detail API ────────────────────────────────────────────────
    detail = {}
    if xn_id:
        try:
            detail = _xn_post('/ai/recruitments/detail', {"_id": xn_id}) or {}
            meta["detail_api_ok"] = True
        except RuntimeError as e:
            meta["detail_error"] = str(e)
    else:
        meta["detail_error"] = 'user has no xn_user_id'

    # ── 3. Signature ─────────────────────────────────────────────────
    signature_bytes = None
    signature_url = _v(detail.get('signature_url'))
    if signature_url:
        signature_bytes, _ = _download(signature_url, timeout=30)
        meta["signature_ok"] = bool(signature_bytes)

    # ── 4. Document list ─────────────────────────────────────────────
    documents = []
    if xn_id or email:
        payload = {"_id": xn_id} if xn_id else {"email": email}
        try:
            doc_data = _xn_post('/ai/recruitments/user-document-list', payload)
            if isinstance(doc_data, list):
                documents = doc_data
            elif isinstance(doc_data, dict):
                documents = doc_data.get('documents') or []
            meta["documents_ok"] = True
        except RuntimeError as e:
            meta["documents_error"] = str(e)

    # ── 5. Application form -> Gemini ────────────────────────────────
    extracted = {}
    app_doc = None
    for d in documents:
        name = _v(d.get('document_type_name')).lower()
        if name == 'application' or 'application form' in name:
            if _v(d.get('url')):
                app_doc = d
                break

    if app_doc:
        meta["appform_found"] = True
        meta["appform_url"] = _v(app_doc.get('url'))
        raw, ctype = _download(meta["appform_url"])
        if raw:
            extracted, gerr = _extract_appform_via_gemini(raw, ctype, meta["appform_url"])
            meta["gemini_ok"] = bool(extracted)
            meta["gemini_error"] = gerr
        else:
            meta["gemini_error"] = 'application form download failed (404 or unreachable)'

    # ── 6. ID documents to tick ──────────────────────────────────────
    id_docs = {label: False for label in _ID_DOC_TYPES}

    if extracted.get('id_documents_provided'):
        # Trust the ticks already on the extracted Application form
        meta["id_docs_source"] = 'application_form'
        for raw_label in extracted['id_documents_provided']:
            low = _v(raw_label).lower()
            for canonical, aliases in _ID_DOC_ALIASES.items():
                if any(a in low for a in aliases):
                    id_docs[canonical] = True
    elif not meta["appform_found"]:
        # No Application form -> tick based on what's in the document list
        meta["id_docs_source"] = 'document_list'
        for d in documents:
            if not _v(d.get('url')):
                continue
            low = _v(d.get('document_type_name')).lower()
            for canonical, aliases in _ID_DOC_ALIASES.items():
                if any(a in low for a in aliases):
                    id_docs[canonical] = True

    # ── 7. Merge: mongo <- detail API <- Gemini ──────────────────────
    def pick(gemini_key, detail_key, mongo_key=None):
        return _v(_first(extracted.get(gemini_key),
                         detail.get(detail_key),
                         user.get(mongo_key or detail_key)))

    first_name = _first(detail.get('first_name'), user.get('first_name'))
    last_name = _first(detail.get('last_name'), user.get('last_name'))
    full_name = _v(_first(extracted.get('full_name'),
                          user.get('name'),
                          f"{_v(first_name)} {_v(last_name)}".strip()))

    sub_types = _first(detail.get('user_sub_type_ids'),
                       user.get('user_sub_type_ids'), []) or []
    if isinstance(sub_types, str):
        sub_types = [sub_types]

    references = _first(detail.get('references'), user.get('references'), []) or []

    # ── Derived fields the template needs ────────────────────────────
    role = _v(_first(extracted.get('role'),
                     detail.get('user_type'),
                     user.get('designation')))

    # "Yes; Stamp 4" — permission to work plus the visa type, as printed
    # on the template.
    _perm = _yn(_first(detail.get('permission_to_work'),
                       user.get('permission_to_work')))
    _visa = _v(_lookup_name('visa_types', _first(detail.get('visa_type_id'),
                                                user.get('visa_type_id'))))
    visa_status = _v(extracted.get('work_permit_visa_status'))
    if not visa_status:
        visa_status = '; '.join(x for x in (_perm, _visa) if x)

    # Section 3 tick — Nurse (NMBI) vs HCA (QQI L5)
    _role_low = role.lower()
    if extracted.get('is_nurse_nmbi') is not None or \
       extracted.get('is_hca_qqi') is not None:
        is_nurse = bool(extracted.get('is_nurse_nmbi'))
        is_hca = bool(extracted.get('is_hca_qqi'))
    else:
        is_nurse = 'nurse' in _role_low or 'nursing' in _role_low
        is_hca = (('healthcare assistant' in _role_low) or
                  ('health care assistant' in _role_low) or
                  _re.search(r'\bhca\b', _role_low) is not None or
                  bool(_v(detail.get('qqi_status_number'))))
        if is_hca:
            is_nurse = False

    # Vaccinated — Yes when any immunisation is recorded on the profile
    _vax_flags = [
        _first(detail.get('tuberculosis_vaccine'), user.get('tuberculosis_vaccine')),
        _first(detail.get('hepatitis_antibody'), user.get('hepatitis_antibody')),
        _first(detail.get('mmr_vaccine'), user.get('mmr_vaccine')),
        _first(detail.get('covid_19_vaccine'), user.get('covid_19_vaccine')),
    ]
    vaccinated = _v(extracted.get('vaccinated'))
    if not vaccinated:
        vaccinated = 'Yes' if any(_truthy(f) for f in _vax_flags) else 'No'

    # Address is stored multi-line in mongo; the template prints one line.
    address = _v(_first(extracted.get('address'),
                        detail.get('address'), user.get('address')))
    address = _re.sub(r'\s*\n\s*', ', ', address)
    address = _re.sub(r'\s{2,}', ' ', address).strip(' ,')

    ctx = {
        "meta": meta,
        "signature_bytes": signature_bytes,
        "id_docs": id_docs,
        "documents": documents,

        # Personal
        "full_name":        full_name,
        "first_name":       _v(first_name),
        "last_name":        _v(last_name),
        "dob":              _fmt_date(_first(extracted.get('date_of_birth'),
                                             detail.get('dob'), user.get('dob'))),
        "gender":           _v(_first(extracted.get('gender'),
                                      _lookup_name('genders',
                                                _first(detail.get('gender_id'),
                                                       user.get('gender_id'))))),
        "email":            pick('email', 'email'),
        "phone":            _v(_first(extracted.get('mobile_number'),
                                      detail.get('phone_number'),
                                      user.get('phone'))),
        "address":          address,
        "eir_code":         pick('eir_code', 'eir_code'),
        "county":           _v(_first(_lookup_name('counties',
                                                _first(detail.get('county_id'),
                                                       user.get('county_id'))))),
        "country":          _v(_first(extracted.get('nationality'),
                                      _lookup_name('countries',
                                                _first(detail.get('country_id'),
                                                       user.get('country_id'))))),
        "pps_number":       pick('pps_number', 'pps_number'),

        # Position
        "role":             role,
        "position":         role,
        "visa_status":      visa_status,
        "is_nurse":         is_nurse,
        "is_hca":           is_hca,
        "vaccinated":       vaccinated,
        "specialities":     _v(_first(extracted.get('specialities'),
                                      ', '.join(_v(s) for s in sub_types))),
        "total_experience": _v(_first(extracted.get('total_years_of_experience'),
                                      _exp(_first(detail.get('experience_year'),
                                                  user.get('experience_year')),
                                           _first(detail.get('experience_month'),
                                                  user.get('experience_month'))))),
        "masters":          _yn(_first(detail.get('masters'), user.get('masters'))),
        "uniform_size":     pick('uniform_size', 'uniform_size'),
        "own_transport":    _v(_first(extracted.get('own_transport'),
                                      _TRAVEL_MODES.get(
                                          _first(detail.get('travel_mode'),
                                                 user.get('travel_mode')), ''))),

        # Most recent employment
        "company_name":     pick('employer', 'company_name'),
        "job_title":        pick('job_title', 'job_title'),
        "company_phone":    _v(f"{_v(_first(detail.get('company_dial_code'), user.get('company_dial_code')))} "
                               f"{_v(_first(detail.get('company_phone'), user.get('company_phone')))}").strip(),
        "company_county":   _lookup_name('counties',
                                                _first(detail.get('company_county_id'),
                                                       user.get('company_county_id'))),
        "company_experience": _exp(_first(detail.get('last_company_experience_year'),
                                          user.get('last_company_experience_year')),
                                   _first(detail.get('last_company_experience_month'),
                                          user.get('last_company_experience_month'))),

        # Right to work
        "permission_to_work": _yn(_first(extracted.get('right_to_work'),
                                         detail.get('permission_to_work'),
                                         user.get('permission_to_work'))),
        "visa_type":          _visa,
        "work_permit_exemption": _yn(_first(extracted.get('work_permit_exemption'),
                                            detail.get('work_permit_exemption'),
                                            user.get('work_permit_exemption'))),

        # Health
        "tuberculosis_vaccine": _yn(_first(extracted.get('tuberculosis_vaccine'),
                                           detail.get('tuberculosis_vaccine'),
                                           user.get('tuberculosis_vaccine'))),
        "hepatitis_antibody":   _yn(_first(extracted.get('hepatitis_antibody'),
                                           detail.get('hepatitis_antibody'),
                                           user.get('hepatitis_antibody'))),
        "mmr_vaccine":          _yn(_first(extracted.get('mmr_vaccine'),
                                           detail.get('mmr_vaccine'),
                                           user.get('mmr_vaccine'))),
        "covid_19_vaccine":     _yn(_first(extracted.get('covid_19_vaccine'),
                                           detail.get('covid_19_vaccine'),
                                           user.get('covid_19_vaccine'))),
        "criminal_conviction":  _yn(extracted.get('criminal_conviction')),

        # Repeating sections
        "qualifications":     extracted.get('qualifications') or [],
        "employment_history": extracted.get('employment_history') or [],
        "references":         _merge_references(references,
                                                extracted.get('references') or []),

        "declaration_date": _v(_first(extracted.get('declaration_date'),
                                      datetime.utcnow().strftime('%d/%m/%Y'))),
        "notes": _v(extracted.get('notes')),
    }

    return ctx, ''


_TRAVEL_MODES = {
    1: 'Own car',
    2: 'Public transport',
    3: 'Own car / Public transport',
    '1': 'Own car',
    '2': 'Public transport',
    '3': 'Own car / Public transport',
}


def _merge_references(api_refs, gemini_refs):
    """Prefer the structured API references; fall back to Gemini's."""
    out = []
    for r in (api_refs or []):
        out.append({
            "name":         _v(r.get('name')),
            "job_role":     _v(r.get('job_role')),
            "organization": _v(r.get('organization')),
            "email":        _v(r.get('email')),
            "phone":        _v(f"{_v(r.get('dial_code'))} {_v(r.get('phone'))}").strip(),
            "status":       _v(r.get('status')),
        })
    if out:
        return out
    for r in (gemini_refs or []):
        out.append({
            "name":         _v(r.get('name')),
            "job_role":     _v(_first(r.get('position'), r.get('job_role'))),
            "organization": _v(_first(r.get('organisation'), r.get('organization'))),
            "email":        _v(r.get('email')),
            "phone":        _v(_first(r.get('telephone'), r.get('phone'))),
            "status":       '',
        })
    return out


# ──────────────────────────────────────────────────────────────────────
# PDF rendering — matches the official application_form.docx template
# ──────────────────────────────────────────────────────────────────────

_TITLE_BLUE = '#2E5C8A'     # document title
_HEADING_BLUE = '#5B8DB8'   # "Section 1 – Personal Details"
_TICK_FILL = '#8DB3CE'      # checked box fill
_TICK_MARK = '#1E7B34'      # green check
_BOX_LINE = '#4A4A4A'

# Fixed declaration wording, verbatim from the template.
_DECLARATION = (
    'I declare that the information provided in this application form is true '
    'and accurate to the best of my knowledge. I understand that any false or '
    'misleading information may result in the withdrawal of an offer of '
    'employment or termination of employment.'
)


_FONT_CANDIDATES = [
    # (regular, bold) — first pair that exists on disk wins
    ('/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
     '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf'),
    ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
     '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'),
]

_FONTS_READY = None


def _register_fonts():
    """
    Register Liberation Sans so the output matches the Word template's
    metrics. Falls back to built-in Helvetica if no TTFs are installed.
    Returns (regular, bold).
    """
    global _FONTS_READY
    if _FONTS_READY is not None:
        return _FONTS_READY

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.fonts import addMapping

    for reg, bold in _FONT_CANDIDATES:
        if not (os.path.exists(reg) and os.path.exists(bold)):
            continue
        try:
            pdfmetrics.registerFont(TTFont('AppSans', reg))
            pdfmetrics.registerFont(TTFont('AppSans-Bold', bold))
            addMapping('AppSans', 0, 0, 'AppSans')
            addMapping('AppSans', 1, 0, 'AppSans-Bold')
            _FONTS_READY = ('AppSans', 'AppSans-Bold')
            return _FONTS_READY
        except Exception:
            continue

    _FONTS_READY = ('Helvetica', 'Helvetica-Bold')
    return _FONTS_READY


def _styles():
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib import colors

    ss = getSampleStyleSheet()
    regular, bold = _register_fonts()
    return {
        'title': ParagraphStyle(
            'title', parent=ss['Normal'], fontName=bold,
            fontSize=13, leading=17, alignment=TA_CENTER,
            textColor=colors.HexColor(_TITLE_BLUE), spaceAfter=14),
        'heading': ParagraphStyle(
            'heading', parent=ss['Normal'], fontName=bold,
            fontSize=12.5, leading=16,
            textColor=colors.HexColor(_HEADING_BLUE),
            spaceBefore=8, spaceAfter=2),
        'field': ParagraphStyle(
            'field', parent=ss['Normal'], fontName=regular,
            fontSize=10, leading=13, spaceAfter=7),
        'plain': ParagraphStyle(
            'plain', parent=ss['Normal'], fontName=regular,
            fontSize=10, leading=13, spaceAfter=7),
        'body': ParagraphStyle(
            'body', parent=ss['Normal'], fontName=regular,
            fontSize=10, leading=14, spaceAfter=8),
    }


def _tick(checked, size=10):
    """
    A tick-box drawn as vector graphics.

    The template uses the ☑ / ☐ glyphs, but reportlab's core fonts have no
    box characters, so these are drawn instead: an outlined square, filled
    and given a green check when ticked.
    """
    from reportlab.graphics.shapes import Drawing, Rect, PolyLine
    from reportlab.lib import colors

    d = Drawing(size, size)
    if checked:
        d.add(Rect(0.5, 0.5, size - 1, size - 1,
                   fillColor=colors.HexColor(_TICK_FILL),
                   strokeColor=colors.HexColor(_TICK_FILL),
                   strokeWidth=0.6))
        d.add(PolyLine(
            [size * 0.22, size * 0.52,
             size * 0.42, size * 0.28,
             size * 0.80, size * 0.74],
            strokeColor=colors.HexColor(_TICK_MARK),
            strokeWidth=1.7, strokeLineCap=1, strokeLineJoin=1))
    else:
        d.add(Rect(0.5, 0.5, size - 1, size - 1,
                   fillColor=None,
                   strokeColor=colors.HexColor(_BOX_LINE),
                   strokeWidth=0.7))
    return d


def _tick_row(items, st, col_widths):
    """A single row of `(label, checked)` tick-boxes."""
    from reportlab.platypus import Table, TableStyle, Paragraph

    cells, widths = [], []
    for (label, checked), w in zip(items, col_widths):
        cells.append(_tick(checked))
        cells.append(Paragraph(label, st['plain']))
        widths.append(13)
        widths.append(w - 13)

    t = Table([cells], colWidths=widths)
    t.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 1),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
    ]))
    return t


def _logo_flowable(max_w=115, max_h=31):
    """The Xpress Health logo, right-aligned as in the template."""
    from reportlab.platypus import Image, Spacer

    data = None
    path = os.environ.get('APPFORM_LOGO_PATH', '')
    if path and os.path.exists(path):
        try:
            with open(path, 'rb') as fh:
                data = fh.read()
        except Exception:
            data = None
    if data is None:
        try:
            from .appform_assets import XPRESS_LOGO_BYTES
            data = XPRESS_LOGO_BYTES
        except Exception:
            try:
                from appform_assets import XPRESS_LOGO_BYTES
                data = XPRESS_LOGO_BYTES
            except Exception:
                return Spacer(1, max_h)

    try:
        img = Image(_io.BytesIO(data))
        ratio = (img.imageHeight / img.imageWidth) if img.imageWidth else 0.27
        img.drawWidth = max_w
        img.drawHeight = min(max_h, max_w * ratio)
        img.hAlign = 'RIGHT'
        return img
    except Exception:
        return Spacer(1, max_h)


def _build_appform_pdf(ctx):
    """
    Render the merged context as the Xpress Health Application Form,
    reproducing the layout of application_form.docx.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Image, KeepTogether)

    st = _styles()
    buf = _io.BytesIO()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.0 * inch, rightMargin=1.0 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        title=f"Xpress Health Application Form — {ctx['full_name']}",
        author='Xpress Health',
        subject='Application Form',
    )
    W = doc.width
    story = []

    def field(label, value):
        """**Label:** value — the template's field style."""
        val = _v(value)
        story.append(Paragraph(
            f"<b>{label}</b>{(' ' + val) if val else ''}", st['field']))

    # ── Logo ─────────────────────────────────────────────────────────
    story.append(_logo_flowable())
    story.append(Spacer(1, 18))

    # ── Title ────────────────────────────────────────────────────────
    story.append(Paragraph('Xpress Health Application Form', st['title']))

    # ── Section 1 — Personal Details ─────────────────────────────────
    story.append(Paragraph('Section 1 &#8211; Personal Details', st['heading']))
    field('Full Name:', ctx['full_name'])
    field('Email:', ctx['email'])
    field('Role:', ctx['role'])
    field('Address:', ctx['address'])
    field('Eircode/Postcode:', ctx['eir_code'])
    field('Mobile Number:', ctx['phone'])
    field('Work Permit / Visa Status:', ctx['visa_status'])
    field('PPS Number (if applicable)', ctx['pps_number'])

    # ── Section 2 — Identity Verification ────────────────────────────
    story.append(Paragraph('Section 2 &#8211; Identity Verification', st['heading']))
    story.append(Paragraph('ID Proof:', st['plain']))
    id_docs = ctx['id_docs']
    story.append(_tick_row(
        [('Passport', id_docs.get('Passport', False)),
         ('Birth Certificate', id_docs.get('Birth Certificate', False)),
         ('Driving Licence', id_docs.get('Driving Licence', False)),
         ('Proof of Address', id_docs.get('Proof of Address', False))],
        st, [W * 0.20, W * 0.28, W * 0.26, W * 0.26]))
    story.append(Spacer(1, 4))

    # ── Section 3 — Qualification and Experience ─────────────────────
    story.append(Paragraph('Section 3 &#8211; Qualification and Experience',
                           st['heading']))
    story.append(_tick_row(
        [('Nurse (NMBI):', ctx['is_nurse']),
         ('HCA (QQI L5):', ctx['is_hca'])],
        st, [W * 0.30, W * 0.70]))
    story.append(Spacer(1, 6))
    field('Total years of experience:', ctx['total_experience'])

    # ── References 1–3 ───────────────────────────────────────────────
    refs = list(ctx['references'])[:3]
    while len(refs) < 3:
        refs.append({})

    for i, ref in enumerate(refs, start=1):
        block = [
            Paragraph(f'Reference {i}', st['heading']),
            Paragraph(f"<b>Name:</b>{(' ' + _v(ref.get('name'))) if _v(ref.get('name')) else ''}", st['field']),
            Paragraph(f"<b>Position:</b>{(' ' + _v(ref.get('job_role'))) if _v(ref.get('job_role')) else ''}", st['field']),
            Paragraph(f"<b>Organisation:</b>{(' ' + _v(ref.get('organization'))) if _v(ref.get('organization')) else ''}", st['field']),
            Paragraph(f"<b>Telephone:</b>{(' ' + _v(ref.get('phone'))) if _v(ref.get('phone')) else ''}", st['field']),
            Paragraph(f"<b>Email:</b>{(' ' + _v(ref.get('email'))) if _v(ref.get('email')) else ''}", st['field']),
        ]
        story.append(KeepTogether(block))

    # ── Vaccinated ───────────────────────────────────────────────────
    story.append(Spacer(1, 14))
    story.append(Paragraph('Vaccinated:', st['plain']))
    story.append(Paragraph(f"Yes / No: {ctx['vaccinated']}", st['plain']))

    # ── Declaration + signature ──────────────────────────────────────
    decl = [Paragraph(_DECLARATION, st['body']),
            Paragraph('<b>Applicant Signature:</b>', st['field'])]

    if ctx.get('signature_bytes'):
        try:
            img = Image(_io.BytesIO(ctx['signature_bytes']))
            ratio = (img.imageHeight / img.imageWidth) if img.imageWidth else 0.38
            img.drawWidth = 1.6 * inch
            img.drawHeight = min(0.75 * inch, 1.6 * inch * ratio)
            img.hAlign = 'LEFT'
            decl.append(Spacer(1, 4))
            decl.append(img)
        except Exception:
            decl.append(Spacer(1, 0.55 * inch))
    else:
        decl.append(Spacer(1, 0.55 * inch))

    story.append(KeepTogether(decl))

    doc.build(story)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────
# Routes
# ──────────────────────────────────────────────────────────────────────


def _staff_id_for(user):
    """
    The HSE upload API validates staff_id against the user service, so it
    wants the staff/XN id — not the local Mongo _id. Prefer the explicit
    fields, fall back to the Mongo id as a last resort.
    """
    for key in ('staff_id', 'xn_user_id', 'xn_staff_id', 'user_service_id'):
        if _v(user.get(key)):
            return _v(user.get(key))
    return _v(user.get('_id'))


def _appform_filename(ctx, fallback='applicant'):
    safe = (ctx.get('full_name') or fallback).replace(' ', '_').replace('/', '_')
    return f"ApplicationForm_{safe}.pdf"


def _resolve_user_id(user_id):
    """Allow the id via path, query string, or JSON body."""
    if user_id:
        return user_id
    body = request.get_json(silent=True) or {}
    return _v(body.get('_id') or body.get('user_id') or
              request.args.get('_id') or request.args.get('user_id'))


def _render(user_id, disposition):
    uid = _resolve_user_id(user_id)
    if not uid:
        return jsonify({"success": False,
                        "error": "Provide a user id in the path, ?_id=, or JSON body"}), 400

    try:
        ctx, err = _gather_appform_context(uid)
    except Exception as e:
        return jsonify({"success": False, "error": f"Context build failed: {e}"}), 500

    if ctx is None:
        return jsonify({"success": False, "error": err}), 404

    try:
        pdf_bytes = _build_appform_pdf(ctx)
    except Exception as e:
        return jsonify({"success": False, "error": f"PDF build failed: {e}"}), 500

    filename = _appform_filename(ctx)

    return Response(pdf_bytes, mimetype='application/pdf', headers={
        "Content-Disposition": f'{disposition}; filename="{filename}"',
        "Cache-Control":       "no-store",
        "X-Appform-Source":    'application_form' if ctx['meta']['appform_found'] else 'profile_data',
    })


@admin_bp.route('/users/<user_id>/appform/preview', methods=['GET'])
@admin_bp.route('/appform/preview/<user_id>', methods=['GET', 'POST'])
@admin_bp.route('/appform/preview', methods=['GET', 'POST'])
@admin_required
def appform_preview(user_id=None):
    """Inline PDF preview — renders in the browser / an <iframe>."""
    try:
        return _render(user_id, 'inline')
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/users/<user_id>/appform/download', methods=['GET'])
@admin_bp.route('/appform/download/<user_id>', methods=['GET', 'POST'])
@admin_bp.route('/appform/download', methods=['GET', 'POST'])
@admin_required
def appform_download(user_id=None):
    """Same PDF, but as a file download."""
    try:
        return _render(user_id, 'attachment')
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@admin_bp.route('/users/<user_id>/appform/json', methods=['GET'])
@admin_bp.route('/appform/data/<user_id>', methods=['GET', 'POST'])
@admin_bp.route('/appform/data', methods=['GET', 'POST'])
@admin_required
def appform_data(user_id=None):
    """
    JSON view of the merged context — shows what came from where
    (detail API, document list, Gemini) before the PDF is rendered.
    """
    uid = _resolve_user_id(user_id)
    if not uid:
        return jsonify({"success": False, "error": "Provide a user id"}), 400

    try:
        ctx, err = _gather_appform_context(uid)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

    if ctx is None:
        return jsonify({"success": False, "error": err}), 404

    payload = {k: v for k, v in ctx.items()
               if k not in ('signature_bytes', 'documents')}
    payload['signature_downloaded'] = bool(ctx.get('signature_bytes'))
    payload['document_type_names'] = [
        _v(d.get('document_type_name')) for d in ctx.get('documents', [])
    ]

    return jsonify({"success": True, "data": payload})


@admin_bp.route('/users/<user_id>/appform/upload', methods=['POST'])
@admin_bp.route('/appform/upload/<user_id>', methods=['POST'])
@admin_bp.route('/appform/upload', methods=['POST'])
@admin_required
def appform_upload(user_id=None):
    """
    Build the Application Form PDF and push it to the HSE document upload API
    as hse_document_type = application_form.

    ?staff_id=... (or JSON body staff_id) overrides the staff id sent.
    """
    uid = _resolve_user_id(user_id)
    if not uid:
        return jsonify({"success": False,
                        "error": "Provide a user id in the path, ?_id=, or JSON body"}), 400

    try:
        user = _find_user(uid)
        if not user:
            return jsonify({"success": False,
                            "error": f"No user found in 'users' for id/email: {uid}"}), 404

        ctx, err = _gather_appform_context(uid)
        if ctx is None:
            return jsonify({"success": False, "error": err}), 404

        body = request.get_json(silent=True) or {}
        staff_id = (_v(request.args.get('staff_id'))
                    or _v(body.get('staff_id'))
                    or _staff_id_for(user))
        if not staff_id:
            return jsonify({"success": False,
                            "error": "No staff id available for this user"}), 400

        filename  = _appform_filename(ctx)
        pdf_bytes = _build_appform_pdf(ctx)

        ok, result = upload_hse_document(pdf_bytes, filename, staff_id,
                                         HSE_APPLICATION_FORM)
        if not ok:
            return jsonify({
                "success":  False,
                "user_id":  str(user.get('_id')),
                "staff_id": staff_id,
                "filename": filename,
                "upload":   result,
            }), result.get('status_code') or 502

        return jsonify({
            "success":           True,
            "user_id":           str(user.get('_id')),
            "staff_id":          staff_id,
            "filename":          filename,
            "hse_document_type": HSE_APPLICATION_FORM,
            "source":            ('application_form'
                                   if ctx['meta']['appform_found'] else 'profile_data'),
            "upload":            result,
        }), 200
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500