"""
live_staffs_reg_record.py
══════════════════════════
Independent module — Xpress Health "Professional Registration Record" generator.

Lets an admin fill a form with a staff member's professional registration
details (registration number, division/category, registration date,
renewal date, etc.), then generates a downloadable PDF and DOCX in
Xpress Health's own branding, storing both in Google Cloud Storage.

This is an internal Xpress Health admin record — not a reproduction of
any external regulator's certificate. Full CRUD: create, read, update,
delete saved records.

Collection: live_staff_reg_records
  {
    _id, staff_id, staff_name, email,
    registration_number, holder_name,
    address_lines: [str, ...],
    division, date_of_initial_registration, renewed_until,
    signatory_name, signatory_title, issue_date,
    pdf_gcs_blob, docx_gcs_blob,
    created_at, updated_at
  }

Routes:
    GET    /live-staffs/reg-record/list                     — list all saved records (+ search)
    GET    /live-staffs/reg-record/get/<rec_id>              — get one record
    POST   /live-staffs/reg-record/save                      — create or update a record (CRUD: create/update)
    POST   /live-staffs/reg-record/delete/<rec_id>            — delete a record           (CRUD: delete)
    GET    /live-staffs/reg-record/download/<rec_id>/pdf      — download generated PDF
    GET    /live-staffs/reg-record/download/<rec_id>/docx     — download generated DOCX
    GET    /live-staffs/reg-record/prefill/<staff_id>         — prefill form fields from staff record
"""

from flask import request, jsonify, Response, render_template
from bson import ObjectId
from datetime import datetime
import os
import io

from database import db
from . import admin_bp
from admin.views import admin_required


# ── Page ──────────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/reg-record')
@admin_required
def live_staff_reg_record_page():
    """Render the Professional Registration Record admin page."""
    return render_template('admin/live_staffs_reg_record.html')


# ── Helpers ───────────────────────────────────────────────────────────

def _v(val):
    if val is None:
        return ''
    return str(val).strip()


def _staffs_col():
    return db.live_staffs


def _reg_record_col():
    return db.live_staff_reg_records


def _gcs_upload(blob_name, data_bytes, content_type='application/octet-stream'):
    from admin.live_staffs import _gcs_upload as _f
    return _f(blob_name, data_bytes, content_type)


def _gcs_download(blob_name):
    from admin.live_staffs import _gcs_download as _f
    return _f(blob_name)


def _serialize(doc):
    if isinstance(doc, list):
        return [_serialize(i) for i in doc]
    if isinstance(doc, dict):
        return {k: _serialize(v) for k, v in doc.items()}
    if isinstance(doc, ObjectId):
        return str(doc)
    if isinstance(doc, datetime):
        return doc.isoformat()
    return doc


# ── Prefill from staff record ─────────────────────────────────────────

@admin_bp.route('/live-staffs/reg-record/prefill/<staff_id>')
@admin_required
def live_staff_reg_record_prefill(staff_id):
    """Return suggested field values pulled from the staff's existing profile."""
    try:
        staff = _staffs_col().find_one({"_id": ObjectId(staff_id)})
        if not staff:
            return jsonify({"success": False, "error": "Staff not found"}), 404

        s1 = staff.get('section_1_personal_details') or {}
        s3 = staff.get('section_3_professional_registration') or {}

        full_name = _v(s1.get('full_name') or '')
        email     = _v(staff.get('email') or s1.get('email_address') or '')
        address   = _v(s1.get('address') or '')
        address_lines = [l.strip() for l in address.split(',') if l.strip()] if address else []

        reg_number = _v(s3.get('registration_number_pin') or staff.get('nmbi_number') or '')
        divisions  = s3.get('divisions_registered_in') or []
        division   = divisions[0] if divisions else ''
        reg_exp    = _v(s3.get('registration_expiry_date') or '')

        return jsonify({
            "success": True,
            "staff_id": staff_id,
            "staff_name": full_name,
            "email": email,
            "registration_number": reg_number,
            "holder_name": full_name,
            "address_lines": address_lines,
            "division": division,
            "renewed_until": reg_exp,
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── List / Read ────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/reg-record/list')
@admin_required
def live_staff_reg_record_list():
    """List all saved registration records, optionally filtered by search term."""
    search = (request.args.get('search') or '').strip()
    query  = {}
    if search:
        import re
        pattern = re.compile(re.escape(search), re.IGNORECASE)
        query = {"$or": [
            {"staff_name": pattern},
            {"email": pattern},
            {"registration_number": pattern},
        ]}
    recs = list(_reg_record_col().find(query).sort("updated_at", -1))
    return jsonify({"success": True, "records": _serialize(recs)})


@admin_bp.route('/live-staffs/reg-record/get/<rec_id>')
@admin_required
def live_staff_reg_record_get(rec_id):
    """Get a single saved registration record by id."""
    try:
        rec = _reg_record_col().find_one({"_id": ObjectId(rec_id)})
        if not rec:
            return jsonify({"success": False, "error": "Record not found"}), 404
        return jsonify({"success": True, "record": _serialize(rec)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Create / Update ──────────────────────────────────────────────────

@admin_bp.route('/live-staffs/reg-record/save', methods=['POST'])
@admin_required
def live_staff_reg_record_save():
    """
    Create a new registration record, or update an existing one if rec_id
    is provided. Always (re)generates the PDF and DOCX and uploads both
    to GCS, refreshing the stored blob paths.

    Body (JSON):
      {
        "rec_id": "..."              (omit/blank to create new),
        "staff_id": "...",
        "registration_number": "...",
        "holder_name": "...",
        "address_lines": ["...", "..."],
        "division": "...",
        "date_of_initial_registration": "DD/MM/YYYY",
        "renewed_until": "DD/MM/YYYY",
        "signatory_name": "...",
        "signatory_title": "...",
        "issue_date": "DD/MM/YYYY"
      }
    """
    data = request.get_json(silent=True) or {}

    rec_id    = (data.get('rec_id') or '').strip()
    staff_id  = (data.get('staff_id') or '').strip()

    fields = {
        "registration_number":          _v(data.get('registration_number')),
        "holder_name":                  _v(data.get('holder_name')),
        "address_lines":                [_v(l) for l in (data.get('address_lines') or []) if _v(l)],
        "division":                     _v(data.get('division')),
        "date_of_initial_registration": _v(data.get('date_of_initial_registration')),
        "renewed_until":                _v(data.get('renewed_until')),
        "signatory_name":               _v(data.get('signatory_name') or 'Betsy Daniel'),
        "signatory_title":              _v(data.get('signatory_title') or 'Head of Recruitment'),
        "issue_date":                   _v(data.get('issue_date')) or datetime.utcnow().strftime('%d/%m/%Y'),
    }

    if not fields["registration_number"] or not fields["holder_name"]:
        return jsonify({"success": False, "error": "registration_number and holder_name are required"}), 400

    staff_name = fields["holder_name"]
    email      = ''
    if staff_id:
        try:
            staff = _staffs_col().find_one({"_id": ObjectId(staff_id)})
            if staff:
                s1 = staff.get('section_1_personal_details') or {}
                staff_name = _v(s1.get('full_name') or staff_name)
                email      = _v(staff.get('email') or s1.get('email_address') or '')
        except Exception:
            pass

    try:
        col = _reg_record_col()
        now = datetime.utcnow()

        if rec_id:
            existing = col.find_one({"_id": ObjectId(rec_id)})
            if not existing:
                return jsonify({"success": False, "error": "Record not found"}), 404
            object_id = existing["_id"]
        else:
            object_id = None

        # ── Build the documents ────────────────────────────────────────
        docx_bytes = _build_reg_record_docx(fields, staff_name)
        pdf_bytes  = _build_reg_record_pdf(fields, staff_name)

        safe_name  = (staff_name or 'staff').replace(' ', '_').replace('/', '_')
        docx_blob  = f"reg_records/{safe_name}_{fields['registration_number']}.docx"
        pdf_blob   = f"reg_records/{safe_name}_{fields['registration_number']}.pdf"

        _gcs_upload(docx_blob, docx_bytes,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        _gcs_upload(pdf_blob, pdf_bytes, content_type='application/pdf')

        record = {
            "staff_id":     staff_id,
            "staff_name":   staff_name,
            "email":        email,
            **fields,
            "docx_gcs_blob": docx_blob,
            "pdf_gcs_blob":  pdf_blob,
            "updated_at":    now,
        }

        if object_id:
            col.update_one({"_id": object_id}, {"$set": record})
            saved_id = str(object_id)
        else:
            record["created_at"] = now
            saved_id = str(col.insert_one(record).inserted_id)

        return jsonify({
            "success": True,
            "rec_id": saved_id,
            "staff_name": staff_name,
            "docx_gcs_blob": docx_blob,
            "pdf_gcs_blob":  pdf_blob,
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Delete ────────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/reg-record/delete/<rec_id>', methods=['POST'])
@admin_required
def live_staff_reg_record_delete(rec_id):
    """Delete a saved registration record (does not delete GCS blobs)."""
    try:
        result = _reg_record_col().delete_one({"_id": ObjectId(rec_id)})
        if result.deleted_count == 0:
            return jsonify({"success": False, "error": "Record not found"}), 404
        return jsonify({"success": True, "message": "Deleted"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Download ──────────────────────────────────────────────────────────

@admin_bp.route('/live-staffs/reg-record/download/<rec_id>/<fmt>')
@admin_required
def live_staff_reg_record_download(rec_id, fmt):
    """Download the generated PDF or DOCX for a saved record. fmt = 'pdf' or 'docx'."""
    if fmt not in ('pdf', 'docx'):
        return jsonify({"success": False, "error": "fmt must be pdf or docx"}), 400
    try:
        rec = _reg_record_col().find_one({"_id": ObjectId(rec_id)})
        if not rec:
            return jsonify({"success": False, "error": "Record not found"}), 404

        blob_field = 'pdf_gcs_blob' if fmt == 'pdf' else 'docx_gcs_blob'
        blob = rec.get(blob_field)
        if not blob:
            return jsonify({"success": False, "error": f"No {fmt} stored for this record"}), 404

        file_bytes = _gcs_download(blob)
        safe_name  = (rec.get('staff_name') or 'staff').replace(' ', '_').replace('/', '_')
        reg_no     = rec.get('registration_number', '')
        filename   = f"Registration_Record_{safe_name}_{reg_no}.{fmt}"

        mimetype = 'application/pdf' if fmt == 'pdf' else \
                   'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        return Response(
            file_bytes,
            mimetype=mimetype,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── DOCX builder — Xpress Health branded ──────────────────────────────

def _build_reg_record_docx(fields, staff_name):
    """Build an Xpress Health 'Professional Registration Record' DOCX."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY  = RGBColor(0x1B, 0x3A, 0x6B)
    GREEN = RGBColor(0x2E, 0x9E, 0x44)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    d = Document()
    for sec in d.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    normal = d.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    def _run(para, text, bold=False, size=11, color=None, italic=False):
        r = para.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        r.font.color.rgb = color if color else BLACK
        return r

    def _set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def sp(pts=6):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = Pt(pts)

    # ── Logo (same pattern as other Xpress Health generated docs) ─────
    logo_added = False
    for logo_path in [
        'static/image/logo.png',
        'static/img/logo.png',
        os.path.join(os.path.dirname(__file__), '..', 'static', 'image', 'logo.png'),
    ]:
        if os.path.exists(logo_path):
            try:
                p_logo = d.add_paragraph()
                p_logo.paragraph_format.space_before = Pt(0)
                p_logo.paragraph_format.space_after  = Pt(4)
                p_logo.add_run().add_picture(logo_path, width=Inches(1.6))
                logo_added = True
                break
            except Exception:
                continue
    if not logo_added:
        p_logo = d.add_paragraph()
        r_logo = p_logo.add_run('XPRESS HEALTH')
        r_logo.bold = True; r_logo.font.size = Pt(14)
        r_logo.font.color.rgb = NAVY; r_logo.font.name = 'Calibri'

    sp(8)

    # ── Title ─────────────────────────────────────────────────────────
    p_title = d.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after  = Pt(10)
    _run(p_title, 'Professional Registration Record', bold=True, size=20, color=NAVY)

    def field_line(label, value):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, label + ': ', bold=True, size=11)
        _run(p, value, size=11)

    field_line('Registration Number', fields.get('registration_number', ''))
    field_line('Name', fields.get('holder_name', staff_name))

    address_lines = fields.get('address_lines') or []
    if address_lines:
        p_addr_label = d.add_paragraph()
        p_addr_label.paragraph_format.space_before = Pt(2)
        p_addr_label.paragraph_format.space_after  = Pt(2)
        _run(p_addr_label, 'Address: ', bold=True, size=11)
        _run(p_addr_label, address_lines[0], size=11)
        for line in address_lines[1:]:
            p_addr = d.add_paragraph()
            p_addr.paragraph_format.space_before = Pt(0)
            p_addr.paragraph_format.space_after  = Pt(0)
            p_addr.paragraph_format.left_indent  = Cm(2.2)
            _run(p_addr, line, size=11)

    sp(10)

    # ── Registration details table ─────────────────────────────────────
    tbl = d.add_table(rows=2, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers_row = ['Division / Category', 'Date of Initial Registration', 'Renewed Until']
    values_row  = [fields.get('division', ''),
                   fields.get('date_of_initial_registration', ''),
                   fields.get('renewed_until', '')]
    col_widths  = [Cm(5.0), Cm(6.0), Cm(5.0)]

    for ci, (hdr, w) in enumerate(zip(headers_row, col_widths)):
        cell = tbl.rows[0].cells[ci]
        cell.width = w
        _set_cell_bg(cell, '1B3A6B')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, hdr, bold=True, size=10, color=WHITE)

    for ci, (val, w) in enumerate(zip(values_row, col_widths)):
        cell = tbl.rows[1].cells[ci]
        cell.width = w
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, val, size=10)

    sp(20)

    # ── Signatory block ──────────────────────────────────────────────
    field_line('Authorised Signatory', fields.get('signatory_name', ''))
    field_line('Designation', fields.get('signatory_title', ''))
    field_line('Date', fields.get('issue_date', ''))

    sp(10)
    p_note = d.add_paragraph()
    p_note.paragraph_format.space_before = Pt(8)
    _run(
        p_note,
        'This record confirms the professional registration details held on file by Xpress Health. '
        'It is an internal administrative record and is not a substitute for the official registration '
        'certificate issued by the relevant professional body.',
        size=9, color=RGBColor(0x55, 0x55, 0x55), italic=True,
    )

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ── PDF builder — Xpress Health branded ───────────────────────────────

def _build_reg_record_pdf(fields, staff_name):
    """Build a matching PDF for the registration record using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    NAVY = HexColor('#1B3A6B')
    GRAY = HexColor('#555555')

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
        topMargin=2.0*cm, bottomMargin=2.0*cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TitleNavy', parent=styles['Title'],
        textColor=NAVY, fontSize=18, leading=22, spaceAfter=10, alignment=0,
    )
    label_style = ParagraphStyle(
        'LabelBold', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=11, spaceAfter=2,
    )
    value_style = ParagraphStyle(
        'ValueNormal', parent=styles['Normal'],
        fontSize=11, spaceAfter=8,
    )
    addr_style = ParagraphStyle(
        'AddrIndent', parent=styles['Normal'],
        fontSize=11, leftIndent=1*cm, spaceAfter=0,
    )
    note_style = ParagraphStyle(
        'Note', parent=styles['Normal'],
        textColor=GRAY, fontSize=8, fontName='Helvetica-Oblique', spaceBefore=8,
    )

    elements = []

    # Logo if available
    for logo_path in [
        'static/image/logo.png',
        'static/img/logo.png',
        os.path.join(os.path.dirname(__file__), '..', 'static', 'image', 'logo.png'),
    ]:
        if os.path.exists(logo_path):
            try:
                elements.append(Image(logo_path, width=4*cm, height=1.4*cm, kind='proportional'))
                elements.append(Spacer(1, 8))
                break
            except Exception:
                continue

    elements.append(Paragraph('Professional Registration Record', title_style))

    elements.append(Paragraph('Registration Number', label_style))
    elements.append(Paragraph(fields.get('registration_number', ''), value_style))

    elements.append(Paragraph('Name', label_style))
    elements.append(Paragraph(fields.get('holder_name', staff_name), value_style))

    address_lines = fields.get('address_lines') or []
    if address_lines:
        elements.append(Paragraph('Address', label_style))
        for line in address_lines:
            elements.append(Paragraph(line, addr_style))
        elements.append(Spacer(1, 8))

    table_data = [
        ['Division / Category', 'Date of Initial Registration', 'Renewed Until'],
        [fields.get('division', ''),
         fields.get('date_of_initial_registration', ''),
         fields.get('renewed_until', '')],
    ]
    tbl = Table(table_data, colWidths=[5.5*cm, 6.5*cm, 5*cm])
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), NAVY),
        ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
        ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0, 0), (-1, -1), 10),
        ('GRID',       (0, 0), (-1, -1), 0.5, colors.HexColor('#999999')),
        ('VALIGN',     (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING',  (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING',   (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING',(0, 0), (-1, -1), 6),
    ]))
    elements.append(tbl)
    elements.append(Spacer(1, 28))

    elements.append(Paragraph('Authorised Signatory', label_style))
    elements.append(Paragraph(fields.get('signatory_name', ''), value_style))
    elements.append(Paragraph('Designation', label_style))
    elements.append(Paragraph(fields.get('signatory_title', ''), value_style))
    elements.append(Paragraph('Date', label_style))
    elements.append(Paragraph(fields.get('issue_date', ''), value_style))

    elements.append(Paragraph(
        'This record confirms the professional registration details held on file by Xpress Health. '
        'It is an internal administrative record and is not a substitute for the official registration '
        'certificate issued by the relevant professional body.',
        note_style
    ))

    doc.build(elements)
    return buf.getvalue()
